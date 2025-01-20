import streamlit as st
import pandas as pd
import openai
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from markdown_pdf import MarkdownPdf, Section
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import re
import base64
import random
import streamlit.components.v1 as components
from bs4 import BeautifulSoup
from streamlit_modal import Modal 
import os
import json
import pymongo
import bcrypt
from bson.objectid import ObjectId
from dotenv import load_dotenv
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import mimetypes
import stripe
from pymongo.errors import ServerSelectionTimeoutError, ConfigurationError, DuplicateKeyError
import matplotlib.pyplot as plt
import plotly.express as px
import logging
from authlib.integrations.requests_client import OAuth2Session
import streamlit_authenticator as stauth

import jwt
import datetime
from jwt import InvalidTokenError
from math import ceil

from typing import Union, List

import tiktoken


# Charger les variables d'environnement
load_dotenv()


# Configuration des Logs
logging.basicConfig(filename='app.log', level=logging.INFO,
                    format='%(asctime)s:%(levelname)s:%(message)s')


# Configuration MongoDB  ToV6fdovgdT89T66  knkjoseph

MONGO_URI=st.secrets["MONGO_URI"]

@st.cache_resource(show_spinner=False)
def get_mongo_client():
    try:
        client = pymongo.MongoClient(MONGO_URI, serverSelectionTimeoutMS=5000)
        client.server_info()  # Force la connexion pour vérifier la disponibilité
        return client
    except ServerSelectionTimeoutError as err:
        logging.error(f"Échec de la connexion à MongoDB : {err}")
        # st.error(f"Échec de la connexion à MongoDB : {err}")
        return None
    except ConfigurationError as err:
        logging.error(f"Erreur de configuration MongoDB : {err}")
        # st.error(f"Erreur de configuration MongoDB : {err}")
        return None
    except Exception as err:
        logging.error(f"Erreur inattendue lors de la connexion à MongoDB : {err}")
        # st.error(f"Erreur inattendue lors de la connexion à MongoDB : {err}")
        return None

client = get_mongo_client()

if client is None:
    st.title("🚫 Connexion à internet a Échouée")
    st.error("Impossible de se connecter. Veuillez vérifier votre connexion Internet ou contacter l'administrateur.")
    if st.button("Réessayer"):
        st.rerun()
    st.stop()  # Arrête l'exécution pour éviter les erreurs suivantes


db = client['ishai']
users_collection = db['users']
consumption_collection = db['token_consumption']
payments_collection = db['payments']

# Créer un index unique sur l'email pour optimiser les recherches
# Créer des index pour optimiser les recherches
# Créer des index pour optimiser les recherches
try:
    users_collection.create_index("email", unique=True)
    payments_collection.create_index("user_id")
    payments_collection.create_index("timestamp")
except Exception as e:
    logging.error(f"Erreur lors de la création des index : {e}")
    st.title("🚫 Connexion à internet a Échouée")
    st.error("Impossible de se connecter. Veuillez vérifier votre connexion Internet ou contacter l'administrateur.")
    if st.button("Réessayer"):
        st.rerun()
    st.stop()    
    
SMTP_SERVER = st.secrets["SMTP_SERVER"]
SMTP_PORT = st.secrets["SMTP_PORT"]
EMAIL_ADDRESS = st.secrets["EMAIL_ADDRESS"]
EMAIL_PASSWORD = st.secrets["EMAIL_PASSWORD"]

# Configuration des paramètres de pièces jointes
MAX_ATTACHMENT_SIZE = 10 * 1024 * 1024  # 10 Mo
ALLOWED_EXTENSIONS = ['jpg', 'jpeg', 'png']

# Configuration Stripe
stripe.api_key = st.secrets["STRIPE_API_KEY"]
STRIPE_PUBLIC_KEY = st.secrets["STRIPE_PUBLIC_KEY"]

# Informations Administrateur
ADMIN_PASSWORD = st.secrets["ADMIN_PASSWORD"]
ADMIN_EMAIL = st.secrets["ADMIN_EMAIL"]

# Configuration Google OAuth
GOOGLE_CLIENT_ID = st.secrets["GOOGLE_CLIENT_ID"]
GOOGLE_CLIENT_SECRET = st.secrets["GOOGLE_CLIENT_SECRET"]
GOOGLE_AUTHORIZATION_ENDPOINT = st.secrets["GOOGLE_AUTHORIZATION_ENDPOINT"]
GOOGLE_TOKEN_ENDPOINT = st.secrets["GOOGLE_TOKEN_ENDPOINT"]
GOOGLE_USERINFO_ENDPOINT = st.secrets["GOOGLE_USERINFO_ENDPOINT"]

# Configuration JWT
JWT_SECRET_KEY = st.secrets["JWT_SECRET_KEY"]
JWT_ALGORITHM = st.secrets["JWT_ALGORITHM"]
JWT_EXP_DELTA_SECONDS = st.secrets["JWT_EXP_DELTA_SECONDS"]

# Gestion des Tokens
DEFAULT_TOKENS_PURCHASED = st.secrets["DEFAULT_TOKENS_PURCHASED"]
TOKEN_NUMBER = st.secrets["TOKEN_NUMBER"]


# Configuration de l'API OpenAI
api_key = st.secrets["API_KEY"]
openai.api_key = api_key




   
    
    
def load_credentials_from_mongo():
    """
    Récupère les utilisateurs depuis MongoDB et les formate pour streamlit-authenticator.
    """
    credentials_dict = {"usernames": {}}

    try:
        # Extraire tous les utilisateurs depuis la base
        all_users = list(users_collection.find())
        for user in all_users:
            # Vérifie les champs nécessaires
            email = user.get("email")
            nom = user.get("nom", "Utilisateur")
            hashed_password = user.get("mot_de_passe", "")
            role = user.get("role", "user")

            if not email or not hashed_password:
                # Ignore les utilisateurs incomplets
                continue

            # Ajoute l'utilisateur au dictionnaire de credentials
            credentials_dict["usernames"][email] = {
                "name": nom,
                "password": hashed_password,
                "email": email,
                "role": [role],  # Enveloppe dans une liste pour la compatibilité
            }

    except Exception as e:
        #st.error(f"Erreur lors du chargement des credentials depuis MongoDB : {e}")
        st.title("🚫 Connexion à internet a Échouée")
        st.error("Impossible de se connecter. Veuillez vérifier votre connexion Internet ou contacter l'administrateur.")
        if st.button("Réessayer"):
            st.rerun()
        st.stop()

    return credentials_dict





# ---------------------- FONCTIONS TOKEN COUNTER ----------------------
@st.cache_resource(show_spinner=False)
def get_encoding_for_model(model_name: str):
    """
    Retourne l'encodage approprié pour un modèle donné.
    Si le modèle n'est pas reconnu, utilise 'cl100k_base' par défaut.
    """
    try:
        encoding = tiktoken.encoding_for_model(model_name)
    except KeyError:
        st.warning(f"Modèle '{model_name}' non reconnu. Utilisation de l'encodage 'cl100k_base' par défaut.")
        encoding = tiktoken.get_encoding("cl100k_base")
    return encoding

def count_tokens(texts: Union[str, List[str]], model_name: str) -> Union[int, List[int]]:
    """
    Compte le nombre de tokens dans une chaîne de texte ou une liste de chaînes de texte.

    Args:
        texts (str ou List[str]): Le(s) texte(s) à compter.
        model_name (str): Le nom du modèle OpenAI pour déterminer l'encodage.

    Returns:
        int ou List[int]: Le nombre de tokens pour chaque texte.
    """
    encoding = get_encoding_for_model(model_name)
    
    if isinstance(texts, list):
        return [len(encoding.encode(text)) for text in texts]
    else:
        return len(encoding.encode(texts))




# ----------------------------------------------------------------------------
# Business Model 
# ----------------------------------------------------------------------------
def update_nom_produit(index_produit):
    # Mettre à jour le nom du produit dans st.session_state
    st.session_state["produits_data"][index_produit]["nom_produit"] = st.session_state[f"nom_produit_{index_produit}"]
    # Forcer la réexécution de l'application pour mettre à jour le selectbox
    st.rerun()


def collect_persona_pme(index_produit):
    """
    Collecte et/ou met à jour les informations de la Persona
    pour le produit index_produit, en gérant 3 types :
    B2C, B2B et Ménage.
    """

    st.subheader("Persona")

    # Récupération du dictionnaire produit et de son sous-dictionnaire "persona"
    produit = st.session_state["produits_data"][index_produit]
    persona_data = produit.get("persona", {})
    
        # Formulaire dédié à la mise à jour du nom du produit
    with st.form(key=f"form_{index_produit}"):
        col1, col2 = st.columns([4, 1])

        with col1:
            nom_produit = st.text_input(
                f"Nom du Produit (Index {index_produit+1}):",
                value=produit.get("nom_produit", f"Produit_{index_produit+1}"),
                key=f"nom_produit_{index_produit}"
            )

        with col2:
            st.write("")
            submit_button = st.form_submit_button("✅ Valider")


        if submit_button:
            st.session_state["produits_data"][index_produit]["nom_produit"] = nom_produit
            st.success(f"Nom du produit mis à jour : {nom_produit}")
            st.rerun()
            # Optionnel : Maintenir la sélection actuelle
            # st.session_state["selected_idx_produit"] = index_produit

    # Sélecteur du Type de Persona (hors du formulaire)
    type_persona = st.selectbox(
        "Sélectionnez le type de Persona",
        ["B2C", "B2B", "Ménage"],
        index=["B2C", "B2B", "Ménage"].index(persona_data.get("type", "B2C")),
        key=f"type_persona_selectbox_{index_produit}"
    )
    #st.write(f"**Type de Persona sélectionné :** {type_persona}")

    # Mise à jour du champ "type" dans le dictionnaire
    persona_data["type"] = type_persona


    # ----- CAS 1 : PERSONA B2C -----
    if type_persona == "B2C":
        st.subheader("Persona - B2C")

        # Données Démographiques B2C
        st.subheader("Données Démographiques")
        b2c_age = st.number_input(
            "Âge (B2C)",
            min_value=18,
            max_value=100,
            value=persona_data.get("âge", 18),
            key=f"b2c_age_{index_produit}",
            help=(
        "Entrez l'âge du persona en années. "
        "Assurez-vous que l'âge correspond au profil cible de votre produit ou service. "
        "Par exemple, si vous ciblez les jeunes adultes, un âge entre 18 et 35 pourrait être approprié."
    )
        )
        b2c_sexe = st.selectbox(
            "Sexe (B2C)",
            ["", "Homme", "Femme"],
            index=["","Homme","Femme"].index(persona_data.get("sexe","")) if persona_data.get("sexe","") in ["","Homme","Femme"] else 0,
            key=f"b2c_sexe_{index_produit}",
            help=(
        "Sélectionnez le sexe du persona. "
        "Cela peut influencer les préférences et les comportements d'achat. "
        "Si le sexe du persona n'est pas pertinent pour votre produit, choisissez vide."
    )
        )
        b2c_localisation = st.text_input(
            "Localisation Géographique (ex: Goma, Rdcongo)",
            placeholder="Goma, Rdcongo",
            value=persona_data.get("localisation",""),
            key=f"b2c_localisation_{index_produit}",
            help=(
        "Indiquez la localisation géographique du persona. "
        "Précisez la ville et le pays, par exemple 'Goma, Rdcongo'. "
        "Cette information aide à comprendre les spécificités culturelles, économiques et logistiques."
    )
        )
        b2c_education = st.text_input(
            "Niveau d'Éducation (B2C)",
            placeholder="Baccalauréat, Licence, Master, Doctorat, Autre",
            value=persona_data.get("éducation",""),
            key=f"b2c_education_{index_produit}",
            help=(
        "Sélectionnez le niveau d'éducation atteint par le persona. "
        "Cela peut influencer le type de produits ou services auxquels il est exposé, "
        "ses attentes en matière de qualité et son pouvoir d'achat."
    )
            
        )
        b2c_profession = st.text_input(
            "Profession (B2C)",
            placeholder="Ex: Ingénieur, Designer",
            value=persona_data.get("profession",""),
            key=f"b2c_profession_{index_produit}",
            help=(
        "Entrez la profession actuelle du persona. "
        "Connaître la profession permet de mieux comprendre le quotidien, les besoins spécifiques "
        "et le pouvoir d'achat du persona. Par exemple, un ingénieur pourrait avoir des besoins différents "
        "d'un designer."
    )
            
        )
        b2c_revenu = st.number_input(
            "Revenu Mensuel Moyen ($)",
            min_value=0.0,
            value=float(persona_data.get("revenu_moyen", 0.0)),
            key=f"b2c_revenu_{index_produit}",
            help=(
        "Entrez le revenu mensuel moyen du persona en dollars. "
        "Cette information est cruciale pour évaluer le pouvoir d'achat et la sensibilité au prix. "
        "Assurez-vous que le revenu est réaliste par rapport au niveau d'éducation et à la profession."
    )
        )

        st.subheader("Paramètres Comportementaux (B2C)")
        b2c_sensibilite_prix = st.selectbox(
            "Sensibilité au Prix (B2C)",
            ["","Très Faible", "Faible", "Modérée", "Élevée", "Très Élevée"],
            index=["","Très Faible", "Faible", "Modérée", "Élevée", "Très Élevée"].index(persona_data.get("sensibilite_prix","")) if persona_data.get("sensibilite_prix","") in ["","Très Faible", "Faible", "Modérée", "Élevée", "Très Élevée"] else 0,
            key=f"b2c_sensibilite_prix_{index_produit}",
            help=(
        "Indiquez la sensibilité au prix du persona. "
        "Un persona avec une 'Très Élevée' sensibilité au prix privilégiera les produits abordables, "
        "tandis qu'un persona avec une 'Très Faible' sensibilité pourrait privilégier la qualité ou la marque."
    )
            
        )
        b2c_frequence_achat = st.text_input(
            "Fréquence d'Achat (B2C)",
            placeholder="Rarement, Mensuellement, Hebdomadairement",
            value=persona_data.get("frequence_achat",""),
            key=f"b2c_frequence_achat_{index_produit}",
            help=(
            "Sélectionnez la fréquence à laquelle le persona achète des produits ou services similaires. "
            "Cela permet d'estimer la régularité des revenus et d'ajuster les stratégies marketing en conséquence."
        )
            
        )
        b2c_volume_achat = st.text_input(
            "Volume d'Achat (B2C)",
            value=persona_data.get("volume_achat",""),
            key=f"b2c_volume_achat_{index_produit}",
            help=(
        "Entrez le volume d'achat habituel du persona. "
        "Cela peut représenter le nombre d'articles achetés par transaction ou la quantité totale dépensée. "
        "Cette information aide à prévoir les besoins en stock et les campagnes de vente."
    )
            
        )
        b2c_perception_qualite = st.text_area(
            "Perception de la Qualité (B2C)",
            placeholder="Décrivez la perception de la qualité...",
            value=persona_data.get("perception_qualite",""),
            key=f"b2c_perception_qualite_{index_produit}",
            help=(
        "Décrivez comment le persona perçoit la qualité de votre produit ou service. "
        "Par exemple, le persona pourrait valoriser la durabilité, le design, ou la fonctionnalité. "
        "Comprendre cette perception aide à aligner les caractéristiques du produit avec les attentes des clients."
    )
            
        )
        b2c_utilisation_tech = st.text_area(
            "Quelles technologies le client B2C utilise",
            placeholder="Décrivez les technologies (smartphone, internet...)",
            value=persona_data.get("utilisation_tech",""),
            key=f"b2c_utilisation_tech_{index_produit}",
            help=(
        "Décrivez les technologies auxquelles le persona a accès. "
        "Cela inclut les appareils (Smartphone, ordinateur), la connectivité Internet, les plateformes sociales, etc. "
        "Ces informations sont essentielles pour déterminer les canaux de communication et de distribution appropriés."
    )
            
        )
        b2c_acces_transport = st.text_area(
            "Accessibilité (Transport) (B2C)",
            value=persona_data.get("acces_transport",""),
            key=f"b2c_acces_transport_{index_produit}",
            help=(
        "Décrivez les moyens de transport que le persona utilise régulièrement. "
        "Cela peut influencer la logistique de livraison, l'emplacement des points de vente, "
        "ou la manière dont les services sont offerts (en ligne vs. physique)."
    )
            
        )
        b2c_temps_disponible = st.text_area(
            "Temps Disponible (B2C)",
            value=persona_data.get("temps_disponible",""),
            key=f"b2c_temps_disponible_{index_produit}",
            help=(
        "Décrivez le temps que le persona peut consacrer à votre produit ou service. "
        "Par exemple, s'il s'agit d'une application mobile, le temps disponible pour l'utiliser quotidiennement est crucial. "
        "Cette information aide à adapter la complexité et la convivialité du produit."
    )
            
        )
        b2c_besoins_specifiques = st.text_area(
            "Besoins Spécifiques (B2C)",
            value=persona_data.get("besoins_specifiques",""),
            key=f"b2c_besoins_specifiques_{index_produit}",
            help=(
            "Décrivez les besoins spécifiques du persona que votre produit ou service vise à satisfaire. "
            "Cela peut inclure des besoins fonctionnels, émotionnels, ou sociaux. "
            "Comprendre ces besoins permet de mieux cibler les fonctionnalités et les avantages proposés."
        )
            
        )
        b2c_motivations = st.text_area(
            "Motivations (B2C)",
            value=persona_data.get("motivations",""),
            key=f"b2c_motivations_{index_produit}",
            help=(
        "Décrivez les motivations qui poussent le persona à utiliser votre produit ou service. "
        "Cela peut inclure le désir de gain de temps, l'amélioration de la qualité de vie, "
        "l'atteinte d'objectifs personnels ou professionnels, etc. "
        "Identifier ces motivations aide à créer des messages marketing efficaces."
    )
            
        )

        st.subheader("Capacité d’Adoption de l’Innovation (B2C)")
        b2c_familiarite_tech = st.text_area(
            "Familiarité Technologique (B2C)",
            value=persona_data.get("familiarite_tech",""),
            key=f"b2c_familiarite_tech_{index_produit}",
            help=(
        "Décrivez le niveau de familiarité technologique du persona. "
        "Par exemple, utilise-t-il régulièrement des smartphones, des applications spécifiques, "
        "ou des plateformes en ligne? Cette information est essentielle pour développer des produits intuitifs "
        "et compatibles avec les habitudes technologiques du persona."
    )
            
        )
        b2c_ouverture_changement = st.text_input(
            "Ouverture au Changement (B2C)",
            value=persona_data.get("ouverture_changement",""),
            key=f"b2c_ouverture_changement_{index_produit}",
            help=(
        "Indiquez le niveau d'ouverture au changement du persona. "
        "Un persona avec une 'Élevée' ouverture sera plus enclin à essayer de nouveaux produits ou services, "
        "tandis qu'un persona avec une 'Faible' ouverture pourrait préférer les solutions éprouvées et familières."
    )
            
        )
        b2c_barrieres = st.text_area(
            "Barrières Psychologiques/Culturelles (B2C)",
            value=persona_data.get("barrieres",""),
            key=f"b2c_barrieres_{index_produit}",
            help=(
        "Décrivez les barrières psychologiques ou culturelles qui pourraient empêcher le persona d'adopter "
        "votre produit ou service. Cela peut inclure des croyances, des habitudes, des normes sociales, "
        "ou des craintes spécifiques. Comprendre ces barrières permet de les adresser dans votre stratégie marketing."
    )
            
        )

        # Mise à jour du dictionnaire "persona_data"
        persona_data.update({
            "type": "B2C",
            "âge": b2c_age,
            "sexe": b2c_sexe,
            "localisation": b2c_localisation,
            "éducation": b2c_education,
            "profession": b2c_profession,
            "revenu_moyen": b2c_revenu,
            "sensibilite_prix": b2c_sensibilite_prix,
            "frequence_achat": b2c_frequence_achat,
            "volume_achat": b2c_volume_achat,
            "perception_qualite": b2c_perception_qualite,
            "utilisation_tech": b2c_utilisation_tech,
            "acces_transport": b2c_acces_transport,
            "temps_disponible": b2c_temps_disponible,
            "besoins_specifiques": b2c_besoins_specifiques,
            "motivations": b2c_motivations,
            "familiarite_tech": b2c_familiarite_tech,
            "ouverture_changement": b2c_ouverture_changement,
            "barrieres": b2c_barrieres
        })

    # ----- CAS 2 : PERSONA B2B -----
    elif type_persona == "B2B":
        st.subheader("Persona - B2B")

        # Données Démographiques B2B
        st.subheader("Données Démographiques")
        b2b_taille_entreprise = st.selectbox(
            "Taille de l'Entreprise (B2B)",
            ["","PME", "Grande Entreprise", "Multinationale"],
            index=["","PME", "Grande Entreprise", "Multinationale"].index(persona_data.get("taille_entreprise","")) if persona_data.get("taille_entreprise","") in ["","PME", "Grande Entreprise", "Multinationale"] else 0,
            key=f"b2b_taille_entreprise_{index_produit}",
            help=(
        "Sélectionnez la taille de l'entreprise du persona. "
        "La taille de l'entreprise influence les besoins, les ressources disponibles et les processus décisionnels. "
        "Par exemple, une grande entreprise peut avoir des besoins plus complexes et des cycles d'achat plus longs que "
        "une petite entreprise."
    )
            
        )
        b2b_secteur_activite = st.text_input(
            "Secteur d'Activité (B2B)",
            placeholder="Ex: Technologie, Santé",
            value=persona_data.get("secteur_activite",""),
            key=f"b2b_secteur_activite_{index_produit}",
            help=(
        "Sélectionnez le secteur d'activité de l'entreprise du persona. "
        "Le secteur d'activité détermine les spécificités, les réglementations et les tendances auxquelles l'entreprise est soumise. "
        "Par exemple, une entreprise du secteur de la santé aura des besoins et des contraintes différents de celles du secteur technologique."
    )
            
        )
        b2b_localisation_ent = st.text_input(
            "Localisation Entreprise (B2B)",
            value=persona_data.get("localisation_entreprise",""),
            key=f"b2b_localisation_entreprise_{index_produit}",
            help=(
        "Indiquez la localisation géographique de l'entreprise du persona. "
        "Précisez la ville et le pays, par exemple 'Goma, Rdcongo'. "
        "Cette information est essentielle pour comprendre les spécificités économiques, culturelles et logistiques locales."
        )
            
        )
        b2b_chiffre_affaires = st.number_input(
            "Chiffre d'Affaires ($) (B2B)",
            min_value=0.0,
            value=float(persona_data.get("chiffre_affaires",0.0)),
            key=f"b2b_chiffre_affaires_{index_produit}",
            help=(
        "Entrez le chiffre d'affaires annuel de l'entreprise en dollars. "
        "Le chiffre d'affaires donne une idée de la taille financière de l'entreprise et de sa capacité d'investissement. "
        "Assurez-vous que le chiffre d'affaires est réaliste par rapport à la taille et au secteur d'activité de l'entreprise."
        )
            
        )
        b2b_nombre_employes = st.number_input(
            "Nombre d'Employés (B2B)",
            min_value=1,
            step=1,
            value=int(persona_data.get("nombre_employes",1)),
            key=f"b2b_nombre_employes_{index_produit}",
             help=(
        "Entrez le nombre total d'employés de l'entreprise. "
        "Le nombre d'employés peut influencer la structure organisationnelle, les besoins en ressources humaines et les capacités opérationnelles."
        )
            
        )

        st.subheader("Décideurs et Influenceurs (B2B)")
        b2b_role_decideur = st.text_input(
            "Rôle du Décideur (B2B)",
            value=persona_data.get("role_decideur",""),
            key=f"b2b_role_decideur_{index_produit}",
            help=(
        "Entrez le rôle du principal décideur au sein de l'entreprise. "
        "Connaître le rôle du décideur aide à cibler les communications et les arguments de vente de manière plus efficace. "
        "Par exemple, un Directeur des Achats aura des priorités différentes de celles d'un Directeur Technique."
        )
        )
        
        b2b_influenceur = st.text_input(
            "Influenceurs Internes (B2B)",
            value=persona_data.get("influenceur",""),
            key=f"b2b_influenceur_{index_produit}",
            help=(
        "Listez les principaux influenceurs internes qui peuvent influencer la décision d'achat. "
        "Cela peut inclure des équipes spécifiques, des départements ou des individus clés. "
        "Comprendre les influenceurs internes permet de mieux adresser les besoins et les préoccupations de l'ensemble des parties prenantes."
        )
            
        )

        # Paramètres Comportementaux B2B
        st.subheader("Paramètres Comportementaux (B2B)")
        b2b_sensibilite_prix = st.selectbox(
            "Sensibilité au Prix (B2B)",
            ["","Faible","Moyenne","Élevée"],
            index=["","Faible","Moyenne","Élevée"].index(persona_data.get("sensibilite_prix","")) if persona_data.get("sensibilite_prix","") in ["","Faible","Moyenne","Élevée"] else 0,
            key=f"b2b_sensibilite_prix_{index_produit}",
            help=(
        "Indiquez la sensibilité au prix de l'entreprise persona. "
        "Une sensibilité 'Très Élevée' signifie que l'entreprise accorde une grande importance au coût, "
        "tandis qu'une sensibilité 'Très Faible' indique une plus grande priorité sur la qualité ou les fonctionnalités, même à un coût plus élevé."
        )
            
        )
        b2b_cycle_achat = st.text_input(
            "Cycle d'Achat (B2B) - Long, Moyen, Court",
            value=persona_data.get("cycle_achat",""),
            key=f"b2b_cycle_achat_{index_produit}",
            help=(
        "Sélectionnez la durée typique du cycle d'achat de l'entreprise. "
        "Un cycle d'achat 'Long' implique des processus décisionnels plus complexes et des délais plus étendus, "
        "tandis qu'un cycle 'Court' indique des décisions plus rapides et moins de formalités."
        )
            
        )
        b2b_volume_achat = st.text_input(
            "Volume d'Achat (B2B) - Faible, Moyen, Élevé",
            value=persona_data.get("volume_achat",""),
            key=f"b2b_volume_achat_{index_produit}",
            help=(
        "Entrez le volume d'achat habituel de l'entreprise. "
        "Cela peut représenter le nombre de licences logicielles achetées, la quantité de matériel commandée, "
        "ou toute autre métrique pertinente selon votre produit ou service. "
        "Cette information aide à prévoir les besoins en stock et à adapter les offres commerciales."
        )
            
        )
        b2b_perception_qualite = st.text_area(
            "Perception de la Qualité (B2B)",
            value=persona_data.get("perception_qualite",""),
            key=f"b2b_perception_qualite_{index_produit}",
            help=(
        "Décrivez comment l'entreprise persona perçoit la qualité de votre produit ou service. "
        "Par exemple, l'accent peut être mis sur la durabilité, la fiabilité, le support client, ou l'innovation technologique. "
        "Comprendre cette perception aide à aligner votre offre avec les attentes du client."
        )
        )
        b2b_besoins_specifiques = st.text_area(
            "Besoins Spécifiques (B2B)",
            value=persona_data.get("besoins_specifiques",""),
            key=f"b2b_besoins_specifiques_{index_produit}",
            help=(
        "Décrivez les besoins spécifiques de l'entreprise que votre produit ou service vise à satisfaire. "
        "Cela peut inclure des besoins fonctionnels, opérationnels, stratégiques ou réglementaires. "
        "Identifier ces besoins permet de mieux cibler les fonctionnalités et les avantages proposés."
        )
            
        )
        b2b_motivations = st.text_area(
            "Motivations (B2B)",
            value=persona_data.get("motivations",""),
            key=f"b2b_motivations_{index_produit}",
            help=(
        "Décrivez les motivations qui poussent l'entreprise persona à utiliser votre produit ou service. "
        "Cela peut inclure l'amélioration de l'efficacité opérationnelle, la réduction des coûts, l'innovation, "
        "l'amélioration de la satisfaction client, etc. Identifier ces motivations aide à créer des messages marketing efficaces."
        )
            
        )

        st.subheader("Capacité d’Adoption de l’Innovation (B2B)")
        b2b_familiarite_tech = st.text_area(
            "Familiarité Technologique (B2B)",
            value=persona_data.get("familiarite_tech",""),
            key=f"b2b_familiarite_tech_{index_produit}",
            help=(
        "Décrivez le niveau de familiarité technologique de l'entreprise persona. "
        "Par exemple, utilise-t-elle des systèmes ERP avancés, des plateformes de cloud computing, des outils de collaboration spécifiques, etc. "
        "Cette information est essentielle pour développer des produits compatibles et adaptés aux habitudes technologiques de l'entreprise."
        )
            
        )
        b2b_ouverture_changement = st.text_input(
            "Ouverture au Changement (B2B) - Faible, Moyenne, Élevée",
            value=persona_data.get("ouverture_changement",""),
            key=f"b2b_ouverture_changement_{index_produit}",
            help=(
        "Indiquez le niveau d'ouverture au changement de l'entreprise persona. "
        "Une ouverture 'Élevée' signifie que l'entreprise est plus encline à essayer de nouveaux produits ou services, "
        "tandis qu'une ouverture 'Faible' peut indiquer une préférence pour les solutions traditionnelles et éprouvées."
        )
            
        )
        
        b2b_barrieres = st.text_area(
            "Barrières Psychologiques/Culturelles (B2B)",
            value=persona_data.get("barrieres",""),
            key=f"b2b_barrieres_{index_produit}",
            help=(
        "Décrivez les barrières psychologiques ou culturelles qui pourraient empêcher l'entreprise persona d'adopter "
        "votre produit ou service. Cela peut inclure des croyances organisationnelles, des habitudes établies, "
        "des normes industrielles ou des craintes spécifiques. Comprendre ces barrières permet de les adresser dans votre stratégie marketing."
        )
            
        )

        # Mise à jour du dictionnaire persona_data
        persona_data.update({
            "type": "B2B",
            "taille_entreprise": b2b_taille_entreprise,
            "secteur_activite": b2b_secteur_activite,
            "localisation_entreprise": b2b_localisation_ent,
            "chiffre_affaires": b2b_chiffre_affaires,
            "nombre_employes": b2b_nombre_employes,
            "role_decideur": b2b_role_decideur,
            "influenceur": b2b_influenceur,
            "sensibilite_prix": b2b_sensibilite_prix,
            "cycle_achat": b2b_cycle_achat,
            "volume_achat": b2b_volume_achat,
            "perception_qualite": b2b_perception_qualite,
            "besoins_specifiques": b2b_besoins_specifiques,
            "motivations": b2b_motivations,
            "familiarite_tech": b2b_familiarite_tech,
            "ouverture_changement": b2b_ouverture_changement,
            "barrieres": b2b_barrieres
        })

    # ----- CAS 3 : PERSONA MÉNAGE -----
    elif type_persona == "Ménage":
        st.subheader("Persona - Ménage")

        if "taille_menage" not in persona_data:
            # On initialise certains champs si absents
            persona_data["taille_menage"] = 1

        # Données Démographiques Ménage
        st.subheader("Données Démographiques")
        menage_taille = st.number_input(
            "Nombre de Personnes dans le Ménage",
            min_value=1,
            value=int(persona_data.get("taille_menage",1)),
            key=f"menage_taille_{index_produit}",
            help=(
        "Entrez le nombre total de personnes vivant dans le ménage. "
        "Cela inclut tous les membres de la famille ou les colocataires. "
        "Le nombre de personnes peut influencer les besoins en produits, les habitudes de consommation et le budget familial."
        )
            
        )
        menage_revenu = st.number_input(
            "Revenu Mensuel du Ménage ($)",
            min_value=0.0,
            value=float(persona_data.get("revenu_menage",0.0)),
            key=f"menage_revenu_{index_produit}",
            help=(
        "Entrez le revenu mensuel total du ménage en dollars. "
        "Cette information est cruciale pour évaluer le pouvoir d'achat, la sensibilité au prix et les priorités financières du ménage. "
        "Assurez-vous que le revenu est réaliste par rapport à la localisation géographique et au nombre de personnes dans le ménage."
        )
            
        )
        menage_localisation = st.text_input(
            "Localisation Géographique (ex: Goma, Rdcongo)",
            value=persona_data.get("localisation_menage",""),
            key=f"menage_localisation_{index_produit}",
            help=(
        "Indiquez la localisation géographique du ménage. "
        "Précisez la ville et le pays, par exemple 'Goma, Rdcongo'. "
        "Cette information aide à comprendre les spécificités culturelles, économiques et logistiques qui peuvent influencer les habitudes de consommation."
        )
            
        )
        menage_type_logement = st.text_input(
            "Type de Logement",
            value=persona_data.get("type_logement",""),
            key=f"menage_type_logement_{index_produit}",
            help=(
        "Sélectionnez le type de logement du ménage. "
        "Le type de logement peut influencer les besoins en mobilier, appareils électroménagers, espace de vie et habitudes de consommation. "
        "Par exemple, un appartement peut nécessiter des solutions de rangement optimisées, tandis qu'une maison peut avoir des besoins différents."
        )
            
        )

        st.subheader("Paramètres Comportementaux (Ménage)")
        menage_sensibilite_prix = st.selectbox(
            "Sensibilité au Prix (Ménage)",
            ["","Faible", "Moyenne", "Élevée"],
            index=["","Faible","Moyenne","Élevée"].index(persona_data.get("sensibilite_prix","")) if persona_data.get("sensibilite_prix","") in ["","Faible","Moyenne","Élevée"] else 0,
            key=f"menage_sensibilite_prix_{index_produit}",
            help=(
        "Indiquez la sensibilité au prix du ménage. "
        "Un ménage avec une 'Très Élevée' sensibilité au prix privilégiera les produits abordables, tandis qu'un ménage avec une 'Très Faible' sensibilité pourrait privilégier la qualité ou les marques, même à un coût plus élevé."
        )
            
        )
        menage_frequence_achat = st.text_input(
            "Fréquence d'Achat (Ménage) - Rarement, Mensuellement, Hebdomadairement",
            value=persona_data.get("frequence_achat",""),
            key=f"menage_frequence_achat_{index_produit}",
            help=(
        "Sélectionnez la fréquence à laquelle le ménage achète des produits ou services similaires. "
        "Cela permet d'estimer la régularité des dépenses et d'ajuster les stratégies marketing en conséquence."
        )
            
        )
        
        menage_volume_achat = st.text_input(
            "Volume d'Achat (Ménage) - Faible, Moyen, Élevé",
            value=persona_data.get("volume_achat",""),
            key=f"menage_volume_achat_{index_produit}",
            help=(
        "Entrez le volume d'achat habituel du ménage. "
        "Cela peut représenter le nombre d'articles achetés par transaction ou la quantité totale dépensée. "
        "Cette information aide à prévoir les besoins en stock et à adapter les campagnes de vente."
        )
            
        )
        menage_perception_qualite = st.text_area(
            "Perception de la Qualité (Ménage)",
            value=persona_data.get("perception_qualite",""),
            key=f"menage_perception_qualite_{index_produit}",
            help=(
        "Décrivez comment le ménage perçoit la qualité des produits ou services qu'il achète. "
        "Par exemple, le ménage pourrait valoriser la durabilité, le design, ou la fonctionnalité. "
        "Comprendre cette perception aide à aligner les caractéristiques du produit avec les attentes des clients."
        )
            
        )
        menage_utilisation_tech = st.text_area(
            "Utilisation Technologique (Ménage)",
            value=persona_data.get("utilisation_tech",""),
            key=f"menage_utilisation_tech_{index_produit}",
            help=(
        "Décrivez l'utilisation technologique au sein du ménage. "
        "Cela inclut les appareils utilisés (Smartphones, tablettes, ordinateurs, etc.), les plateformes en ligne préférées, et le niveau de compétence technologique. "
        "Ces informations sont essentielles pour déterminer les canaux de communication et les types de produits technologiques adaptés."
        )
            
        )
        menage_acces_transport = st.text_area(
            "Accessibilité (Transport) (Ménage)",
            value=persona_data.get("acces_transport",""),
            key=f"menage_acces_transport_{index_produit}",
            help=(
        "Décrivez les moyens de transport que le ménage utilise régulièrement. "
        "Cela peut influencer la logistique de livraison, l'emplacement des points de vente, ou la manière dont les services sont offerts (en ligne vs. physique). "
        "Par exemple, un ménage sans accès facile à une voiture peut préférer les livraisons à domicile."
        )
            
        )
        menage_temps_disponible = st.text_area(
            "Temps Disponible (Ménage)",
            value=persona_data.get("temps_disponible",""),
            key=f"menage_temps_disponible_{index_produit}",
            help=(
        "Décrivez le temps que le ménage peut consacrer à l'utilisation de votre produit ou service. "
        "Par exemple, s'il s'agit d'une application mobile, le temps disponible pour l'utiliser quotidiennement est crucial. "
        "Cette information aide à adapter la complexité et la convivialité du produit."
        )
            
        )
        menage_besoins_specifiques = st.text_area(
            "Besoins Spécifiques (Ménage)",
            value=persona_data.get("besoins_specifiques",""),
            key=f"menage_besoins_specifiques_{index_produit}",
            help=(
        "Décrivez les besoins spécifiques du ménage que votre produit ou service vise à satisfaire. "
        "Cela peut inclure des besoins fonctionnels, émotionnels, ou sociaux. "
        "Comprendre ces besoins permet de mieux cibler les fonctionnalités et les avantages proposés."
        )
            
        )
        menage_motivations = st.text_area(
            "Motivations (Ménage)",
            value=persona_data.get("motivations",""),
            key=f"menage_motivations_{index_produit}",
            help=(
        "Décrivez les motivations qui poussent le ménage à utiliser votre produit ou service. "
        "Cela peut inclure le désir de gain de temps, l'amélioration de la qualité de vie, l'atteinte d'objectifs personnels ou familiaux, etc. "
        "Identifier ces motivations aide à créer des messages marketing efficaces."
        )
            
        )

        st.subheader("Capacité d’Adoption de l’Innovation (Ménage)")
        menage_familiarite_tech = st.text_area(
            "Familiarité avec certaines Technologies (Ménage)",
            value=persona_data.get("familiarite_tech",""),
            key=f"menage_familiarite_tech_{index_produit}",
            help=(
        "Décrivez le niveau de familiarité technologique du ménage. "
        "Par exemple, utilise-t-il régulièrement des smartphones, des applications spécifiques, des plateformes en ligne, etc. "
        "Cette information est essentielle pour développer des produits intuitifs et compatibles avec les habitudes technologiques du ménage."
        )
            
        )
        menage_ouverture_changement = st.text_input(
            "Ouverture au Changement (Ménage) - Faible, Moyenne, Élevée",
            value=persona_data.get("ouverture_changement",""),
            key=f"menage_ouverture_changement_{index_produit}",
            help=(
        "Indiquez le niveau d'ouverture au changement du ménage. "
        "Un ménage avec une 'Élevée' ouverture sera plus enclin à essayer de nouveaux produits ou services, tandis qu'un ménage avec une 'Faible' ouverture pourrait préférer les solutions éprouvées et familières."
        )
            
        )
        menage_barrieres = st.text_area(
            "Barrières Psychologiques/Culturelles (Ménage)",
            value=persona_data.get("barrieres",""),
            key=f"menage_barrieres_{index_produit}",
            help=(
        "Décrivez les barrières psychologiques ou culturelles qui pourraient empêcher le ménage d'adopter votre produit ou service. "
        "Cela peut inclure des croyances, des habitudes, des normes sociales, ou des craintes spécifiques. "
        "Comprendre ces barrières permet de les adresser dans votre stratégie marketing."
        )
            
        )

        # Mise à jour du dictionnaire persona_data
        persona_data.update({
            "type": "Ménage",
            "taille_menage": menage_taille,
            "revenu_menage": menage_revenu,
            "localisation_menage": menage_localisation,
            "type_logement": menage_type_logement,
            "sensibilite_prix": menage_sensibilite_prix,
            "frequence_achat": menage_frequence_achat,
            "volume_achat": menage_volume_achat,
            "perception_qualite": menage_perception_qualite,
            "utilisation_tech": menage_utilisation_tech,
            "acces_transport": menage_acces_transport,
            "temps_disponible": menage_temps_disponible,
            "besoins_specifiques": menage_besoins_specifiques,
            "motivations": menage_motivations,
            "familiarite_tech": menage_familiarite_tech,
            "ouverture_changement": menage_ouverture_changement,
            "barrieres": menage_barrieres
        })

    # Enfin, on sauvegarde ce persona_data dans le produit
    produit["persona"] = persona_data
    st.session_state["produits_data"][index_produit] = produit



def collect_arbre_probleme(index_produit):
    """
    Collecte et/ou met à jour l'Arbre à Problème pour le produit index_produit.
    Pré-remplit les champs avec les données existantes si disponibles.
    """
    st.header("Arbre à Problème")
    
    # Récupération des données existantes
    produit = st.session_state["produits_data"][index_produit]
    data = produit.get("problem_tree", {})
    
    # 1. Contexte
    st.subheader("Contexte")
    contexte = st.text_area(
        "Décrire le Contexte",
        placeholder="Exemple : Dans le domaine de la santé publique dans les zones rurales...",
        value=data.get("contexte", ""),
        help=(
            "Décrivez le contexte général dans lequel le problème se situe. "
            "Incluez des informations pertinentes sur le secteur d'activité, la localisation géographique, "
            "les conditions socio-économiques, les tendances actuelles, et tout autre élément contextuel important. "
            "Cette section doit fournir une vue d'ensemble qui aide à comprendre les facteurs externes et internes influençant le problème. "
            "Par exemple : 'Dans le domaine de la santé publique dans les zones rurales de la République Démocratique du Congo, l'accès aux soins de santé est limité en raison de l'insuffisance des infrastructures médicales et du manque de personnel qualifié.'"
        )
    )
    
    st.markdown("---")  # Séparateur visuel
    
    # 2. Problème Principal
    st.subheader("Problème Principal")
    probleme_principal = st.text_area(
        "Décrire le Problème Principal",
        placeholder="Exemple : Un accès limité aux soins de santé de base.",
        value=data.get("probleme_principal", ""),
        help=(
            "Formulez clairement le problème central que vous souhaitez aborder. "
            "Assurez-vous que la description est précise, concise et reflète bien la nature et l'ampleur du problème. "
            "Cette section doit identifier le cœur du défi ou de la difficulté rencontrée. "
            "Par exemple : 'Un accès limité aux soins de santé de base dans les zones rurales, entraînant une augmentation des taux de mortalité infantile et une détérioration de la santé générale de la population.'"
        )
    )
    
    st.markdown("---")  # Séparateur visuel
    
    # 3. Causes Principales
    st.subheader("Causes Principales")
    causes = data.get("causes_principales", ["", "", ""])
    cause1 = st.text_input(
        "Cause 1",
        placeholder="Exemple : Manque d'infrastructures médicales...",
        value=causes[0] if len(causes) > 0 else "",
        help=(
            "Identifiez et décrivez la première cause principale du problème. "
            "Cette cause doit être directement liée au problème principal et contribuer de manière significative à sa persistance. "
            "Analysez les facteurs sous-jacents qui génèrent ou exacerbent le problème. "
            "Par exemple : 'Manque d'infrastructures médicales adéquates dans les zones rurales, rendant difficile l'accès aux services de santé essentiels.'"
        )
    )
    cause2 = st.text_input(
        "Cause 2",
        placeholder="Exemple : Faible financement gouvernemental...",
        value=causes[1] if len(causes) > 1 else "",
        help=(
            "Identifiez et décrivez la deuxième cause principale du problème. "
            "Assurez-vous qu'elle est également directement liée au problème principal et qu'elle joue un rôle important dans sa manifestation. "
            "Par exemple : 'Faible financement gouvernemental alloué aux services de santé ruraux, limitant les ressources disponibles pour le personnel médical et les équipements nécessaires.'"
        )
    )
    cause3 = st.text_input(
        "Cause 3 (Facultatif)",
        placeholder="Exemple : Isolement géographique...",
        value=causes[2] if len(causes) > 2 else "",
        help=(
            "Identifiez et décrivez une troisième cause principale du problème, si nécessaire. "
            "Ce champ est facultatif et peut être laissé vide si une troisième cause n'est pas pertinente ou nécessaire. "
            "Par exemple : 'Isolement géographique des communautés rurales, compliquant la distribution des médicaments et la mobilité des professionnels de santé.'"
        )
    )
    
    st.markdown("---")  # Séparateur visuel
    
    # 4. Impact
    st.subheader("Impact")
    impact = st.text_area(
        "Décrire l'Impact",
        placeholder="Exemple : Augmentation de la mortalité infantile, retard de développement économique.",
        value=data.get("impact", ""),
        help=(
            "Décrivez les conséquences du problème principal sur les parties prenantes et l'environnement. "
            "Incluez les impacts sociaux, économiques, environnementaux, et sanitaires. "
            "Cette section doit illustrer l'ampleur et la gravité des effets négatifs engendrés par le problème. "
            "Par exemple : 'Augmentation de la mortalité infantile, retard de développement économique, détérioration de la qualité de vie des résidents, et surmenage des services de santé existants.'"
        )
    )
    
    st.markdown("---")  # Séparateur visuel
    
    # 5. Parties Prenantes
    st.subheader("Parties Prenantes")
    parties = data.get("parties_prenantes", ["", "", ""])
    partie1 = st.text_input(
        "Partie Prenante 1",
        placeholder="Exemple : Gouvernement local...",
        value=parties[0] if len(parties) > 0 else "",
        help=(
            "Identifiez et décrivez la première partie prenante impliquée ou affectée par le problème. "
            "Les parties prenantes peuvent être des individus, des groupes, des organisations ou des institutions ayant un intérêt direct ou indirect dans le problème. "
            "Par exemple : 'Gouvernement local responsable de la gestion des services de santé dans les zones rurales.'"
        )
    )
    partie2 = st.text_input(
        "Partie Prenante 2",
        placeholder="Exemple : ONG...",
        value=parties[1] if len(parties) > 1 else "",
        help=(
            "Identifiez et décrivez la deuxième partie prenante impliquée ou affectée par le problème. "
            "Ces parties peuvent jouer un rôle clé dans la résolution du problème ou être affectées par ses conséquences. "
            "Par exemple : 'Organisations non gouvernementales (ONG) travaillant à améliorer les services de santé dans les régions défavorisées.'"
        )
    )
    partie3 = st.text_input(
        "Partie Prenante 3 (Facultatif)",
        placeholder="Exemple : Résidents des zones rurales...",
        value=parties[2] if len(parties) > 2 else "",
        help=(
            "Identifiez et décrivez une troisième partie prenante impliquée ou affectée par le problème, si nécessaire. "
            "Ce champ est facultatif et peut être laissé vide si une troisième partie prenante n'est pas pertinente ou nécessaire. "
            "Par exemple : 'Résidents des zones rurales ayant un accès limité aux services de santé essentiels.'"
        )
    )
    
    st.markdown("---")  # Séparateur visuel
    
    # 6. Opportunités
    st.subheader("Opportunités")
    opportunites = data.get("opportunites", ["", "", ""])
    opportunite1 = st.text_input(
        "Opportunité 1",
        placeholder="Exemple : Introduction de cliniques mobiles...",
        value=opportunites[0] if len(opportunites) > 0 else "",
        help=(
            "Identifiez et décrivez la première opportunité pour résoudre le problème. "
            "Les opportunités peuvent provenir de changements technologiques, de nouvelles politiques, de financements disponibles, ou de partenariats potentiels. "
            "Par exemple : 'Introduction de cliniques mobiles pour fournir des services de santé directement dans les communautés rurales isolées.'"
        )
    )
    opportunite2 = st.text_input(
        "Opportunité 2",
        placeholder="Exemple : Formation d'agents de santé communautaire...",
        value=opportunites[1] if len(opportunites) > 1 else "",
        help=(
            "Identifiez et décrivez la deuxième opportunité pour résoudre le problème. "
            "Assurez-vous que cette opportunité est réalisable et alignée avec les ressources et les capacités disponibles. "
            "Par exemple : 'Formation d'agents de santé communautaire locaux pour assurer une présence continue et un soutien médical dans les villages reculés.'"
        )
    )
    opportunite3 = st.text_input(
        "Opportunité 3 (Facultatif)",
        placeholder="Exemple : Partenariats avec des organisations internationales...",
        value=opportunites[2] if len(opportunites) > 2 else "",
        help=(
            "Identifiez et décrivez une troisième opportunité pour résoudre le problème, si nécessaire. "
            "Ce champ est facultatif et peut être laissé vide si une troisième opportunité n'est pas pertinente ou nécessaire. "
            "Par exemple : 'Établissement de partenariats avec des organisations internationales pour obtenir des financements et des ressources supplémentaires.'"
        )
    )
    
    # Collecte des données
    problem_tree = {
        "contexte": contexte,
        "probleme_principal": probleme_principal,
        "causes_principales": [cause for cause in [cause1, cause2, cause3] if cause.strip()],
        "impact": impact,
        "parties_prenantes": [partie for partie in [partie1, partie2, partie3] if partie.strip()],
        "opportunites": [opp for opp in [opportunite1, opportunite2, opportunite3] if opp.strip()]
    }
    
    # Stockage des données dans la session
    st.session_state["produits_data"][index_produit]["problem_tree"] = problem_tree

def collect_analyse_marche_pme(index_produit):
    """
    Collecte et/ou met à jour l'Analyse du Marché pour le produit index_produit.
    Pré-remplit les champs avec les données existantes si disponibles.
    """
    st.header("Analyse du Marché - PME")
    
    # Récupération des données existantes
    produit = st.session_state["produits_data"][index_produit]
    data = produit.get("analyse_marche", {})
    
    # Taille du Marché
    st.subheader("Taille du Marché")
    taille_marche = st.text_area(
        "Taille du Marché", 
        placeholder="Décrivez la taille du marché, les segments et la valeur totale.",
        value=data.get("taille_marche", ""),
        help=(
            "Décrivez la taille du marché que vous visez. Incluez des informations sur la portée géographique, les segments de clientèle, "
            "et la valeur totale du marché. Cette section doit fournir une vue d'ensemble quantitative et qualitative du marché potentiel. "
            "Par exemple : 'Le marché de la santé numérique en Afrique de l'Ouest représente environ 2 milliards de dollars avec un taux de croissance annuel de 15%.'"
        )
    )
    
    # Segments du Marché
    st.subheader("Segments du Marché")
    segments_marche = st.text_area(
        "Segments du Marché",
        placeholder="Décrivez les segments du marché...",
        value=data.get("segments_marche", ""),
        help=(
            "Décrivez les différents segments de marché que vous ciblez. Identifiez les sous-groupes spécifiques au sein du marché global, "
            "basés sur des critères tels que la démographie, la géographie, le comportement, ou les besoins spécifiques. "
            "Par exemple : 'Segments incluant les jeunes professionnels urbains, les familles avec enfants, et les seniors recherchant des solutions de santé à domicile.'"
        )
    )
    
    # Valeur Totale du Marché ($)
    st.subheader("Valeur Totale du Marché ($)")
    valeur_totale = st.text_area(
        "Valeur Totale du Marché ($)", 
        placeholder="Décrivez la valeur totale du marché...",
        value=data.get("valeur_totale", ""),
        help=(
            "Entrez la valeur totale estimée du marché en dollars. Cette estimation doit refléter la taille financière du marché que vous ciblez. "
            "Utilisez des données secondaires fiables, des études de marché ou des rapports industriels pour appuyer votre estimation. "
            "Par exemple : 'La valeur totale du marché de l'e-commerce en Europe est estimée à 700 milliards de dollars en 2024.'"
        )
    )
    
    # Offres Concurrentes
    st.subheader("Offres Concurrentes")
    offres_concurrentes = st.text_area(
        "Offres Concurrentes", 
        placeholder="Décrivez les offres concurrentes...",
        value=data.get("offres_concurrentes", ""),
        help=(
            "Décrivez les principales offres concurrentes sur le marché. Identifiez les entreprises ou les produits qui répondent déjà aux besoins de vos segments cibles. "
            "Analysez leurs forces et faiblesses, leurs parts de marché, leurs stratégies de prix, et leurs propositions de valeur. "
            "Par exemple : 'Les principaux concurrents incluent HealthPlus et MedTech Solutions, offrant des applications de gestion de santé avec des fonctionnalités similaires mais à des prix plus élevés.'"
        )
    )
    
    # Niveau de Satisfaction
    st.subheader("Niveau de Satisfaction")
    niveau_satisfaction = st.text_area(
        "Niveau de Satisfaction", 
        placeholder="Décrivez le niveau de satisfaction...",
        value=data.get("niveau_satisfaction", ""),
        help=(
            "Évaluez le niveau de satisfaction actuel des clients vis-à-vis des offres concurrentes. Utilisez des enquêtes, des avis clients, ou des études de satisfaction pour recueillir ces informations. "
            "Identifiez les points forts et les lacunes des offres existantes du point de vue des clients. "
            "Par exemple : 'Les clients apprécient la facilité d'utilisation des applications concurrentes, mais signalent un manque de support client réactif et de fonctionnalités personnalisables.'"
        )
    )
    
    # Tendances du Marché
    st.subheader("Tendances du Marché")
    tendances = st.text_area(
        "Tendances du Marché", 
        placeholder="Décrivez les tendances du marché...",
        value=data.get("tendances", ""),
        help=(
            "Décrivez les principales tendances qui influencent actuellement le marché et qui pourraient avoir un impact futur. "
            "Cela peut inclure des évolutions technologiques, des changements réglementaires, des variations des comportements des consommateurs, etc. "
            "Par exemple : 'Croissance de l'utilisation des smartphones, augmentation de la demande pour des solutions de santé personnalisées, et renforcement des régulations sur la protection des données de santé.'"
        )
    )
    
    # Innovations Émergentes
    st.subheader("Innovations Émergentes")
    innovations = st.text_area(
        "Innovations Émergentes", 
        placeholder="Décrivez les innovations émergentes...",
        value=data.get("innovations", ""),
        help=(
            "Identifiez et décrivez les innovations émergentes dans votre secteur qui pourraient créer de nouvelles opportunités ou menacer les acteurs existants. "
            "Cela peut inclure des technologies disruptives, des nouveaux modèles d'affaires, ou des avancées scientifiques. "
            "Par exemple : 'Développement de l'intelligence artificielle pour le diagnostic médical, adoption croissante de la télémédecine, et utilisation de la blockchain pour sécuriser les données de santé.'"
        )
    )
    
    # Comportements Émergents
    st.subheader("Comportements Émergents")
    comportements_emergents = st.text_area(
        "Comportements Émergents", 
        placeholder="Décrivez les comportements émergents...",
        value=data.get("comportements_emergents", ""),
        help=(
            "Analysez les changements dans les comportements des consommateurs qui pourraient influencer la demande sur le marché. "
            "Cela peut inclure des préférences accrues pour des produits durables, une tendance vers l'achat en ligne, ou une demande pour des services personnalisés. "
            "Par exemple : 'Les consommateurs montrent une préférence croissante pour les solutions de santé intégrées et personnalisées, avec une utilisation accrue des applications mobiles pour le suivi de la santé.'"
        )
    )
    
    # Collecte des données
    analyse_marche = {
        "taille_marche": taille_marche,
        "segments_marche": segments_marche,
        "valeur_totale": valeur_totale,
        "offres_concurrentes": offres_concurrentes,
        "niveau_satisfaction": niveau_satisfaction,
        "tendances": tendances,
        "innovations": innovations,
        "comportements_emergents": comportements_emergents
    }
    
    # Stockage des données dans la session
    st.session_state["produits_data"][index_produit]["analyse_marche"] = analyse_marche

def collect_facteurs_limitants_pme(index_produit):
    st.subheader("Facteurs Limitants")
    data = st.session_state["produits_data"][index_produit]["facteurs_limitants"]

    techno_description = st.text_area("Contraintes Technologiques", 
                                      value=data.get("contraintes_technologiques",""),
                                      key=f"techno_description_{index_produit}",
                                      help=(
        "Décrivez les contraintes technologiques qui peuvent limiter le développement ou la mise en œuvre de votre produit ou service. "
        "Cela peut inclure des limitations en termes d'infrastructure, de capacités techniques, de disponibilité des technologies nécessaires, "
        "ou des défis liés à l'intégration avec des systèmes existants. "
        "Par exemple : 'Manque d'accès à des technologies de pointe pour le développement de logiciels personnalisés, rendant difficile la création de fonctionnalités avancées.'"
    ))
    economiques_description = st.text_area("Contraintes Économiques",
                                           value=data.get("contraintes_economiques",""),
                                           key=f"economiques_description_{index_produit}",
                                           help=(
        "Décrivez les contraintes économiques qui peuvent affecter votre projet ou entreprise. "
        "Cela peut inclure des limitations budgétaires, des fluctuations de marché, des coûts élevés de production, "
        "ou des difficultés à obtenir des financements. "
        "Par exemple : 'Budget limité pour le marketing, ce qui restreint la portée des campagnes publicitaires et la visibilité du produit sur le marché.'"
    ))
    culturelles_description = st.text_area("Contraintes Culturelles",
                                           value=data.get("contraintes_culturelles",""),
                                           key=f"culturelles_description_{index_produit}",
                                           help=(
        "Décrivez les contraintes culturelles qui peuvent influencer votre projet ou entreprise. "
        "Cela peut inclure des différences culturelles, des normes sociales, des préférences des consommateurs, ou des résistances au changement. "
        "Par exemple : 'Préférences culturelles pour des produits traditionnels, rendant difficile l'adoption de nouvelles solutions innovantes.'"
    ))
    psych_phys_description = st.text_area("Contraintes Psychologiques et Physiologiques",
                                          value=data.get("contraintes_psych_phys",""),
                                          key=f"psych_phys_description_{index_produit}",
                                          help=(
        "Décrivez les contraintes psychologiques et physiologiques qui peuvent impacter votre projet ou entreprise. "
        "Cela peut inclure des facteurs tels que la résistance au changement, le stress des employés, des limitations physiques liées à la santé, "
        "ou des défis liés à la motivation et à la productivité. "
        "Par exemple : 'Résistance des employés à l'adoption de nouveaux outils technologiques en raison de la peur de la perte d'emploi.'"
    ))
    regulatoires_description = st.text_area("Contraintes Réglementaires",
                                            value=data.get("contraintes_reglementaires",""),
                                            key=f"regulatoires_description_{index_produit}",
                                            help=(
        "Décrivez les contraintes réglementaires qui peuvent affecter votre projet ou entreprise. "
        "Cela peut inclure des lois, des régulations, des normes industrielles, des exigences de conformité, ou des barrières administratives. "
        "Par exemple : 'Conformité stricte aux régulations de protection des données, nécessitant des investissements supplémentaires en sécurité informatique.'"
    ))

    st.session_state["produits_data"][index_produit]["facteurs_limitants"] = {
        "contraintes_technologiques": techno_description,
        "contraintes_economiques": economiques_description,
        "contraintes_culturelles": culturelles_description,
        "contraintes_psych_phys": psych_phys_description,
        "contraintes_reglementaires": regulatoires_description
    }

if 'competitors' not in st.session_state:
    st.session_state.competitors = []

# Fonction de dialogue pour ajouter un concurrent
@st.dialog ("Ajouter un Concurrent", width="large")
def add_competitor_dialog():
    with st.form("add_competitor_form", clear_on_submit=True):
        # Création de deux colonnes
        col1, col2 = st.columns(2)
        
        with col1:
            concurrent_nom = st.text_input(
                "Nom du Concurrent", 
                placeholder="Nom du concurrent...",
                help="Entrez le nom complet du concurrent."
            )
        
        with col2:
            type_concurrent = st.selectbox(
                "Type de Concurrent", 
                options=["Directe", "Indirecte"],
                help="Sélectionnez le type de concurrent."
            )
        # Création de deux colonnes pour le slider et le bouton de soumission
        col3, col4 = st.columns([2, 2])
        with col3:
            force = st.text_area(
                "Forces du Concurrent", 
                placeholder="Décrivez les forces...",
                help="Décrivez les forces de ce concurrent."
            )
        with col4:
            faiblesse = st.text_area(
                "Faiblesses du Concurrent", 
                placeholder="Décrivez les faiblesses...",
                help="Décrivez les faiblesses de ce concurrent."
            )
        # Création de trois colonnes pour la perception et le bouton de soumission
        col5, col6, col7 = st.columns([4, 2, 1])
        
        with col5:
            perception =  st.text_area(
                "Perception de l’offre de la concurrence (Niveau de Satisfaction)", 
                placeholder="Décrivez la satisfaction des clients envers ce concurrent... (par ex: Produit de bonne qualité mais prix élevé)",
                help="Évaluez la satisfaction des clients envers ce concurrent."
            )
        
        with col7:
            submit = st.form_submit_button("Ajouter")
            if submit:
                if concurrent_nom.strip() == "":
                    st.error("Le nom du concurrent ne peut pas être vide.")
                else:
                    competitor = {
                        "Nom": concurrent_nom,
                        "Forces": force,
                        "Faiblesses": faiblesse,
                        "Type": type_concurrent,
                        "Perception": perception
                    }
                    st.session_state.competitors.append(competitor)
                    st.success(f"Concurrent **{concurrent_nom}** ajouté avec succès!")
                    st.rerun()


# Fonction de dialogue pour modifier un concurrent
@st.dialog ("Modifier un Concurrent", width="large")
def edit_competitor_dialog(index):
    competitor = st.session_state.competitors[index]
    with st.form(f"edit_competitor_form_{index}", clear_on_submit=False):
        col1, col2 = st.columns(2)
        
        with col1:
            concurrent_nom = st.text_input(
                "Nom du Concurrent", 
                value=competitor['Nom'],
                help="Entrez le nom complet du concurrent."
            )
        with col2:
            type_concurrent = st.selectbox(
                "Type de Concurrent", 
                options=["Directe", "Indirecte"], 
                index=["Directe", "Indirecte"].index(competitor['Type']),
                help="Sélectionnez le type de concurrent."
            )
        col3, col4 = st.columns([2, 2])
        
        with col3:    
            force = st.text_area(
                "Forces du Concurrent", 
                value=competitor['Forces'],
                help="Décrivez les forces de ce concurrent."
            )
            
        with col4:
            faiblesse = st.text_area(
                "Faiblesses du Concurrent", 
                value=competitor['Faiblesses'],
                help="Décrivez les faiblesses de ce concurrent."
            )
        
        # Création de trois colonnes pour la perception et le bouton de soumission
        col5, col6, col7 = st.columns([4, 2, 1])

        with col5:
            perception = st.text_area(
                "Perception de l’offre de la concurrence (Niveau de Satisfaction)", 
                competitor['Perception'],
                help="Évaluez la satisfaction des clients envers ce concurrent."
            )
            
        with col7:    
            submit = st.form_submit_button("Mettre à jour")
            if submit:
                if concurrent_nom.strip() == "":
                    st.error("Le nom du concurrent ne peut pas être vide.")
                else:
                    st.session_state.competitors[index] = {
                        "Nom": concurrent_nom,
                        "Forces": force,
                        "Faiblesses": faiblesse,
                        "Type": type_concurrent,
                        "Perception": perception
                    }
                    st.success(f"Concurrent **{concurrent_nom}** mis à jour avec succès!")
                    st.rerun()  # Ferme automatiquement le modal

# Fonction pour supprimer un concurrent
def delete_competitor(index):
    competitor = st.session_state.competitors.pop(index)
    st.success(f"Concurrent **{competitor['Nom']}** supprimé avec succès!")
    st.rerun()  # Rafraîchir l'application pour mettre à jour la liste

# Fonction pour tronquer le texte
def tronquer_texte(texte, max_caracteres=100):
    if len(texte) > max_caracteres:
        return texte[:max_caracteres] + "..."
    else:
        return texte

# Fonction pour collecter la concurrence
def collect_concurrence_pme_multi():
    """
    Ancienne fonction qui affiche la liste des concurrents et permet d'ajouter, modifier ou supprimer.
    Utilise st.session_state.competitors.
    Retourne la liste mise à jour des concurrents.
    """
    st.subheader("Liste des Concurrents")
    if st.session_state.competitors:
        # Créer une ligne d'en-tête
        header_cols = st.columns([2, 2, 2, 2, 2, 1, 1])
        header_cols[0].markdown("**Nom**")
        header_cols[1].markdown("**Forces**")
        header_cols[2].markdown("**Faiblesses**")
        header_cols[3].markdown("**Type**")
        header_cols[4].markdown("**Perception**")
        header_cols[5].markdown("**Act**")
        header_cols[6].markdown("**Act**")
        
        # Afficher les boutons d'action pour chaque concurrent
        for index, competitor in enumerate(st.session_state.competitors):
            cols = st.columns([2, 2, 2, 2, 2, 1, 1])  # Ajuster les proportions selon vos besoins
            nom_tronque = tronquer_texte(competitor['Nom'], max_caracteres=15)
            forces_tronque = tronquer_texte(competitor['Forces'], max_caracteres=15)
            faiblesse_tronque = tronquer_texte(competitor['Faiblesses'], max_caracteres=15)
            perception_tronque = tronquer_texte(competitor['Perception'], max_caracteres=15)
            
            cols[0].write(nom_tronque)
            cols[1].write(forces_tronque)
            cols[2].write(faiblesse_tronque)
            cols[3].write(competitor['Type'])
            cols[4].write(perception_tronque)
            # Bouton Modifier
            if cols[5].button("✏️", key=f"modify_{index}"):
                edit_competitor_dialog(index)
            # Bouton Supprimer
            if cols[6].button("🗑️", key=f"delete_{index}"):
                delete_competitor(index)
    else:
        st.info("Aucun concurrent ajouté pour le moment.")
    
    # Bouton pour ouvrir le dialogue d'ajout en dehors du formulaire
    if st.button("Ajouter un Concurrent"):
        add_competitor_dialog()
    
    # Retourner la liste des concurrents
    return st.session_state.competitors


def collect_concurrence_pme(index_produit):
    """
    Collecte et/ou met à jour la Concurrence pour le produit index_produit.
    Synchronise la liste des concurrents spécifiques au produit avec la logique existante.
    """
    # 1. Récupérer le produit courant
    produit = st.session_state["produits_data"][index_produit]
    
    # 2. Initialiser la liste des concurrents si absente
    if "competitors" not in produit:
        produit["competitors"] = []
    
    # 3. Synchroniser la liste spécifique au produit avec la variable globale
    st.session_state.competitors = produit["competitors"]
    
    # 4. Appeler la fonction legacy qui manipule st.session_state.competitors
    collect_concurrence_pme_multi()
    
    # 5. Re-synchroniser la liste mise à jour dans le produit
    produit["competitors"] = st.session_state.competitors
    
    # 6. Mettre à jour la session
    st.session_state["produits_data"][index_produit] = produit


# ----------------------------------------------------------------------------
# 2) Fonctions de collecte des données pour Startups
# ----------------------------------------------------------------------------

def collect_persona_startup():
    st.header("Persona - Startup")
    
    # Données Démographiques
    st.subheader("Données Démographiques")
    age = st.number_input("Âge", min_value=18, max_value=100, value=30)
    sexe = st.text_input("Sexe", "Homme/Femme/Autre")
    localisation_detail = st.text_input("Localisation Géographique (ex: Paris, France)", "Paris, France")
    education = st.text_input("Niveau d'Éducation", "Ex: Licence, Master")
    profession = st.text_input("Profession", "Ex: Ingénieur, Designer")
    revenu_moyen = st.number_input("Revenu Moyen ($)", min_value=0, step=100, value=1000)
    
    # Paramètres Comportementaux
    st.subheader("Paramètres Comportementaux")
    
    sensibilite_prix = st.text_input("Sensibilité au Prix", placeholder="Décrivez la sensibilité au prix...")
    frequence_achat = st.text_input("Fréquence d'Achat", placeholder="Décrivez la fréquence d'achat...")
    volume_achat = st.text_input("Volume d'Achat", placeholder="Décrivez le volume d'achat...")
    perception_qualite = st.text_area("Perception de la Qualité", placeholder="Décrivez la perception de la qualité...")
    utilisation_tech = st.text_area("Utilisation Technologique", placeholder="Décrivez l'utilisation technologique...")
    acces_transport = st.text_area("Accessibilité (Transport)", placeholder="Décrivez l'accessibilité via le transport...")
    temps_disponible = st.text_area("Temps Disponible", placeholder="Décrivez le temps disponible...")
    besoins_specifiques = st.text_area("Besoins Spécifiques",placeholder= "Décrivez les besoins spécifiques...")
    motivations = st.text_area("Motivations", placeholder="Décrivez les motivations des clients...")
    
    # Capacité d’Adoption de l’Innovation
    st.subheader("Capacité d’Adoption de l’Innovation")
    
    familiarite_tech = st.text_area("Familiarité avec certaines Technologies", placeholder="Décrivez la familiarité technologique...")
    ouverture_changement = st.text_input("Ouverture au Changement", placeholder="Faible/Moyenne/Élevée")
    barrières = st.text_area("Barrières Psychologiques/Culturelles", placeholder="Décrivez les barrières psychologiques ou culturelles...")
    
    persona = {
        "âge": age,
        "sexe": sexe,
        "localisation": localisation_detail,
        "éducation": education,
        "profession": profession,
        "revenu_moyen": revenu_moyen,
        "sensibilite_prix": sensibilite_prix,
        "frequence_achat": frequence_achat,
        "volume_achat": volume_achat,
        "perception_qualite": perception_qualite,
        "utilisation_tech": utilisation_tech,
        "acces_transport": acces_transport,
        "temps_disponible": temps_disponible,
        "besoins_specifiques": besoins_specifiques,
        "motivations": motivations,
        "familiarite_tech": familiarite_tech,
        "ouverture_changement": ouverture_changement,
        "barrieres": barrières
    }
    
    return persona

def collect_analyse_marche_startup():
    st.header("Analyse du Marché - Startup")
    
    # Taille du Marché
    st.subheader("Taille du Marché")
    taille_marche = st.text_area("Taille du Marché", 
                                 placeholder="Décrivez la taille du marché, les segments et la valeur totale.",
                                 help=(
        "Décrivez la taille du marché que vous visez. Incluez des informations sur la portée géographique, les segments de clientèle, "
        "et la valeur totale du marché. Cette section doit fournir une vue d'ensemble quantitative et qualitative du marché potentiel. "
        "Par exemple : 'Le marché de la santé numérique en Afrique de l'Ouest représente environ 2 milliards de dollars avec un taux de croissance annuel de 15%.'"
    ))
    
    # Segments du Marché
    st.subheader("Segments du Marché")
    segments_marche = st.text_area("Segments du Marché",
                                   placeholder="Décrivez les segments du marché...",
                                   help=(
        "Décrivez les différents segments de marché que vous ciblez. Identifiez les sous-groupes spécifiques au sein du marché global, "
        "basés sur des critères tels que la démographie, la géographie, le comportement, ou les besoins spécifiques. "
        "Par exemple : 'Segments incluant les jeunes professionnels urbains, les familles avec enfants, et les seniors recherchant des solutions de santé à domicile.'"
    ))
    
    # Valeur Totale du Marché ($)
    st.subheader("Valeur Totale du Marché ($)")
    valeur_totale = st.text_area("Valeur Totale du Marché ($)", 
                                 placeholder="Décrivez la valeur totale du marché...",
                                 help=(
        "Entrez la valeur totale estimée du marché en dollars. Cette estimation doit refléter la taille financière du marché que vous ciblez. "
        "Utilisez des données secondaires fiables, des études de marché ou des rapports industriels pour appuyer votre estimation. "
        "Par exemple : 'La valeur totale du marché de l'e-commerce en Europe est estimée à 700 milliards de dollars en 2024.'"
    ))
    
    # Offres Concurrentes
    st.subheader("Offres Concurrentes")
    offres_concurrentes = st.text_area("Offres Concurrentes", 
                                       placeholder="Décrivez les offres concurrentes...",
                                       help=(
        "Décrivez les principales offres concurrentes sur le marché. Identifiez les entreprises ou les produits qui répondent déjà aux besoins de vos segments cibles. "
        "Analysez leurs forces et faiblesses, leurs parts de marché, leurs stratégies de prix, et leurs propositions de valeur. "
        "Par exemple : 'Les principaux concurrents incluent HealthPlus et MedTech Solutions, offrant des applications de gestion de santé avec des fonctionnalités similaires mais à des prix plus élevés.'"
    ))
    
    # Niveau de Satisfaction
    st.subheader("Niveau de Satisfaction")
    niveau_satisfaction = st.text_area("Niveau de Satisfaction", 
                                       placeholder="Décrivez le niveau de satisfaction...",
                                       help=(
        "Évaluez le niveau de satisfaction actuel des clients vis-à-vis des offres concurrentes. Utilisez des enquêtes, des avis clients, ou des études de satisfaction pour recueillir ces informations. "
        "Identifiez les points forts et les lacunes des offres existantes du point de vue des clients. "
        "Par exemple : 'Les clients apprécient la facilité d'utilisation des applications concurrentes, mais signalent un manque de support client réactif et de fonctionnalités personnalisables.'"
    ))
    
    # Tendances du Marché
    st.subheader("Tendances du Marché")
    tendances = st.text_area("Tendances du Marché", 
                             placeholder="Décrivez les tendances du marché...",
                            help=(
        "Décrivez les principales tendances qui influencent actuellement le marché et qui pourraient avoir un impact futur. "
        "Cela peut inclure des évolutions technologiques, des changements réglementaires, des variations des comportements des consommateurs, etc. "
        "Par exemple : 'Croissance de l'utilisation des smartphones, augmentation de la demande pour des solutions de santé personnalisées, et renforcement des régulations sur la protection des données de santé.'"
    ))
    
    # Innovations Émergentes
    st.subheader("Innovations Émergentes")
    innovations = st.text_area("Innovations Émergentes", 
                               help=(
        "Identifiez et décrivez les innovations émergentes dans votre secteur qui pourraient créer de nouvelles opportunités ou menacer les acteurs existants. "
        "Cela peut inclure des technologies disruptives, des nouveaux modèles d'affaires, ou des avancées scientifiques. "
        "Par exemple : 'Développement de l'intelligence artificielle pour le diagnostic médical, adoption croissante de la télémédecine, et utilisation de la blockchain pour sécuriser les données de santé.'"
    ))
    
    # Comportements Émergents
    st.subheader("Comportements Émergents")
    comportements_emergents = st.text_area("Comportements Émergents", 
                                           placeholder="Décrivez les comportements émergents...",
                                           help=(
        "Analysez les changements dans les comportements des consommateurs qui pourraient influencer la demande sur le marché. "
        "Cela peut inclure des préférences accrues pour des produits durables, une tendance vers l'achat en ligne, ou une demande pour des services personnalisés. "
        "Par exemple : 'Les consommateurs montrent une préférence croissante pour les solutions de santé intégrées et personnalisées, avec une utilisation accrue des applications mobiles pour le suivi de la santé.'"
    ))
    
    analyse_marche = {
        "taille_marche": taille_marche,
        "segments_marche": segments_marche,
        "valeur_totale": valeur_totale,
        "offres_concurrentes": offres_concurrentes,
        "niveau_satisfaction": niveau_satisfaction,
        "tendances": tendances,
        "innovations": innovations,
        "comportements_emergents": comportements_emergents
    }
    
    return analyse_marche

def collect_facteurs_limitants_startup():
    st.header("Facteurs Limitants - Startup")
    
    # Contraintes Technologiques
    st.subheader("Contraintes Technologiques")
    contraintes_techno = st.text_area("Contraintes Technologiques", 
                                      placeholder="Décrivez les contraintes technologiques...",
                                      help=(
        "Décrivez les contraintes technologiques qui peuvent limiter le développement ou la mise en œuvre de votre produit ou service. "
        "Cela peut inclure des limitations en termes d'infrastructure, de capacités techniques, de disponibilité des technologies nécessaires, "
        "ou des défis liés à l'intégration avec des systèmes existants. "
        "Par exemple : 'Manque d'accès à des technologies de pointe pour le développement de logiciels personnalisés, rendant difficile la création de fonctionnalités avancées.'"
    ))
    
    # Contraintes Économiques
    st.subheader("Contraintes Économiques")
    contraintes_economiques = st.text_area("Contraintes Économiques", 
                                           placeholder="Décrivez les contraintes économiques...",
                                           help=(
        "Décrivez les contraintes économiques qui peuvent affecter votre projet ou entreprise. "
        "Cela peut inclure des limitations budgétaires, des fluctuations de marché, des coûts élevés de production, "
        "ou des difficultés à obtenir des financements. "
        "Par exemple : 'Budget limité pour le marketing, ce qui restreint la portée des campagnes publicitaires et la visibilité du produit sur le marché.'"
    ))
    
    # Contraintes Culturelles
    st.subheader("Contraintes Culturelles")
    contraintes_culturelles = st.text_area("Contraintes Culturelles", 
                                           placeholder="Décrivez les contraintes culturelles...",
                                           help=(
        "Décrivez les contraintes culturelles qui peuvent influencer votre projet ou entreprise. "
        "Cela peut inclure des différences culturelles, des normes sociales, des préférences des consommateurs, ou des résistances au changement. "
        "Par exemple : 'Préférences culturelles pour des produits traditionnels, rendant difficile l'adoption de nouvelles solutions innovantes.'"
    ))
    
    # Contraintes Psychologiques et Physiologiques
    st.subheader("Contraintes Psychologiques et Physiologiques")
    contraintes_psych_phys = st.text_area("Contraintes Psychologiques et Physiologiques", 
                                          placeholder="Décrivez ces contraintes...",
                                          help=(
        "Décrivez les contraintes psychologiques et physiologiques qui peuvent impacter votre projet ou entreprise. "
        "Cela peut inclure des facteurs tels que la résistance au changement, le stress des employés, des limitations physiques liées à la santé, "
        "ou des défis liés à la motivation et à la productivité. "
        "Par exemple : 'Résistance des employés à l'adoption de nouveaux outils technologiques en raison de la peur de la perte d'emploi.'"
    ))
    
    # Contraintes Réglementaires
    st.subheader("Contraintes Réglementaires")
    contraintes_reglementaires = st.text_area("Contraintes Réglementaires", 
                                              placeholder="Décrivez les contraintes réglementaires...",
                                              help=(
        "Décrivez les contraintes réglementaires qui peuvent affecter votre projet ou entreprise. "
        "Cela peut inclure des lois, des régulations, des normes industrielles, des exigences de conformité, ou des barrières administratives. "
        "Par exemple : 'Conformité stricte aux régulations de protection des données, nécessitant des investissements supplémentaires en sécurité informatique.'"
    ))
    
    facteurs_limitants = {
        "contraintes_technologiques": contraintes_techno,
        "contraintes_economiques": contraintes_economiques,
        "contraintes_culturelles": contraintes_culturelles,
        "contraintes_psych_phys": contraintes_psych_phys,
        "contraintes_reglementaires": contraintes_reglementaires
    }
    
    return facteurs_limitants

def collect_concurrence_startup():
    st.header("Évaluation de la Concurrence - Startup")
    
    # Concurrents Directs
    concurrents_directs = st.text_area("Concurrents Directs", placeholder="Listez les concurrents directs...")
    
    # Concurrents Indirects
    concurrents_indirects = st.text_area("Concurrents Indirects", placeholder="Listez les concurrents indirects...")
    
    # Forces des Concurrents
    forces_concurrents = st.text_area("Forces des Concurrents", placeholder="Décrivez les forces des concurrents...")
    
    # Faiblesses des Concurrents
    faiblesses_concurrents = st.text_area("Faiblesses des Concurrents", placeholder="Décrivez les faiblesses des concurrents...")
    
    # Niveau de Satisfaction des Clients envers les Concurrents
    satisfaction_concurrence = st.slider("Satisfaction des Clients envers les Concurrents", 0, 10, 5)
    
    # Niveau de Confiance des Clients envers les Concurrents
    confiance_concurrence = st.slider("Confiance des Clients envers les Concurrents", 0, 10, 5)
    
    concurrence = {
        "concurrents_directs": concurrents_directs,
        "concurrents_indirects": concurrents_indirects,
        "forces_concurrents": forces_concurrents,
        "faiblesses_concurrents": faiblesses_concurrents,
        "satisfaction_concurrence": satisfaction_concurrence,
        "confiance_concurrence": confiance_concurrence
    }
    
    return concurrence




def get_metaprompt(type_entreprise):
    """
    Retourne un metaprompt spécifique basé sur le type d'entreprise.
    """
    metaprompts = {
        "PME": """**Méta-Prompt pour l’Élaboration d’un Business Model pour PME Traditionnelle (Intégrant des Innovations Low-Tech et Adaptées aux Contextes Africains ou Émergents)**

        **Votre Rôle :**  
        Vous êtes un expert en stratégie d’entreprise, marketing, UX, innovation frugale (low-tech et éventuellement high-tech), et élaboration de Business Models. Vous devez générer un Business Model complet, clair, chiffré, cohérent et innovant, adapté à une PME qui opère dans un environnement local (par exemple en Afrique ou dans d’autres pays émergents) où les réalités technologiques, économiques, culturelles et réglementaires diffèrent des contextes occidentaux fortement numérisés.  
        L’innovation ne sera pas seulement technologique de pointe (high-tech), mais aussi low-tech (solutions simples, robustes, faciles d’entretien, peu consommatrices de ressources), et tenant compte des infrastructures limitées, des préférences culturelles, de la disponibilité intermittente de l’électricité, du coût de la connectivité, de l’importance du lien social, etc.

        Votre tâche s’organise en trois phases :  
        1. Configuration Initiale (Collecte et Structuration des Données)  
        2. Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)  
        3. Production Finale (Business Model Canvas)

        Pour chaque phase, suivez les instructions et veillez à :  
        - Prendre en compte la persona (données démographiques, comportementales, capacités d’adoption de l’innovation)
        - Pendre en compte l'arbre à problemes(Problème Principal,Causes Principales,Impact,Parties Prenantes , Opportunités)  
        - Analyser le marché (taille, segments, offres existantes formelles et informelles, niveau de satisfaction, tendances locales, disponibilité de ressources, logistique)  
        - Intégrer les facteurs limitants (technologiques, économiques, culturels, psychologiques, physiologiques, réglementaires, infrastructures limitées)  
        - Évaluer la concurrence (locale, informelle, substituts traditionnels), comprendre les niveaux de satisfaction et de confiance  
        - Comprendre le parcours client (avant, pendant, après), intégrer la carte d’empathie, identifier les gains et souffrances spécifiques au contexte (par exemple : importance du bouche-à-oreille, confiance interpersonnelle, exigence de robustesse, maintenance locale)  
        - Vérifier systématiquement la cohérence, proposer des optimisations et ajustements  
        - Avant d’introduire une innovation (low-tech ou high-tech), s’assurer que la persona est prête à l’adopter, en tenant compte de l’accessibilité, du coût, de la simplicité et de la réputation  
        - Produire un Business Model Canvas complet (9 blocs), avec des méta-prompts spécifiques pour chacun des blocs, adaptés au contexte local

        ---

        ### Phase 1 : Configuration Initiale (Entrée de Données)

        1. **Recueille et structure les informations suivantes :**  
        - **Persona** :  
            - Données démographiques : âge, sexe, localisation (zones urbaines, péri-urbaines, rurales), niveau d’éducation (alphabétisation, langues parlées), profession (artisans, commerçants, agriculteurs, employés, indépendants), revenu moyen.  
            - Paramètres comportementaux : sensibilité au prix (budgets limités, nécessité de micro-paiements), fréquence et volume d’achat (achats ponctuels, saisonniers, hebdomadaires), perception de la qualité (fiabilité, durabilité), utilisation technologique (téléphones basiques, smartphones d’entrée de gamme, accès limité à Internet), accessibilité (distance aux points de vente, transport limité), temps disponible (horaires de travail, saison des récoltes), besoins spécifiques (ex : accès à l’eau, énergie, outils agricoles, services financiers de base, éducation des enfants, soins de santé).  
            - Capacité d’adoption de l’innovation : Familiarité avec certaines technologies (mobile money, radios communautaires, solutions solaires), ouverture au changement dépendant de la preuve sociale, de la confiance dans la communauté, de la simplicité et robustesse du produit/service. Barrières psychologiques/culturelles (méfiance envers les nouvelles solutions étrangères, préférence pour le contact humain, importance de la recommandation de la famille ou du chef de village).  
        
        - **Arbre à Problème** :
            - Contexte:Description générale du domaine ou de la situation actuelle (secteur d'activité, environnement géographique, tendances actuelles du marché), Facteurs externes influençant la situation (réglementations, conditions économiques, technologies émergentes),Facteurs internes pertinents (ressources disponibles, compétences clés, structure organisationnelle).
            - Problème Principal : Identification du défi ou de l'obstacle central (nature du problème, circonstances spécifiques),Impact immédiat sur l'organisation ou le projet (effets sur les opérations, la performance financière, la réputation).
            - Causes Principales :Causes internes contribuant au problème (processus inefficaces, manque de compétences, ressources limitées),Causes externes contribuant au problème (concurrence accrue, changements de marché, évolutions technologiques),Interaction entre les causes internes et externes (comment elles se renforcent mutuellement).
            - Impact:Conséquences financières du problème (pertes de revenus, augmentation des coûts, rentabilité réduite),Effets opérationnels (délai dans les projets, baisse de productivité, qualité des services ou produits affectée),Impact sur les parties prenantes (satisfaction des clients, moral des employés, relations avec les partenaires).
            - Parties Prenantes :Identification des acteurs concernés ou impactés (clients, employés, fournisseurs, investisseurs, communauté locale),Intérêts et attentes de chaque partie prenante vis-à-vis du problème (besoins spécifiques, priorités, préoccupations).
            - Opportunités :Pistes d’amélioration ou de résolution du problème (solutions innovantes, meilleures pratiques),Stratégies pour atténuer les causes principales (formation, réorganisation, investissement technologique),Actions pour maximiser les impacts positifs (exploitation des forces, diversification, partenariats stratégiques),

        - **Analyse du Marché** :  
            - Taille du marché local : estimer la population concernée, le pouvoir d’achat moyen, les infrastructures disponibles.  
            - Segments : populations urbaines vs rurales, artisans, commerçants, coopératives, PME locales, secteur informel.  
            - Offres concurrentes existantes : solutions traditionnelles (artisanales, informelles), importations bas de gamme, programmes d’ONG, concurrents locaux ou étrangers, modèles low-cost.  
            - Niveau de satisfaction actuel : Les clients sont-ils satisfaits des solutions actuelles ? Y a-t-il un manque de fiabilité, de formation, de SAV ?  
            - Tendances : adoption progressive du mobile money, sensibilisation croissante à l’énergie solaire, émergence de petites coopératives, engouement pour des solutions durables et réparables.  
            - Innovations et comportements émergents : réemploi, économie circulaire, mise en commun de ressources, augmentation des transferts d’argent via mobile.  
        
        - **Facteurs Limitants** :  
            - Contraintes technologiques : faible accès à l’électricité stable, couverture internet inégale, outils technologiques rudimentaires, importance de solutions low-tech (pompes manuelles, panneaux solaires simples, systèmes de filtration d’eau basiques).  
            - Contraintes économiques : revenus limités, volatilité des prix, accès restreint au crédit, nécessité d’étaler les paiements (micro-paiements, crédit rotatif, tontines).  
            - Contraintes culturelles : langues locales, importance de la confiance interpersonnelle, réticence à adopter des produits inconnus sans démonstration ou validation par la communauté.  
            - Contraintes psychologiques et physiologiques : besoin de solutions simples d’utilisation, ergonomiques, adaptées aux conditions climatiques (chaleur, poussière), faible taux d’alphabétisation nécessitant des modes d’emploi visuels.  
            - Contraintes réglementaires : normes locales, barrières douanières, absence de normes formelles dans certains secteurs, difficulté à obtenir des certifications officielles.  

        **Après avoir recueilli ces données, effectue une première analyse critique** :  
        - Vérifie la cohérence des informations.  
        - Identifie les lacunes (par exemple, manque d’informations sur le pouvoir d’achat réel, sur le réseau de distribution informel, sur le rôle des leaders d’opinion locaux).  
        - Propose des compléments ou ajustements pour optimiser la qualité des données (ajouter des données sur la saisonnalité du marché, l’influence des ONG, l’impact des conditions climatiques, la présence ou non de microfinance).

        ---

        ### Phase 2 : Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d’Empathie** :  
        - Décris le parcours client (avant, pendant, après) en tenant compte des conditions locales :  
            - Avant : Le client prend conscience de son besoin par le bouche-à-oreille, via un ami, un voisin, un leader communautaire, ou en écoutant la radio locale. Il compare avec les solutions déjà connues (artisan local, réparations informelles, solutions importées). Il évalue la confiance, le prix, la disponibilité.  
            - Pendant : Achat sur un marché local, essai d’une démonstration concrète (démonstration en conditions réelles, sur une place de village), informations données par un vendeur itinérant ou un agent de confiance. Utilisation d’un mode de paiement adapté (cash, mobile money).  
            - Après : Suivi du produit, entretien, besoin de pièces détachées, possibilité de contact direct avec l’entreprise (ligne téléphonique, point de service local), échange d’expériences avec d’autres utilisateurs, éventuel SAV simplifié (réparations locales, pièces détachées bon marché).  
        
        - Identifie les points de contact (marchés, boutiques, intermédiaires locaux, radios communautaires, SMS informatifs), obstacles (faible connectivité, manque d’informations détaillées, barrières linguistiques), moments de vérité (premier essai du produit, première panne et réactivité du SAV), frustrations (produit pas adapté, manuel incompréhensible, manque de fiabilité).  
        
        - Intègre les contraintes physiologiques, psychologiques, économiques, culturelles, technologiques, réglementaires : par exemple, l’importance de la simplicité et de la robustesse pour réduire la crainte d’une technologie trop complexe, la nécessité de support en langue locale, la possibilité de s’adapter aux normes informelles.  
        
        - Crée une carte d’empathie :  
            - Pensées : « Est-ce que cette solution est fiable, reconnue par ma communauté ? Est-ce que je vais perdre mon argent si ça ne marche pas ? »  
            - Sentiments : Méfiance, curiosité, besoin de réassurance, fierté s’il s’agit d’une innovation locale valorisée.  
            - Actions : Demande de conseils à d’autres, observation d’exemples concrets, volonté d’essayer avant d’acheter.

        3. **Gains et Souffrances** :  
        - Liste les gains : par exemple, accès facilité à un service vital (eau, énergie, outil de gestion commerciale simple), réduction du temps et de l’effort, robustesse (moins de pannes), accès à un SAV local, meilleure rentabilité ou productivité.  
        - Liste les souffrances : manque de solutions adaptées, problèmes de maintenance, coûts initiaux trop élevés sans option de paiement flexible, manque de formation pour utiliser correctement le produit.

        4. **Élaboration de la Carte de Valeur** :  
        - Définis la mission de consommation principale : répondre à un besoin fondamental (ex : un outil agricole robuste, une solution d’éclairage solaire fiable, un service financier simple via mobile, un appareil domestique low-tech adapté aux pannes d’électricité).  
        - Identifie les gains déjà fournis par les offres actuelles (ex : disponibilité locale, prix bas) et les souffrances non adressées (faible qualité, pas de SAV, pas d’adaptation aux conditions réelles).  
        - Esquisse une proposition de valeur préliminaire adaptée à la capacité d’adoption de l’innovation par la persona :  
            - Une solution simple, robuste, facilement compréhensible, qui peut être testée avant achat.  
            - Un modèle de distribution local (agents sur le terrain), un SAV accessible, un support en langue locale, des options de paiement flexible (mobile money, tontines, microcrédit).  
            - Intégration progressive d’innovations low-tech (p. ex. appareils mécaniques robustes, panneaux solaires portables) ou high-tech simple (SMS, USSD, application mobile légère) si l’utilisateur est prêt.

        5. **Détermination du Segment de Clients** :  
        - Choisis le type de relation (B2C direct, B2B via des coopératives, B2B2C via des distributeurs locaux).  
        - Priorise les segments qui correspondent le mieux :  
            - Par exemple, petits commerçants urbains ayant un pouvoir d’achat limité mais stables, agriculteurs nécessitant un outil fiable en milieu rural, coopératives d’artisans prêts à adopter une solution pour améliorer leur productivité.  
        - Tient compte de leur sensibilité au prix, de leur ouverture à l’innovation, de leur capacité à comprendre et utiliser la solution, de la nécessité de formation.

        6. **Analyse des Problèmes et Solutions (Canvas de Problème)** :  
        - Identifie les problèmes majeurs : par exemple, la difficulté à accéder à un produit fiable, le manque d’informations, la complexité du produit, le coût trop élevé d’une solution importée haut de gamme.  
        - Associe chaque problème à une solution :  
            - Problème : manque de SAV → Solution : réseau de réparateurs locaux formés.  
            - Problème : prix élevé d’entrée → Solution : offres en micro-paiements, location-vente, partenariats avec microfinance.  
            - Problème : manque de confiance → Solution : démonstrations, témoignages de pairs, communication via radios locales et leaders d’opinion.  
        - Justifie en quoi les solutions sont meilleures que l’existant : plus adaptées, plus abordables, plus simples, prenant en compte la réalité du terrain (faible infrastructure, besoin de résilience, faible taux d’alphabétisation).

        **Après ces étapes, fais une analyse intermédiaire** :  
        - Vérifie la cohérence du contexte, du parcours client, des solutions proposées.  
        - Assure-toi que les innovations (low-tech, partenariats locaux, solutions de paiement flexible) sont compréhensibles et adoptables par la persona.  
        - Propose des ajustements stratégiques : simplification du produit, ajustement du prix, ajout d’un canal de distribution plus local, formation des utilisateurs, partenariats avec des ONG ou des radios locales.

        ---

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses précédentes, génère un Business Model Canvas complet. Utilise les méta-prompts suivants pour chaque bloc, en tenant compte du contexte local, des solutions low-tech et des infrastructures limitées :

        1. **Segments de Clients**  
        Méta-Prompt :  
        « Définis précisément les segments de clients ciblés, en tenant compte :  
        - De leurs caractéristiques sociodémographiques (âge, sexe, localisation, niveau d’éducation, profession, revenu, langue).  
        - De leurs comportements d’achat (fréquence, volume, sensibilité au prix, recours au crédit informel, canaux de confiance : marchés locaux, revendeurs informels, chefs de village, radios).  
        - De leur maturité technologique (téléphones basiques, usage de SMS/USSD, familiarité avec le mobile money, radio, bouche-à-oreille, rencontres physiques).  
        - De leur capacité d’adoption de l’innovation (ouverture au changement si démonstration concrète, barrières culturelles, besoin de preuves, préférences pour du low-tech robuste plutôt que du high-tech fragile).  
        - De leurs contraintes (faible pouvoir d’achat, saisons de récolte, temps de disponibilité, accès difficile à l’électricité ou à internet).  
        Intègre également des scénarios évolutifs (si l’économie se dégrade, réduction de l’achat ou passage à des solutions plus frugales ; si la technologie progresse, adoption graduelle de services numériques simples).  
        Justifie pourquoi ces segments sont retenus : potentiel de rentabilité, facilité d’accès via des canaux locaux, réceptivité à la proposition de valeur (améliorer leur vie de façon concrète, fiable, abordable). »

        2. **Proposition de Valeur**  
        Méta-Prompt :  
        « Détaille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux (eau, énergie, information, outils productifs, services financiers simples).  
        - Les souffrances clientes (manque de fiabilité, difficulté d’entretien, complexité des produits, méfiance) et comment elles sont résolues (simplicité, robustesse, support local, preuves sociales).  
        - Les gains fournis (amélioration de la productivité, économies de temps, durabilité, réduction de la dépendance à des systèmes complexes, meilleure gestion financière) et inclure les bénéfices émotionnels (confiance, fierté, reconnaissance sociale).  
        - La différenciation par rapport aux offres concurrentes : intégration dans le tissu local, formation d’agents locaux, facilité d’entretien, pricing adapté, low-tech combiné avec technologie simple (mobile money), SAV local.  
        - L’introduction progressive de l’innovation : démonstrations pratiques, formation sur le terrain, tutoriels en langue locale, partenariat avec leaders communautaires.  
        - Variantes selon les segments : option premium (un meilleur SAV, une maintenance plus poussée) pour les clients plus solvables, version ultra-simplifiée pour les segments plus conservateurs ou à très faible pouvoir d’achat. »

        3. **Canaux de Distribution**  
        Méta-Prompt :  
        « Définis les canaux par lesquels les clients seront informés, convaincus, achèteront et utiliseront le produit/service, en tenant compte des réalités locales :  
        - Canaux hors ligne : marchés locaux, boutiques physiques, vente itinérante, radios communautaires, affichages, démonstrations sur place, coopératives agricoles, leaders religieux ou communautaires.  
        - Canaux digitaux légers : SMS, USSD, appels téléphoniques, WhatsApp, Facebook local, mobile money.  
        - Nécessité d’omnicanalité adaptée au contexte : cohérence entre communication radio, démonstration physique, et suivi par téléphone.  
        - Simplicité d’accès et besoin d’accompagnement pédagogique (formation dans les marchés, brochures visuelles, tutoriels audio).  
        - Adaptabilité des canaux si le marché évolue (ex: introduction progressive d’une application mobile si la connectivité s’améliore).  
        Justifie chaque canal (coût, accessibilité, confiance) et comment il réduit les obstacles à l’adoption, améliore la satisfaction, et s’intègre dans le parcours client local. »

        4. **Relations Clients**  
        Méta-Prompt :  
        « Décris la nature et la qualité des relations établies avec les clients :  
        - Personnalisation via un réseau d’agents locaux qui connaissent la langue, la culture, et les besoins.  
        - Communauté : création de groupes d’utilisateurs, d’associations locales, de rencontres de démonstration, événements communautaires où les clients échangent leurs expériences.  
        - Automatisation : mise en place d’un service SMS de rappel, d’une hotline téléphonique simple, d’un chatbot vocal si la technologie le permet (ou service d’appels humains en langue locale).  
        - Fidélisation : réductions pour clients fidèles, options de maintenance préventive, accès à des mises à jour techniques simples, partenariats avec des ONG pour aider à la formation continue.  
        - Gestion des plaintes et retours : politique claire de SAV, réparation locale, garantie adaptée, délais de réponse rapides.  
        Intègre la dimension culturelle (contact humain valorisé), psychologique (confiance, besoin de réassurance), réglementaire (respect des règles locales, si existantes). Explique comment ces relations évoluent au fil du temps et renforcent la CLV dans un contexte de marché volatile. »

        5. **Sources de Revenus**  
        Méta-Prompt :
        « Détaille les mécanismes de génération de revenus :  
        - Modèles de tarification : vente directe à prix abordable, options de micro-paiements échelonnés, crédit via partenaire de microfinance, location-vente, abonnement léger (maintenance), freemium (démonstration gratuite, paiement pour les pièces détachées).  
        - Justification des prix : aligner le prix sur le pouvoir d’achat, offrir un excellent rapport qualité/durabilité/prix, tenir compte des référentiels locaux (si les concurrents informels sont très bon marché, justifier la valeur par la fiabilité).  
        - Réductions des freins économiques : essai avant achat, garantie satisfait ou remboursé, partenariats avec ONG ou institutions locales.  
        - Diversification des revenus : ventes croisées (pièces détachées, formation), partenariats B2B (ventes en gros à des coopératives), publicité locale, sponsorisation par des institutions de développement.  
        - Adaptation aux changements : si le marché se contracte, proposer des modèles encore plus frugaux, si la réglementation change, s’adapter par des produits conformes.  
        Explique comment cette structure de revenus soutient la viabilité à long terme et reste cohérente avec la proposition de valeur et la sensibilité au prix de la persona. »

        6. **Ressources Clés**  
        Méta-Prompt :  
        « Identifie toutes les ressources indispensables :  
        - Ressources Humaines : agents locaux (formés aux langues et contextes locaux), réparateurs, formateurs, personnels de SAV.  
        - Ressources Technologiques : outils de communication simples (téléphones basiques, logiciels légers), systèmes de paiement mobile, éventuellement une plateforme centralisée mais légère.  
        - Ressources Intellectuelles : savoir-faire sur l’adaptation du produit au contexte local, guides visuels, partenariats de R&D avec des instituts techniques locaux.  
        - Ressources Matérielles : pièces détachées robustes, matériaux durables, équipements simples qui ne nécessitent pas une infrastructure complexe.  
        - Ressources Financières : capital initial, fonds de roulement, accès à la microfinance ou à des investisseurs sociaux, trésorerie pour faire face aux saisons difficiles.  
        - Ressources Relationnelles : liens solides avec les communautés, chefs traditionnels, radios locales, ONG, institutions de développement.  
        Pour chaque ressource, justifie pourquoi elle est critique (ex. sans agents locaux, pas de confiance ; sans matériaux robustes, produit inutilisable), et comment ces ressources assurent un avantage concurrentiel durable. »

        7. **Activités Clés**  
        Méta-Prompt :  
        « Décris les activités indispensables :  
        - Développement & Innovation : adapter le produit aux conditions locales (climat, langue), améliorer la durabilité, simplifier l’usage.  
        - Production & Livraison : fabrication locale ou semi-locale, contrôle de la qualité, approvisionnement en pièces robustes, logistique simple (transport par camions, motos, ânes si nécessaire).  
        - Marketing & Ventes : communication via radios communautaires, démonstrations physiques, formation d’agents, distribution de brochures visuelles.  
        - Relation Client & Support : formation du personnel de SAV, mise en place d’une hotline téléphonique, ateliers pratiques, visites régulières sur le terrain.  
        - Partenariats & Négociations : conclure des partenariats avec ONG, coopératives, associations villageoises, négocier des conditions avantageuses avec fournisseurs locaux.  
        Intègre une perspective adaptative : si la demande fluctue, ajuster les stocks, si une nouvelle réglementation apparaît, adapter le produit. Justifie comment chaque activité soutient la proposition de valeur. »

        8. **Partenaires Clés**  
        Méta-Prompt :  
        « Liste et justifie les partenaires stratégiques :  
        - Fournisseurs locaux : garantissant disponibilité et qualité des matières premières.  
        - Distributeurs locaux et intermédiaires informels : accès direct à la clientèle, réduction des coûts d’acquisition.  
        - Partenaires technologiques locaux ou ONG : formation, maintenance, R&D adaptée.  
        - Organismes de certification locaux, influenceurs communautaires, médias (radios, journaux locaux) : augmentent la crédibilité et la confiance.  
        - Institutions financières (microfinance) : faciliter l’accès au crédit, au paiement échelonné.  
        Anticipe les risques (un partenaire clé fait défaut, troubles politiques, pénuries) et prévois des alternatives (autres fournisseurs, diversification géographique). Explique comment ces partenariats renforcent la proposition de valeur et l’efficacité opérationnelle. »

        9. **Structure de Coûts**  
        Méta-Prompt :  
        « Détaille les coûts :  
        - Coûts fixes : salaires des agents locaux, loyers de petits entrepôts, licences minimales, amortissement de matériel de base.  
        - Coûts variables : achat des matières premières, commission aux revendeurs, campagnes radio, formation continue, SAV.  
        - Coûts liés à l’innovation : R&D pour adapter le produit, formation des équipes, tests terrain.  
        Analyse la rentabilité :  
        - Le modèle de revenus couvre-t-il ces coûts ?  
        - Possibilités de réduire les coûts (sourcing local moins cher, économies d’échelle, recyclage, revente de pièces usagées).  
        - Stratégies pour faire face aux fluctuations (augmenter la part de services, moduler les prix, limiter le stock).  
        Explique comment la structure de coûts reste en ligne avec la proposition de valeur, le niveau de vie local, et comment elle assure la pérennité financière à long terme. »

        ---

        **Instructions Finales** :  
        Après avoir utilisé ces méta-prompts pour chaque bloc du Business Model Canvas, effectue une dernière vérification :  
        - Assure-toi que tous les blocs sont cohérents et alignés avec la proposition de valeur, le parcours client et les réalités locales.  
        - Vérifie que l’innovation (low-tech ou high-tech adaptée) est réellement adoptable par la persona, apporte un avantage concurrentiel durable, et que les contraintes (culturelles, économiques, réglementaires, infrastructurelles) sont prises en compte.  
        - Contrôle la rentabilité, la viabilité à long terme, et la flexibilité face aux changements (variations saisonnières, crises économiques, évolution des réglementations ou de la pénétration technologique).  
        - Ajuste les éléments (segments, prix, canaux, partenariats) si nécessaire pour améliorer la robustesse du modèle.  
        - Fournis un récapitulatif global du Business Model, mettant en avant la logique, la cohérence, la proposition de valeur différenciante et quelques chiffres (taille du marché estimée, coûts, revenus, marge, etc.) pour valider la viabilité économique.

        Le résultat final doit être un Business Model clair, complet, adapté au contexte local, prêt à être testé ou implémenté, avec une feuille de route pour l’adoption progressive de l’innovation et une vision claire des points de différenciation face aux solutions traditionnelles ou informelles existantes.
        """,
        
        
        "Startup": """ Tu es un assistant expert en stratégie d’entreprise, marketing, UX, innovation et élaboration de Business Models. Ton rôle est de générer un Business Model complet, clair, chiffré, cohérent et innovant, en suivant trois phases : Configuration Initiale, Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation) et Production Finale (Business Model Canvas).

        Tout au long du processus, tu dois :
        - Prendre en compte la persona (données démographiques, comportementales, capacités d’adoption de l’innovation).
        - Pendre en compte l'arbre à problemes(Problème Principal,Causes Principales,Impact,Parties Prenantes , Opportunités)  
        - Analyser le marché (taille, segments, offres existantes, niveau de satisfaction, tendances).
        - Intégrer les facteurs limitants (technologiques, économiques, culturels, psychologiques, physiologiques, réglementaires).
        - Évaluer la concurrence et comprendre le niveau de satisfaction actuel.
        - Comprendre le parcours client (avant, pendant, après), la carte d’empathie, les gains et souffrances.
        - Vérifier systématiquement la cohérence, proposer des optimisations et ajustements.
        - Avant d’introduire une innovation, t’assurer que la persona est prête à l’adopter.
        - Produire un Business Model Canvas complet (9 blocs), avec des meta-prompts spécifiques pour chacun des blocs.

        Voici les étapes :

        ### Phase 1 : Configuration Initiale (Entrée de Données)

        1. Recueille et structure les informations suivantes :
        - **Persona :**
            - Données démographiques : Âge, sexe, localisation, niveau d’éducation, profession, revenu.
            - Paramètres comportementaux : Sensibilité au prix, budget, fréquence et volume d’achat, perception de la qualité, utilisation technologique, accessibilité, temps disponible, besoins, motivations.
            - Capacité d’adoption de l’innovation : Familiarité technologique, ouverture au changement, barrières psychologiques ou culturelles.
        
        - **Arbre à Problème** :
            - Contexte:Description générale du domaine ou de la situation actuelle (secteur d'activité, environnement géographique, tendances actuelles du marché), Facteurs externes influençant la situation (réglementations, conditions économiques, technologies émergentes),Facteurs internes pertinents (ressources disponibles, compétences clés, structure organisationnelle).
            - Problème Principal : Identification du défi ou de l'obstacle central (nature du problème, circonstances spécifiques),Impact immédiat sur l'organisation ou le projet (effets sur les opérations, la performance financière, la réputation).
            - Causes Principales :Causes internes contribuant au problème (processus inefficaces, manque de compétences, ressources limitées),Causes externes contribuant au problème (concurrence accrue, changements de marché, évolutions technologiques),Interaction entre les causes internes et externes (comment elles se renforcent mutuellement).
            - Impact:Conséquences financières du problème (pertes de revenus, augmentation des coûts, rentabilité réduite),Effets opérationnels (délai dans les projets, baisse de productivité, qualité des services ou produits affectée),Impact sur les parties prenantes (satisfaction des clients, moral des employés, relations avec les partenaires).
            - Parties Prenantes :Identification des acteurs concernés ou impactés (clients, employés, fournisseurs, investisseurs, communauté locale),Intérêts et attentes de chaque partie prenante vis-à-vis du problème (besoins spécifiques, priorités, préoccupations).
            - Opportunités :Pistes d’amélioration ou de résolution du problème (solutions innovantes, meilleures pratiques),Stratégies pour atténuer les causes principales (formation, réorganisation, investissement technologique),Actions pour maximiser les impacts positifs (exploitation des forces, diversification, partenariats stratégiques),
        
        - **Analyse du Marché :**
            - Taille du marché, segments, valeur totale.
            - Offres concurrentes, niveau de satisfaction, tendances, innovations, comportements émergents.
        - **Facteurs Limitants :**
            - Contraintes technologiques, économiques, culturelles, réglementaires, physiologiques, psychologiques.
        
        Après avoir recueilli ces données, effectue une première analyse critique :
        - Vérifie la cohérence des informations.
        - Identifie les lacunes.
        - Propose des compléments ou ajustements pour optimiser la qualité des données.

        ### Phase 2 : Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d’Empathie :**
        - Décris le parcours client (avant, pendant, après consommation).
        - Identifie les points de contact, obstacles, moments de vérité, frustrations.
        - Intègre les contraintes physiologiques, psychologiques, économiques, culturelles, technologiques, réglementaires.
        - Crée une carte d’empathie (pensées, sentiments, actions) pour comprendre l’expérience du client à chaque étape.

        3. **Gains et Souffrances :**
        - À partir du parcours client et de la carte d’empathie, liste les gains (bénéfices, réassurance, sentiment de compétence) et les souffrances (problèmes non résolus, frustrations, coûts d’opportunité).

        4. **Élaboration de la Carte de Valeur :**
        - Définis la mission de consommation principale (besoin fondamental).
        - Identifie les gains déjà fournis par les offres actuelles.
        - Mets en évidence les souffrances non adressées.
        - Esquisse une proposition de valeur préliminaire, adaptée à la capacité d’adoption de l’innovation par la persona.

        5. **Détermination du Segment de Clients :**
        - Choisis le type de relation (B2C, B2B, B2B2C…).
        - Priorise les segments (taille, pouvoir d’achat, sensibilité au prix, ouverture à l’innovation, contraintes) qui correspondent le mieux à la proposition de valeur.

        6. **Analyse des Problèmes et Solutions (Canvas de Problème) :**
        - Identifie clairement les problèmes majeurs à résoudre.
        - Associe chaque problème à une solution spécifique, justifie en quoi elle est meilleure que les offres existantes.

        Après ces étapes, effectue une analyse intermédiaire :
        - Vérifie la cohérence du contexte, du parcours client, des solutions proposées.
        - Assure-toi que les innovations sont compréhensibles, utiles et adoptables par la persona.
        - Propose des ajustements stratégiques (simplification de l’offre, ajustement du prix, sélection de segments plus pertinents, etc.) si nécessaire.

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses précédentes, génère un Business Model Canvas complet. Utilise les meta-prompts suivants pour chaque bloc :

        1. **Segments de Clients**  
        Méta-Prompt :  
        « Définis précisément les segments de clients ciblés, en tenant compte :  
        - De leurs caractéristiques sociodémographiques (âge, sexe, localisation, niveau d’éducation, profession, revenu).  
        - De leurs comportements d’achat (fréquence, volume, sensibilité au prix, critères de qualité) et de leur maturité technologique (utilisation d’outils numériques, appareils connectés, plateformes en ligne).  
        - De leur capacité d’adoption de l’innovation (ouverture au changement, barrières psychologiques, éventuelle réticence culturelle).  
        - De leurs contraintes physiologiques (accessibilité, ergonomie), psychologiques (stress, anxiété, besoin de réassurance), économiques (pouvoir d’achat, rapport qualité/prix), culturelles (normes, tabous) et réglementaires (normes légales, certifications).  
        Intègre également des scénarios évolutifs :  
        - Si la technologie évolue, comment ce segment réagit-il ?  
        - S’il y a une crise économique, ces clients réduisent-ils leur consommation ?  
        - Une partie du segment est-elle prête à payer plus pour des options premium ?  
        Justifie pourquoi ces segments sont retenus, comment ils se distinguent de segments non ciblés, et comment leur potentiel de rentabilité, leur facilité d’accès, et leur réceptivité à la proposition de valeur justifient leur inclusion. »

        2. **Proposition de Valeur**  
        Méta-Prompt :  
        « Détaille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux adressés (mission de consommation principale).  
        - Les souffrances clientes (manque de temps, complexité, mauvaise qualité, manque de confiance, crainte face à l’innovation) et comment elles sont résolues.  
        - Les gains fournis (gain de temps, économie d’argent, facilité d’utilisation, statut social, tranquillité d’esprit), y compris les bénéfices émotionnels et symboliques.  
        - La différenciation par rapport aux offres concurrentes (qualité supérieure, innovation plus accessible, prix compétitifs, service client exemplaire, partenariats de prestige).  
        - L’intégration de l’innovation : montre comment elle est introduite progressivement, comment l’éducation ou la formation du client est assurée, et comment les barrières à l’adoption sont levées (essais gratuits, démonstrations, tutoriels, certifications reconnues).  
        - Prévois des variantes de proposition de valeur en fonction des segments, si nécessaire (une version premium pour les early adopters innovants, une version simplifiée pour les plus conservateurs). »

        3. **Canaux de Distribution**  
        Méta-Prompt :  
        « Définis les canaux par lesquels les clients seront informés, convaincus, achèteront et utiliseront le produit/service. Considère :  
        - Les canaux en ligne (site web, application mobile, plateformes e-learning, réseaux sociaux, partenariats avec marketplaces, influenceurs, SEO, SEA).  
        - Les canaux hors ligne (magasins physiques, salons professionnels, conférences, revendeurs, agents sur le terrain).  
        - La nécessité de cohérence entre les points de contact (omnicanal), la simplicité d’accès, le besoin d’accompagnement pédagogique (webinaires, tutoriels vidéo), et les contraintes technologiques de la persona (faible bande passante, préférence pour un canal mobile vs desktop).  
        - L’adaptabilité des canaux si les conditions du marché changent (pénurie d’un canal, évolution légale, concurrence d’un nouveau canal).  
        Justifie pourquoi chaque canal est choisi, comment il s’intègre dans le parcours client, comment il favorise l’adoption de l’innovation, et comment il est optimisé pour réduire les coûts d’acquisition et améliorer la satisfaction. »

        4. **Relations Clients**  
        Méta-Prompt :  
        « Décris la nature et la qualité des relations que l’entreprise établira avec ses clients :  
        - Personnalisation : existe-t-il un accompagnement individuel, des conseils sur mesure, une assistance humaine ou une IA conversationnelle ?  
        - Communauté : les clients peuvent-ils interagir entre eux (forums, réseaux sociaux, clubs, rencontres physiques) pour renforcer leur sentiment d’appartenance et échanger des expériences ?  
        - Automatisation : y a-t-il des éléments de self-service, de chatbots, de bases de connaissances en ligne ? Est-ce adapté aux cibles moins technophiles ?  
        - Fidélisation : cartes de fidélité, programmes de récompenses, contenus exclusifs, mises à jour gratuites, offres spéciales pour clients fidèles.  
        - Gestion des plaintes et retours : procédures de remboursement, garantie de satisfaction, SLA pour répondre aux demandes critiques.  
        Intègre la dimension psychologique (rassurer les clients sur l’innovation), culturelle (certains clients préfèrent un contact humain), réglementaire (besoin de conformité avec les lois sur la protection des données).  
        Explique comment ces relations évoluent au fil du temps (du premier contact à la fidélisation), comment elles améliorent la CLV, et comment elles s’adaptent aux changements de marché (nouveaux concurrents, crises économiques). »

        5. **Sources de Revenus**  
        Méta-Prompt :  
        « Détaille les mécanismes de génération de revenus :  
        - Modèle de tarification : abonnement mensuel, paiement à l’usage, achat unique, freemium avec options premium, licences, commissions.  
        - Justification des prix : comment le prix reflète-t-il la valeur perçue par le client ? Est-il aligné avec le pouvoir d’achat du segment, la concurrence, la qualité et l’innovation proposée ?  
        - Options de réduction des freins économiques : essais gratuits, garantie satisfait ou remboursé, paiement échelonné, remises pour les early adopters.  
        - Diversification des revenus : ventes croisées, upselling, partenariats, publicité, formation complémentaire, monétisation de données (en respectant la réglementation).  
        - Adaptation à des changements de contexte : si le marché se contracte, proposer un modèle plus flexible ? Si une réglementation limite certains types de revenus, anticiper une alternative ?  
        Explique comment cette structure de revenus soutient la croissance, la rentabilité, et s’intègre avec les coûts prévus. Vérifie la cohérence avec la proposition de valeur et la sensibilité au prix de la persona. »

        6. **Ressources Clés**  
        Méta-Prompt :  
        « Identifie toutes les ressources indispensables :  
        - Ressources Humaines : équipes multidisciplinaires (ingénieurs, designers UX, experts marketing, formateurs, support client multilingue) nécessaires à la création, maintenance, amélioration de l’offre.  
        - Ressources Technologiques : plateformes e-learning, serveurs, logiciels de personnalisation, outils d’IA, applications mobiles, infrastructure IT sécurisée.  
        - Ressources Intellectuelles : brevets, marques, contenus propriétaires, méthodologies exclusives, licences de tiers, données clients protégées.  
        - Ressources Financières : capitaux nécessaires au lancement, trésorerie pour résister à une période de faible demande, fonds pour R&D.  
        - Ressources Relationnelles : partenariats stratégiques, accès à un réseau d’influenceurs, certification par des organismes reconnus.  
        Explique pour chaque ressource pourquoi elle est critique, comment elle se combine avec les autres pour délivrer la proposition de valeur, soutenir l’adoption de l’innovation, et maintenir un avantage concurrentiel. Prends en compte la robustesse de la chaîne d’approvisionnement, la résilience face aux crises, et la propriété intellectuelle. »

        7. **Activités Clés**  
        Méta-Prompt :  
        « Décris les activités indispensables pour que le Business Model fonctionne :  
        - Développement & Innovation : R&D, amélioration continue, intégration de nouvelles fonctionnalités, veille concurrentielle, tests utilisateurs.  
        - Production & Livraison : création de contenu, mise à jour régulière, gestion du stock (si produit physique), maintenance technique, logistique.  
        - Marketing & Ventes : campagnes publicitaires, référencement, webinaires de démonstration, éducation du marché, gestion des promotions.  
        - Relation Client & Support : formation du personnel du support, chatbots, assistance multicanal, traitement des plaintes, suivi de la satisfaction.  
        - Partenariats & Négociations : recherche, signature et entretien des partenariats clés, mise en place de conditions avantageuses.  
        Intègre une perspective adaptative :  
        - Quelles activités mener si la demande fluctue fortement ?  
        - Comment réallouer les ressources si une nouvelle réglementation émerge ?  
        Justifie comment chaque activité soutient la proposition de valeur, favorise l’adoption de l’innovation, et contribue à la rentabilité globale. »

        8. **Partenaires Clés**  
        Méta-Prompt :  
        « Liste et justifie les partenaires stratégiques critiques :  
        - Fournisseurs : apportant des ressources rares, de haute qualité ou à un coût avantageux.  
        - Distributeurs : offrant un accès facilité à certains segments, réduisant les coûts d’acquisition, améliorant la visibilité.  
        - Partenaires technologiques : fournissant une infrastructure fiable, des outils d’IA performants, ou des solutions complémentaires (API, intégrations).  
        - Organismes de certification, influenceurs, médias spécialisés : augmentant la crédibilité, validant la qualité, rassurant sur l’innovation.  
        - Associations professionnelles, clusters, écosystèmes sectoriels : permettant de suivre les tendances, d’anticiper les changements réglementaires, d’échanger les bonnes pratiques.  
        Explique comment ces partenariats renforcent la proposition de valeur, améliorent la confiance du client, augmentent l’efficacité opérationnelle, réduisent les coûts ou les risques, et soutiennent la stratégie à long terme. Anticipe les risques : et si un partenaire clé fait défaut ? Quels sont les plans B ? »

        9. **Structure de Coûts**  
        Méta-Prompt :  
        « Détaille tous les coûts engendrés par les ressources, activités et partenariats clés :  
        - Coûts fixes (salaires, loyers, licences, amortissement de l’infrastructure).  
        - Coûts variables (marketing, support client, acquisition de nouveaux outils, commission aux partenaires).  
        - Coûts liés à l’innovation (R&D, tests, formations du personnel), et comment ils sont amortis dans le temps.  
        Analyse la rentabilité :  
        - Le modèle de revenus couvre-t-il ces coûts ?  
        - Quelles mesures de réduction de coûts sont possibles (automatisation, sourcing moins cher, économies d’échelle) ?  
        - Comment réagir face à des fluctuations du marché (baisse de la demande, hausse des prix des ressources) ?  
        Explique comment la structure de coûts reste alignée avec la proposition de valeur, les segments, et les moyens financiers de l’entreprise. Justifie la pérennité financière en montrant que les marges sont satisfaisantes, que le CAC est raisonnable par rapport à la CLV, et que le modèle reste rentable même en cas de stress. »

        ### Instructions Finales

        Après avoir utilisé ces méta-prompts pour chaque bloc du Business Model Canvas, effectue une dernière vérification :

        - Assure-toi que tous les blocs sont cohérents entre eux et s’alignent parfaitement avec la proposition de valeur et le parcours client.
        - Vérifie que l’innovation proposée est bien adoptable par la persona, qu’elle apporte un avantage concurrentiel durable, et que les contraintes sont gérées.  
        - Contrôle la rentabilité, la viabilité à long terme, et la flexibilité pour s’adapter aux changements de marché.
        - Ajuste les éléments (segments, prix, canaux, partenariats) si nécessaire pour améliorer la robustesse du modèle.

        Le résultat final doit être un Business Model clair, complet, et prêt à être testé ou implémenté, avec une feuille de route pour l’adoption de l’innovation et une vision claire des points de différenciation face à la concurrence.


        Enfin, fournis un récapitulatif global du Business Model, mettant en avant la logique, la cohérence, et la proposition de valeur différenciante. Indique, si possible, des chiffres (taille du marché, CAC, CLV, taux de conversion, CA projeté) pour valider la viabilité économique.""",
       
        "Autre": "Fournissez une approche générale adaptée à votre entreprise."
    }
    return metaprompts.get(type_entreprise, metaprompts["Autre"])



def obtenir_business_model(nom_entreprise, type_entreprise, montant_projet ,previousdata, rubriques, ameliorations, generation=1):
    
    """
    Interroge ChatGPT (API OpenAI) pour générer le contenu textuel
    des différents blocs du Business Model Canvas.
    'type_entreprise' peut être "PME", "Startup", "Grande Entreprise", etc.
    'previousdata' peut etre du contenue html generer precedement par chatgpt
    """
    MODEL="gpt-4o"
    MAX_TOKENS_PER_REQUEST = 150
    utilisateur = get_current_user()
    
    # Calculer les tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed
    
    
    # Récupérer le metaprompt basé sur le type d'entreprise
    metaprompt = get_metaprompt(type_entreprise)

    
    if generation == 1:
        # Première génération avec les nouvelles rubriques
        prompt = f"""
        {metaprompt}
        
        Mener la reflexions du generation du business modele sur base des indications(Méta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres données sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des données collecter sur internet 
        Génère le contenu d'un Business Model Canvas en format HTML et CSS encapsulé dans des blocs de code sans aucun autre texte pour une entreprise nommée '{nom_entreprise}'.
        Le type d'entreprise est : {type_entreprise}.
        Le montant que le proprietaire est en mesure d'investir dans le projet est:{montant_projet} veuillez considerer ce montant lors de generation de ressources clés, Canaux de Distribution, Relation client 
        
        Utilisez les données ci après(dans la rubriques) comme données collecté lors de la Phase 1 (Configuration Initiale (Entrée de Données)): {rubriques}
        Certains partie du rubriques peuvent etre vide, si c'est les cas generer les données manquantes. les chiffres entrer pour l'utilisateur doivent etre imperativement tenue en compte
        
        À faire impérativement :
        Je veux impérativement 9 blocs distincts, rédigés en français, avec les titres en gras et des listes à puces si nécessaire :
          - Partenaires clés
          - Activités clés
          - Offre (proposition de valeur)
          - Relation client
          - Segments de clientèle
          - Ressources clés
          - Canaux de distribution
          - Structure des coûts
          - Sources de revenus
        Fournissez 5 à 10 points ou éléments (phrases) par bloc pour un contenu riche et adapté, soyez concis.
        """
    else:
        # Deuxième génération (amélioration) en utilisant le BMC précédent et les nouvelles rubriques
        # Prompt ajusté sans numérotation dans les titres
        prompt = f"""
        {metaprompt}
        
        Voici les autres recommandations pour generer un business model ameliorer:{ameliorations}.
        
        Voici le business model à ameliorer generer precedement {previousdata}.
        
        Ameliorer ces business modeles modeles sur bases de metaprompt 
        Mener la reflexions du generation du business modele sur base des indications(Méta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres données sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des données collecter sur internet 
        Génère le contenu d'un Business Model Canvas en format HTML pour une entreprise nommée '{nom_entreprise}'.
        Le type d'entreprise est : {type_entreprise}.
        Le montant que le proprietaire est en mesure d'investir dans le projet est:{montant_projet} veuillez considerer ce montant lors de generation de ressources clés, Canaux de Distribution, Relation client 
        
        
        
        sachant que les données qui ont permit la generation du precedent business model sont: {rubriques}.
        si l'utlisateur a donner les données complementaires, veuillez en tenir compte dans la generation, et ca doit etre imperativement prioritaire.
        Si dans un bloque un utilisateur n'as pas donner des informations (elements), veuillez generer,
        Si l'utilisateur à donné des elements que vous juger peu, generer d'autres et les ajoutées à ce que l'utlisateur à fournit.
        
        à faire imperativement est:
        Je veux impérativement 9 blocs distincts, rédigés en français, avec les titres en gras et des listes à puces si nécessaire :
        - Partenaires clés
        - Activités clés
        - Offre (proposition de valeur)
        - Relation client
        - Segments de clientèle
        - Ressources clés
        - Canaux de distribution
        - Structure des coûts
        - Sources de revenus
        Fournis 5 à 10 points ou élements(phrases) , meme plus pour chacun afin d'avoir un contenu riche et adapté, soyez concis.
        """
    # Calculer les tokens nécessaires (entrée + réponse prévue)
    # Calculer le nombre de tokens dans l'entrée utilisateur
    
    tokens_in_input = count_tokens(nom_entreprise+""+type_entreprise+""+previousdata+""+json.dumps(rubriques) +""+ameliorations, MODEL)
    tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
    # Vérifier si l'utilisateur a assez de tokens
    if tokens_remaining < tokens_needed:
        st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
        return
    else:
        try:
            response = openai.ChatCompletion.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "Tu es un assistant expert en génération de business  et business plan."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=5000,
                temperature=0.7
            )
            html_genere = response.choices[0].message.content.strip()
            tokens_utilises = response['usage']['total_tokens']
                    # Consommer les tokens
            success, message = consommer_tokens(st.session_state['user_info'], tokens_utilises)
            
            if success:
                st.success(message)
            else:
                st.error(f"Erreur")
            return html_genere
        except Exception as e:
            st.error(f"Erreur lors de la génération du contenu  {e}")
            return ""




# ----------------------------------------------------------------------------
# 3) CHAIN OF THOUGTHTS
# ----------------------------------------------------------------------------


def get_metaprompt_chain_of_thougtht(type_chain_of_thougtht):
    """
    Retourne un metaprompt spécifique basé sur le type de Chain of Thought.
    """
    metaprompts = {
                "Creativite":"""Méta-Prompt : Chain of Thought pour la Vérification de Cohérence 
                1. Contexte et Typologie (B2B / B2C / B2B2C)
                Identifier le type de cible : B2B, B2C ou B2B2C.
                Méthode : Vérifier si la solution s’adresse à des entreprises, à des clients finaux, ou à une chaîne mixte.
                Vérification : Le document ou la matrice de conception mentionne-t-il explicitement la nature de la relation commerciale ?
                Vérifier l’adaptation au contexte choisi (enjeux de volume, personnalisation, cycle de décision…).
                Suggestion si incohérence :
                Ajuster la persona (profil acheteur / utilisateur) pour qu’elle reflète la bonne cible (ex. acheteur d’entreprise, consommateur final…).
                2. Analyse de la Persona avec l’Approche Sinus Milieu
                Identifier le(s) milieu(x) Sinus le plus proche de la persona
                Objectif : Situer la persona sur un axe socioculturel précis (ex. “Traditionnel”, “Hédoniste”, “Moderne matérialiste”, “Pragmatique adaptatif”…).
                Méthode : Analyser les valeurs, le style de vie, les motivations profondes décrites dans la fiche persona.
                Vérification : La persona est-elle cohérente avec un segment Sinus (pouvoir d’achat, rapport à l’innovation, etc.) ?
                Adapter la proposition de valeur aux valeurs du milieu Sinus
                Objectif : Vérifier la compatibilité entre les valeurs du milieu (ex. sécurité, hédonisme, éthique, recherche de statut) et les bénéfices proposés.
                Suggestion si incohérence :
                Revoir la concurrence : y a-t-il un rival déjà bien ancré dans ce milieu ?
                Réajuster la persona : la persona peut appartenir à un autre milieu Sinus (ex. plus “progressiste” ou plus “traditionnel”) si les données terrain le justifient.
                Chiffrer la taille du milieu Sinus visé
                Objectif : Évaluer la portée de la cible (quel % de la population ou du marché total).
                Suggestion d’itération :
                Si ce milieu représente moins de 5% du marché, envisager un ciblage complémentaire (deuxième milieu Sinus proche) ou affiner la proposition.
                3. Arbre à Problèmes et Pertinence du Besoin
                Racines du Problème : Les causes profondes du besoin sont-elles bien définies ?
                Méthode : Faire correspondre chaque branche de l’arbre à un aspect concret de la “douleur client” (en lien avec les motivations Sinus).
                Vérification : Les causes sont-elles validées par des retours terrain (sondages, interviews) ?
                Si incohérence (branches peu justifiées, besoin trop faible) :
                Liste de points à creuser :
                Ajuster la persona (réévaluer ses motivations ou son budget) si la “douleur” n’est pas si forte pour ce milieu Sinus.
                Vérifier les contraintes (techniques, légales) : la solution proposée pourrait être trop ambitieuse ou sous-dimensionnée.
                Examiner la concurrence : un acteur répond-il déjà mieux à ce problème ?
                4. Contraintes Limitantes (Légales, Techniques, Budgétaires)
                Clarté et réalisme des contraintes
                Méthode : Associer chaque contrainte (budgétaire, légale…) aux fonctionnalités clés.
                Vérification : Compatibilité avec le milieu Sinus visé (pouvoir d’achat, ouverture à l’innovation, etc.).
                Si incohérence (conflit entre besoin et contrainte) :
                Liste d’actions :
                Réduire la portée du produit si le budget dépasse trop la capacité de paiement de la cible Sinus.
                Ajuster la persona : si les utilisateurs ciblés ont un budget plus élevé, resegmenter le milieu Sinus.
                5. Analyse de la Concurrence
                Comparaison fonctionnelle et tarifs
                Méthode : Vérifier qu’il existe un tableau comparant fonctionnalités, prix, positionnement.
                Vérification : La concurrence vise-t-elle le même milieu Sinus ou un autre ? Les arguments marketing sont-ils comparables ?
                Si incohérence (prix trop élevé par rapport au milieu Sinus, attributs redondants) :
                Liste d’investigations :
                Revoir la persona (et donc le milieu Sinus) : surestime-t-on le pouvoir d’achat ?
                Ajuster l’arbre à problèmes : la proposition de valeur est-elle unique ?
                Revoir les contraintes budgétaires (peut-on proposer une version allégée pour être compétitif ?).
                6. Caractéristiques du Produit / Service / Processus
                Clarté de la Proposition de Valeur
                Méthode : Établir la correspondance “pain point -> fonctionnalité -> bénéfice” en tenant compte du style de consommation du milieu Sinus.
                Vérification : Les fonctionnalités répondent-elles réellement aux valeurs/besoins du milieu ?
                Priorisation (Kano ou autre)
                Vérification : Les “must-have” sont-ils alignés avec les besoins profonds de la cible Sinus ? Les “attractives” sont-elles pertinentes culturellement ?
                Si incohérence :
                Revoir la concurrence (les must-have sont-ils déjà standardisés chez les concurrents ?).
                Retester la persona (motivation vs. nouvelles options de solution).
                7. Évaluation de la Cohérence Globale (Inclusion Sinus Milieu)
                Catégories d’évaluation

                On évalue 6 catégories :
                Persona & Sinus Milieu
                Arbre à Problèmes
                Contraintes
                Concurrence
                Caractéristiques du produit (valeur ajoutée, priorisation)
                Adéquation Sinus Milieu (styles de vie, valeurs, budget)
                Noter chaque catégorie de 1 à 5

                1 = Très faible cohérence
                5 = Cohérence complète
                Calcul du Taux de Cohérence

                Taux de coheˊrence=6×5∑(notes sur 5)×100%
                Exemple : Si la somme des notes est 24/30, la cohérence = (24/30)×100% = 80%.
                Interprétation :

                ≥80% : Cohérence satisfaisante. Ajustements mineurs.
                60-79% : Cohérence moyenne. Des interventions ciblées sont nécessaires.
                ≤59% : Cohérence insuffisante. Il faut prioritairement revoir la persona/Sinus Milieu ou l’arbre à problèmes.
                8. Instructions en cas de Cohérence Faible ou Moyenne
                Ajuster la Persona & le Milieu Sinus
                Ex. : Affiner le revenu, le degré d’ouverture à l’innovation, la sensibilité aux arguments “status / sécurité / fun”.
                Compléter ou Corriger l’Arbre à Problèmes
                Ex. : Ajouter des données terrain, réévaluer l’impact réel du problème, distinguer causes primaires/secondaires.
                Modifier les Contraintes
                Ex. : Revoir le budget ou la faisabilité technique si en contradiction avec les attentes du milieu Sinus.
                Actualiser la Concurrence
                Ex. : Mettre à jour le benchmark (nouveaux entrants, changements de prix) ou réévaluer le positionnement.
                9. Synthèse et Relance du Processus de Créativité
                Mettre en œuvre les Ajustements Autorisés
                Persona (et milieu Sinus), Arbre à problèmes, Contraintes, Concurrence.
                Recalculer la Note de Cohérence
                Vérifier si elle dépasse maintenant 80%.
                Boucler la Mise à Jour
                Valider que les fonctionnalités et la proposition de valeur restent en accord avec les nouvelles informations.
                Obtenir la Version Finalisée
                Énoncer clairement la proposition de valeur, la cible Sinus, les contraintes, et la position concurrentielle mises à jour.
                """,
                
                "BM":"""Chain of Thought pour Analyser et Vérifier le Business Model
                1. Identification des 9 blocs du Business Model
                Pour rappel, les 9 blocs du Business Model Canvas sont :

                Segments de clientèle
                Proposition de valeur
                Canaux de distribution
                Relations client
                Flux de revenus
                Ressources clés
                Activités clés
                Partenaires clés
                Structure de coûts
                But : S’assurer que chaque bloc est clairement décrit et chiffré si nécessaire (équipements, ressources humaines, coûts estimés, etc.).

                2. Vérifier la cohérence entre les “résultats de créativité” et le Business Model
                Alignement avec la Persona

                Méthode : Vérifier que le segment de clientèle (bloc 1) correspond à la persona issue de la créativité (habitudes, pouvoir d’achat, milieu socioculturel, etc.).
                Si incohérence :
                Revoir la Persona (ex. si le segment de clientèle est trop large ou trop restreint par rapport aux insights de la phase de créativité).
                Ajuster la Proposition de valeur (bloc 2) pour qu’elle répondre aux vrais besoins de la persona.
                Intégration de l’Arbre à problèmes / Empathy Map

                Méthode : Chaque douleur (pain point) identifiée doit trouver une réponse claire dans la Proposition de valeur et dans les Activités clés.
                Si incohérence :
                Compléter ou réviser la Proposition de valeur (bloc 2) pour couvrir les problèmes fondamentaux repérés.
                Vérifier que les Activités clés (bloc 7) et Ressources clés (bloc 6) permettent concrètement de résoudre ces problèmes.
                Priorisation des fonctionnalités (Kano) et impact sur le BM

                Méthode : Les fonctionnalités “Must-be” (basique) et “One-Dimensional” (attendues) doivent apparaître en priorité dans les Activités clés et Ressources clés. Les fonctionnalités “Attractives” peuvent être prévues mais doivent être budgétées avec prudence.
                Si incohérence :
                Vérifier la Structure de coûts (bloc 9) pour éviter un investissement disproportionné dans des fonctionnalités “attractives” non essentielles.
                Réduire le périmètre initial du produit (MVP) pour limiter les investissements et assurer une rentabilité progressive.
                3. Analyse et Validation de chaque Bloc
                Passons en revue les 9 blocs pour un contrôle de cohérence plus approfondi.

                3.1. Segments de clientèle (Bloc 1)
                Objectif : Vérifier la précision du ou des segments cibles (B2C, B2B, B2B2C, sous-segments, etc.).
                Vérification :
                Les segments sont-ils clairement reliés à la persona ?
                Les volumes (nombre d’entreprises, nombre de clients particuliers) et la localisation sont-ils estimés ?
                Si incohérence :
                Points à chercher :
                Ajuster la persona (profil, pouvoir d’achat, localisation).
                Vérifier l’arbre à problèmes (la douleur est-elle trop différente d’un segment à l’autre ?)
                3.2. Proposition de valeur (Bloc 2)
                Objectif : S’assurer que la solution (produit/service/processus) décrit bien comment elle résout les problèmes identifiés.
                Vérification :
                Texte clair et concis : “Nous résolvons [problème X] pour [segment Y] grâce à [caractéristique unique Z].”
                Inclus les retours de la phase de créativité (pains/gains).
                Si incohérence :
                Points à préciser :
                Manque d’éléments différenciateurs ?
                Proposition trop vague vs. attentes concrètes de la persona ?
                Alignement insuffisant avec la priorisation Kano (focus sur des features non-essentielles).
                3.3. Canaux de distribution (Bloc 3)
                Objectif : Vérifier comment la proposition de valeur est livrée au client (physique, digital, vente directe, etc.).
                Vérification :
                Les canaux sont-ils cohérents avec la persona (milieu urbain/rural, habileté numérique, budget pub) ?
                Est-ce que le coût logistique (transport, plateforme e-commerce) est estimé ?
                Si incohérence :
                Points à rechercher :
                Persona mal définie (ex. utilisation d’un canal digital alors que la cible est peu connectée).
                Contradiction avec la structure de coûts (oublier les frais de distribution).
                3.4. Relations client (Bloc 4)
                Objectif : Évaluer la stratégie de gestion de la relation (service client, SAV, self-service, communauté).
                Vérification :
                Les formes de relation sont-elles adaptées au volume de clients (B2B vs. B2C) et aux attentes repérées (arbre à problèmes) ?
                Si incohérence :
                Points à éclaircir :
                Besoin de SAV important ? Est-il budgété ?
                La concurrence propose-t-elle un meilleur accompagnement ?
                3.5. Flux de revenus (Bloc 5)
                Objectif : Contrôler le modèle de monétisation (abonnement, vente à l’unité, location, freemium…).
                Vérification :
                Prix fixés en accord avec le pouvoir d’achat du segment, le positionnement (premium/entrée de gamme), et la concurrence.
                Les montants de CA prévisionnels sont-ils plausibles (basés sur un volume de ventes réaliste) ?
                Si incohérence :
                Points à corriger :
                Pricing trop ambitieux vs. revenu moyen de la persona.
                Sous-estimation du coût d’acquisition client (marketing, pub).
                3.6. Ressources clés (Bloc 6)
                Objectif : Lister précisément les ressources (humaines, matérielles, financières, intellectuelles) nécessaires.
                Vérification :
                Les équipements sont-ils nommés (ex. “2 imprimantes 3D de marque X”, “licence CRM Y”…) ?
                Le budget RH (salaires, compétences requises) est-il cohérent avec la taille du projet ?
                Si incohérence :
                Points à rechercher :
                Manque d’équipement critique (laboratoire, logiciel, entrepôt).
                Contradiction avec le modèle Kano : trop de ressources pour des fonctionnalités “attractives” mais non prioritaires.
                3.7. Activités clés (Bloc 7)
                Objectif : Quelles sont les actions indispensables pour créer et délivrer la proposition de valeur (production, R&D, marketing, maintenance…) ?
                Vérification :
                Activités en adéquation avec les fonctionnalités “must-have” et “one-dimensional” (Kano).
                Planning / roadmap (et/ou phasage MVP) permettant de maîtriser les coûts.
                Si incohérence :
                Points à clarifier :
                Activités non nécessaires (gaspillage de ressources).
                Absence de marketing / support alors que l’arbre à problèmes montre un besoin de forte éducation du marché.
                3.8. Partenaires clés (Bloc 8)
                Objectif : Identifier tous les acteurs externes qui renforcent la chaîne de valeur (fournisseurs, distributeurs, co-développeurs).
                Vérification :
                Contrats (ou MoU) mentionnés ? Rôle précis de chaque partenaire ?
                Délégation d’activités (sous-traitance) clairement estimée dans la structure de coûts ?
                Si incohérence :
                Points à vérifier :
                Un partenaire existe-t-il réellement ou n’est-ce qu’une hypothèse non validée ?
                Y a-t-il un risque de dépendance majeure ?
                3.9. Structure de coûts (Bloc 9)
                Objectif : Dresser la liste de tous les coûts (fixes, variables), incluant les équipements, salaires, marketing, etc.
                Vérification :
                Alignement avec les Ressources clés et les Activités clés.
                Ratio coût/revenu : la marge est-elle suffisante pour être viable ?
                Si incohérence :
                Points à rectifier :
                Oubli de certains postes de dépenses (maintenance, licences logicielles).
                Sur-investissement dans des fonctionnalités “attractives” au détriment de la rentabilité court terme.
                4. Établir un Score de Cohérence (en %)
                Méthode de notation : Évaluer chacun des 9 blocs sur une échelle de 1 à 5 (1 = faible cohérence, 5 = parfaite cohérence).
                Exemple de calcul : Taux de coheˊrence=9×5∑(notes sur 5 pour chaque bloc)×100%
                Si la somme des notes est 35/45 (9 blocs × 5 = 45 points), alors la cohérence = (35/45)×100 = 78%.
                Interprétation :
                ≥80% : Cohérence satisfaisante ; quelques optimisations mineures.
                60-79% : Cohérence moyenne ; réviser les points faibles identifiés.
                ≤59% : Cohérence insuffisante ; retravailler fortement Persona / Proposition de Valeur / Coûts.
                5. Guidelines en cas d’Incohérences
                Si le Taux de cohérence est faible ou moyen, proposez des pistes d’action, en vous limitant (selon les règles) à :

                Persona
                Ajuster la segmentation, le pouvoir d’achat, le comportement d’achat.
                Arbre à problèmes
                Réexaminer les causes profondes, vérifier si la proposition de valeur répond vraiment aux souffrances majeures.
                Contraintes
                Réduire l’ambition technique ou le scope initial pour réduire les coûts (MVP).
                Concurrence
                Actualiser le benchmark, revoir l’argumentaire de différenciation ou la tarification.
                But : Rendre le Business Model réaliste et conforme aux insights de la créativité (notamment le modèle de Kano, pour éviter les surinvestissements inutiles).

                6. Optimisation du Business Model
                Limiter les investissements superflus
                Approche Kano : se concentrer d’abord sur les fonctionnalités “must-have” et “one-dimensional”.
                Clarifier le chiffrage
                Exemple : Nommer précisément les équipements (“Imprimante 3D Stratasys, 2 machines CNC, 3 licences CRM Salesforce, etc.”) et leur coût.
                Budgéter la masse salariale (nombre de postes, salaires annuels, charges sociales).
                Planifier une stratégie MVP
                Objectif : Tester le marché avec un investissement minimal, prouver la traction avant de déployer des fonctionnalités “attractives”.
                7. Conclusion de la Chain of Thought
                Application :
                Parcourir chaque bloc du Business Model en vérifiant l’intégration des résultats de la créativité (persona, pains, priorisation Kano).
                Noter la cohérence sur 9 blocs (sur 5 points chacun) pour obtenir un score en %.
                Si le score <80%, orienter l’utilisateur vers des actions correctrices précises (sur la persona, l’arbre à problèmes, les contraintes, la concurrence).
                Mettre à jour les coûts, ressources et planning de manière à obtenir un Business Model viable et cohérent.
                """,
                
                "Analisis":"""Chain of Thought pour Analyser les Analyses Financières
                
                1. Vérification du Contexte et des Hypothèses Initiales
                Alignement avec le Business Model
                Objectif : Contrôler que les hypothèses financières (volumes de vente, prix, coûts fixes/variables, etc.) découlent du Business Model validé.
                Vérification :
                Les segments de clientèle et la proposition de valeur (ex. prix moyen, cycle d’achat) sont-ils bien traduits en chiffres ?
                Les hypothèses de volume/ventes sont-elles cohérentes avec le marché ciblé et la concurrence ?
                Référence à la Phase de Créativité
                Objectif : Vérifier que la connaissance des pains et gains clients, issue de la phase de créativité (persona, arbre à problèmes), se reflète dans la politique tarifaire et dans la structure de coûts.
                Si incohérence détectée :
                Liste d’informations à chercher :
                Revoir la persona (niveau de revenu, budget moyen, fréquence d’achat) ;
                Recalibrer la taille du marché ou la part de marché visée ;
                Vérifier la stratégie (B2B/B2C) vs. les hypothèses de prix ou volumes.
                2. Identification et Chiffrage des Ressources
                Noms des Équipements et Ressources
                Objectif : L’analyse financière doit mentionner clairement les équipements (machines, serveurs, véhicules, etc.) et ressources humaines (profils, salaires).
                Vérification :
                Existe-t-il une liste détaillée des ressources (ex. 3 postes de développeurs, 1 local de 50 m², 2 machines CNC…) ?
                Les coûts associés (achat, location, maintenance, salaires) sont-ils chiffrés de manière réaliste (devis, références marché) ?
                Optimisation via le Modèle de Kano
                Objectif : Vérifier que l’on ne surinvestit pas dans des ressources pour des fonctionnalités “indifférentes” ou “mineures”.
                Si incohérence détectée :
                Liste d’actions :
                Réévaluer les fonctionnalités “attractives” vs. “indifférentes” et limiter l’achat d’équipements s’il s’agit d’une option non prioritaire ;
                Réattribuer le budget aux “must-have” identifiés dans le modèle de Kano.
                3. Plan de Trésorerie et Financement
                Chronologie des Flux
                Objectif : Vérifier que les dépenses (investissements initiaux, coûts récurrents) et les encaissements (ventes, subventions, levées de fonds) sont planifiés dans un calendrier réaliste.
                Vérification :
                Y a-t-il un tableau mensuel/trimestriel des flux de trésorerie (cash in / cash out) ?
                Le besoin en fonds de roulement (BFR) est-il estimé en tenant compte des délais de paiement (B2B) ou de la saisonnalité (B2C) ?
                Plan de Financement
                Objectif : Confirmer la source des financements (apport personnel, prêt bancaire, investisseurs, crowdfunding) et leur correspondance avec les besoins chiffrés.
                Si incohérence détectée :
                Liste de vérifications :
                Les montants alloués couvrent-ils vraiment l’investissement nécessaire à la mise en place des fonctionnalités “must-have” ?
                Les investissements “facultatifs” (features attractives mais secondaires) peuvent-ils être reportés ou financés plus tard ?
                4. Compteur de Résultat Prévisionnel (CRP) et Structure des Coûts
                Estimation des Ventes
                Objectif : S’assurer que les recettes attendues sont cohérentes avec les hypothèses de volume/prix et l’étude de marché.
                Vérification : Les quantités annuelles (ou mensuelles) de ventes sont-elles crédibles (ex. 10 000 ventes/an) au regard de la cible ?
                Charges Variables et Fixes
                Objectif : Vérifier le niveau de charges (matières premières, sous-traitance, marketing, etc.).
                Si incohérence détectée :
                Liste d’actions :
                Revoir la tarification si la marge unitaire est trop faible ou trop élevée ;
                Réduire certaines charges liées aux fonctionnalités jugées “non essentielles” (Kano) pour ne pas alourdir les coûts.
                5. Bilan Prévisionnel et Amortissements
                Clarté sur les Équipements
                Objectif : Retrouver, dans l’actif du bilan, les biens d’équipement cités (ordinateur, serveurs, machines).
                Vérification : Les amortissements sont-ils calculés sur une durée réaliste ? Les taux d’amortissement sont-ils cohérents avec les normes comptables (ex. 3 ans, 5 ans) ?
                Patrimoine et Fonds Propres
                Objectif : S’assurer que la structure du financement (capitaux propres vs. dettes) est en adéquation avec la capacité de l’entreprise à rembourser.
                Si incohérence détectée :
                Liste de vérifications :
                Les capitaux propres initiaux du porteur de projet sont-ils sous-estimés ou surestimés ?
                Les immobilisations superflues (cf. Kano) peuvent-elles être allégées ou différées ?
                6. Seuil de Rentabilité et Point Mort
                Calcul du Seuil de Rentabilité
                Objectif : Identifier le chiffre d’affaires minimal à atteindre pour couvrir tous les coûts (fixes + variables).
                Vérification : La méthode de calcul (marge sur coûts variables ou autre) est-elle explicitée ?
                Temps d’Atteinte du Point Mort
                Objectif : Vérifier la durée nécessaire pour être rentable (6 mois, 1 an, 2 ans ?).
                Si incohérence détectée :
                Liste de questions :
                Les projections de ventes sont-elles trop optimistes ?
                Existe-t-il une solution pour réduire les coûts (et donc le point mort) en sélectionnant moins de fonctionnalités ?
                7. Analyse de Sensibilité et Scénarios
                Scénarios Pessimiste / Moyen / Optimiste
                Objectif : Mesurer la robustesse du projet face à des variations (ventes 20% plus faibles, coûts 10% plus élevés…).
                Vérification : L’analyse financière présente-t-elle 2 ou 3 scénarios ? Les impacts sur le cash-flow ou la rentabilité sont-ils clairs ?
                Mesures d’Ajustement
                Objectif : Anticiper des plans B (ex. reporter un investissement, réduire les effectifs).
                Si incohérence détectée :
                Liste de contrôles :
                Les ressources “non prioritaires” (Kano) peuvent-elles être réduites en cas de scénario pessimiste ?
                Les partenariats (coûteux) sont-ils flexibles ?
                8. Cohérence avec les Ressources du Porteur de Projet
                Vérifier l’Adéquation avec le Profil du Porteur
                Objectif : Comparer les besoins financiers totaux (investissements, BFR) avec les ressources réelles du porteur (apports, capacités d’endettement).
                Vérification :
                Les apports personnels sont-ils en phase avec le plan de financement ?
                L’endettement prévu (taux, durée) est-il réaliste au vu du niveau de risque ?
                Si incohérence :
                Liste d’actions :
                Réduire les dépenses initiales en priorisant les fonctionnalités via Kano.
                Chercher un co-investisseur ou des subventions.
                9. Calcul du Taux de Cohérence
                Pour aboutir à un score global, on peut évaluer 7 catégories financières (chacune notée de 1 à 5) :

                Hypothèses de Chiffre d’Affaires (alignement marché, volumes, prix)
                Coûts Variables et Fixes (réalisme, cohérence avec ressources)
                Plan de Trésorerie (flux mensuels/trimestriels, BFR, timing)
                Financement et Fonds Propres (source, niveau d’endettement)
                Amortissements et Bilan Prévisionnel (durées, coûts d’équipement)
                Seuil de Rentabilité / Point Mort (réalisme, timing)
                Analyse de Sensibilité (scénarios, options de réduction via Kano)
                Notation : 1 = très faible, 5 = très bonne cohérence.
                Formule : Taux de coheˊrence=7×5∑(notes sur 5)×100%
                Exemple : si la somme des notes est 26/35, le taux = (26/35)×100% ≈ 74%.
                Interprétation :

                ≥80% : Analyses financières cohérentes. Ajustements mineurs.
                60-79% : Cohérence moyenne. Des ajustements ciblés sont requis.
                ≤59% : Cohérence insuffisante. Il faut revoir en profondeur les hypothèses, la structure des coûts ou les ressources du porteur.
                10. Suggestions si Cohérence Insuffisante
                Hypothèses de Ventes
                Revoir la persona et la taille de marché. Valider la disposition à payer via un sondage rapide (budget ~2% du total).
                Modèle de Kano
                Supprimer ou reporter les fonctionnalités “indifférentes” pour réduire l’investissement initial.
                Ressources / Équipements
                Lister clairement les besoins. Vérifier s’il existe une solution moins coûteuse (ex. location vs. achat).
                Financement
                Vérifier si les apports et les prêts sont suffisants pour couvrir la trésorerie pendant x mois.
                Réalisme des Marges
                S’assurer que la marge unitaire permet de couvrir les frais fixes à moyen terme.
                11. Conclusion et Itération Finale
                Mettre à Jour le Dossier Financier
                Ajuster les données (prix, volume, ressources, plan de trésorerie) en fonction des incohérences détectées.
                Recalculer le Taux de Cohérence
                Vérifier si le score s’améliore (objectif : ≥80%).
                Validation Globale
                S’assurer que la phase de créativité (pains/gains, persona), le business model (segments, ressources, coûts) et les analyses financières (réalistes, chiffrées, optimisées) forment un ensemble cohérent.
                Utilisation de cette Chain of Thought
                Appliquer les étapes (1 à 11) pour analyser l’existant (documents financiers, hypothèses de recettes, coûts, plan de financement, etc.).
                Attribuer une note de 1 à 5 à chaque catégorie (hypothèses de CA, plan de trésorerie, etc.).
                Calculer le Taux de Cohérence final en %.
                Si le score est <80%, suivre la liste de suggestions (priorisation via Kano, révision de la persona, etc.).
                Revalider jusqu’à ce que les analyses financières soient suffisamment réalistes, claires et ajustées aux ressources du porteur et aux stratégies du Business Model.
                """,
                
                "BP":""" Chain of Thought pour Analyser le Plan d’Affaires Global
                1. Vérification du Cadre Global et du Contexte du Projet
                Clarifier la portée du plan d’affaires
                Objectif : Savoir si ce plan couvre un nouveau produit, un service, une activité globale ou l’ensemble de l’entreprise.
                Vérification : L’introduction ou le sommaire mentionne-t-il explicitement le périmètre (ex. “Développement d’une nouvelle gamme B2C” ou “Plan global de l’entreprise”)?
                Alignement avec les analyses macro-micro
                Objectif : S’assurer que les diagnostics PESTEL & PORTER (environnement politique, économique, social, technologique, environnemental, légal, ainsi que concurrence) sont bien pris en compte.
                Vérification : Les menaces et opportunités identifiées (ex. cadre réglementaire, forces concurrentielles) sont-elles intégrées à la stratégie finale du plan ?
                But : Confirmer que le plan d’affaires s’ancre dans un contexte clairement défini et prend en compte l’environnement macro et micro.

                2. Validation de la Cohérence Générale avec les Phases Précédentes
                Résultats de la Créativité
                Objectif : Vérifier que les insights clés (persona, arbre à problèmes, proposition de valeur, priorisations via Kano) sont bien retranscrits dans le plan.
                Vérification : Les pains et gains clients, les fonctionnalités must-have ou attractive sont-elles présentes dans la description du produit/service ?
                Business Model et Analyses Financières
                Objectif : Contrôler la cohérence entre le Business Model (9 blocs) et les documents financiers (hypothèses de ventes, structure de coûts, plan de financement).
                Vérification : Les chiffres mentionnés (CA prévisionnel, budget marketing, coûts de production) correspondent-ils à ce qui apparaît dans les annexes financières ?
                But : Éviter tout décalage entre les blocs de création de valeur et les tableaux financiers.

                3. Vérification de la Planification et des Objectifs SMART
                Nature et Formulation des Objectifs
                Objectif : Les objectifs (de vente, de rentabilité, de développement…) sont-ils Spécifiques, Mesurables, Atteignables, Réalistes et Temporellement définis ?
                Vérification : Existe-t-il un plan d’action avec des indicateurs quantitatifs (ex. “Atteindre 100k€ de CA au 2e trimestre” ou “Conquérir 10% de part de marché local en un an”) ?
                Cohérence avec les Paramètres du Projet
                Objectif : Vérifier que la taille de l’entreprise, le niveau de ressources, et le stade de développement (start-up, PME, scale-up) sont en accord avec ces objectifs.
                Si incohérence détectée :
                Liste d’actions :
                Réduire ou échelonner l’ambition si les ressources ne suffisent pas.
                Détailler plus clairement le calendrier de réalisation (ex. plan Gantt).
                But : S’assurer que les objectifs ne sont pas seulement ambitieux, mais également réalistes et planifiés.

                4. Vérification des Informations Critiques et Identifiants de l’Entreprise
                Informations Légales & Forme Juridique
                Objectif : S’assurer que le plan inclut la forme juridique (SARL, SA, SAS, etc.), le numéro RCCM, le numéro d’Impôt, le numéro d’identification nationale, etc.
                Vérification : Ces informations sont-elles présentes ? Si l’entreprise n’est pas encore créée, y a-t-il une mention sur la forme prévue et les formalités à venir ?
                Identité Bancaire et Dates Clés
                Objectif : Vérifier que le plan renseigne (ou prévoit) les comptes bancaires, la date de création, les évolutions majeures (levées de fonds, changements de statuts).
                Si incohérence détectée :
                Liste de données manquantes :
                Absence d’IBAN ou de compte dédié ?
                Prochaines échéances d’immatriculation non précisées ?
                But : Garantir que toutes les mentions légales et coordonnées nécessaires sont présentes pour un dossier professionnel et complet.

                5. Vérification de la Section “Marché, Concurrence et Marketing”
                Précisions sur le Marché
                Objectif : Confirmer la taille du marché, la segmentation, la dynamique (croissance, saisonnalité) et les preuves associées (études, sondages).
                Vérification : Les données sont-elles réelles ou simplement estimées ? Y a-t-il des sources citées ?
                Concurrence et Stratégie Marketing
                Objectif : Vérifier que le plan mentionne les acteurs concurrents (directs, indirects) et la stratégie pour se positionner (prix, distribution, communication).
                Si incohérence détectée :
                Liste de questions :
                Les concurrents majeurs identifiés lors de l’analyse de l’environnement (Porter) apparaissent-ils ici ?
                Les 4P (Produit, Prix, Place, Promotion) sont-ils détaillés ?
                But : S’assurer que la vision du marché est fondée, et que la stratégie marketing est adaptée et étayée.

                6. Vérification de la Section “Organisation, Personnel et Associés”
                Présentation du Personnel Clé
                Objectif : Vérifier que le plan nomme les dirigeants, les fondateurs, les associés stratégiques, leurs rôles et leurs compétences.
                Vérification : Les CV, expériences, ou atouts sont-ils brièvement évoqués ? Les postes manquants (recrutements à venir) sont-ils identifiés ?
                Coût de la Main-d’Œuvre et Fiches de Poste
                Objectif : S’assurer que le plan financier inclut le coût du personnel et la répartition des rôles (ex. 2 commerciaux, 1 community manager, etc.).
                Si incohérence détectée :
                Liste d’actions :
                Ajouter une section décrivant chaque poste clé, son salaire, ses missions.
                Vérifier la formation ou l’accompagnement éventuel (coûts supplémentaires ?).
                But : Garantir la clarté sur l’organisation humaine du projet et la réalité des coûts salariaux.

                7. Vérification des Aspects “Production et Aménagements”
                Production / Prestations
                Objectif : Détailler les moyens de production, la chaîne d’approvisionnement, la logistique.
                Vérification : Les équipements (machines, locaux, serveurs) décrits sont-ils en cohérence avec le Business Model et les analyses financières (CAPEX/OPEX) ?
                Conformité Juridique / Réglementaire
                Objectif : S’assurer que le plan intègre les éventuels certificats, licences, normes à respecter selon le secteur (alimentaire, cosmétique, numérique, etc.).
                Si incohérence détectée :
                Liste de vérifications :
                Les risques réglementaires identifiés dans la PESTEL sont-ils traités (DGCCRF, autorisations, douanes) ?
                Les coûts induits (inspection, homologation) sont-ils budgétisés ?
                But : Vérifier qu’aucun angle mort technique ou réglementaire ne compromet l’ensemble.

                8. Analyse des Risques et Cohérence avec PESTEL & PORTER
                Liste des Risques Identifiés
                Objectif : Le plan d’affaires inclut-il une liste de risques (concurrence agressive, évolution réglementaire, évolution technologique, etc.) ?
                Vérification : Les facteurs issus de la PESTEL (ex. instabilité politique, nouvelles lois écologiques) et de la concurrence (5 Forces de Porter) sont-ils mentionnés ?
                Stratégie de Gestion des Risques
                Objectif : Les mesures pour atténuer/éviter ces risques sont-elles définies (plans B, assurance, R&D alternée) ?
                Si incohérence détectée :
                Liste d’ajustements :
                Ajouter un tableau de risques croisant la probabilité et l’impact.
                Prévoir des scénarios (pessimiste/modéré/optimiste) pour la croissance ou la réglementation.
                But : Garantir que la stratégie tient compte des risques réels, en lien direct avec l’environnement macro et micro analysé.

                9. Réalité et Faisabilité du Plan (Cohérence Finale)
                Appréciation Globale
                Objectif : Vérifier la faisabilité au regard des ressources (humaines, financières, techniques) et du calendrier proposé.
                Méthode : Relire la synthèse et les annexes (tableaux financiers, organisation, partenariats) pour un contrôle croisé.
                Estimation du Taux de Cohérence
                Proposition : Noter 7 critères (marché & concurrence, organisation & personnel, production & aménagements, aspects juridiques, objectifs SMART, analyses financières, intégration macro-micro) de 1 à 5.
                Formule : Taux de coheˊrence=7×5∑(notes sur 5)×100%
                Exemple : Si la somme atteint 30/35, on obtient (30/35) × 100% ≈ 85%.
                But : Fournir une mesure synthétique pour évaluer la maturité globale du plan d’affaires.

                10. Proposition d’Ajustements ou Corrections
                (En cas d’incohérence ou de note <80%.)

                Étape : Marché et Concurrence
                Amélioration : Apporter des données chiffrées issues d’études récentes ; préciser la stratégie face aux concurrents directs.
                Étape : Personnel et Associés
                Amélioration : Clarifier l’identité des associés, le rôle de chacun, la politique de rémunération.
                Étape : Analyse des Risques
                Amélioration : Ajouter un tableau de suivi des risques (facteur, impact, probabilité, solutions).
                Étape : Formalisation et Identifiants
                Amélioration : Indiquer la forme juridique, le RCCM, le numéro d’Impôt, etc. ou préciser la procédure pour les obtenir.
                Étape : Objectifs SMART
                Amélioration : Rendre les objectifs plus quantifiables (ex. “Atteindre 1 000 utilisateurs actifs d’ici 6 mois”), lier chaque objectif à un responsable et une deadline.
                But : Permettre à l’utilisateur de corriger précisément les points faibles, afin de finaliser un plan d’affaires complet et crédible.

                Synthèse de la Méthodologie
                Balayage Général : S’assurer que le plan reprend toutes les étapes précédentes (créativité, BM, finances) et les met en cohérence.
                Contrôle de la Complétude : Vérifier la présence des informations légales, des données de marché, de la structure organisationnelle, etc.
                Analyse de la Cohérence : Checker la cohésion interne via la notation (voir Taux de cohérence).
                Ajustements : Proposer une liste structurée d’améliorations à réaliser si le score de cohérence est insuffisant.
                De cette manière, on assure que le Business Plan final est :

                Réaliste : en phase avec les moyens (humains, financiers, techniques).
                Conforme aux Résultats de la Créativité : propositions de valeur cohérentes, fonctionnalités justifiées.
                Aligné sur l’Environnement Macro-Micro : PESTEL et PORTER intégrés.
                Fiable : avec des objectifs SMART, des mentions légales, une organisation claire, une analyse des risques documentée. """,
                "Autre": ""
    }
    return metaprompts.get(type_chain_of_thougtht, metaprompts["Autre"])



def chain_of_thougtht(type_chain_of_thougtht, montant_projet, nom_entreprise, previousdata, rubriques, Analisis, previousbp, generation=1):
    
    """
    Interroge ChatGPT (API OpenAI) pour générer le contenu textuel
    des de verifier la coherence, generer precedement par chatgpt. 
    """
    
    MODEL="gpt-4o"
    MAX_TOKENS_PER_REQUEST = 150
    utilisateur = get_current_user()
    
    # Calculer les tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    # Récupérer le metaprompt basé sur le type d'entreprise
    metaprompt = get_metaprompt_chain_of_thougtht(type_chain_of_thougtht)

   
    if generation == 1:
        # Première génération avec les nouvelles rubriques
        prompt = f"""
        {metaprompt}
        voici les données à considerer: {rubriques}
        Format :
            - Résumer l’analyse en expliquant chaque point de vérification.
            - Dire si c'est coherent ou pas.
            - Proposer un plan d’action (liste structurée).
        """
    elif generation == 2:
        # Deuxième génération (amélioration) en utilisant le BMC précédent et les nouvelles rubriques
        # Prompt ajusté sans numérotation dans les titres
        prompt = f"""
        voici le metaprompt:{metaprompt}
        
        
        Voici le business model genener precedement: {previousdata}
        Ameliorer ces business modeles modeles sur bases de metaprompt
        Mener la reflexions du generation du business modele sur base des indications(Méta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres données sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des données collecter sur internet 
        Génère le contenu d'un Business Model Canvas en format HTML pour une entreprise nommée '{nom_entreprise}'.
        
        sachant que les données qui ont permit la generation du precedent business model sont: {rubriques}.
        si l'utlisateur a donner les données complementaires, veuillez en tenir compte dans la generation, et ca doit etre imperativement prioritaire.
        Si dans un bloque un utilisateur n'as pas donner des informations (elements), veuillez generer,
        Si l'utilisateur à donné des elements que vous juger peu, generer d'autres et les ajoutées à ce que l'utlisateur à fournit.
        
        à faire imperativement est:
        Format:
            - Résumer l’analyse en expliquant chaque point de vérification.
            - Dire si c'est coherent ou pas.
            - Proposer un plan d’action (liste structurée) en se basant sur le metaprompt.
        
        """
    elif generation == 3:
        # Première génération avec les nouvelles rubriques
        prompt = f"""
        {metaprompt}
        voici les données à qu'on a utliser pour generer le business model: {rubriques}.
        voici le business model generer: {previousdata} .
        voici le données des analyses financier : {Analisis} .
        
        Format :
            - Résumer l’analyse en expliquant chaque point de vérification.
            - Dire si c'est coherent ou pas.
            - Proposer un plan d’action (liste structurée).
        """
    else :
        # Deuxième génération (amélioration) en utilisant le BMC précédent et les nouvelles rubriques
        # Prompt ajusté sans numérotation dans les titres
        prompt = f"""
        voici le metaprompt:{metaprompt}
        
        voici les données à qu'on a utliser pour generer le business model: {rubriques}.
        voici le business model generer: {previousdata} .
        voici le données des analyses financier : {Analisis} .
        voici le business plan generer pecedement: {previousbp} .
        
        Ameliorer ces business plan sur bases de metaprompt
        Mener la reflexions du generation du business modele sur base des indications(Méta-Prompt) precedents du metaprompts; 
        Chercher les chiffres et autres données sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des données collecter sur internet 
        Génère le contenu d'un Business Model Canvas en format HTML pour une entreprise nommée '{nom_entreprise}'.
        
        sachant que les données qui ont permit la generation du precedent business model sont: {rubriques}.
        si l'utlisateur a donner les données complementaires, veuillez en tenir compte dans la generation, et ca doit etre imperativement prioritaire.
        Si dans un bloque un utilisateur n'as pas donner des informations (elements), veuillez generer,
        Si l'utilisateur à donné des elements que vous juger peu, generer d'autres et les ajoutées à ce que l'utlisateur à fournit.
        
        à faire imperativement est:
        Format:
            - Résumer l’analyse en expliquant chaque point de vérification.
            - Dire si c'est coherent ou pas.
            - Proposer un plan d’action (liste structurée) en se basant sur le metaprompt.
        
        """
    # Calculer les tokens nécessaires (entrée + réponse prévue)
    # Calculer le nombre de tokens dans l'entrée utilisateur
    tokens_in_input = count_tokens(type_chain_of_thougtht+""+json.dumps(previousdata)+""+json.dumps(rubriques)+""+Analisis+""+previousbp, MODEL)
    tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
    # Vérifier si l'utilisateur a assez de tokens
    if tokens_remaining < tokens_needed:
        st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
        return
    else:
        try:
            response = openai.ChatCompletion.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "Tu es un assistant expert en génération de business  et business plan."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=5000,
                temperature=0.7
            )
            html_genere = response.choices[0].message.content.strip()
            tokens_utilises = response['usage']['total_tokens']
            # Consommer les tokens
            success, message = consommer_tokens(st.session_state['user_info'], tokens_utilises)
            return html_genere
        except Exception as e:
            st.error(f"Erreur lors de la génération du contenu : {e}")
            return ""


# ----------------------------------------------------------------------------
# 2) Fonction pour créer le fichier Word (format tableau) avec python-docx
# ----------------------------------------------------------------------------

def get_produit(index_produit):
    if 0 <= index_produit < len(st.session_state["produits_data"]):
        return st.session_state["produits_data"][index_produit]
    else:
        st.error("Index du produit invalide.")
        return None


def get_business_model_by_product_index(product_idx):
    """
    Récupère le Business Model (BMC) d'un produit donné par son index.

    Parameters:
    - product_idx (int): L'index du produit dans st.session_state["produits_data"].

    Returns:
    - dict: Le Business Model du produit, contenant 'titre', 'description', et 'date_bmc'.
    - None: Si aucun Business Model n'est trouvé ou si l'index est invalide.
    """
    # Vérifier que 'produits_data' existe dans session_state
    if "produits_data" not in st.session_state:
        st.error("Aucune donnée de produit trouvée. Veuillez d'abord charger les données.")
        return None
    
    produits = st.session_state["produits_data"]
    
    # Vérifier que l'index est un entier
    if not isinstance(product_idx, int):
        st.error("L'index du produit doit être un entier.")
        return None
    
    # Vérifier que l'index est dans la plage valide
    if product_idx < 0 or product_idx >= len(produits):
        st.error(f"Index du produit invalide. Il doit être entre 0 et {len(produits) - 1}.")
        return None
    
    produit = produits[product_idx]
    nom_produit = produit.get("nom_produit", f"Produit_{product_idx}")
    
    # Vérifier si le produit possède un Business Model
    business_models = produit.get("business_models", [])
    
    if len(business_models) == 0:
        st.info(f"Aucun Business Model trouvé pour le produit '{nom_produit}'.")
        return None
    
    # Puisqu'il y a au maximum un BMC, on accède au premier élément
    business_model = business_models[0]
    titre = business_model.get("titre", "Titre non disponible")
    description = business_model.get("description", "Description non disponible")
    date_bmc = business_model.get("date_bmc", "Date non disponible")
    
    # Retourner le Business Model
    return {
        "titre": titre,
        "description": description,
        "date_bmc": date_bmc
    }

def add_business_model(index_produit, business_model):
    produit = get_produit(index_produit)
    if produit is not None:
        if "business_models" not in produit:
            produit["business_models"] = []
            st.warning(f"'business_models' initialisé pour le produit '{produit['nom_produit']}'")
        if len(produit["business_models"]) >= 1:
            # Remplacer le BMC existant
            produit["business_models"][0] = business_model
            st.success("Business Model remplacé avec succès!")
        else:
            # Ajouter un nouveau BMC
            produit["business_models"].append(business_model)
            st.success("Business Model ajouté avec succès!")

def update_business_model(index_produit, business_model):
    produit = get_produit(index_produit)
    if produit is not None:
        if "business_models" in produit and len(produit["business_models"]) >= 1:
            produit["business_models"][0] = business_model
            st.success("Business Model mis à jour avec succès!")
        else:
            st.error("Aucun Business Model existant à mettre à jour.")

def delete_business_model(index_produit):
    produit = get_produit(index_produit)
    if produit is not None:
        if "business_models" in produit and len(produit["business_models"]) >= 1:
            bm = produit["business_models"].pop(0)
            st.success(f"Business Model '{bm['titre']}' supprimé avec succès!")
            
            # Gestion de 'business_model_precedent'
            if 'business_model_precedent' in st.session_state:
                bm_p = st.session_state['business_model_precedent'].get('product_idx')
                
                if bm_p == index_produit:
                    # Le Business Model supprimé appartenait au produit sélectionné
                    st.session_state['business_model_precedent'] = {}
        else:
            st.error("Aucun Business Model à supprimer.")





def generer_docx_business_modelss(nom_entreprise, date_bmc, contenu_business_model, doc, value=1):
    """
    Construit un document Word reproduisant un tableau avec la disposition souhaitée
    pour le Business Model Canvas. La mise en forme inclut des titres en gras et
    des listes à puces.
    'contenu_business_model' : le contenu HTML renvoyé par ChatGPT,
    qu'on découpe ensuite pour remplir chaque bloc.
    """
    # Créer un nouveau document Word
    if value == 1:
        doc = Document()

    # Définir les styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titre principal
    titre = doc.add_heading(level=1)
    titre_run = titre.add_run(f"Business Model Canvas de {nom_entreprise}")
    titre_run.bold = True
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date : {date_bmc}")
    date_run.bold = True
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter un saut de ligne
    doc.add_paragraph("")

    # Créer un tableau de 6 lignes × 5 colonnes
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    # Ajuster les largeurs des colonnes (en pouces)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.8)  # Ajustez selon vos besoins

    # 1) Ligne 0 : Titre (fusion des 5 colonnes)
    cell00 = table.cell(0, 0)
    cell00_merge = cell00.merge(table.cell(0, 4))
    cell00_merge.text = f"Business Model Canvas de {nom_entreprise}"
    for paragraph in cell00_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2) Ligne 1 : Nom de l'entreprise et Date (fusion des colonnes)
    cell10 = table.cell(1, 0)
    cell10_merge = cell10.merge(table.cell(1, 2))
    cell10_merge.text = f"**Nom de l'entreprise**: {nom_entreprise}"
    for paragraph in cell10_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    cell13 = table.cell(1, 3)
    cell13_merge = cell13.merge(table.cell(1, 4))
    cell13_merge.text = f"**Date**: {date_bmc}"
    for paragraph in cell13_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3) Ligne 2 : Headers des 5 blocs
    headers = ["Partenaires clés", "Activités clés", "Offre (proposition de valeur)", 
               "Relation client", "Segments de clientèle"]
    for idx, header in enumerate(headers):
        cell = table.cell(2, idx)
        paragraphe = cell.paragraphs[0]
        run = paragraphe.add_run(header)
        run.bold = True
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4) Ligne 3 : Contenus des 5 blocs
    # Initialiser les blocs
    blocs = {
        "Partenaires clés": "",
        "Activités clés": "",
        "Offre (proposition de valeur)": "",
        "Relation client": "",
        "Segments de clientèle": ""
    }

    # Utiliser BeautifulSoup pour parser le HTML
    soup = BeautifulSoup(contenu_business_model, 'html.parser')

    # Fonction pour trouver le bon header tag (h3 par défaut, avec flexibilité)
    def trouver_header(soup, header):
        # Regex pour capturer optionnellement des numéros suivis de points et espaces
        pattern = rf"^(?:\d+\.\s*)?{re.escape(header)}$"
        # Chercher dans les balises h3
        header_tag = soup.find(['h2', 'h3', 'h4', 'h5', 'h6'], text=re.compile(pattern, re.IGNORECASE))
        return header_tag

    # Extraire chaque bloc
    for header in blocs.keys():
        h_tag = trouver_header(soup, header)
        if h_tag:
            content = []
            for sibling in h_tag.find_next_siblings():
                if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                    break  # Arrêter si un nouveau header est trouvé
                if sibling.name == 'ul':
                    for li in sibling.find_all('li'):
                        content.append(f"- {li.get_text(strip=True)}")
                elif sibling.name == 'p':
                    content.append(sibling.get_text(strip=True))
                elif isinstance(sibling, str):
                    text = sibling.strip()
                    if text:
                        content.append(text)
            blocs[header] = '\n'.join(content)

    # Debug: Afficher les blocs extraits (à désactiver en production)
    # st.write("Blocs extraits :", blocs)

    # Fonction pour ajouter du contenu formaté dans une cellule
    def ajouter_contenu(cell, titre, contenu):
        """
        Ajoute du contenu formaté dans une cellule Word.
        Le titre est en gras, suivi de listes à puces si nécessaire.
        """
        # Supprimer le texte initial (par défaut) dans la cellule
        cell.text = ""

        # Ajouter le titre en gras
        paragraphe = cell.add_paragraph()
        run = paragraphe.add_run(titre)
        run.bold = True

        # Ajouter le contenu
        # Diviser le contenu par les sauts de ligne
        lignes = contenu.split('\n')
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            # Vérifier si la ligne commence par '-', '+', '•' pour une liste à puces
            if re.match(r'^[-+•]\s+', ligne):
                # Ajouter une puce
                item = re.sub(r'^[-+•]\s+', '', ligne)
                p = cell.add_paragraph(item, style='List Bullet')
            else:
                # Ajouter un paragraphe normal
                p = cell.add_paragraph(ligne)

    # Remplir les cellules de la ligne 3
    ordre_blocs = [
        "Partenaires clés", "Activités clés", "Offre (proposition de valeur)",
        "Relation client", "Segments de clientèle"
    ]

    for idx, bloc in enumerate(ordre_blocs):
        cell = table.cell(3, idx)
        ajouter_contenu(cell, bloc, blocs[bloc])

    # 5) Ligne 4 : Structure de coûts (fusion 3 cols) et Sources de revenus (fusion 2 cols)
    # Fusionner les cellules pour "Structure de coûts" (colonnes 0-2)
    cell40 = table.cell(4, 0)
    cell40_merge = cell40.merge(table.cell(4, 2))
    cell40_merge.text = f"**Structure de coûts**:\n\n"

    # Fusionner les cellules pour "Sources de revenus" (colonnes 3-4)
    cell43 = table.cell(4, 3)
    cell43_merge = cell43.merge(table.cell(4, 4))
    cell43_merge.text = f"**Sources de revenus**:\n\n"

    # Extraire les contenus pour ces blocs
    structure_couts = ""
    sources_revenus = ""

    # Structure des coûts
    strong_tag = trouver_header(soup, "Structure des coûts")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        structure_couts = '\n'.join(content)

    # Sources de revenus
    strong_tag = trouver_header(soup, "Sources de revenus")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        sources_revenus = '\n'.join(content)

    # Remplir les cellules fusionnées
    ajouter_contenu(cell40_merge, "Structure de coûts", structure_couts)
    ajouter_contenu(cell43_merge, "Sources de revenus", sources_revenus)

    # Ajuster les paragraphes existants
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Ajouter un saut de ligne à la fin
    doc.add_paragraph("")

    # Convertir le document en binaire pour téléchargement via Streamlit
    fichier_io = BytesIO()
    doc.save(fichier_io)
    fichier_io.seek(0)
    return fichier_io


def generer_docx_business_models(nom_entreprise, date_bmc, contenu_business_model, nom_produit, doc=None, value=1): 
    """
    Construit un document Word reproduisant un tableau avec la disposition souhaitée
    pour le Business Model Canvas. La mise en forme inclut des titres en gras et
    des listes à puces.
    'contenu_business_model' : le contenu HTML renvoyé par ChatGPT,
    qu'on découpe ensuite pour remplir chaque bloc.
    """
    # Créer un nouveau document Word
    if value == 1 or doc is None:
        doc = Document()

    # Définir les styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titre principal
    titre = doc.add_heading(level=1)
    titre_run = titre.add_run(f"Business Model Canvas de {nom_entreprise}")
    titre_run.bold = True
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading(f"Produit : {nom_produit}", level=2)

    # Date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date : {date_bmc}")
    date_run.bold = True
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter un saut de ligne
    doc.add_paragraph("")

    # Créer un tableau de 7 lignes × 5 colonnes
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'

    # Ajuster les largeurs des colonnes (en pouces)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.8)  # Ajustez selon vos besoins

    # 1) Ligne 0 : Titre (fusion des 5 colonnes)
    cell00 = table.cell(0, 0)
    cell00_merge = cell00.merge(table.cell(0, 4))
    cell00_merge.text = f"Business Model Canvas de {nom_entreprise}"
    for paragraph in cell00_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2) Ligne 1 : Nom de l'entreprise et Date (fusion des colonnes)
    cell10 = table.cell(1, 0)
    cell10_merge = cell10.merge(table.cell(1, 2))
    cell10_merge.text = f"**Nom de l'entreprise**: {nom_entreprise}"
    for paragraph in cell10_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    cell13 = table.cell(1, 3)
    cell13_merge = cell13.merge(table.cell(1, 4))
    cell13_merge.text = f"**Date**: {date_bmc}"
    for paragraph in cell13_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3) Ligne 2 : Headers des 5 blocs
    headers = ["Partenaires clés", "Activités clés", "Offre (proposition de valeur)", 
               "Relation client", "Segments de clientèle"]
    for idx, header in enumerate(headers):
        cell = table.cell(2, idx)
        paragraphe = cell.paragraphs[0]
        run = paragraphe.add_run(header)
        run.bold = True
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4) Ligne 3 : Contenus des 5 blocs
    # Initialiser les blocs
    blocs = {
        "Partenaires clés": "",
        "Activités clés": "",
        "Offre (proposition de valeur)": "",
        "Relation client": "",
        "Segments de clientèle": "",
        "Ressources Clés": "",
        "Canaux de Distribution": ""
    }

    # Utiliser BeautifulSoup pour parser le HTML
    soup = BeautifulSoup(contenu_business_model, 'html.parser')

    # Fonction pour trouver le bon header tag (h3 par défaut, avec flexibilité)
    def trouver_header(soup, header):
        # Regex pour capturer optionnellement des numéros suivis de points et espaces
        pattern = rf"^(?:\d+\.\s*)?{re.escape(header)}$"
        # Chercher dans les balises h2 à h6
        header_tag = soup.find(['h2', 'h3', 'h4', 'h5', 'h6'], text=re.compile(pattern, re.IGNORECASE))
        return header_tag

    # Extraire chaque bloc
    for header in blocs.keys():
        h_tag = trouver_header(soup, header)
        if h_tag:
            content = []
            for sibling in h_tag.find_next_siblings():
                if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                    break  # Arrêter si un nouveau header est trouvé
                if sibling.name == 'ul':
                    for li in sibling.find_all('li'):
                        content.append(f"- {li.get_text(strip=True)}")
                elif sibling.name == 'p':
                    content.append(sibling.get_text(strip=True))
                elif isinstance(sibling, str):
                    text = sibling.strip()
                    if text:
                        content.append(text)
            blocs[header] = '\n'.join(content)

    # Fonction pour ajouter du contenu formaté dans une cellule
    def ajouter_contenu(cell, titre, contenu):
        """
        Ajoute du contenu formaté dans une cellule Word.
        Le titre est en gras, suivi de listes à puces si nécessaire.
        """
        # Supprimer le texte initial (par défaut) dans la cellule
        cell.text = ""

        # Ajouter le titre en gras
        paragraphe = cell.add_paragraph()
        run = paragraphe.add_run(titre)
        run.bold = True

        # Ajouter le contenu
        # Diviser le contenu par les sauts de ligne
        lignes = contenu.split('\n')
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            # Vérifier si la ligne commence par '-', '+', '•' pour une liste à puces
            if re.match(r'^[-+•]\s+', ligne):
                # Ajouter une puce
                item = re.sub(r'^[-+•]\s+', '', ligne)
                p = cell.add_paragraph(item, style='List Bullet')
            else:
                # Ajouter un paragraphe normal
                p = cell.add_paragraph(ligne)

    # Remplir les cellules de la ligne 3
    ordre_blocs = [
        "Partenaires clés", 
        "Activités clés", 
        "Offre (proposition de valeur)",
        "Relation client", 
        "Segments de clientèle"
    ]

    for idx, bloc in enumerate(ordre_blocs):
        cell = table.cell(3, idx)
        ajouter_contenu(cell, bloc, blocs[bloc])

    # Ajouter les nouvelles rubriques dans la ligne 4
    # "Ressources Clés" sous "Activités Clés" (colonne 1)
    cell_ressources = table.cell(4, 1)
    ajouter_contenu(cell_ressources, "Ressources Clés", blocs["Ressources Clés"])

    # "Canaux de Distribution" sous "Relation Client" (colonne 3)
    cell_canaux = table.cell(4, 3)
    ajouter_contenu(cell_canaux, "Canaux de Distribution", blocs["Canaux de Distribution"])

    # 5) Ligne 5 : Structure de coûts (fusion 3 cols) et Sources de revenus (fusion 2 cols)
    # Fusionner les cellules pour "Structure de coûts" (colonnes 0-2)
    cell50 = table.cell(5, 0)
    cell50_merge = cell50.merge(table.cell(5, 2))
    cell50_merge.text = f"**Structure de coûts**:\n\n"

    # Fusionner les cellules pour "Sources de revenus" (colonnes 3-4)
    cell53 = table.cell(5, 3)
    cell53_merge = cell53.merge(table.cell(5, 4))
    cell53_merge.text = f"**Sources de revenus**:\n\n"

    # Extraire les contenus pour ces blocs
    structure_couts = ""
    sources_revenus = ""

    # Structure des coûts
    strong_tag = trouver_header(soup, "Structure des coûts")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        structure_couts = '\n'.join(content)

    # Sources de revenus
    strong_tag = trouver_header(soup, "Sources de revenus")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        sources_revenus = '\n'.join(content)

    # Remplir les cellules fusionnées
    ajouter_contenu(cell50_merge, "Structure de coûts", structure_couts)
    ajouter_contenu(cell53_merge, "Sources de revenus", sources_revenus)

    # Ajuster les paragraphes existants
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Ajouter un saut de ligne à la fin
    doc.add_paragraph("")

    # Convertir le document en binaire pour téléchargement via Streamlit
    fichier_io = BytesIO()
    doc.save(fichier_io)
    fichier_io.seek(0)
    return fichier_io

def generer_docx_business_model(nom_entreprise, date_bmc, contenu_business_model, nom_produit, doc=None, value=1): 
    """
    Construit un document Word reproduisant un tableau avec la disposition souhaitée
    pour le Business Model Canvas. La mise en forme inclut des titres en gras et
    des listes à puces.
    
    Parameters:
    - nom_entreprise (str): Nom de l'entreprise.
    - date_bmc (str): Date du Business Model Canvas.
    - contenu_business_model (str): Contenu HTML renvoyé par ChatGPT.
    - nom_produit (str): Nom du produit.
    - doc (Document, optional): Document Word existant à modifier.
    - value (int, optional): Indicateur pour créer un nouveau document ou non.
    
    Returns:
    - BytesIO: Document Word en binaire pour téléchargement.
    """
    
    # Créer un nouveau document Word si nécessaire
    if value == 1 or doc is None:
        doc = Document()

    # Définir les styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titre principal
    titre = doc.add_heading(level=1)
    titre_run = titre.add_run(f"Business Model Canvas de {nom_entreprise}")
    titre_run.bold = True
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Titre du produit
    doc.add_heading(f"Produit : {nom_produit}", level=2)

    # Date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date : {date_bmc}")
    date_run.bold = True
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter un saut de ligne
    doc.add_paragraph("")

    # Créer un tableau de 7 lignes × 5 colonnes
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'

    # Ajuster les largeurs des colonnes (en pouces)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.8)  # Ajustez selon vos besoins

    # 1) Ligne 0 : Titre (fusion des 5 colonnes)
    cell00 = table.cell(0, 0)
    cell00_merge = cell00.merge(table.cell(0, 4))
    cell00_merge.text = f"Business Model Canvas de {nom_entreprise}"
    for paragraph in cell00_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2) Ligne 1 : Nom de l'entreprise et Date (fusion des colonnes)
    cell10 = table.cell(1, 0)
    cell10_merge = cell10.merge(table.cell(1, 2))
    cell10_merge.text = f"Nom de l'entreprise: {nom_entreprise}"
    for paragraph in cell10_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    cell13 = table.cell(1, 3)
    cell13_merge = cell13.merge(table.cell(1, 4))
    cell13_merge.text = f"Date: {date_bmc}"
    for paragraph in cell13_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3) Ligne 2 : Headers des 5 blocs
    headers = ["Partenaires clés", "Activités clés", "Offre (proposition de valeur)", 
               "Relation client", "Segments de clientèle"]
    for idx, header in enumerate(headers):
        cell = table.cell(2, idx)
        paragraphe = cell.paragraphs[0]
        run = paragraphe.add_run(header)
        run.bold = True
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4) Ligne 3 : Contenus des 5 blocs
    # Initialiser les blocs
    blocs = {
        "Partenaires clés": "",
        "Activités clés": "",
        "Offre (proposition de valeur)": "",
        "Relation client": "",
        "Segments de clientèle": "",
        "Ressources clés": "",
        "Canaux de distribution": ""
    }

    # Utiliser BeautifulSoup pour parser le HTML
    soup = BeautifulSoup(contenu_business_model, 'html.parser')

    # Fonction pour trouver le bon header tag (h2 par défaut, avec flexibilité)
    def trouver_header(soup, header):
        """
        Trouve le tag d'en-tête correspondant au bloc spécifié, en nettoyant les astérisques.

        Parameters:
        - soup (BeautifulSoup): Objet BeautifulSoup du contenu HTML.
        - header (str): Nom du bloc à trouver.

        Returns:
        - Tag or None: Tag BeautifulSoup correspondant à l'en-tête ou None.
        """
        # Regex pour capturer optionnellement des numéros suivis de points et espaces
        pattern = rf"^(?:\d+\.\s*)?{re.escape(header)}$"
        # Chercher dans les balises h2 à h6 sans les astérisques
        for tag in soup.find_all(['h2', 'h3', 'h4', 'h5', 'h6']):
            # Nettoyer le texte en enlevant les astérisques
            tag_text = tag.get_text(strip=True).replace('**', '')
            if re.match(pattern, tag_text, re.IGNORECASE):
                return tag
        return None

    # Extraire chaque bloc
    for header in blocs.keys():
        h_tag = trouver_header(soup, header)
        if h_tag:
            content = []
            for sibling in h_tag.find_next_siblings():
                if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                    break  # Arrêter si un nouveau header est trouvé
                if sibling.name == 'ul':
                    for li in sibling.find_all('li'):
                        content.append(f"- {li.get_text(strip=True)}")
                elif sibling.name == 'p':
                    content.append(sibling.get_text(strip=True))
                elif isinstance(sibling, str):
                    text = sibling.strip()
                    if text:
                        content.append(text)
            blocs[header] = '\n'.join(content)

    # Fonction pour ajouter du contenu formaté dans une cellule
    def ajouter_contenu(cell, titre, contenu):
        """
        Ajoute du contenu formaté dans une cellule Word.
        Le titre est en gras, suivi de listes à puces si nécessaire.

        Parameters:
        - cell (Cell): Cellule du tableau Word.
        - titre (str): Titre du bloc.
        - contenu (str): Contenu du bloc.
        """
        # Supprimer le texte initial dans la cellule
        cell.text = ""

        # Ajouter le titre en gras
        paragraphe = cell.add_paragraph()
        run = paragraphe.add_run(titre)
        run.bold = True

        # Ajouter le contenu
        lignes = contenu.split('\n')
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            # Vérifier si la ligne commence par '-', '+', '•' pour une liste à puces
            if re.match(r'^[-+•]\s+', ligne):
                # Ajouter une puce
                item = re.sub(r'^[-+•]\s+', '', ligne)
                p = cell.add_paragraph(item, style='List Bullet')
            else:
                # Ajouter un paragraphe normal
                p = cell.add_paragraph(ligne)

    # Remplir les cellules de la ligne 3
    ordre_blocs = [
        "Partenaires clés", 
        "Activités clés", 
        "Offre (proposition de valeur)",
        "Relation client", 
        "Segments de clientèle"
    ]

    for idx, bloc in enumerate(ordre_blocs):
        cell = table.cell(3, idx)
        ajouter_contenu(cell, bloc, blocs.get(bloc, ""))

    # Ajouter les nouvelles rubriques dans la ligne 4
    # "Ressources clés" sous "Activités clés" (colonne 1)
    cell_ressources = table.cell(4, 1)
    ajouter_contenu(cell_ressources, "Ressources clés", blocs.get("Ressources clés", ""))

    # "Canaux de distribution" sous "Relation client" (colonne 3)
    cell_canaux = table.cell(4, 3)
    ajouter_contenu(cell_canaux, "Canaux de distribution", blocs.get("Canaux de distribution", ""))

    # 5) Ligne 5 : Structure de coûts (fusion 3 cols) et Sources de revenus (fusion 2 cols)
    # Fusionner les cellules pour "Structure de coûts" (colonnes 0-2)
    cell50 = table.cell(5, 0)
    cell50_merge = cell50.merge(table.cell(5, 2))
    cell50_merge.text = f"Structure de coûts:\n\n"

    # Fusionner les cellules pour "Sources de revenus" (colonnes 3-4)
    cell53 = table.cell(5, 3)
    cell53_merge = cell53.merge(table.cell(5, 4))
    cell53_merge.text = f"Sources de revenus:\n\n"

    # Extraire les contenus pour ces blocs
    structure_couts = ""
    sources_revenus = ""

    # Structure des coûts
    strong_tag = trouver_header(soup, "Structure des coûts")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        structure_couts = '\n'.join(content)

    # Sources de revenus
    strong_tag = trouver_header(soup, "Sources de revenus")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        sources_revenus = '\n'.join(content)

    # Remplir les cellules fusionnées
    ajouter_contenu(cell50_merge, "Structure de coûts", structure_couts)
    ajouter_contenu(cell53_merge, "Sources de revenus", sources_revenus)

    # Ajuster les paragraphes existants
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Ajouter un saut de ligne à la fin
    doc.add_paragraph("")

    # Convertir le document en binaire pour téléchargement via Streamlit
    fichier_io = BytesIO()
    doc.save(fichier_io)
    fichier_io.seek(0)
    return fichier_io








def page_generer_business_model():
    st.header("Étape 2 : Générer le Business Model Canvas")
    if st.session_state.get('business_model_precedent'):
        produit_selectionne = st.selectbox(
            "Sélectionnez un Produit",
            options=[prod["nom_produit"] for prod in st.session_state["produits_data"]],
            key="produit_selectionne_affichage_bmc"
        )
        # Trouver l'index du produit sélectionné
        index_produit_affichage = next((index for (index, d) in enumerate(st.session_state["produits_data"]) if d["nom_produit"] == produit_selectionne), None)
        st.write("Le Business Model Canvas initial a été généré. Vous pouvez le télécharger ci-dessous ou procéder à son amélioration.")
        if st.button("Ameliorer Business Model"):
            with st.spinner("Amelioration en cours..."):
                html_content = chain_of_thougtht("BM",st.session_state.montant_projet,st.session_state.nom_entreprise,get_business_model_by_product_index(index_produit_affichage).get("description"), st.session_state["produits_data"][index_produit_affichage], Analisis='', previousbp='', generation=2)
                st.write(html_content)
                if html_content:
                    with st.spinner("Generation en cours..."):                        
                        contenu_bmc_ameliore = obtenir_business_model(
                        nom_entreprise=st.session_state.nom_entreprise,
                        montant_projet=st.session_state.montant_projet,
                        previousdata=get_business_model_by_product_index(index_produit_affichage).get("description"),
                        type_entreprise=st.session_state.type_entreprise,
                        rubriques= st.session_state["produits_data"][index_produit_affichage],
                        ameliorations=html_content,
                        generation=2)
                        
                        updated_bm = {
                            "titre": get_business_model_by_product_index(index_produit_affichage)['titre'],
                            "description": contenu_bmc_ameliore,
                            "date_bmc": get_business_model_by_product_index(index_produit_affichage)['date_bmc']
                                }
                        # Stocker le BMC initial dans la session pour la deuxième générationst.session_state.business_model_precedent = contenu_bmc_ameliore
                        update_business_model(index_produit_affichage, updated_bm)

        if index_produit_affichage is not None:
            business_models = st.session_state["produits_data"][index_produit_affichage].get("business_models", [])
            if business_models:
                bm_titles = [bm["titre"] for bm in business_models]
                selected_bm_idx = st.selectbox(
                    "Sélectionnez un Business Model",
                    options=list(range(len(bm_titles))),
                    format_func=lambda x: bm_titles[x],
                    key="selected_bm_idx_affichage"
                )
                
                st.session_state['business_model_precedent'] = {
                     'product_idx': index_produit_affichage
                     }

                selected_bm = business_models[selected_bm_idx]  
                     
                st.write(f"### {selected_bm['titre']}")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Modifier", key=f"modify_bm_{index_produit_affichage}_{selected_bm_idx}"):
                        with st.expander(f"Modifier Business Model {selected_bm['titre']}"):
                            nouveau_titre = st.text_input("Titre du Business Model", value=selected_bm['titre'], key=f"nouveau_titre_bm_{selected_bm_idx}")
                            nouvelle_description = st.text_area("Description", value=selected_bm['description'], key=f"nouvelle_description_bm_{selected_bm_idx}")
                            if st.button("Mettre à Jour", key=f"mettre_a_jour_bm_{selected_bm_idx}"):
                                updated_bm = {
                                    "titre": nouveau_titre,
                                    "description": nouvelle_description,
                                    "date_bmc": selected_bm.get("date_bmc", "")
                                }
                                update_business_model(index_produit_affichage, updated_bm)
                                st.rerun()
                with col2:
                    if st.button("Supprimer", key=f"delete_bm_{index_produit_affichage}_{selected_bm_idx}"):
                        delete_business_model(index_produit_affichage)
                        st.rerun()
                
                
                
                html_content = selected_bm['description']

                # Encoder le contenu HTML en Base64
                encoded_html = base64.b64encode(html_content.encode('utf-8')).decode('utf-8')

                # Créer l'URL de données
                data_url = f"data:text/html;base64,{encoded_html}"

                st.markdown(
                    f"""
                    <iframe src="{data_url}" width="100%" height="1500" frameborder="0" scrolling="yes"></iframe>
                    """,
                    unsafe_allow_html=True
                )
            
        # Vous pouvez également proposer d'autres actions ici si nécessaire
    else:
        st.info("Veuillez d'abord collecter toutes les données et générer le Business Model Canvas initial dans l'onglet 'Collecte des Données'.")


def ct_model():
    """
    Fonction principale de l'application Streamlit.
    """
    st.header("Étape 2 : Amelioration à effectuer")
    #st.write(st.session_state.business_model_precedent )
    #st.write(st.session_state["produits_data"])
    if st.session_state.get('business_model_precedent'):
        st.write("Amelioration à effectuer pour une bonne coherence du Business ")
        if st.button("Quoi ameliorer"):
            with st.spinner("Detection des Ameliorations en cours..."):
                html_content = chain_of_thougtht("Analysis", st.session_state.montant_projet,st.session_state.nom_entreprise, json.dumps(st.session_state.business_model_precedent) , st.session_state["produits_data"], Analisis='', previousbp='', generation=2)
                st.write(html_content)
                


def afficher_informations_cv_document(cv_data, query="Pouvez-vous résumer ce CV ? si oui si dans le resumer mentionner ses qualificetions les formations faites, son numero de telephone, adresse mail et son adresse"):
    """
    Fonction principale pour traiter un CV passé directement comme document et retourner des informations générées.

    Args:
        cv_data (UploadedFile): Le fichier téléchargé contenant le CV.
        query (str): La question ou le résumé demandé pour le CV.

    Returns:
        str: Les informations générées à partir du CV.
    """
    if not cv_data:
        raise ValueError("Aucun document fourni pour le traitement.")

    try:
        # Enregistrer le fichier temporairement
        file_path = "uploaded_document.pdf"
        with open(file_path, "wb") as f:
            f.write(cv_data.read())

        # Charger et diviser les documents
        documents = load_and_split_documents(file_path)
        
        # Fusionner tous les textes des documents pour le comptage des tokens
        full_text = " ".join([doc.page_content for doc in documents])

        # Définir le modèle et l'encodage correspondant
        MODEL="gpt-4o"  # Remplacez par le modèle que vous utilisez

        # Définir le nombre maximal de tokens par requête (réponse prévue)
        MAX_TOKENS_REPONSE = 300

        # Compter les tokens dans la requête (query + full_text)
        tokens_in_input = count_tokens(query + "full_text", MODEL)

        # Définir le nombre total de tokens nécessaires (entrée + réponse prévue)
        tokens_needed = tokens_in_input + MAX_TOKENS_REPONSE

        # Récupérer les informations de l'utilisateur
        utilisateur = st.session_state.get('user_info', None)
        if not utilisateur:
            st.error("Vous devez être connecté pour effectuer cette action.")
            return None

        tokens_purchased = utilisateur.get('tokens_purchased', 0)
        tokens_consumed = utilisateur.get('tokens_consumed', 0)
        tokens_remaining = tokens_purchased - tokens_consumed

        # Vérifier si l'utilisateur a assez de tokens
        if tokens_remaining < tokens_needed:
            st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
            return None

        db = create_faiss_db(documents)
        
        if db:
            success, message = consommer_tokens(st.session_state['user_info'], tokens_in_input)
            if not success:
                return ""

            

        # Configurer le modèle conversationnel
        llm = ChatOpenAI(temperature=0.7, openai_api_key=api_key, max_tokens=300)
        memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
        qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory)

        # Poser une question ou générer un résumé
        reponse = qa_chain.run({"question": query})
        tokens_in_inputs = count_tokens(query + ""+full_text+""+ reponse, MODEL)

        success, message = consommer_tokens(st.session_state['user_info'], tokens_in_inputs)
        if not success:
            return ""
        
        # Supprimer le fichier temporaire
        #os.remove(file_path)
        return reponse

    except Exception as e:
        raise ValueError(f"Erreur lors du traitement du CV : {str(e)}")

def serialize_membres(membres):
    serializable_membres = []
    for membre in membres:
        membre_copy = {k: v for k, v in membre.items() if k != 'cv'}
        serializable_membres.append(membre_copy)
    return serializable_membres

def ajouter_informations_personnel():
    """
    Fonction permettant de gérer les informations des membres de l'entreprise
    et de générer et enregistrer automatiquement les résumés des CV.
    """
    st.header("Gestion des membres du personnel")
    st.write("Ajoutez les informations des membres dans des conteneurs, jusqu'à un maximum de 5.")

    # Initialisation des données
    if "membres" not in st.session_state:
        st.session_state["membres"] = []  # Liste des membres ajoutés

    max_conteneurs = 5

    # Ajouter un nouveau conteneur
    if len(st.session_state["membres"]) < max_conteneurs:
        if st.button("Ajouter un membre"):
            st.session_state["membres"].append({
                "fonction": "",
                "utiliser_cv": False,
                "cv": None,
                "informations": "",
                "resume_cv": None  # Nouveau champ pour enregistrer le résumé généré
            })

    # Afficher les conteneurs existants
    for idx, membre in enumerate(st.session_state["membres"]):
        with st.container():
            col1, col2, col3, col4 = st.columns([2, 1, 3, 1])

            with col1:
                membre["fonction"] = st.text_input(
                    f"Fonction au sein de l'entreprise (Membre {idx + 1})",
                    value=membre["fonction"],
                    key=f"fonction_{idx}",
                    placeholder="Entrez la fonction"
                )

            with col2:
                membre["utiliser_cv"] = st.checkbox(
                    f"Utiliser un CV ",
                    value=membre["utiliser_cv"],
                    key=f"checkbox_cv_{idx}"
                )

            with col3:
                if membre["utiliser_cv"]:
                    cv_data = st.file_uploader(
                        f"Téléchargez le CV (Membre {idx + 1})",
                        type=["pdf", "docx"],
                        key=f"uploader_cv_{idx}"
                    )

                    # Vérification de la taille du fichier
                    if cv_data is not None and cv_data.size > 3 * 1024 * 1024:  # 3 MB
                        st.error("La taille du fichier dépasse 3 MB. Veuillez télécharger un fichier plus petit.")
                        membre["cv"] = None     # Réinitialiser si la taille est dépassée
                    else:
                        membre["cv"] = cv_data  # Enregistrer le fichier s'il est valide

                        # Générer et enregistrer automatiquement le résumé
                        if membre["cv"] and not membre.get("resume_cv"):
                            membre["resume_cv"] = afficher_informations_cv_document(membre["cv"])
                else:
                    membre["informations"] = st.text_area(
                        f"Informations (Membre {idx + 1})",
                        value=membre["informations"],
                        key=f"text_area_{idx}",
                        placeholder="Entrez les informations manuellement",
                        height=100
                    )

            with col4:
                if st.button("Supprimer", key=f"supprimer_{idx}"):
                    st.session_state["membres"].pop(idx)
                    # Mise à jour dynamique sans rechargement
                    st.rerun()

    # Afficher un message si la limite est atteinte
    if len(st.session_state["membres"]) >= max_conteneurs:
        st.warning("Vous avez atteint le nombre maximum de membres (5).")

    # Afficher la liste des membres enregistrés
    st.subheader("Résumé des membres ajoutés")
    if st.session_state["membres"]:
        for idx, membre in enumerate(st.session_state["membres"]):
            st.write(f"**Membre {idx + 1}**")
            st.write(f"- **Fonction :** {membre.get('fonction', 'Non défini')}")
            if membre.get("utiliser_cv"):
                if membre.get("cv"):
                    # Afficher uniquement le nom du fichier
                    st.write(f"- **CV :** {membre['cv'].name}")
                    # Afficher le résumé généré
                    st.write(f"- **Résumé généré :** {membre.get('resume_cv', 'Résumé non disponible')}")
                else:
                    st.write("- **CV :** Non valide ou non téléchargé")
            else:
                st.write(f"- **Informations :** {membre.get('informations', 'Non fourni')}")




def page_collecte_donnees():
    st.subheader("Collecte des Données")

    col1, col2, col3 = st.columns([6, 1, 6])
    with col1:
        ui_val = st.number_input(
            "Combien de produits afficher ?",
            min_value=st.session_state["nb_products"],
            step=1,
            value=st.session_state["nb_products"],
            key="ui_nb_products"
        )
        # Synchronisation : si l'utilisateur change la valeur

        if ui_val > st.session_state["nb_products"]:
            # Mettre à jour uniquement si la valeur a augmenté
            st.session_state["nb_products"] = ui_val
        elif ui_val < st.session_state["nb_products"]:
            # Si l'utilisateur diminue la valeur (non autorisé), afficher un message
            st.warning("Utilisez uniquement les boutons pour augmenter le nombre de produits.")
                # 2) Bouton "Ajouter un produit"
            
    with col3:
        # 3) Bouton "Supprimer un produit"
        with st.container():
            col1, col2 = st.columns([6, 1])  # Première colonne plus large que la deuxième

            # Colonne 1 : Selectbox pour choisir un produit
            with col1:
                if len(st.session_state["produits_data"]) >= 0:
                    suppr_index = st.selectbox(
                        "Choisissez un produit à supprimer",
                        options=["Aucun"] + list(range(len(st.session_state["produits_data"]))),
                        format_func=lambda x: (
                            "Aucun"
                            if x == "Aucun"
                            else f"{x+1} - {st.session_state['produits_data'][x].get('nom_produit', '(inconnu)')}"
                        ),
                        key="suppr_index_select",
                    )

            # Colonne 2 : Bouton de suppression
            with col2:
                if len(st.session_state["produits_data"]) > 0 and suppr_index != "Aucun":
                    st.markdown(
                        """
                        <div style="height: 25px;"></div>
                        """,
                        unsafe_allow_html=True,
                    )

                    if st.button("🗑️", key="confirm_delete_button"):
                        # Supprimer le produit sélectionné
                        st.session_state["produits_data"].pop(suppr_index)

                        # Ajuster le nombre de produits
                        st.session_state["nb_products"] = len(st.session_state["produits_data"])

                        # Supprimer le business_model_precedent associé si nécessaire
                        if "business_model_precedent" in st.session_state:
                            bm_p = st.session_state["business_model_precedent"].get("product_idx")
                            if bm_p == suppr_index:
                                st.session_state["business_model_precedent"] = {}

                        st.success("Produit supprimé.")
                        st.rerun()
         
         




    # 4) Si nb_products > taille, on crée des produits vides
    if len(st.session_state["produits_data"]) < st.session_state["nb_products"]:
        diff = st.session_state["nb_products"] - len(st.session_state["produits_data"])
        for _ in range(diff):
            i_new = len(st.session_state["produits_data"])
            st.session_state["produits_data"].append({
                "nom_produit": f"Produit_{i_new+1}",
                "persona": {},
                "problem_tree": {},
                "analyse_marche": {},
                "facteurs_limitants": {},
                "concurrence": {}
            })

    # 5) Affichage
    if len(st.session_state["produits_data"])==0:
        st.info("Aucun produit.")
        return

    nb_affiche = min(st.session_state["nb_products"], len(st.session_state["produits_data"]))
    label_list = []
    for i, p in enumerate(st.session_state["produits_data"]):
        nomp = p.get("nom_produit", f"Produit_{i+1}")
        #label_list.append(f"{i+1} - {nomp}")
    label_list = [f"{i+1} - {p['nom_produit']}" for i, p in enumerate(st.session_state["produits_data"])]

    # Générer la liste des labels après mise à jour
    label_list = [f"{i+1} - {p['nom_produit']}" for i, p in enumerate(st.session_state["produits_data"])]

    # Créer le selectbox en utilisant la sélection sauvegardée
    selected_idx = st.selectbox(
        "Sélectionnez un produit à modifier",
        options=list(range(len(st.session_state["produits_data"]))),
        format_func=lambda x: label_list[x],
        index=st.session_state["selected_idx_produit"],
        key="selected_idx_produit"
    )
    
    # Mettre à jour la sélection dans st.session_state
    # st.session_state["selected_idx_produit"] = selected_idx
    # Détection des changements de sélection
    if st.session_state['selected_idx_produit'] != st.session_state['previous_selected_idx_produit']:
        st.session_state['previous_selected_idx_produit'] = st.session_state['selected_idx_produit']
        #st.success(f"Produit sélectionné changé : {label_list[selected_idx]}")


    # Sous-onglets : Persona, Arbre à Problème, Analyse du Marché, Facteurs Limitants, Concurrence
    tabs = st.tabs([
        "Persona",
        "Arbre à Problème",
        "Analyse du Marché",
        "Facteurs Limitants",
        "Concurrence"
    ])
    with tabs[0]:
        if 'type_entreprise' not in st.session_state:
            st.warning("Veuillez sélectionner le type d'entreprise dans la barre latérale.")
        else:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                collect_persona_pme(selected_idx)
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                collect_persona_pme(selected_idx)
                
    with tabs[1]:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                collect_arbre_probleme(selected_idx)
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                collect_arbre_probleme(selected_idx)
                        
    with tabs[2]:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                collect_analyse_marche_pme(selected_idx)
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                collect_analyse_marche_pme(selected_idx)

    with tabs[3]:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                collect_facteurs_limitants_pme(selected_idx)
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                collect_facteurs_limitants_pme(selected_idx)
        
    with tabs[4]:
            if st.session_state.type_entreprise == "PME":
                st.subheader("Collecte de Persona pour PME")
                collect_concurrence_pme(selected_idx)
            elif st.session_state.type_entreprise == "Startup":
                st.subheader("Collecte de Persona pour Startup")
                collect_concurrence_pme(selected_idx)
            # Bouton pour Générer le BMC Initial après avoir collecté toutes les données
            with st.form("form_generate_initial"):
                st.write("Après avoir collecté toutes les données, cliquez sur le bouton ci-dessous pour générer le Business Model Canvas initial.")
                submit_generate_initial = st.form_submit_button("Générer BMC Initial")
                submit_generate_ct = st.form_submit_button("Verifier coherence")
                
                
            if submit_generate_ct:
                  
                if "produits_data" in st.session_state and st.session_state["produits_data"]:
                    with st.spinner("Génération en cours..."):
                        html_content = chain_of_thougtht("Creativite",
                                                         st.session_state.montant_projet,
                                                         st.session_state.nom_entreprise, 
                                                         "", 
                                                         st.session_state["produits_data"][selected_idx], 
                                                         Analisis='',
                                                         previousbp='', 
                                                         generation=1)
                        st.markdown(html_content)

                            
                else:
                    st.info("Aucun produit ajouté pour le moment.")
                    
                    
            if submit_generate_initial:   
                if "produits_data" in st.session_state and st.session_state["produits_data"]:
                    with st.spinner("Génération en cours..."):
                        date_bmc = st.date_input("Date du BMC", value=datetime.date.today(), key="date_bmc_generate")
                        date_bmc_str = date_bmc.isoformat()
                        # Générer le premier BMC
                        contenu_bmc_initial = obtenir_business_model(                          
                            nom_entreprise=st.session_state.nom_entreprise,
                            type_entreprise=st.session_state.type_entreprise,
                            montant_projet=st.session_state.montant_projet,
                            rubriques=st.session_state["produits_data"][selected_idx],
                            previousdata="",
                            ameliorations="",
                            generation=1
                        )
                        #st.write(contenu_bmc_initial)
                        #st.write(st.session_state["produits_data"][selected_idx])
                        st.subheader("Contenu Initial Généré Ishai")
                        encoded_html = base64.b64encode(contenu_bmc_initial.encode('utf-8')).decode('utf-8')

                        # Créer l'URL de données
                        data_url = f"data:text/html;base64,{encoded_html}"

                        st.markdown(
                            f"""
                            <iframe src="{data_url}" width="100%" height="1500" frameborder="0" scrolling="yes"></iframe>
                            """,
                            unsafe_allow_html=True
                        )
                        
                        if not contenu_bmc_initial:
                            st.error("Erreur lors de la génération du contenu initial. Veuillez réessayer ou selectionner d'abord un produit.")
                        else:
                            # Générer le document Word en mémoire
                            # Enregistrer le Business Model dans session_state
                            business_model = {
                                "titre": f"BMC_{nomp}_{date_bmc}",
                                "description": contenu_bmc_initial ,
                                "date_bmc": date_bmc_str
                            }
                            add_business_model(selected_idx, business_model)
                            
                            doc = Document()
                            docx_bytes_initial = generer_docx_business_model(
                                nom_entreprise=st.session_state.nom_entreprise,
                                date_bmc=date_bmc.strftime("%d %B %Y"),
                                contenu_business_model=contenu_bmc_initial,
                                nom_produit=nomp,
                                doc=doc,
                                value=1
                            )
                            
                            st.success("Business Model Canvas initial généré avec succès !")
                            
                            # Proposer le téléchargement du document Word
                            st.download_button(
                                label="Télécharger le Business Model Canvas Initial (Word)",
                                data=docx_bytes_initial,
                                file_name=f"BMC_Initial_{st.session_state.nom_entreprise.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            # Stocker le BMC initial dans la session pour la deuxième génération
                            st.session_state.business_model_precedent = contenu_bmc_initial
                            
                            # Optionnel : Afficher le contenu généré pour vérification
                            
                            #st.markdown(contenu_bmc_initial, unsafe_allow_html=True)

                            
                else:
                    st.info("Aucun produit ajouté pour le moment.")
        
   


    # Pour debug : afficher le contenu actuel
    #st.write("Produits en mémoire :", st.session_state["produits_data"])

# ----------------------------------------------------------------------------
# Business plan 
# ----------------------------------------------------------------------------




# Initialiser le dictionnaire principal dans session_state
if "data" not in st.session_state:
    st.session_state["data"] = {}
    
# Section 1 : Informations Générales
def page_informations_generales():
    st.title("Informations Générales")
    
    # Accès au dictionnaire principal
    data = st.session_state["data"]
    
    # Collecte des entrées et stockage dans le dictionnaire principal
    data["informations_generales"] = data.get("informations_generales", {})
    info = data["informations_generales"]
    info["prenom_nom"] = st.text_input("Prénom, nom :", value=info.get("prenom_nom", ""))
    info["intitule_projet"] = st.text_input("Intitulé de votre projet :", value=info.get("intitule_projet", ""))
    info["statut_juridique"] = st.selectbox(
        "Votre statut juridique :",
        ["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"],
        index=["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"].index(info.get("statut_juridique", "Micro-entreprise"))
    )
    info["telephone"] = st.text_input("Votre numéro de téléphone :", value=info.get("telephone", ""))
    info["email"] = st.text_input("Votre adresse e-mail :", value=info.get("email", ""))
    info["ville"] = st.text_input("Votre ville ou commune d'activité :", value=info.get("ville", ""))
    info["type_vente"] = st.selectbox(
        "Vente de marchandises ou de services ?",
        ["Marchandises", "Services", "Mixte"],
        index=["Marchandises", "Services", "Mixte"].index(info.get("type_vente", "Marchandises"))
    )
    
    # Mise à jour des données dans le dictionnaire principal
    st.session_state["data"]["informations_generales"] = info

def page_besoins_demarrage():
    st.title("Besoins de Démarrage")
    
    # Accès au dictionnaire principal
    data = st.session_state.get("data", {})
    
    # Liste des besoins réorganisée
    besoins = [
        "Frais d’établissement", 
        "Logiciels, formations",
        "Matériel professionnel",
        "Matériel autre",
        "Matériel de bureau",
        "Stock de matières et produits",
        "Enseigne et éléments de communication",
        "Véhicule",
        "Frais de dossier",
        "Frais de notaire",
        "Trésorerie de départ",
        "Frais d’ouverture de compteurs",
        "Dépôt de marque",
        "Droits d’entrée",
        "Achat fonds de commerce ou parts",
        "Droit au bail",
        "Caution ou dépôt de garantie"
    ]
    
    # Initialiser le dictionnaire pour stocker les besoins
    data["besoins_demarrage"] = data.get("besoins_demarrage", {})
    besoins_demarrage = data["besoins_demarrage"]
    
    total_besoins = 0.0
    
    for besoin in besoins:
        montant = st.number_input(
            f"{besoin} ($)",
            min_value=0.0,
            key=f"besoin_{besoin}",
            value=besoins_demarrage.get(besoin, 0.0)
        )
        besoins_demarrage[besoin] = montant
        total_besoins += montant
    
    data["total_besoins"] = total_besoins
    
    st.write("---")
    st.markdown(f"**Total des Besoins de Démarrage :** {total_besoins:.2f} $")
    
    # Durée d'amortissement
    data["duree_amortissement"] = st.number_input(
        "Durée d'amortissement des investissements (en années) :",
        min_value=1,
        key="duree_amortissement",
        value=data.get("duree_amortissement", 3)
    )
    
    # Mise à jour des données dans le dictionnaire principal
    st.session_state["data"] = data



def calculer_pret_interet_fixe(montant, taux_annuel, duree_mois):
    """
    Calcule les détails d'un prêt avec intérêts fixes par mois.

    Args:
        montant (float): Montant du prêt en euros.
        taux_annuel (float): Taux d'intérêt annuel en pourcentage.
        duree_mois (int): Durée du prêt en mois.

    Returns:
        dict: Détails du prêt incluant mensualité, total à rembourser, principal mensuel,
              intérêts totaux et intérêts par année.
    """
    if duree_mois <= 0:
        return {
            "mensualite": 0.0,
            "total_a_rembourser": 0.0,
            "principal_mensuel": 0.0,
            "interet_mensuel": 0.0,
            "interets_totaux": 0.0,
            "interets_annee1": 0.0,
            "interets_annee2": 0.0,
            "interets_annee3": 0.0
        }

    taux_mensuel = taux_annuel / 100 / 12

    # Calcul de la mensualité en utilisant la formule PMT
    try:
        mensualite = (taux_mensuel * montant) / (1 - (1 + taux_mensuel) ** (-duree_mois))
    except ZeroDivisionError:
        mensualite = 0.0

    # Principal mensuel fixe
    principal_mensuel = montant / duree_mois

    # Intérêt mensuel
    interet_mensuel = mensualite - principal_mensuel

    # Total à rembourser
    total_a_rembourser = mensualite * duree_mois

    # Intérêts totaux
    interets_totaux = interet_mensuel * duree_mois

    # Intérêts par année, limités à 12 mois maximum
    interets_annee1 = interet_mensuel * min(duree_mois, 12)
    interets_annee2 = interet_mensuel * min(max(duree_mois - 12, 0), 12)
    interets_annee3 = interet_mensuel * min(max(duree_mois - 24, 0), 12)

    return {
        "mensualite": round(mensualite, 2),
        "total_a_rembourser": round(total_a_rembourser, 2),
        "principal_mensuel": round(principal_mensuel, 2),
        "interet_mensuel": round(interet_mensuel, 2),
        "interets_totaux": round(interets_totaux, 2),
        "interets_annee1": round(interets_annee1, 2),
        "interets_annee2": round(interets_annee2, 2),
        "interets_annee3": round(interets_annee3, 2)
    }  

def page_financement():
    st.title("Financement des Besoins de Démarrage")
    
    data = st.session_state.get("data", {})
    
    # Initialiser la section des financements
    if "financements" not in data:
        data["financements"] = {}
    
    financements_dict = data["financements"]
    
    total_financement = 0.0
    
    st.subheader("Apports")
    
    # Apport personnel ou familial
    apport_personnel = st.number_input(
        "Apport personnel ou familial ($)",
        min_value=0.0,
        key="financement_apport_personnel",
        value=financements_dict.get("Apport personnel ou familial", 0.00)
    )
    financements_dict["Apport personnel ou familial"] = apport_personnel
    total_financement += apport_personnel
    
    # Apports en nature (en valeur)
    apport_nature = st.number_input(
        "Apports en nature (en valeur) ($)",
        min_value=0.0,
        key="financement_apport_nature",
        value=financements_dict.get("Apports en nature (en valeur)", 0.00)
    )
    financements_dict["Apports en nature (en valeur)"] = apport_nature
    total_financement += apport_nature
    
    st.subheader("Prêts")
    
    # Nombre de prêts (maximum 3)
    num_prets = 3  # Limité à 3 prêts comme demandé
    
    interets_prets = {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    }
    
    for i in range(1, num_prets + 1):
        st.markdown(f"#### Prêt {i}")
        pret_name = st.text_input(
            f"Nom du prêt {i}",
            value=financements_dict.get(f"Prêt {i}", {}).get("nom", f"Prêt {i}"),
            key=f"pret_{i}_nom"
        )
        pret_montant = st.number_input(
            f"Montant du {pret_name} ($)",
            min_value=0.0,
            value=financements_dict.get(f"Prêt {i}", {}).get("montant", 0.0),
            key=f"pret_{i}_montant"
        )
        pret_taux = st.number_input(
            f"Taux du {pret_name} (%)",
            min_value=0.0,
            max_value=100.0,
            value=financements_dict.get(f"Prêt {i}", {}).get("taux", 0.0),
            key=f"pret_{i}_taux"
        )
        pret_duree = st.number_input(
            f"Durée du {pret_name} (en mois)",
            min_value=1,
            value=financements_dict.get(f"Prêt {i}", {}).get("duree", 12),
            key=f"pret_{i}_duree"
        )
        
        # Stocker les détails du prêt
        financements_dict[f"Prêt {i}"] = {
            "nom": pret_name,
            "montant": pret_montant,
            "taux": pret_taux,
            "duree": pret_duree
        }
        total_financement += pret_montant
        
        # Calculer les détails du remboursement du prêt
        if pret_montant > 0 and pret_taux > 0 and pret_duree > 0:
            pret_info = calculer_pret_interet_fixe(pret_montant, pret_taux, pret_duree)
            # Stocker les résultats du calcul
            financements_dict[f"Prêt {i}"].update(pret_info)
            # Ajouter les intérêts par année
            interets_prets["annee1"] += pret_info["interets_annee1"]
            interets_prets["annee2"] += pret_info["interets_annee2"]
            interets_prets["annee3"] += pret_info["interets_annee3"]
            
            # Afficher les détails du prêt pour vérification
            st.write(f"**Détails du {pret_name}:**")
            st.write(f"Mensualité : {pret_info['mensualite']:.2f} $")
            st.write(f"Total à rembourser : {pret_info['total_a_rembourser']:.2f} $")
            st.write(f"Principal mensuel : {pret_info['principal_mensuel']:.2f} $")
            st.write(f"Intérêt mensuel : {pret_info['interet_mensuel']:.2f} $")
            st.write(f"Intérêts totaux : {pret_info['interets_totaux']:.2f} $")
            st.write(f"Intérêts Année 1 : {pret_info['interets_annee1']:.2f} $")
            st.write(f"Intérêts Année 2 : {pret_info['interets_annee2']:.2f} $")
            st.write(f"Intérêts Année 3 : {pret_info['interets_annee3']:.2f} $")
            st.write("---")
    
    st.subheader("Subventions")
    
    # Nombre de subventions (maximum 2)
    num_subventions = 2  # Limité à 2 subventions comme demandé
    
    for i in range(1, num_subventions + 1):
        st.markdown(f"#### Subvention {i}")
        subvention_name = st.text_input(
            f"Nom de la subvention {i}",
            value=financements_dict.get(f"Subvention {i}", {}).get("nom", f"Subvention {i}"),
            key=f"subvention_{i}_nom"
        )
        subvention_montant = st.number_input(
            f"Montant de {subvention_name} ($)",
            min_value=0.0,
            value=financements_dict.get(f"Subvention {i}", {}).get("montant", 0.0),
            key=f"subvention_{i}_montant"
        )
        # Stocker les détails de la subvention
        financements_dict[f"Subvention {i}"] = {
            "nom": subvention_name,
            "montant": subvention_montant
        }
        total_financement += subvention_montant
    
    st.subheader("Autres Financements")
    
    # Autre financement
    autre_financement = st.number_input(
        "Autre financement ($)",
        min_value=0.0,
        key="financement_autre",
        value=financements_dict.get("Autre financement", 0.00)
    )
    financements_dict["Autre financement"] = autre_financement
    total_financement += autre_financement
    
    st.write("---")
    st.markdown(f"**Total des Financements :** {total_financement:,.2f} $")
    
    # Validation du total des financements
    besoin_total = data.get("besoins", 0.0)  # Assurez-vous que cette clé existe dans vos données
    if besoin_total > 0 and total_financement != besoin_total:
        st.error(f"Le total des financements ({total_financement:,.2f} $) ne correspond pas au besoin total ({besoin_total:,.2f} $). Veuillez ajuster les montants.")
    elif besoin_total > 0:
        st.success(f"Le total des financements correspond au besoin total ({besoin_total:,.2f} $).")
    
    # Stocker les données dans la session
    data["financements"] = financements_dict
    data["total_financement"] = total_financement
    data["interets_prets"] = interets_prets  # Stocker les intérêts des prêts
    
    st.session_state["data"] = data


def page_charges_fixes():
    st.title("Charges Fixes sur 3 Années")
    
    data = st.session_state.get("data", {})
    
    charges_fixes = [
        "Assurances véhicule et RC pro", "Téléphone, internet", "Autres abonnements",
        "Carburant", "Frais de déplacement / hébergement", "Eau, électricité, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et matériel",
        "Nettoyage des locaux", "Budget publicité et communication", "Emplacements",
        "Expert comptable, avocats", "Frais bancaires et terminal carte bleue", "Taxes, CFE"
    ]
    
    # Initialisation des charges fixes si non présentes
    if "charges_fixes" not in data:
        data["charges_fixes"] = {"annee1": {}, "annee2": {}, "annee3": {}}
        for charge in charges_fixes:
            data["charges_fixes"]["annee1"][charge] = 0.0
            data["charges_fixes"]["annee2"][charge] = 0.0
            data["charges_fixes"]["annee3"][charge] = 0.0
    charges_fixes_dict = data["charges_fixes"]
    
    # Initialisation des charges supplémentaires si non présentes
    if "charges_supplementaires" not in data:
        data["charges_supplementaires"] = []
    
    # Fonctions de mise à jour
    def update_year1(charge):
        year1_key = f"charge_{charge}_annee1"
        year2_key = f"charge_{charge}_annee2"
        year3_key = f"charge_{charge}_annee3"
        
        year1_val = st.session_state.get(year1_key, 0.0)
        
        # Mettre à jour année 2 et 3 seulement si l'utilisateur n'a pas déjà modifié ces champs
        if st.session_state.get(f"updated_{year2_key}", False) == False:
            st.session_state[year2_key] = year1_val
            charges_fixes_dict["annee2"][charge] = year1_val
        if st.session_state.get(f"updated_{year3_key}", False) == False:
            st.session_state[year3_key] = year1_val
            charges_fixes_dict["annee3"][charge] = year1_val

    def update_year2(charge):
        year2_key = f"charge_{charge}_annee2"
        year3_key = f"charge_{charge}_annee3"
        
        year2_val = st.session_state.get(year2_key, 0.0)
        
        # Mettre à jour année 3 seulement si l'utilisateur n'a pas déjà modifié ce champ
        if st.session_state.get(f"updated_{year3_key}", False) == False:
            st.session_state[year3_key] = year2_val
            charges_fixes_dict["annee3"][charge] = year2_val

    def update_year3(charge):
        # Indiquer que l'année 3 a été mise à jour manuellement
        year3_key = f"charge_{charge}_annee3"
        st.session_state[f"updated_{year3_key}"] = True

    st.subheader("Charges Fixes par Défaut")
    for charge in charges_fixes:
        col1, col2, col3 = st.columns(3)
        with col1:
            year1_key = f"charge_{charge}_annee1"
            if year1_key not in st.session_state:
                st.session_state[year1_key] = charges_fixes_dict["annee1"].get(charge, 0.0)
            montant1 = st.number_input(
                f"{charge} - Année 1 ($)",
                min_value=0.0,
                key=year1_key,
                on_change=update_year1,
                args=(charge,),
                value=st.session_state[year1_key]
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            year2_key = f"charge_{charge}_annee2"
            if year2_key not in st.session_state:
                st.session_state[year2_key] = charges_fixes_dict["annee2"].get(charge, 0.0)
                st.session_state[f"updated_{year2_key}"] = False
            montant2 = st.number_input(
                f"{charge} - Année 2 ($)",
                min_value=0.0,
                key=year2_key,
                on_change=update_year2,
                args=(charge,),
                value=st.session_state[year2_key]
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            year3_key = f"charge_{charge}_annee3"
            if year3_key not in st.session_state:
                st.session_state[year3_key] = charges_fixes_dict["annee3"].get(charge, 0.0)
                st.session_state[f"updated_{year3_key}"] = False
            montant3 = st.number_input(
                f"{charge} - Année 3 ($)",
                min_value=0.0,
                key=year3_key,
                on_change=update_year3,
                args=(charge,),
                value=st.session_state[year3_key]
            )
            charges_fixes_dict["annee3"][charge] = montant3
        
    # Charges supplémentaires
    st.write("---")
    st.subheader("Ajouter des Charges Supplémentaires")
    
    nouvelle_charge = st.text_input("Nom de la nouvelle charge :", key="nouvelle_charge")
    
    if st.button("Ajouter la charge"):
        nouvelle_charge = nouvelle_charge.strip()
        if nouvelle_charge and nouvelle_charge not in data["charges_supplementaires"]:
            data["charges_supplementaires"].append(nouvelle_charge)
            charges_fixes_dict["annee1"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee2"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee3"][nouvelle_charge] = 0.0
            # Réinitialiser le champ de texte
            st.session_state["nouvelle_charge"] = ""
    
    for charge in data["charges_supplementaires"]:
        col1, col2, col3 = st.columns(3)
        with col1:
            year1_key = f"charge_{charge}_supp_annee1"
            if year1_key not in st.session_state:
                st.session_state[year1_key] = charges_fixes_dict["annee1"].get(charge, 0.0)
            montant1 = st.number_input(
                f"{charge} - Année 1 ($)",
                min_value=0.0,
                key=year1_key,
                on_change=update_year1,
                args=(charge,),
                value=st.session_state[year1_key]
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            year2_key = f"charge_{charge}_supp_annee2"
            if year2_key not in st.session_state:
                st.session_state[year2_key] = charges_fixes_dict["annee2"].get(charge, 0.0)
                st.session_state[f"updated_{year2_key}"] = False
            montant2 = st.number_input(
                f"{charge} - Année 2 ($)",
                min_value=0.0,
                key=year2_key,
                on_change=update_year2,
                args=(charge,),
                value=st.session_state[year2_key]
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            year3_key = f"charge_{charge}_supp_annee3"
            if year3_key not in st.session_state:
                st.session_state[year3_key] = charges_fixes_dict["annee3"].get(charge, 0.0)
                st.session_state[f"updated_{year3_key}"] = False
            montant3 = st.number_input(
                f"{charge} - Année 3 ($)",
                min_value=0.0,
                key=year3_key,
                on_change=update_year3,
                args=(charge,),
                value=st.session_state[year3_key]
            )
            charges_fixes_dict["annee3"][charge] = montant3
    
    # Calcul des totaux
    total_annee1 = sum(charges_fixes_dict["annee1"].values())
    total_annee2 = sum(charges_fixes_dict["annee2"].values())
    total_annee3 = sum(charges_fixes_dict["annee3"].values())
    
    data["total_charges_fixes_annee1"] = total_annee1
    data["total_charges_fixes_annee2"] = total_annee2
    data["total_charges_fixes_annee3"] = total_annee3
    
    st.write("---")
    st.markdown(f"**Total Charges Fixes Année 1 :** {total_annee1:.2f} $")
    st.markdown(f"**Total Charges Fixes Année 2 :** {total_annee2:.2f} $")
    st.markdown(f"**Total Charges Fixes Année 3 :** {total_annee3:.2f} $")
    
    st.session_state["data"] = data

def page_chiffre_affaires():
    st.title("Chiffre d'Affaires Prévisionnel")
    
    data = st.session_state.get("data", {})
    type_vente = data.get("informations_generales", {}).get("type_vente", "Marchandises")
    
    data["chiffre_affaires"] = data.get("chiffre_affaires", {})
    chiffre_affaires_dict = data["chiffre_affaires"]
    
    mois = [f"Mois {i}" for i in range(1, 13)]
    
    # Fonctions de mise à jour
    def update_jours_travailles(nom_vente):
        key_jours_mois1 = f"{nom_vente}_Mois 1_jours"
        new_val = st.session_state.get(key_jours_mois1, 0)
        for mois_nom in mois[1:]:
            key = f"{nom_vente}_{mois_nom}_jours"
            if not st.session_state.get(f"updated_{key}", False):
                st.session_state[key] = new_val
                chiffre_affaires_dict[key] = new_val

    def update_ca_moyen_jour(nom_vente):
        key_ca_mois1 = f"{nom_vente}_Mois 1_ca_moyen"
        new_val = st.session_state.get(key_ca_mois1, 0.0)
        for mois_nom in mois[1:]:
            key = f"{nom_vente}_{mois_nom}_ca_moyen"
            if not st.session_state.get(f"updated_{key}", False):
                st.session_state[key] = new_val
                chiffre_affaires_dict[key] = new_val

    def mark_updated(key):
        st.session_state[f"updated_{key}"] = True

    def calcul_chiffre_affaires(nom_vente):
        mois_list = [f"Mois {i}" for i in range(1, 13)]
        data_ca = []
        
        st.subheader(f"Année 1 - {nom_vente}")
        for mois_nom in mois_list:
            col1, col2, col3 = st.columns(3)
            key_jours = f"{nom_vente}_{mois_nom}_jours"
            key_ca_moyen = f"{nom_vente}_{mois_nom}_ca_moyen"
            key_ca = f"{nom_vente}_{mois_nom}_ca"
            
            with col1:
                if mois_nom == "Mois 1":
                    montant_jours = st.number_input(
                        f"{mois_nom} - Nombre de jours travaillés",
                        min_value=0,
                        key=key_jours,
                        value=chiffre_affaires_dict.get(key_jours, 0),
                        on_change=update_jours_travailles,
                        args=(nom_vente,)
                    )
                else:
                    montant_jours = st.number_input(
                        f"{mois_nom} - Nombre de jours travaillés",
                        min_value=0,
                        key=key_jours,
                        value=chiffre_affaires_dict.get(key_jours, 0),
                        on_change=lambda key=key_jours: mark_updated(key)
                    )
                chiffre_affaires_dict[key_jours] = montant_jours
            
            with col2:
                if mois_nom == "Mois 1":
                    montant_ca_moyen = st.number_input(
                        f"{mois_nom} - Chiffre d'affaires moyen / jour ($)",
                        min_value=0.0,
                        key=key_ca_moyen,
                        value=chiffre_affaires_dict.get(key_ca_moyen, 0.0),
                        on_change=update_ca_moyen_jour,
                        args=(nom_vente,)
                    )
                else:
                    montant_ca_moyen = st.number_input(
                        f"{mois_nom} - Chiffre d'affaires moyen / jour ($)",
                        min_value=0.0,
                        key=key_ca_moyen,
                        value=chiffre_affaires_dict.get(key_ca_moyen, 0.0),
                        on_change=lambda key=key_ca_moyen: mark_updated(key)
                    )
                chiffre_affaires_dict[key_ca_moyen] = montant_ca_moyen
            
            ca_mensuel = montant_jours * montant_ca_moyen
            chiffre_affaires_dict[key_ca] = ca_mensuel
            data_ca.append({
                "mois": mois_nom,
                "jours_travailles": montant_jours,
                "ca_moyen_jour": montant_ca_moyen,
                "ca_mensuel": ca_mensuel
            })
            
            with col3:
                st.write(f"CA mensuel: {ca_mensuel:.2f} $")
        
        df_ca = pd.DataFrame(data_ca)
        total_ca_annee1 = df_ca["ca_mensuel"].sum()
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee1"] = total_ca_annee1
        
        st.write("---")
        st.markdown(f"**Total Chiffre d'Affaires Année 1 ({nom_vente}) :** {total_ca_annee1:.2f} $")
        
        # Pourcentages d'augmentation
        key_aug_annee2 = f"{nom_vente}_augmentation_annee2"
        key_aug_annee3 = f"{nom_vente}_augmentation_annee3"
        pourcentage_augmentation_annee2 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'année 1 et l'année 2 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee2,
            value=chiffre_affaires_dict.get(key_aug_annee2, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee2] = pourcentage_augmentation_annee2
        pourcentage_augmentation_annee3 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'année 2 et l'année 3 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee3,
            value=chiffre_affaires_dict.get(key_aug_annee3, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee3] = pourcentage_augmentation_annee3
        
        total_ca_annee2 = total_ca_annee1 * (1 + pourcentage_augmentation_annee2 / 100)
        total_ca_annee3 = total_ca_annee2 * (1 + pourcentage_augmentation_annee3 / 100)
        
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee2"] = total_ca_annee2
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee3"] = total_ca_annee3
        
        st.markdown(f"**Total Chiffre d'Affaires Année 2 ({nom_vente}) :** {total_ca_annee2:.2f} $")
        st.markdown(f"**Total Chiffre d'Affaires Année 3 ({nom_vente}) :** {total_ca_annee3:.2f} $")
    
    if type_vente in ["Marchandises", "Mixte"]:
        calcul_chiffre_affaires("Marchandises")
    if type_vente in ["Services", "Mixte"]:
        calcul_chiffre_affaires("Services")
    
    # Calcul du total CA toutes ventes
    total_ca_annee1 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee1", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee2 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee2", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee3 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee3", 0.0) for type in ["Marchandises", "Services"]
    )
    
    data["total_chiffre_affaires_annee1"] = total_ca_annee1
    data["total_chiffre_affaires_annee2"] = total_ca_annee2
    data["total_chiffre_affaires_annee3"] = total_ca_annee3
    
    st.write("---")
    st.markdown(f"**Total Chiffre d'Affaires Année 1 (toutes ventes) :** {total_ca_annee1:.2f} $")
    st.markdown(f"**Total Chiffre d'Affaires Année 2 (toutes ventes) :** {total_ca_annee2:.2f} $")
    st.markdown(f"**Total Chiffre d'Affaires Année 3 (toutes ventes) :** {total_ca_annee3:.2f} $")
    
    st.session_state["data"] = data
# Section 6 : Charges Variables
def page_charges_variables():
    st.title("Charges Variables")
    
    data = st.session_state["data"]
    type_vente = data["informations_generales"].get("type_vente", "Marchandises")
    
    if type_vente in ["Marchandises", "Mixte"]:
        st.markdown("""
        ### Vos charges variables
        Les charges variables sont liées au niveau d’activité ou à la production. 
        Il s’agit des achats de marchandises destinées à être revendues, des achats de matières destinées à être transformées, 
        des commissions versées à des agents commerciaux.
        """)
        
        data["charges_variables"] = data.get("charges_variables", {})
        charges_variables = data["charges_variables"]
        
        # Coût d'achat des marchandises en %
        cout_achat_marchandises_pct = st.number_input(
            "Quel est, en % du prix de vente, le coût d'achat de vos marchandises ? (concerne uniquement le chiffre d'affaires vente de marchandises)",
            min_value=0.0,
            max_value=100.0,
            format="%.2f",
            key="cout_achat_marchandises_pct",
            value=charges_variables.get("cout_achat_marchandises_pct", 0.0)
        )
        charges_variables["cout_achat_marchandises_pct"] = cout_achat_marchandises_pct
        
        st.write(f"Coût d'achat des marchandises : {cout_achat_marchandises_pct:.2f}% du prix de vente")
        
        total_ca_marchandises_annee1 = data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0)
        total_charges_variables = total_ca_marchandises_annee1 * cout_achat_marchandises_pct / 100.0
        
        data["total_charges_variables"] = total_charges_variables
        
        st.write(f"Total des Charges Variables Année 1 : {total_charges_variables:.2f} $")
        
    else:
        st.info("Cette section est uniquement applicable si vous vendez des marchandises ou des services mixtes.")
        data["total_charges_variables"] = 0.0
    
    st.session_state["data"] = data

# Section 7 : Fonds de Roulement
def page_fonds_roulement():
    st.title("Votre Besoin en Fonds de Roulement")
    
    data = st.session_state["data"]
    
    st.markdown("""
    ### Déterminez votre besoin en fonds de roulement
    Le fonds de roulement représente le montant nécessaire pour financer le cycle d'exploitation de votre entreprise.
    """)
    
    data["fonds_roulement"] = data.get("fonds_roulement", {})
    fonds_roulement = data["fonds_roulement"]
    
    duree_credits_clients = st.number_input(
        "Durée moyenne des crédits accordés aux clients (en jours) :",
        min_value=0,
        help="Temps moyen qu'un client met pour vous payer.",
        key="duree_credits_clients",
        value=fonds_roulement.get("duree_credits_clients", 0)
    )
    fonds_roulement["duree_credits_clients"] = duree_credits_clients
    
    duree_dettes_fournisseurs = st.number_input(
        "Durée moyenne des dettes fournisseurs (en jours) :",
        min_value=0,
        help="Temps moyen que vous mettez pour payer vos fournisseurs.",
        key="duree_dettes_fournisseurs",
        value=fonds_roulement.get("duree_dettes_fournisseurs", 0)
    )
    fonds_roulement["duree_dettes_fournisseurs"] = duree_dettes_fournisseurs
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    
    bfr = (total_ca_annee1 * duree_credits_clients / 360) - (total_charges_variables * duree_dettes_fournisseurs / 360)
    fonds_roulement["bfr"] = bfr
    
    st.write("---")
    st.markdown(f"**Durée moyenne des crédits clients :** {duree_credits_clients} jours")
    st.markdown(f"**Durée moyenne des dettes fournisseurs :** {duree_dettes_fournisseurs} jours")
    st.markdown(f"**Besoin en Fonds de Roulement (BFR) Année 1 :** {bfr:.2f} $")
    
    st.session_state["data"] = data

# Section 8 : Salaires
def page_salaires():
    st.title("Salaires Employés et Rémunération Chef d'Entreprise")
    
    data = st.session_state["data"]
    data["salaires"] = data.get("salaires", {"employes": {}, "dirigeants": {}})
    salaires = data["salaires"]
    
    st.markdown("""
    ### Saisissez les salaires et rémunérations pour les 3 années
    Veuillez entrer les chiffres annuels pour les salaires des employés et la rémunération nette des dirigeants.
    """)
    
    st.subheader("Salaires Employés (Net)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["employes"][key] = st.number_input(
            f"Salaires Employés Année {annee} ($)",
            min_value=0.0,
            key=f"salaires_employes_annee_{annee}",
            value=salaires["employes"].get(key, 0.0)
        )
    
    st.subheader("Rémunération Nette Dirigeant(s)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["dirigeants"][key] = st.number_input(
            f"Rémunération Dirigeant Année {annee} ($)",
            min_value=0.0,
            key=f"remuneration_dirigeant_annee_{annee}",
            value=salaires["dirigeants"].get(key, 0.0)
        )
    
    st.write("---")
    accre = st.selectbox(
        "Le(s) dirigeant(s) bénéficient-ils de l'ACRE ?",
        options=["Oui", "Non"],
        key="accre",
        index=["Oui", "Non"].index(data.get("accre", "Non")),
        help="Veuillez sélectionner 'Oui' si les dirigeants bénéficient de l'ACRE. Cette question est obligatoire."
    )
    data["accre"] = accre
    
    total_salaires_annee1 = salaires["employes"]["annee1"] + salaires["dirigeants"]["annee1"]
    total_salaires_annee2 = salaires["employes"]["annee2"] + salaires["dirigeants"]["annee2"]
    total_salaires_annee3 = salaires["employes"]["annee3"] + salaires["dirigeants"]["annee3"]
    
    data["total_salaires_annee1"] = total_salaires_annee1
    data["total_salaires_annee2"] = total_salaires_annee2
    data["total_salaires_annee3"] = total_salaires_annee3
    
    st.write("---")
    st.markdown(f"**Total Salaires et Rémunération Année 1 :** {total_salaires_annee1:.2f} $")
    st.markdown(f"**Total Salaires et Rémunération Année 2 :** {total_salaires_annee2:.2f} $")
    st.markdown(f"**Total Salaires et Rémunération Année 3 :** {total_salaires_annee3:.2f} $")
    
    st.session_state["data"] = data

# Section 9 : Contrôle de Rentabilité
def page_rentabilite():
    st.title("Contrôle de Rentabilité")
    
    data = st.session_state["data"]
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_chiffre_affaires = data.get("total_chiffre_affaires_annee1", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    
    if total_chiffre_affaires > 0:
        marge_brute = ((total_chiffre_affaires - total_charges_variables) / total_chiffre_affaires) * 100.0
    else:
        marge_brute = 0.0
    
    charges_fixes_totales = total_charges_fixes_annee1 + total_salaires_annee1
    if marge_brute > 0:
        seuil_rentabilite = charges_fixes_totales / (marge_brute / 100.0)
    else:
        seuil_rentabilite = 0.0
    
    if total_chiffre_affaires >= seuil_rentabilite:
        rentabilite = "Rentable"
        message_rentabilite = "L'entreprise est rentable."
        couleur_rentabilite = "green"
    else:
        rentabilite = "Non rentable"
        message_rentabilite = "L'entreprise n'est pas rentable. Il faut augmenter le chiffre d'affaires ou réduire les charges."
        couleur_rentabilite = "red"
    
    data["marge_brute"] = marge_brute
    data["seuil_rentabilite"] = seuil_rentabilite
    data["rentabilite"] = rentabilite
    
    st.write("---")
    st.markdown(f"**Marge Brute :** {marge_brute:.2f} %")
    st.markdown(f"**Seuil de Rentabilité :** {seuil_rentabilite:.2f} $")
    st.markdown(f"<div style='background-color:{couleur_rentabilite}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{rentabilite}</strong> - {message_rentabilite}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 10 : Trésorerie de Départ
def page_tresorerie():
    st.title("Contrôle du Niveau de votre Trésorerie de Départ")
    data = st.session_state["data"]
    besoins_demarrage=data.get("besoins_demarrage", 0.0)
    tresorerie_depart1 = besoins_demarrage.get("Trésorerie de départ", 0.0)
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    st.markdown(f"**Trésorerie de départ :** {tresorerie_depart1 :.2f} $")
    tresorerie_depart=tresorerie_depart1
    
    data["tresorerie_depart"] = tresorerie_depart
    
    seuil_tresorerie = total_charges_fixes_annee1 / 4.0  # 3 mois de charges fixes
    if tresorerie_depart >= seuil_tresorerie:
        niveau_tresorerie = "Adéquate"
        message_tresorerie = "Votre trésorerie de départ est adéquate pour couvrir les charges initiales."
        couleur_tresorerie = "green"
    else:
        niveau_tresorerie = "Trop faible"
        message_tresorerie = "Votre trésorerie de départ est trop faible. Prévoyez plus de trésorerie pour couvrir les charges."
        couleur_tresorerie = "red"
    
    data["niveau_tresorerie"] = niveau_tresorerie
    
    st.write("---")
    st.markdown(f"### Résultat pour la 1ère année :")
    st.markdown(f"<div style='background-color:{couleur_tresorerie}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{niveau_tresorerie}</strong> - {message_tresorerie}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 11 : Récapitulatif
def page_recapitulatif():
    st.title("Récapitulatif Complet des Données")
    
    data = st.session_state["data"]
    
    st.subheader("1. Informations Générales")
    info = data.get("informations_generales", {})
    st.write(f"Prénom, nom : {info.get('prenom_nom', '')}")
    st.write(f"Intitulé du projet : {info.get('intitule_projet', '')}")
    st.write(f"Statut juridique : {info.get('statut_juridique', '')}")
    st.write(f"Téléphone : {info.get('telephone', '')}")
    st.write(f"Email : {info.get('email', '')}")
    st.write(f"Ville : {info.get('ville', '')}")
    st.write(f"Type de vente : {info.get('type_vente', '')}")
    
    st.subheader("2. Besoins de Démarrage")
    besoins = data.get("besoins_demarrage", {})
    total_besoins = data.get("total_besoins", 0.0)
    for besoin, montant in besoins.items():
        st.write(f"{besoin} : {montant:.2f} $")
    st.write(f"**Total des Besoins de Démarrage : {total_besoins:.2f} $**")
    
    st.title("Récapitulatif des Financements")
    data = st.session_state.get("data", {})
    financements_dict = data.get("financements", {})
    total_financement = data.get("total_financement", 0.0)
    st.subheader("Financements")
    for financement, details in financements_dict.items():
        if isinstance(details, dict):
            montant = details.get("montant", 0.0)
            st.write(f"{details.get('nom', financement)} : {montant:.2f} $")
        else:
            montant = details
            st.write(f"{financement} : {montant:.2f} $")
    
    st.markdown(f"**Total des Financements :** {total_financement:.2f} $")
    
    st.subheader("4. Charges Fixes sur 3 Années")
    charges_fixes_dict = data.get("charges_fixes", {"annee1": {}, "annee2": {}, "annee3": {}})
    total_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_annee2 = data.get("total_charges_fixes_annee2", 0.0)
    total_annee3 = data.get("total_charges_fixes_annee3", 0.0)
    charges_supp = data.get("charges_supplementaires", [])
    
    for charge in charges_fixes_dict["annee1"].keys():
        montant1 = charges_fixes_dict["annee1"].get(charge, 0.0)
        montant2 = charges_fixes_dict["annee2"].get(charge, 0.0)
        montant3 = charges_fixes_dict["annee3"].get(charge, 0.0)
        st.write(f"{charge} - Année 1 : {montant1:.2f} $, Année 2 : {montant2:.2f} $, Année 3 : {montant3:.2f} $")
    
    st.write(f"**Total Charges Fixes Année 1 : {total_annee1:.2f} $**")
    st.write(f"**Total Charges Fixes Année 2 : {total_annee2:.2f} $**")
    st.write(f"**Total Charges Fixes Année 3 : {total_annee3:.2f} $**")
    
    st.subheader("5. Chiffre d'Affaires Prévisionnel")
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_ca_annee2 = data.get("total_chiffre_affaires_annee2", 0.0)
    total_ca_annee3 = data.get("total_chiffre_affaires_annee3", 0.0)
    
    st.write(f"Total Chiffre d'Affaires Année 1 : {total_ca_annee1:.2f} $")
    st.write(f"Total Chiffre d'Affaires Année 2 : {total_ca_annee2:.2f} $")
    st.write(f"Total Chiffre d'Affaires Année 3 : {total_ca_annee3:.2f} $")
    
    st.subheader("6. Charges Variables")
    cout_achat_marchandises_pct = data.get("charges_variables", {}).get("cout_achat_marchandises_pct", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    st.write(f"Coût d'achat des marchandises (% du CA) : {cout_achat_marchandises_pct:.2f} %")
    st.write(f"Total Charges Variables Année 1 : {total_charges_variables:.2f} $")
    
    st.subheader("7. Fonds de Roulement")
    fonds_roulement = data.get("fonds_roulement", {})
    duree_credits_clients = fonds_roulement.get("duree_credits_clients", 0)
    duree_dettes_fournisseurs = fonds_roulement.get("duree_dettes_fournisseurs", 0)
    bfr = fonds_roulement.get("bfr", 0.0)
    st.write(f"Durée moyenne des crédits clients : {duree_credits_clients} jours")
    st.write(f"Durée moyenne des dettes fournisseurs : {duree_dettes_fournisseurs} jours")
    st.write(f"Besoin en Fonds de Roulement (BFR) Année 1 : {bfr:.2f} $")
    
    st.subheader("8. Salaires et Rémunération")
    salaires = data.get("salaires", {})
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires_employes = salaires.get("employes", {}).get(key, 0.0)
        remuneration_dirigeants = salaires.get("dirigeants", {}).get(key, 0.0)
        st.write(f"Année {annee} : Salaires employés : {salaires_employes:.2f} $, Rémunération dirigeants : {remuneration_dirigeants:.2f} $")
        st.write(f"Total Salaires Année {annee} : {(salaires_employes + remuneration_dirigeants):.2f} $")
    
    st.subheader("9. Rentabilité")
    marge_brute = data.get("marge_brute", 0.0)
    seuil_rentabilite = data.get("seuil_rentabilite", 0.0)
    rentabilite = data.get("rentabilite", "Non rentable")
    st.write(f"Marge Brute : {marge_brute:.2f} %")
    st.write(f"Seuil de Rentabilité : {seuil_rentabilite:.2f} $")
    st.write(f"Rentabilité : {rentabilite}")
    
    st.subheader("10. Trésorerie de Départ")
    tresorerie_depart = data.get("tresorerie_depart", 0.0)
    niveau_tresorerie = data.get("niveau_tresorerie", "Trop faible")
    st.write(f"Montant de la Trésorerie Initiale : {tresorerie_depart:.2f} $")
    st.write(f"Niveau de Trésorerie : {niveau_tresorerie}")
    
    st.session_state["data"] = data
    
    
    
    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def page_investissements_et_financements(): 
    st.title("Investissements et Financements")
    
    # Initialiser la clé 'export_data' dans session_state si elle n'existe pas
    if 'export_data' not in st.session_state:
        st.session_state['export_data'] = {}
    
    # Récupérer les données de la session
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Initialiser une liste pour stocker toutes les lignes du tableau
    table_data = []
    
    # Immobilisations Incorporelles
    immobilisations_incorporelles = {
        "Frais d’établissement": data.get("besoins_demarrage", {}).get("Frais d’établissement", 0.0),
        "Frais d’ouverture de compteurs": data.get("besoins_demarrage", {}).get("Frais d’ouverture de compteurs", 0.0),
        "Logiciels, formations": data.get("besoins_demarrage", {}).get("Logiciels, formations", 0.0),
        "Dépôt de marque": data.get("besoins_demarrage", {}).get("Dépôt de marque", 0.0),
        "Droits d’entrée": data.get("besoins_demarrage", {}).get("Droits d’entrée", 0.0),
        "Achat fonds de commerce ou parts": data.get("besoins_demarrage", {}).get("Achat fonds de commerce ou parts", 0.0),
        "Droit au bail": data.get("besoins_demarrage", {}).get("Droit au bail", 0.0),
        "Caution ou dépôt de garantie": data.get("besoins_demarrage", {}).get("Caution ou dépôt de garantie", 0.0),
        "Frais de dossier": data.get("besoins_demarrage", {}).get("Frais de dossier", 0.0),
        "Frais de notaire": data.get("besoins_demarrage", {}).get("Frais de notaire", 0.0),
    }
    total_incorporelles = sum(immobilisations_incorporelles.values())
    table_data.append({
        "Investissements": "Immobilisations incorporelles",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_incorporelles:.2f}"
    })
    for desc, montant in immobilisations_incorporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Immobilisations Corporelles
    immobilisations_corporelles = {
        "Enseigne et éléments de communication": data.get("besoins_demarrage", {}).get("Enseigne et éléments de communication", 0.0),
        "Véhicule": data.get("besoins_demarrage", {}).get("Véhicule", 0.0),
        "Matériel professionnel": data.get("besoins_demarrage", {}).get("Matériel professionnel", 0.0),
        "Matériel autre": data.get("besoins_demarrage", {}).get("Matériel autre", 0.0),
        "Matériel de bureau": data.get("besoins_demarrage", {}).get("Matériel de bureau", 0.0),
    }
    total_corporelles = sum(immobilisations_corporelles.values())
    table_data.append({
        "Investissements": "Immobilisations corporelles",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_corporelles:.2f}"
    })
    for desc, montant in immobilisations_corporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Autres Investissements
    autres_investissements = {
        "Stock de matières et produits": data.get("besoins_demarrage", {}).get("Stock de matières et produits", 0.0),
        "Trésorerie de départ": data.get("besoins_demarrage", {}).get("Trésorerie de départ", 0.0)
    }
    total_autres = sum(autres_investissements.values())
    table_data.append({
        "Investissements": "Autres investissements",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_autres:.2f}"
    })
    for desc, montant in autres_investissements.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # TOTAL BESOINS
    total_besoins = total_incorporelles + total_corporelles + total_autres
    table_data.append({
        "Investissements": "TOTAL BESOINS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_besoins:.2f}"
    })
    
    # Section FINANCEMENT DES INVESTISSEMENTS
    table_data.append({
        "Investissements": "FINANCEMENT DES INVESTISSEMENTS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": ""
    })
    table_data.append({
        "Investissements": "Montant $ hors taxes",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": ""
    })
    
    # Apport Personnel
    financements = data.get("financements", {})
    apport_personnel = {
        "Apport personnel ou familial": financements.get("Apport personnel ou familial", 0.0),
        "Apports en nature (en valeur)": financements.get("Apports en nature (en valeur)", 0.0),
    }
    total_apport_personnel = sum(apport_personnel.values())
    table_data.append({
        "Investissements": "Apport personnel",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_apport_personnel:.2f}"
    })
    for desc, montant in apport_personnel.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant ($)": f"{montant:.2f}"
        })
    
    # Emprunts Dynamiques
    emprunts_keys = ["Prêt 1", "Prêt 2", "Prêt 3"]
    emprunts_list = []
    total_emprunts = 0.0

    # Collecter les détails des emprunts
    for i, key in enumerate(emprunts_keys, start=1):
        pret = financements.get(key, {})
        nom = pret.get("nom", "")
        taux = pret.get("taux", 0.0)
        duree = pret.get("duree", 0)
        montant = pret.get("montant", 0.0)
        
        # Définir le nom de l'emprunt
        emprunt_nom = nom if nom else f"Prêt {i}"
        
        # Ajouter les détails du prêt
        if montant > 0:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": f"{taux:.2f}%",
                "Durée (mois)": duree,
                "Montant ($)": f"{montant:.2f}"
            })
            total_emprunts += montant
        else:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": "-",
                "Durée (mois)": "-",
                "Montant ($)": "0.00"
            })

    # TOTAL EMPRUNTS - placé avant les emprunts individuels
    table_data.append({
        "Investissements": "TOTAL EMPRUNTS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_emprunts:.2f}"
    })

    # Ajouter les emprunts individuels après le total
    for emprunt in emprunts_list:
        table_data.append(emprunt)

     # Subventions Dynamiques
    subventions_keys = ["Subvention 1", "Subvention 2"]
    subventions_list = []
    total_subventions = 0.0
    
    # Calculer le total des subventions d'abord
    for i, key in enumerate(subventions_keys, start=1):
        subv = financements.get(key, {})
        nom = subv.get("nom", "")
        montant = subv.get("montant", 0.0)
        
        # Définir le nom de la subvention
        subvention_nom = nom if nom else f"Subvention {i}"
        
        # Ajouter les détails de la subvention
        if montant > 0:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Durée (mois)": "",
                "Montant ($)": f"{montant:.2f}"
            })
            total_subventions += montant
        else:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Durée (mois)": "",
                "Montant ($)": "0.00"
            })
    
    # TOTAL SUBVENTIONS - placé avant les subventions individuelles
    table_data.append({
        "Investissements": "TOTAL SUBVENTIONS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_subventions:.2f}"
    })
    
    # Ajouter les subventions individuelles après le total
    for subv in subventions_list:
        table_data.append(subv)
    
    # Autre Financement
    autre_financement = financements.get("Autre financement", 0.0)
    table_data.append({
        "Investissements": "Autre financement",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{autre_financement:.2f}"
    })
    
    # TOTAL RESSOURCES
    total_ressources = total_apport_personnel + total_emprunts + total_subventions + autre_financement
    table_data.append({
        "Investissements": "TOTAL RESSOURCES",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant ($)": f"{total_ressources:.2f}"
    })
    
    # Vérification de l'équilibre
    if total_ressources == total_besoins:
        equilibrium_message = "Le total des ressources couvre exactement les besoins."
        equilibrium_type = "success"
    elif total_ressources > total_besoins:
        surplus = total_ressources - total_besoins
        equilibrium_message = f"Les ressources dépassent les besoins de {surplus:.2f} $."
        equilibrium_type = "info"
    else:
        deficit = total_besoins - total_ressources
        equilibrium_message = f"Il manque {deficit:.2f} $ pour couvrir les besoins."
        equilibrium_type = "warning"
    
    if equilibrium_type == "success":
        st.success(equilibrium_message)
    elif equilibrium_type == "info":
        st.info(equilibrium_message)
    else:
        st.warning(equilibrium_message)
    
    st.write("---")
    
    # Créer le DataFrame unique avec les quatre colonnes
    df_unique = pd.DataFrame(table_data, columns=["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"])
    
    # Afficher le tableau dans Streamlit
    st.dataframe(df_unique.style.apply(lambda x: ['background-color: #f0f0f0' if pd.isna(v) else '' for v in x], axis=1))
    
    # Stocker les totaux dans la session
    data["total_investissements"] = total_besoins
    data["total_financements"] = total_ressources
    
    st.session_state["data"] = data
    
    # Stocker les données d'exportation dans la nouvelle session
    st.session_state['export_data_investissements'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": table_data,
        "equilibre": {
            "type": equilibrium_type,
            "message": equilibrium_message
        }
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger en Markdown"):
        markdown_content = f"# Investissements et Financements\n\n**Projet :** {projet}\n\n**Porteur de projet :** {porteur_projet}\n\n"
        
        # Convertir le DataFrame en Markdown
        markdown_content += df_unique.to_markdown(index=False)
        markdown_content += f"\n\n---\n\n{equilibrium_message}\n"
        
        markdown_bytes = markdown_content.encode('utf-8')
        st.download_button(
            label="Télécharger le Markdown",
            data=markdown_bytes,
            file_name="investissements_et_financements.md",
            mime="text/markdown"
        )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger en Word"):
        export_data = st.session_state.get('export_data', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            doc = Document()
            doc.add_heading('Investissements et Financements', level=1)
            doc.add_paragraph(f"**Projet :** {export_data['projet']}")
            doc.add_paragraph(f"**Porteur de projet :** {export_data['porteur_projet']}")
            doc.add_page_break()
            
            # Créer le DataFrame pour Word
            df_word = pd.DataFrame(export_data['table_data'], columns=["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"])
            
            # Ajouter le tableau au document Word
            table = doc.add_table(rows=1, cols=len(df_word.columns))
            table.style = 'Light List Accent 1'  # Choisissez un style approprié
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df_word.columns):
                hdr_cells[i].text = column
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for index, row in df_word.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
                    # Mettre en gras les catégories principales et les totaux
                    if row["Investissements"] in ["INVESTISSEMENTS", "Montant $ hors taxes",
                                                 "Immobilisations incorporelles", "Immobilisations corporelles",
                                                 "Autres investissements", "TOTAL BESOINS",
                                                 "FINANCEMENT DES INVESTISSEMENTS", "Apport personnel",
                                                 "Emprunt", "TOTAL EMPRUNTS", "Subvention",
                                                 "TOTAL SUBVENTIONS", "Autre financement", "TOTAL RESSOURCES"]:
                        run = row_cells[i].paragraphs[0].runs
                        if run:
                            run[0].font.bold = True
            doc.add_paragraph()
            doc.add_paragraph(export_data['equilibre']['message'])
            
            # Enregistrer le document dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="Télécharger le fichier Word",
                data=buffer,
                file_name="investissements_et_financements.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


def page_salaires_charges_sociales():
    st.title("Salaires et Charges Sociales")
    
    # Initialiser la clé 'export_data_salaires_charges_sociales' dans session_state si elle n'existe pas
    if 'export_data_salaires_charges_sociales' not in st.session_state:
        st.session_state['export_data_salaires_charges_sociales'] = {}
    
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    statut_juridique = data.get("informations_generales", {}).get("statut_juridique", "")
    benefice_accre = data.get("accre", "Non")  # Assurez-vous que cette information est bien stockée dans data
    
    # Déterminer le statut social du dirigeant en fonction du statut juridique
    if statut_juridique in ["Entreprise individuelle", "EURL (IS)", "EIRL (IS)", "Micro-entreprise"]:
        statut_social_dirigeant = "Travailleur Non Salarié (TNS)"
    elif statut_juridique in ["SARL (IS)", "SAS (IS)", "SASU (IS)"]:
        statut_social_dirigeant = "Assimilé Salarié"
    else:
        statut_social_dirigeant = "Autre"
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    st.write(f"**Statut juridique :** {statut_juridique}")
    st.write(f"**Bénéfice de l'ACRE :** {benefice_accre}")
    st.write(f"**Statut social du (des) dirigeant(s) :** {statut_social_dirigeant}")
    
    st.write("---")
    
    # Récupérer les données de salaires
    salaires = data.get("salaires", {})
    salaires_dirigeant = salaires.get("dirigeants", {})
    salaires_employes = salaires.get("employes", {})
    
    # Définir les taux de charges sociales en fonction du statut juridique et de l'ACCRE
    taux_charges_dirigeant = {
        # Sans ACCRE
        "Sans ACCRE": {
            "Travailleur Non Salarié (TNS)": 0.45,
            "Assimilé Salarié": 0.80,  # Taux approximatif pour les assimilés salariés
        },
        # Avec ACCRE
        "Avec ACCRE": {
            "Travailleur Non Salarié (TNS)": 0.22,
            "Assimilé Salarié": 0.50,  # Taux réduit pour les assimilés salariés avec ACRE
        }
    }
    
    # Sélection du taux approprié pour le dirigeant
    if benefice_accre.lower() == "oui":
        taux_dirigeant = taux_charges_dirigeant["Avec ACCRE"].get(statut_social_dirigeant, 0.45)
    else:
        taux_dirigeant = taux_charges_dirigeant["Sans ACCRE"].get(statut_social_dirigeant, 0.45)
    
    # Taux de charges sociales pour les employés
    taux_charges_employe = 0.72  # Comme indiqué, multiplier par 0.72 qu'il ait ACCRE ou pas
    
    # Préparation des données pour le tableau
    annees = ["Année 1", "Année 2", "Année 3"]
    remuneration_dirigeant = []
    augmentation_dirigeant = []
    charges_sociales_dirigeant = []
    remuneration_employes = []
    augmentation_employes = []
    charges_sociales_employes = []
    
    for i, annee in enumerate(annees):
        annee_key = f"annee{i+1}"
        # Rémunération du (des) dirigeants
        remu_dirigeant = salaires_dirigeant.get(annee_key, 0.0)
        remuneration_dirigeant.append(remu_dirigeant)
        # % augmentation dirigeant
        if i == 0:
            aug_dirigeant = "-"
        else:
            previous_remu_dirigeant = remuneration_dirigeant[i-1]
            if previous_remu_dirigeant != 0:
                aug_dirigeant_value = ((remu_dirigeant - previous_remu_dirigeant) / previous_remu_dirigeant) * 100
                aug_dirigeant = f"{aug_dirigeant_value:.2f}%"
            else:
                aug_dirigeant = "-"
        augmentation_dirigeant.append(aug_dirigeant)
        # Charges sociales du (des) dirigeant(s)
        charge_sociale_dirigeant = remu_dirigeant * taux_dirigeant
        charges_sociales_dirigeant.append(charge_sociale_dirigeant)
        
        # Salaires des employés
        remu_employes = salaires_employes.get(annee_key, 0.0)
        remuneration_employes.append(remu_employes)
        # % augmentation employés
        if i == 0:
            aug_employes = "-"
        else:
            previous_remu_employes = remuneration_employes[i-1]
            if previous_remu_employes != 0:
                aug_employes_value = ((remu_employes - previous_remu_employes) / previous_remu_employes) * 100
                aug_employes = f"{aug_employes_value:.2f}%"
            else:
                aug_employes = "-"
        augmentation_employes.append(aug_employes)
        # Charges sociales employés
        charge_sociale_employes = remu_employes * taux_charges_employe
        charges_sociales_employes.append(charge_sociale_employes)
    
    # Création du DataFrame pour l'affichage
    df = pd.DataFrame({
        "": ["Rémunération du (des) dirigeants", "% augmentation", "Charges sociales du (des) dirigeant(s)",
             "Salaires des employés", "% augmentation", "Charges sociales employés"],
        "Année 1": [f"{remuneration_dirigeant[0]:.2f} $", augmentation_dirigeant[0], f"{charges_sociales_dirigeant[0]:.2f} $",
                    f"{remuneration_employes[0]:.2f} $", augmentation_employes[0], f"{charges_sociales_employes[0]:.2f} $"],
        "Année 2": [f"{remuneration_dirigeant[1]:.2f} $", augmentation_dirigeant[1], f"{charges_sociales_dirigeant[1]:.2f} $",
                    f"{remuneration_employes[1]:.2f} $", augmentation_employes[1], f"{charges_sociales_employes[1]:.2f} $"],
        "Année 3": [f"{remuneration_dirigeant[2]:.2f} $", augmentation_dirigeant[2], f"{charges_sociales_dirigeant[2]:.2f} $",
                    f"{remuneration_employes[2]:.2f} $", augmentation_employes[2], f"{charges_sociales_employes[2]:.2f} $"]
    })
    
    st.table(df)
    
    # Stocker les charges sociales dans les données pour exportation
    data["charges_sociales"] = {
        "dirigeants": {
            "annee1": charges_sociales_dirigeant[0],
            "annee2": charges_sociales_dirigeant[1],
            "annee3": charges_sociales_dirigeant[2]
        },
        "employes": {
            "annee1": charges_sociales_employes[0],
            "annee2": charges_sociales_employes[1],
            "annee3": charges_sociales_employes[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Salaires et Charges Sociales
    export_table_data = []
    
    # Ajouter les lignes du tableau
    for index, row in df.iterrows():
        export_table_data.append({
            "Description": row[""],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_salaires_charges_sociales'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "statut_juridique": statut_juridique,
        "benefice_accre": benefice_accre,
        "statut_social_dirigeant": statut_social_dirigeant,
        "table_data": export_table_data
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Salaires en Markdown"):
        export_data = st.session_state.get('export_data_salaires_charges_sociales', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Salaires et Charges Sociales\n\n**Projet :** {export_data['projet']}\n\n"
            markdown_content += f"**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += f"**Statut juridique :** {export_data['statut_juridique']}\n\n"
            markdown_content += f"**Bénéfice de l'ACRE :** {export_data['benefice_accre']}\n\n"
            markdown_content += f"**Statut social du (des) dirigeant(s) :** {export_data['statut_social_dirigeant']}\n\n"
            markdown_content += "---\n\n"
            
            # Créer un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += f"\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="salaires_charges_sociales.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Salaires en Word"):
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donnée disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donnée disponible pour l'exportation des Investissements et Financements.")
            return
        
        doc = Document()
        
        # Ajouter la première table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Durée (mois)"]) if row["Durée (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant ($)"]
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme supplémentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxième table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Année 1"]
            row_cells_sal[2].text = row["Année 2"]
            row_cells_sal[3].text = row["Année 3"]
            
            # Mise en forme des lignes spécifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si nécessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter des informations supplémentaires si nécessaire
        doc.add_paragraph()
        doc.add_paragraph("Les charges sociales sont calculées en fonction des taux applicables.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Télécharger le fichier Word Complet",
            data=buffer,
            file_name="investissements_et_salaires_charges_sociales.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def page_detail_amortissements():
    st.title("Détail des Amortissements")
    
    # Initialiser la clé 'export_data_detail_amortissements' dans session_state si elle n'existe pas
    if 'export_data_detail_amortissements' not in st.session_state:
        st.session_state['export_data_detail_amortissements'] = {}
    
    data = st.session_state.get("data", {})
    
    st.write("---")
    
    # Récupérer la durée d'amortissement
    duree_amortissement = data.get("duree_amortissement", 3)  # Par défaut à 3 ans si non défini
    if duree_amortissement <= 0:
        st.warning("La durée d'amortissement doit être supérieure à zéro.")
        return
    
    # Récupérer les montants des investissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Fonction pour calculer les amortissements
    def calcul_amortissements(items):
        amortissements = {}
        total_amort = [0.0, 0.0, 0.0]
        for item in items:
            amount = besoins_demarrage.get(item, 0.0)
            annual_depreciation = amount / duree_amortissement if duree_amortissement > 0 else 0.0
            amortization_years = [0.0, 0.0, 0.0]
            for year in range(3):
                if year < duree_amortissement:
                    amortization_years[year] = annual_depreciation
                    total_amort[year] += annual_depreciation
            amortissements[item] = amortization_years
        return amortissements, total_amort
    
    # Incorporels
    incorporels_items = [
        "Frais d’établissement",
        "Logiciels, formations",
        "Droits d’entrée",
        "Frais de dossier",
        "Frais de notaire"
    ]
    incorporels_amortissements, total_incorporels_amort = calcul_amortissements(incorporels_items)
    
    # Corporels
    corporels_items = [
        "Enseigne et éléments de communication",
        "Véhicule",
        "Matériel professionnel",
        "Matériel autre",
        "Matériel de bureau"
    ]
    corporels_amortissements, total_corporels_amort = calcul_amortissements(corporels_items)
    
    # Total amortissements par année
    total_amortissements = [
        total_incorporels_amort[year] + total_corporels_amort[year] for year in range(3)
    ]
    
    # Création d'un tableau unique
    st.subheader("Amortissements")
    amortissements_data = []
    
    # Ajout des totaux des catégories
    amortissements_data.append({
        "Amortissement": "Amortissements incorporels",
        "Année 1": f"{total_incorporels_amort[0]:.2f}",
        "Année 2": f"{total_incorporels_amort[1]:.2f}",
        "Année 3": f"{total_incorporels_amort[2]:.2f}"
    })
    
    # Ajout d'une ligne vide pour la lisibilité
    amortissements_data.append({
        "Amortissement": "",
        "Année 1": "",
        "Année 2": "",
        "Année 3": ""
    })
    
    # Ajout des détails des amortissements incorporels
    for item in incorporels_items:
        amortization_years = incorporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Année 1": f"{amortization_years[0]:.2f}",
            "Année 2": f"{amortization_years[1]:.2f}",
            "Année 3": f"{amortization_years[2]:.2f}"
        })
    
    # Ajout d'une ligne vide pour la lisibilité
    amortissements_data.append({
        "Amortissement": "",
        "Année 1": "",
        "Année 2": "",
        "Année 3": ""
    })
    amortissements_data.append({
        "Amortissement": "Amortissements corporels",
        "Année 1": f"{total_corporels_amort[0]:.2f}",
        "Année 2": f"{total_corporels_amort[1]:.2f}",
        "Année 3": f"{total_corporels_amort[2]:.2f}"
    })
        # Ajout d'une ligne vide pour la lisibilité
    amortissements_data.append({
        "Amortissement": "",
        "Année 1": "",
        "Année 2": "",
        "Année 3": ""
    })
        
    # Ajout des détails des amortissements corporels
    for item in corporels_items:
        amortization_years = corporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Année 1": f"{amortization_years[0]:.2f}",
            "Année 2": f"{amortization_years[1]:.2f}",
            "Année 3": f"{amortization_years[2]:.2f}"
        })
    
    # Ajout d'une ligne vide pour la lisibilité
    amortissements_data.append({
        "Amortissement": "",
        "Année 1": "",
        "Année 2": "",
        "Année 3": ""
    })
    
    # Total amortissements
    amortissements_data.append({
        "Amortissement": "Total Amortissements",
        "Année 1": f"{total_amortissements[0]:.2f}",
        "Année 2": f"{total_amortissements[1]:.2f}",
        "Année 3": f"{total_amortissements[2]:.2f}"
    })
    
    # Création du DataFrame
    df_amortissements = pd.DataFrame(amortissements_data)
    
    # Affichage du tableau avec des bordures pour plus de clarté
    st.table(df_amortissements.style.set_properties(**{
        'text-align': 'right'
    }).set_table_styles([
        {'selector': 'th', 'props': [('text-align', 'center')]}
    ]))

    
    # Stocker les amortissements dans les données pour exportation
    data["amortissements"] = {
        "incorporels": {
            "annee1": total_incorporels_amort[0],
            "annee2": total_incorporels_amort[1],
            "annee3": total_incorporels_amort[2]
        },
        "corporels": {
            "annee1": total_corporels_amort[0],
            "annee2": total_corporels_amort[1],
            "annee3": total_corporels_amort[2]
        },
        "total": {
            "annee1": total_amortissements[0],
            "annee2": total_amortissements[1],
            "annee3": total_amortissements[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Détail des Amortissements
    export_table_amortissements = []
    for row in amortissements_data:
        export_table_amortissements.append({
            "Amortissement": row["Amortissement"],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_detail_amortissements'] = {
        "amortissements": export_table_amortissements
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Amortissements en Markdown"):
        export_data = st.session_state.get('export_data_detail_amortissements', {})
        if not export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Détail des Amortissements\n\n"
            markdown_content += "---\n\n"
            
            # Amortissements
            markdown_content += "## Amortissements\n\n"
            df_amortissements_md = pd.DataFrame(export_data['amortissements'])
            markdown_content += df_amortissements_md.to_markdown(index=False)
            markdown_content += "\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="detail_amortissements.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Amortissements en Word"):
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donnée disponible pour l'exportation des Amortissements.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donnée disponible pour l'exportation des Investissements et Financements.")
            return
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donnée disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        doc = Document()
        
        # Ajouter la première table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Durée (mois)"]) if row["Durée (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant ($)"]
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme supplémentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxième table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Année 1"]
            row_cells_sal[2].text = row["Année 2"]
            row_cells_sal[3].text = row["Année 3"]
            
            # Mise en forme des lignes spécifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si nécessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la troisième table : Détail des Amortissements
        doc.add_heading('Détail des Amortissements', level=1)
        
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        
        # Vérifier si les données d'amortissements sont disponibles
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donnée disponible pour l'exportation des Amortissements.")
            return
        
        # Créer le tableau Amortissements dans Word
        doc.add_heading('Amortissements', level=2)
        table_word_amort = doc.add_table(rows=1, cols=4)
        table_word_amort.style = 'Light List Accent 1'
        table_word_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_word_amort.rows[0].cells
        headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Amortissements au tableau
        for row in export_data_amortissements['amortissements']:
            row_cells_amort = table_word_amort.add_row().cells
            row_cells_amort[0].text = row["Amortissement"]
            row_cells_amort[1].text = row["Année 1"]
            row_cells_amort[2].text = row["Année 2"]
            row_cells_amort[3].text = row["Année 3"]
        
        # Ajouter des informations supplémentaires si nécessaire
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Télécharger le fichier Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        

def telecharger_document_complet():
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([export_data_investissements.get("table_data"),
                export_data_salaires.get("table_data"),
                export_data_amortissements.get("amortissements"),
                export_data_compte.get("table_data")]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
    doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
    doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_amortissements['amortissements']:
        row_cells = table_amort.add_row().cells
        row_cells[0].text = row.get("Amortissement", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells = table_compte.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le fichier Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def calculer_impot_societes(resultat_avant_impots):
    """
    Calcule l'Impôt sur les Sociétés (IS) selon la formule progressive.

    Args:
        resultat_avant_impots (float): Résultat avant impôts.

    Returns:
        float: Montant de l'IS.
    """
    if resultat_avant_impots < 0:
        return 0.0
    elif resultat_avant_impots > 38120:
        return 38120 * 0.15 + (resultat_avant_impots - 38120) * 0.28
    else:
        return resultat_avant_impots * 0.15

def page_compte_resultats_previsionnel():
    st.title("Compte de résultats prévisionnel sur 3 ans")
    
    # Initialiser la clé 'export_data_compte_resultats_previsionnel' dans session_state si elle n'existe pas
    if 'export_data_compte_resultats_previsionnel' not in st.session_state:
        st.session_state['export_data_compte_resultats_previsionnel'] = {}
    
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Préparation des données
    # Chiffre d'affaires
    ca_marchandises = [
        data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee3", 0.0)
    ]
    ca_services = [
        data["chiffre_affaires"].get("total_ca_Services_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee3", 0.0)
    ]
    total_ca = [
        ca_marchandises[0] + ca_services[0],
        ca_marchandises[1] + ca_services[1],
        ca_marchandises[2] + ca_services[2]
    ]
    
    # Achats consommés (charges variables) - Supposés nuls si pas de marchandises vendues
    data["charges_variables"] = data.get("charges_variables", {})
    charges_variables = data["charges_variables"]
    cout_achat_marchandises_pct=charges_variables.get("cout_achat_marchandises_pct", 0.0)
    
    charges_variables = [ca_marchandises[0]* cout_achat_marchandises_pct / 100.0,
                         ca_marchandises[1]* cout_achat_marchandises_pct / 100.0,
                         ca_marchandises[2]* cout_achat_marchandises_pct / 100.0
                         ]
    
    # charges exploitations (charges exploitations) - Supposés nuls si pas de marchandises vendues
    
    charges_exploitations = charges_variables 

    
    # Marge brute = Total CA - Achats consommés
    marge_brute = [
        total_ca[0] - charges_variables[0],
        total_ca[1] - charges_variables[1],
        total_ca[2] - charges_variables[2]
    ]
    
    # Charges externes (charges fixes)
    charges_fixes_data = data.get("charges_fixes", {})
    charges_fixes_annee1 = charges_fixes_data.get("annee1", {})
    charges_fixes_annee2 = charges_fixes_data.get("annee2", {})
    charges_fixes_annee3 = charges_fixes_data.get("annee3", {})
    
    # Liste des charges externes détaillées
    liste_charges = [
        "Assurances véhicule et RC pro", "Téléphone, internet", "Autres abonnements",
        "Carburant", "Frais de déplacement / hébergement", "Eau, électricité, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et matériel",
        "Nettoyage des locaux", "Budget publicité et communication", "Emplacements",
        "Expert comptable, avocats", "Markting"
    ]
    
    # Récupération des montants pour chaque charge
    charges_detaillees = {}
    total_charges_fixes = [0.0, 0.0, 0.0]
    for charge in liste_charges:
        montant_annee1 = charges_fixes_annee1.get(charge, 0.0)
        montant_annee2 = charges_fixes_annee2.get(charge, 0.0)
        montant_annee3 = charges_fixes_annee3.get(charge, 0.0)
        charges_detaillees[charge] = [montant_annee1, montant_annee2, montant_annee3]
        total_charges_fixes[0] += montant_annee1
        total_charges_fixes[1] += montant_annee2
        total_charges_fixes[2] += montant_annee3
    
    # Valeur ajoutée = Marge brute - Charges externes
    valeur_ajoutee = [
        marge_brute[0] - total_charges_fixes[0],
        marge_brute[1] - total_charges_fixes[1],
        marge_brute[2] - total_charges_fixes[2]
    ]
    
    # Impôts et taxes (ajouter d'autres impôts si nécessaire)
    impots_et_taxes = [
        charges_fixes_annee1.get("Taxes, CFE", 0.0),
        charges_fixes_annee2.get("Taxes, CFE", 0.0),
        charges_fixes_annee3.get("Taxes, CFE", 0.0)
    ]
    
    # Salaires employés
    salaires_employes = [
        data["salaires"]["employes"].get("annee1", 0.0),
        data["salaires"]["employes"].get("annee2", 0.0),
        data["salaires"]["employes"].get("annee3", 0.0)
    ]
    
    # Charges sociales employés
    charges_sociales_employes = [
        data["charges_sociales"]["employes"].get("annee1", 0.0),
        data["charges_sociales"]["employes"].get("annee2", 0.0),
        data["charges_sociales"]["employes"].get("annee3", 0.0)
    ]
    
    # Prélèvement dirigeant(s)
    salaires_dirigeants = [
        data["salaires"]["dirigeants"].get("annee1", 0.0),
        data["salaires"]["dirigeants"].get("annee2", 0.0),
        data["salaires"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Charges sociales dirigeant(s)
    charges_sociales_dirigeants = [
        data["charges_sociales"]["dirigeants"].get("annee1", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee2", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Excédent brut d'exploitation = Valeur ajoutée - Impôts et taxes - Salaires - Charges sociales
    ebe = [
        valeur_ajoutee[0] - impots_et_taxes[0] - salaires_employes[0] - charges_sociales_employes[0] - salaires_dirigeants[0] - charges_sociales_dirigeants[0],
        valeur_ajoutee[1] - impots_et_taxes[1] - salaires_employes[1] - charges_sociales_employes[1] - salaires_dirigeants[1] - charges_sociales_dirigeants[1],
        valeur_ajoutee[2] - impots_et_taxes[2] - salaires_employes[2] - charges_sociales_employes[2] - salaires_dirigeants[2] - charges_sociales_dirigeants[2]
    ]
    
    # Frais bancaires, charges financières
    frais_bancaires = [
        charges_fixes_annee1.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee2.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee3.get("Frais bancaires et terminal carte bleue", 0.0)
    ]
    
    # Intérêts des prêts
    interets_prets = data.get("interets_prets", {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    })
    
    # Ajouter les intérêts des prêts aux autres frais financiers
    frais_financiers = [
        interets_prets.get("annee1", 0.0),
        interets_prets.get("annee2", 0.0),
        interets_prets.get("annee3", 0.0)
    ]
    
    # Total des frais bancaires et charges financières
    total_frais_financiers = [
        frais_bancaires[0] + frais_financiers[0],
        frais_bancaires[1] + frais_financiers[1],
        frais_bancaires[2] + frais_financiers[2]
    ]
    
    # Dotations aux amortissements (supposées nulles si non fournies)
    amortissements = [0.0, 0.0, 0.0]
    
    # Résultat avant impôts = EBE - Frais bancaires - Amortissements
    resultat_avant_impots = [
        ebe[0] - total_frais_financiers[0] - amortissements[0],
        ebe[1] - total_frais_financiers[1] - amortissements[1],
        ebe[2] - total_frais_financiers[2] - amortissements[2]
    ]
    
    # Impôt sur les sociétés (selon la formule progressive)
    impot_societes = [
        calculer_impot_societes(resultat_avant_impots[0]),
        calculer_impot_societes(resultat_avant_impots[1]),
        calculer_impot_societes(resultat_avant_impots[2])
    ]
    
    # Résultat net comptable (résultat de l'exercice)
    resultat_net = [
        resultat_avant_impots[0] - impot_societes[0],
        resultat_avant_impots[1] - impot_societes[1],
        resultat_avant_impots[2] - impot_societes[2]
    ]
    
    # Préparation des données pour le tableau
    tableau = {
        "": [
            "Produits d'exploitation",
            "Chiffre d'affaires HT vente de marchandises",
            "Chiffre d'affaires HT services",
            "",
            "Charges d'exploitation(charge variable)",
            "Achats consommés",
            "",
            "Marge brute",
            "Charges externes",
            ""
        ],
        "Année 1": [
            f"{total_ca[0]:,.2f} $",
            f"{ca_marchandises[0]:,.2f} $",
            f"{ca_services[0]:,.2f} $",
            "",
            f"{charges_exploitations[0]:,.2f} $",
            f"{charges_variables[0]:,.2f} $",
            "",
            f"{marge_brute[0]:,.2f} $",
            "",
            ""
        ],
        "Année 2": [
            f"{total_ca[1]:,.2f} $",
            f"{ca_marchandises[1]:,.2f} $",
            f"{ca_services[1]:,.2f} $",
            "",
            f"{charges_exploitations[1]:,.2f} $",
            f"{charges_variables[1]:,.2f} $",
            "",
            f"{marge_brute[1]:,.2f} $",
            "",
            ""
        ],
        "Année 3": [
            f"{total_ca[2]:,.2f} $",
            f"{ca_marchandises[2]:,.2f} $",
            f"{ca_services[2]:,.2f} $",
            "",
            f"{charges_exploitations[2]:,.2f} $",
            f"{charges_variables[2]:,.2f} $",
            "",
            f"{marge_brute[2]:,.2f} $",
            "",
            ""
        ]
    }
    
    # Ajouter les charges détaillées au tableau
    for charge in liste_charges:
        tableau[""].append(charge)
        tableau["Année 1"].append(f"{charges_detaillees[charge][0]:,.2f} $")
        tableau["Année 2"].append(f"{charges_detaillees[charge][1]:,.2f} $")
        tableau["Année 3"].append(f"{charges_detaillees[charge][2]:,.2f} $")
    
    # Ajouter le total des charges externes
    tableau[""].append("Total Charges externes")
    tableau["Année 1"].append(f"{total_charges_fixes[0]:,.2f} $")
    tableau["Année 2"].append(f"{total_charges_fixes[1]:,.2f} $")
    tableau["Année 3"].append(f"{total_charges_fixes[2]:,.2f} $")
    
    # Continuer à remplir le tableau
    additional_rows = {
        "Valeur ajoutée": valeur_ajoutee,
        "Impôts et taxes": impots_et_taxes,
        "Salaires employés": salaires_employes,
        "Charges sociales employés": charges_sociales_employes,
        "Prélèvement dirigeant(s)": salaires_dirigeants,
        "Charges sociales dirigeant(s)": charges_sociales_dirigeants,
        "Excédent brut d'exploitation": ebe,
        "Frais bancaires, charges financières": total_frais_financiers,
        "Dotations aux amortissements": amortissements,
        "Résultat avant impôts": resultat_avant_impots,
        "Impôt sur les sociétés": impot_societes,
        "Résultat net comptable (résultat de l'exercice)": resultat_net
    }
    
    for key, values in additional_rows.items():
        tableau[""].append(key)
        tableau["Année 1"].append(f"{values[0]:,.2f} $")
        tableau["Année 2"].append(f"{values[1]:,.2f} $")
        tableau["Année 3"].append(f"{values[2]:,.2f} $")
    
    # Créer le DataFrame
    df_resultats = pd.DataFrame(tableau)
    
    # Afficher le tableau
    st.table(df_resultats)
    
    # Ajouter les variables calculées au dictionnaire 'data'
    data["compte_de_resultat"] = {
        "total_ca": total_ca,
        "ca_marchandises": ca_marchandises,
        "ca_services": ca_services,
        "charges_exploitations":charges_exploitations,
        "charges_variables": charges_variables,
        "marge_brute": marge_brute,
        "charges_fixes": total_charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "salaires_employes": salaires_employes,
        "charges_sociales_employes": charges_sociales_employes,
        "salaires_dirigeants": salaires_dirigeants,
        "charges_sociales_dirigeants": charges_sociales_dirigeants,
        "ebe": ebe,
        "frais_bancaires": frais_bancaires,
        "frais_financiers": frais_financiers,
        "total_frais_financiers": total_frais_financiers,
        "amortissements": amortissements,
        "resultat_avant_impots": resultat_avant_impots,
        "impot_societes": impot_societes,
        "resultat_net": resultat_net
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Compte de résultats prévisionnel
    export_table_compte = []
    for index, row in df_resultats.iterrows():
        export_table_compte.append({
            "Description": row[""],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_compte_resultats_previsionnel'] = {
        "table_data": export_table_compte
    }
    
    # Section Export
    st.header("Exporter les données")
    # Bouton pour télécharger le document complet
    st.button("Télécharger le Document Word Complet", on_click=telecharger_document_complet)
    


import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_soldes_intermediaires_de_gestion():
    st.title("Soldes intermédiaires de gestion")
    
    # Récupérer les données de la session
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les données nécessaires pour les calculs
    compte_resultat = data.get("compte_de_resultat", {})
    
    # Fonction pour assurer que les listes contiennent trois éléments et convertir en float
    def get_three_years_data(key):
        values = compte_resultat.get(key, [])
        processed_values = []
        for v in values:
            try:
                processed_values.append(float(v))
            except (ValueError, TypeError):
                processed_values.append(0.0)
        # Compléter avec 0.0 si moins de 3 éléments
        while len(processed_values) < 3:
            processed_values.append(0.0)
        return processed_values[:3]
    
    # Récupération des données avec validation
    total_ca = get_three_years_data("total_ca")
    ca_marchandises = get_three_years_data("ca_marchandises")
    ca_services = get_three_years_data("ca_services")
    achats_consommes = get_three_years_data("charges_variables")  # Actuellement défini à [0.0, 0.0, 0.0]
    charges_fixes = get_three_years_data("charges_fixes")
    impot_societes = get_three_years_data("impot_societes")
    impots_et_taxes = get_three_years_data("impots_et_taxes")
    salaires_employes = get_three_years_data("salaires_employes")
    charges_sociales_employes = get_three_years_data("charges_sociales_employes")
    salaires_dirigeants = get_three_years_data("salaires_dirigeants")
    charges_sociales_dirigeants = get_three_years_data("charges_sociales_dirigeants")
    amortissements = get_three_years_data("amortissements")
    total_frais_financiers = get_three_years_data("total_frais_financiers")
    
    # Calcul des différents soldes intermédiaires
    ventes_production_reelle = [ca_marchandises[i] + ca_services[i] for i in range(3)]
    marge_globale = [ventes_production_reelle[i] - achats_consommes[i] for i in range(3)]
    valeur_ajoutee = [marge_globale[i] - charges_fixes[i] for i in range(3)]
    charges_personnel = [
        salaires_employes[i] + charges_sociales_employes[i] + salaires_dirigeants[i] + charges_sociales_dirigeants[i]
        for i in range(3)
    ]
    ebe = [valeur_ajoutee[i] - impots_et_taxes[i] - charges_personnel[i] for i in range(3)]
    resultat_exploitation = [ebe[i] - amortissements[i] for i in range(3)]
    resultat_financier = [-total_frais_financiers[i] for i in range(3)]
    resultat_courant = [resultat_exploitation[i] + resultat_financier[i] for i in range(3)]
    resultat_exercice = [resultat_courant[i] - impot_societes[i] for i in range(3)]
    capacite_autofinancement = [resultat_exercice[i] + amortissements[i] for i in range(3)]
    
    # Fonction de calcul des pourcentages avec gestion de la division par zéro
    def calculate_percentage(value, ca):
        return (value / ca * 100) if ca != 0 else 0.0
    
    # Préparation des données pour le tableau
    soldes = [
        "Chiffre d'affaires",
        "Ventes + production réelle",
        "Achats consommés",
        "Marge globale",
        "Charges externes",
        "Valeur ajoutée",
        "Impôts et taxes",
        "Charges de personnel",
        "Excédent brut d'exploitation (EBE)",
        "Dotations aux amortissements",
        "Résultat d'exploitation",
        "Charges financières",
        "Résultat financier",
        "Résultat courant",
        "Résultat de l'exercice",
        "Capacité d'autofinancement"
    ]
    
    # Initialiser le data_table avec les soldes
    data_table = {"Soldes intermédiaires de gestion": soldes}
    
    # Ajouter les données pour chaque année et leurs pourcentages
    for year in range(3):
        data_table[f"Année {year+1}"] = [
            total_ca[year],
            ventes_production_reelle[year],
            achats_consommes[year],
            marge_globale[year],
            charges_fixes[year],
            valeur_ajoutee[year],
            impots_et_taxes[year],
            charges_personnel[year],
            ebe[year],
            amortissements[year],
            resultat_exploitation[year],
            total_frais_financiers[year],
            resultat_financier[year],
            resultat_courant[year],
            resultat_exercice[year],
            capacite_autofinancement[year]
        ]
        
        data_table[f"% Année {year+1}"] = [
            100.0,  # Chiffre d'affaires
            100.0,  # Ventes + production réelle
            calculate_percentage(achats_consommes[year], total_ca[year]),
            calculate_percentage(marge_globale[year], total_ca[year]),
            calculate_percentage(charges_fixes[year], total_ca[year]),
            calculate_percentage(valeur_ajoutee[year], total_ca[year]),
            calculate_percentage(impots_et_taxes[year], total_ca[year]),
            calculate_percentage(charges_personnel[year], total_ca[year]),
            calculate_percentage(ebe[year], total_ca[year]),
            calculate_percentage(amortissements[year], total_ca[year]),
            calculate_percentage(resultat_exploitation[year], total_ca[year]),
            calculate_percentage(total_frais_financiers[year], total_ca[year]),
            calculate_percentage(resultat_financier[year], total_ca[year]),
            calculate_percentage(resultat_courant[year], total_ca[year]),
            calculate_percentage(resultat_exercice[year], total_ca[year]),
            calculate_percentage(capacite_autofinancement[year], total_ca[year])
        ]
    
    # Créer le DataFrame avec les données
    df = pd.DataFrame(data_table)
    
    # Définir l'ordre des colonnes alternées entre "Année x" et "%"
    columns_order = ["Soldes intermédiaires de gestion"]
    for year in range(3):
        columns_order.append(f"Année {year+1}")
        columns_order.append(f"% Année {year+1}")
    df = df[columns_order]
    
    # Afficher le tableau avec une mise en forme améliorée
    st.dataframe(
        df.style.format({
            "Année 1": "{:,.2f} $",
            "Année 2": "{:,.2f} $",
            "Année 3": "{:,.2f} $",
            "% Année 1": "{:.2f}%",
            "% Année 2": "{:.2f}%",
            "% Année 3": "{:.2f}%"
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les résultats dans les données pour exportation
    data["soldes_intermediaires_de_gestion"] = {
        "ca": total_ca,
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "marge_globale": marge_globale,
        "charges_externes": charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "charges_personnel": charges_personnel,
        "ebe": ebe,
        "dotations_aux_amortissements": amortissements,
        "resultat_exploitation": resultat_exploitation,
        "charges_financieres": total_frais_financiers,
        "resultat_financier": resultat_financier,
        "resultat_courant": resultat_courant,
        "resultat_exercice": resultat_exercice,
        "capacite_autofinancement": capacite_autofinancement
    }
    
    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Soldes Intermédiaires de Gestion avec % colonnes
    export_table_soldes = []
    for idx, solde in enumerate(soldes):
        export_table_soldes.append({
            "Description": solde,
            "Année 1": data_table["Année 1"][idx],
            "% Année 1": data_table["% Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "% Année 2": data_table["% Année 2"][idx],
            "Année 3": data_table["Année 3"][idx],
            "% Année 3": data_table["% Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_soldes_intermediaires_de_gestion'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_soldes
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Soldes Intermédiaires en Markdown"):
        export_data = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Soldes intermédiaires de gestion\n\n**Projet :** {export_data['projet']}\n\n**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += "---\n\n"
            
            # Créer un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += "\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="soldes_intermediaires_gestion.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Soldes Intermédiaires en Word"):
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_soldes or "table_data" not in export_data_soldes:
            st.error("Aucune donnée disponible pour l'exportation des Soldes intermédiaires de gestion.")
            return
        
        # Vérifiez que toutes les autres sections sont également exportées
        if not all([
            export_data_investissements.get("table_data"),
            export_data_salaires.get("table_data"),
            export_data_amortissements.get("amortissements"),
            export_data_compte.get("table_data")
        ]):
            st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
            return
        
        # Créer un document Word
        doc = Document()
        
        ### 1. Ajouter la section Investissements et Financements ###
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements
        table_inv = doc.add_table(rows=1, cols=4)
        table_inv.style = 'Light List Accent 1'
        table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_investissements['table_data']:
            row_cells = table_inv.add_row().cells
            row_cells[0].text = row.get("Investissements", "")
            row_cells[1].text = row.get("Taux (%)", "")
            row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
            row_cells[3].text = row.get("Montant ($)", "")
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 2. Ajouter la section Salaires et Charges Sociales ###
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales
        table_sal = doc.add_table(rows=1, cols=4)
        table_sal.style = 'Light List Accent 1'
        table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_salaires['table_data']:
            row_cells = table_sal.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 3. Ajouter la section Détail des Amortissements ###
        doc.add_heading('Détail des Amortissements', level=1)
        
        # Créer le tableau Détail des Amortissements
        table_amort = doc.add_table(rows=1, cols=4)
        table_amort.style = 'Light List Accent 1'
        table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_amort.rows[0].cells
        headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_amortissements['amortissements']:
            row_cells = table_amort.add_row().cells
            row_cells[0].text = row.get("Amortissement", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
        
        ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
        doc.add_heading('Compte de Résultats Prévisionnel', level=1)
        
        # Créer le tableau Compte de Résultats Prévisionnel
        table_compte = doc.add_table(rows=1, cols=4)
        table_compte.style = 'Light List Accent 1'
        table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_compte = table_compte.rows[0].cells
        headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_compte):
            hdr_cells_compte[i].text = header
            for paragraph in hdr_cells_compte[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_compte['table_data']:
            row_cells = table_compte.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
        
        ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
        doc.add_heading('Soldes intermédiaires de gestion', level=1)
        
        # Créer le tableau Soldes intermédiaires de gestion avec 7 colonnes
        table_soldes = doc.add_table(rows=1, cols=7)
        table_soldes.style = 'Light List Accent 1'
        table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_soldes = table_soldes.rows[0].cells
        headers_soldes = ["Description", "Année 1", "% Année 1", "Année 2", "% Année 2", "Année 3", "% Année 3"]
        for i, header in enumerate(headers_soldes):
            hdr_cells_soldes[i].text = header
            for paragraph in hdr_cells_soldes[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_soldes['table_data']:
            row_cells = table_soldes.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = f"{row.get('Année 1', 0.0):,.2f} $"
            row_cells[2].text = f"{row.get('% Année 1', 0.0):.2f}%"
            row_cells[3].text = f"{row.get('Année 2', 0.0):,.2f} $"
            row_cells[4].text = f"{row.get('% Année 2', 0.0):.2f}%"
            row_cells[5].text = f"{row.get('Année 3', 0.0):,.2f} $"
            row_cells[6].text = f"{row.get('% Année 3', 0.0):.2f}%"
            
            # Alignement des cellules de pourcentage
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Bouton de téléchargement
        st.download_button(
            label="Télécharger le Document Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Message de confirmation
        st.success("Le document Word complet a été généré avec succès !")



    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    
    
    
    

def telecharger_document_complets():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.0f} $"
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.0f} $"
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.0f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")


def calculate_remboursements_emprunts(financements, years=3):
    """
    Votre fonction de calcul existante...
    """
    remboursements = [0.0 for _ in range(years)]  # Initialiser les remboursements pour chaque année

    for loan_name, loan_info in financements.items():
        # Vérifier que loan_info est un dictionnaire et commence par "Prêt "
        if isinstance(loan_info, dict) and loan_name.startswith("Prêt "):
            required_keys = {"montant", "duree", "taux"}
            if not required_keys.issubset(loan_info.keys()):
                st.warning(f"Le prêt '{loan_name}' est incomplet et sera ignoré.")
                continue  # Ignorer les financements incomplets

            montant = loan_info.get("montant", 0.0)
            duree_mois = loan_info.get("duree", 60)  # Par défaut 60 mois
            taux_annuel = loan_info.get("taux", 5.0)  # Par défaut 5%
            principal_mensuel =  montant / duree_mois if duree_mois > 0 else 0.0

            # Calcul des remboursements par année basés sur principal_mensuel
            # Principal Year 1
            if duree_mois > 12:
                principal_year1 = principal_mensuel * 12
            else:
                principal_year1 = principal_mensuel * duree_mois

            # Principal Year 2
            if duree_mois - 12 < 0:
                principal_year2 = 0.0
            elif duree_mois > 24:
                principal_year2 = principal_mensuel * 12
            else:
                principal_year2 = principal_mensuel * (duree_mois - 12)

            # Principal Year 3
            if duree_mois - 24 < 0:
                principal_year3 = 0.0
            elif duree_mois > 36:
                principal_year3 = principal_mensuel * 12
            else:
                principal_year3 = principal_mensuel * (duree_mois - 24)

            # Ajouter les remboursements principaux au total par année
            remboursements[0] += round(principal_year1, 2)
            if years >= 2:
                remboursements[1] += round(principal_year2, 2)
            if years >= 3:
                remboursements[2] += round(principal_year3, 2)
        else:
            # Ignorer les financements qui ne sont pas des prêts (e.g., Apports, Subventions)
            continue

    return remboursements

def page_capacite_autofinancement():
    """
    Affiche le tableau de Capacité d'Autofinancement en utilisant les données de la session.
    """
    st.title("Capacité d'autofinancement")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    # Récupérer les données de la session
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les soldes intermédiaires de gestion
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # Fonction pour convertir les valeurs en float, remplacer les erreurs par 0.0
    def safe_float_conversion(values):
        return [float(x) if isinstance(x, (int, float)) else 0.0 for x in values]
    
    # Récupérer et convertir les données nécessaires
    resultat_exercice = safe_float_conversion(soldes_intermediaires.get("resultat_exercice", [0.0, 0.0, 0.0]))
    dotations_aux_amortissements = safe_float_conversion(soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0]))
    capacite_autofinancement = safe_float_conversion(soldes_intermediaires.get("capacite_autofinancement", [0.0, 0.0, 0.0]))
    
    # Récupérer les financements
    financements = data.get("financements", {})
    
    # Filtrer uniquement les prêts (dictionnaires) nommés avec "Prêt " pour éviter les subventions
    pret_financements = {
        k: v for k, v in financements.items()
        if isinstance(v, dict) and k.startswith("Prêt ")
    }
    
    # Calculer les remboursements des emprunts
    remboursements_emprunts = calculate_remboursements_emprunts(pret_financements, years=3)
    
    # Autofinancement net = Capacité d'autofinancement - Remboursements des emprunts
    autofinancement_net = [
        capacite_autofinancement[i] - remboursements_emprunts[i]
        for i in range(3)
    ]
    
    # Préparer les valeurs monétaires
    values = {
        "Année 1": [
            resultat_exercice[0],
            dotations_aux_amortissements[0],
            capacite_autofinancement[0],
            remboursements_emprunts[0],
            autofinancement_net[0]
        ],
        "Année 2": [
            resultat_exercice[1],
            dotations_aux_amortissements[1],
            capacite_autofinancement[1],
            remboursements_emprunts[1],
            autofinancement_net[1]
        ],
        "Année 3": [
            resultat_exercice[2],
            dotations_aux_amortissements[2],
            capacite_autofinancement[2],
            remboursements_emprunts[2],
            autofinancement_net[2]
        ]
    }
    
    # Préparer le tableau final avec les labels
    capacite_fonc = [
        "Résultat de l'exercice",
        "+ Dotation aux amortissements",
        "Capacité d'autofinancement",
        "- Remboursements des emprunts",
        "Autofinancement net"
    ]
    
    data_table = {
        "Capacité d'autofinancement": capacite_fonc,
        "Année 1": values["Année 1"],
        "Année 2": values["Année 2"],
        "Année 3": values["Année 3"]
    }
    
    # Créer le DataFrame avec les données
    df = pd.DataFrame(data_table)
    
    # Définir l'ordre des colonnes
    columns_order = ["Capacité d'autofinancement",
                     "Année 1",
                     "Année 2",
                     "Année 3"]
    df = df[columns_order]
    
    # Définir la fonction de formatage
    def format_value(x):
        if x == 0.0:
            return "-"
        else:
            return f"{x:,.2f} $"
    
    # Afficher le tableau avec une mise en forme améliorée
    st.dataframe(
        df.style.format({
            "Année 1": format_value,
            "Année 2": format_value,
            "Année 3": format_value,
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les résultats dans les données
    data["capacite_autofinancement"] = {
        "resultat_exercice": resultat_exercice,
        "dotations_aux_amortissements": dotations_aux_amortissements,
        "capacite_autofinancement": capacite_autofinancement,
        "remboursements_emprunts": remboursements_emprunts,
        "autofinancement_net": autofinancement_net
    }
    
    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Capacité d'Autofinancement
    export_table_capacite = []
    for idx, label in enumerate(capacite_fonc):
        export_table_capacite.append({
            "Description": label,
            "Année 1": values["Année 1"][idx],
            "Année 2": values["Année 2"][idx],
            "Année 3": values["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_capacite_autofinancement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_capacite
    }
    
    # Section Export
    st.header("Exporter les données")
    
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_investissements_et_financements"):
        telecharger_document_complet()

    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_seuil_rentabilite_economique():
    st.title("Seuil de rentabilité économique")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]

    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les données nécessaires avec les clés exactes
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # Récupérer 'ventes_production_reelle' et 'achats_consommes' du 'soldes_intermediaires'
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # Récupérer les charges nécessaires du 'soldes_intermediaires'
    charges_externes = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])
    impots_et_taxes = soldes_intermediaires.get("impots_et_taxes", [0.0, 0.0, 0.0])
    charges_personnel = soldes_intermediaires.get("charges_personnel", [0.0, 0.0, 0.0])
    dotations_aux_amortissements = soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0])
    charges_financieres = soldes_intermediaires.get("charges_financieres", [0.0, 0.0, 0.0])
    
    # Récupérer 'resultat_avant_impots' du 'compte_resultat'
    resultat_avant_impots = compte_resultat.get("resultat_avant_impots", [0.0, 0.0, 0.0])
    
    # Vérifier si les listes ont 3 éléments
    if not (len(ventes_production_reelle) == len(achats_consommes) == len(charges_externes) == len(impots_et_taxes) == len(charges_personnel) == len(dotations_aux_amortissements) == len(charges_financieres) == len(resultat_avant_impots) == 3):
        st.error("Les listes de données ne contiennent pas exactement 3 éléments. Veuillez vérifier les données.")
        return
    
    # Calcul des Coûts fixes pour chaque année
    couts_fixes = []
    for i in range(3):
        cout_fix = (
            charges_externes[i] +
            impots_et_taxes[i] +
            charges_personnel[i] +
            dotations_aux_amortissements[i] +
            charges_financieres[i]
        )
        couts_fixes.append(cout_fix)
   
    # Total des coûts variables = Achats consommés
    total_couts_variables = achats_consommes.copy()
    
    # Marge sur coûts variables
    marge_sur_couts_variables = []
    for i in range(3):
        marge = ventes_production_reelle[i] - total_couts_variables[i]
        marge_sur_couts_variables.append(marge)
    
    # Taux de marge sur coûts variables
    taux_marge_sur_couts_variables = []
    for i in range(3):
        if ventes_production_reelle[i] != 0:
            taux_marge = marge_sur_couts_variables[i] / ventes_production_reelle[i]
        else:
            taux_marge = 0.0
        taux_marge_sur_couts_variables.append(taux_marge)
    
    # Total des charges
    total_charges = []
    for i in range(3):
        total_charge = couts_fixes[i] + total_couts_variables[i]
        total_charges.append(total_charge)
   
    # Seuil de rentabilité (CA)
    seuil_rentabilite_ca = []
    for i in range(3):
        if taux_marge_sur_couts_variables[i] != 0:
            seuil_ca = couts_fixes[i] / taux_marge_sur_couts_variables[i]
        else:
            seuil_ca = 0.0
        seuil_rentabilite_ca.append(seuil_ca)
    
    # Excédent / insuffisance
    excedent_insuffisance = []
    for i in range(3):
        excedent = ventes_production_reelle[i] - seuil_rentabilite_ca[i]
        excedent_insuffisance.append(excedent)
    
    # Point mort
    point_mort_ca_par_jour_ouvre = []
    for i in range(3):
        point_mort = seuil_rentabilite_ca[i] / 250
        point_mort_ca_par_jour_ouvre.append(point_mort)
    
    # Préparation des données pour le tableau
    data_table = {
        "Seuil de rentabilite_economique": [
            "Ventes + Production réelle",
            "Achats consommés",
            "Total des coûts variables",
            "Marge sur coûts variables",
            "Taux de marge sur coûts variables",
            "Coûts fixes",
            "Total des charges",
            "Résultat courant avant impôts",
            "Seuil de rentabilite (chiffre d'affaires)",
            "Excédent / insuffisance",
            "Point mort en chiffre d'affaires par jour ouvré"
        ],
        "Année 1": [
            ventes_production_reelle[0],
            achats_consommes[0],
            total_couts_variables[0],
            marge_sur_couts_variables[0],
            taux_marge_sur_couts_variables[0],
            couts_fixes[0],
            total_charges[0],
            resultat_avant_impots[0],
            seuil_rentabilite_ca[0],
            excedent_insuffisance[0],
            point_mort_ca_par_jour_ouvre[0]
        ],
        "Année 2": [
            ventes_production_reelle[1],
            achats_consommes[1],
            total_couts_variables[1],
            marge_sur_couts_variables[1],
            taux_marge_sur_couts_variables[1],
            couts_fixes[1],
            total_charges[1],
            resultat_avant_impots[1],
            seuil_rentabilite_ca[1],
            excedent_insuffisance[1],
            point_mort_ca_par_jour_ouvre[1]
        ],
        "Année 3": [
            ventes_production_reelle[2],
            achats_consommes[2],
            total_couts_variables[2],
            marge_sur_couts_variables[2],
            taux_marge_sur_couts_variables[2],
            couts_fixes[2],
            total_charges[2],
            resultat_avant_impots[2],
            seuil_rentabilite_ca[2],
            excedent_insuffisance[2],
            point_mort_ca_par_jour_ouvre[2]
        ]
    }
    
    # Créer le DataFrame
    df = pd.DataFrame(data_table)
    
    # Définir "Seuil de rentabilite_economique" comme index
    df.set_index("Seuil de rentabilite_economique", inplace=True)
    
    # Étape 4: Définir une fonction de formatage
    # Étape 4: Définir une fonction de formatage
    def format_row(row):
        if row.name == "Taux de marge sur coûts variables":
            # Formater en pourcentage avec deux décimales
            return row.apply(lambda x: "{:.2f} %".format(x))
        else:
            # Formater en dollars avec séparateurs de milliers et sans décimales
            return row.apply(lambda x: "{:,.0f} $".format(x) if isinstance(x, (int, float)) else x)

    # Étape 5: Appliquer le formatage
    df_formatted = df.apply(format_row, axis=1)

    # Étape 6: Afficher le tableau formaté avec Streamlit
    st.table(df_formatted)
    # Stocker les résultats dans les données
    data["seuil_rentabilite_economique"] = {
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "total_couts_variables": total_couts_variables,
        "marge_sur_couts_variables": marge_sur_couts_variables,
        "taux_marge_sur_couts_variables": taux_marge_sur_couts_variables,
        "couts_fixes": couts_fixes,
        "total_charges": total_charges,
        "resultat_courant_avant_impots": resultat_avant_impots,
        "seuil_rentabilite_ca": seuil_rentabilite_ca,
        "excedent_insuffisance": excedent_insuffisance,
        "point_mort_ca_par_jour_ouvre": point_mort_ca_par_jour_ouvre
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Seuil de rentabilité économique
    export_table_seuil = []
    for idx, label in enumerate(data_table["Seuil de rentabilite_economique"]):
        export_table_seuil.append({
            "Description": label,
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_seuil_rentabilite_economique'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_seuil
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_seuil_rentabilite_economique"):
        telecharger_document_complet()


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.1f} "
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.1f} "
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.1f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Besoin en Fonds de Roulement ###
    doc.add_heading('Besoin en fonds de roulement', level=1)
    
    # Créer le tableau Besoin en Fonds de Roulement dans Word
    table_bfr = doc.add_table(rows=1, cols=5)
    table_bfr.style = 'Light List Accent 1'
    table_bfr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_bfr = table_bfr.rows[0].cells
    headers_bfr = ["Analyse clients / fournisseurs", "Délai jours", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_bfr):
        hdr_cells_bfr[i].text = header
        for paragraph in hdr_cells_bfr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_bfr[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_bfr['table_data']:
        row_cells_bfr = table_bfr.add_row().cells
        row_cells_bfr[0].text = row.get("Analyse clients / fournisseurs", "")
        row_cells_bfr[1].text = row.get("Délai jours", "")
        row_cells_bfr[2].text = row.get("Année 1", "")
        row_cells_bfr[3].text = row.get("Année 2", "")
        row_cells_bfr[4].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")

def page_besoin_fonds_roulement():
    st.title("Besoin en fonds de roulement")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")

    # Récupérer les délais clients et fournisseurs depuis "besoin_fonds_roulement"
    besoin_fonds = data.get("fonds_roulement", {})
    delai_clients = besoin_fonds.get("duree_credits_clients", 0)  # Durée moyenne des crédits accordés aux clients en jours
    delai_fournisseurs = besoin_fonds.get("duree_dettes_fournisseurs", 0)  # Durée moyenne des crédits accordés aux fournisseurs en jours

    st.write("---")
    
    # Récupérer "Ventes + Production réelle" et "Achats consommés" depuis "soldes_intermediaires_de_gestion"
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # Vérifier si les délais sont non nuls
    if delai_clients == 0 or delai_fournisseurs == 0:
        st.error("Les délais de paiement des clients et des fournisseurs ne sont pas renseignés. Veuillez les saisir dans la section 'Votre besoin en fonds de roulement'.")
        return
    
    # Vérifier si les listes ont 3 éléments
    if not (len(ventes_production_reelle) == len(achats_consommes) == 3):
        st.error("Les listes de 'Ventes + Production réelle' ou 'Achats consommés' ne contiennent pas exactement 3 éléments. Veuillez vérifier les données.")
        return
    
    # Calculer le Volume crédit client HT = Ventes + Production réelle / (delai_jours * 365)
    volume_credit_client_ht = []
    for i in range(3):
        vcc_ht = (ventes_production_reelle[i] * delai_clients) / 365
        volume_credit_client_ht.append(vcc_ht)
    
    # Calculer le Volume dettes fournisseurs HT = Achats consommés / (delai_jours * 365)
    volume_dettes_fournisseurs_ht = []
    for i in range(3):
        vdf_ht = (achats_consommes[i] * delai_fournisseurs) / 365
        volume_dettes_fournisseurs_ht.append(vdf_ht)
    
    # Calculer le Besoin en fonds de roulement (BFR) = Volume crédit client HT - Volume dettes fournisseurs HT
    bfr = [volume_credit_client_ht[i] - volume_dettes_fournisseurs_ht[i] for i in range(3)]
    
    # Afficher les résultats intermédiaires
    st.write("### Résultats des Calculs")
    st.write(f"**Volume crédit client HT Année 1** : {volume_credit_client_ht[0]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Année 1** : {volume_dettes_fournisseurs_ht[0]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Année 1** : {bfr[0]:.2f} $")
    st.write(f"**Volume crédit client HT Année 2** : {volume_credit_client_ht[1]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Année 2** : {volume_dettes_fournisseurs_ht[1]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Année 2** : {bfr[1]:.2f} $")
    st.write(f"**Volume crédit client HT Année 3** : {volume_credit_client_ht[2]:.2f} $")
    st.write(f"**Volume dettes fournisseurs HT Année 3** : {volume_dettes_fournisseurs_ht[2]:.2f} $")
    st.write(f"**Besoin en fonds de roulement Année 3** : {bfr[2]:.2f} $")
    
    # Préparer les données pour le tableau
    data_table = {
        "Analyse clients / fournisseurs": [
            "Besoins",
            "Volume crédit client HT",
            "Ressources",
            "Volume dettes fournisseurs HT",
            "Besoin en fonds de roulement"
        ],
        "Délai jours": [
            "",
            f"{delai_clients}",
            "",
            f"{delai_fournisseurs}",
            ""
        ],
        "Année 1": [
            "",
            f"{volume_credit_client_ht[0]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[0]:.2f} $",
            f"{bfr[0]:.2f} $"
        ],
        "Année 2": [
            "",
            f"{volume_credit_client_ht[1]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[1]:.2f} $",
            f"{bfr[1]:.2f} $"
        ],
        "Année 3": [
            "",
            f"{volume_credit_client_ht[2]:.2f} $",
            "",
            f"{volume_dettes_fournisseurs_ht[2]:.2f} $",
            f"{bfr[2]:.2f} $"
        ]
    }
    
    df = pd.DataFrame(data_table)
    
    # Afficher le tableau
    st.write("### Tableau du Besoin en fonds de roulement")
    st.table(df)
    
    # Stocker les résultats dans les données
    data["besoin_fonds_roulement"] = {
        "delai_clients": delai_clients,
        "delai_fournisseurs": delai_fournisseurs,
        "volume_credit_client_ht": volume_credit_client_ht,
        "volume_dettes_fournisseurs_ht": volume_dettes_fournisseurs_ht,
        "bfr": bfr
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Besoin en fonds de roulement
    export_table_bfr = []
    for idx, label in enumerate(data_table["Analyse clients / fournisseurs"]):
        export_table_bfr.append({
            "Analyse clients / fournisseurs": label,
            "Délai jours": data_table["Délai jours"][idx],
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_besoin_fonds_roulement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_bfr
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_besoin_fonds_roulement"):
        telecharger_document_complet()


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.2f} "
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.2f} "
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.2f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Plan de Financement à Trois Ans ###
    doc.add_heading('Plan de financement à trois ans', level=1)
    
    # Créer le tableau Plan de Financement à Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=5)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Plan de financement à trois ans", "")
        row_cells_plan[1].text = row.get("Année 1", "")
        row_cells_plan[2].text = row.get("Année 2", "")
        row_cells_plan[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")




def page_plan_financement_trois_ans(): 
    st.title("Plan de financement à trois ans")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les besoins démarrage
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Calcul des Immobilisations incorporelles et corporelles
    # Définissez quels éléments de "besoins_demarrage" correspondent à chaque catégorie
    immobilisations_inc = sum([
        besoins_demarrage.get("Frais d’établissement", 0),
        besoins_demarrage.get("Frais d’ouverture de compteurs", 0),
        besoins_demarrage.get("Logiciels, formations", 0),
        besoins_demarrage.get("Dépôt de marque", 0),
        besoins_demarrage.get("Droits d’entrée", 0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0),
        besoins_demarrage.get("Droit au bail", 0),
        besoins_demarrage.get("Caution ou dépôt de garantie", 0),
        besoins_demarrage.get("Frais de dossier", 0),
        besoins_demarrage.get("Frais de notaire", 0),
    ])
    
    immobilisations_corp = sum([
        besoins_demarrage.get("Enseigne et éléments de communication", 0),
        besoins_demarrage.get("Véhicule", 0),
        besoins_demarrage.get("Matériel professionnel", 0),
        besoins_demarrage.get("Matériel autre", 0),
        besoins_demarrage.get("Matériel de bureau", 0)
    ])
    
    immobilisations = [
        immobilisations_inc + immobilisations_corp,  # Année 1
        0.0,  # Année 2
        0.0   # Année 3
    ]
    
    # Acquisition des stocks
    acquisition_stocks = [
        besoins_demarrage.get("Stock de matières et produits", 0),
        0.0,  # Année 2
        0.0   # Année 3
    ]
    
    # Variation du Besoin en fonds de roulement (BFR)
    besoin_fonds = data.get("besoin_fonds_roulement", {})
    bfr = besoin_fonds.get("bfr", [0.0, 0.0, 0.0])
    
    # Variation BFR = BFR année n - BFR année n-1
    variation_bfr = [
        bfr[0],                    # Variation en année 1 (BFR année 1 - BFR année 0)
        bfr[1] - bfr[0],           # Variation en année 2
        bfr[2] - bfr[1]            # Variation en année 3
    ]
    
    # Remboursement d'emprunts
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    
    # Total des besoins
    total_besoins = [
        immobilisations[0] + acquisition_stocks[0] + variation_bfr[0] + remboursements_emprunts[0],
        immobilisations[1] + acquisition_stocks[1] + variation_bfr[1] + remboursements_emprunts[1],
        immobilisations[2] + acquisition_stocks[2] + variation_bfr[2] + remboursements_emprunts[2]
    ]
    
    # Apport personnel
    financements = data.get("financements", {})
    apport_personnel = financements.get("Apport personnel ou familial", 0.0)
    apports_nature = financements.get("Apports en nature (en valeur)", 0.0)
    apport_total = apport_personnel + apports_nature
    apport_personnel_list = [apport_total, 0.0, 0.0] 
    
    # Emprunts
    pret_1 = financements.get("Prêt 1", {}).get("montant", 0.0)
    pret_2 = financements.get("Prêt 2", {}).get("montant", 0.0)
    pret_3 = financements.get("Prêt 3", {}).get("montant", 0.0)
    total_emprunts = pret_1 + pret_2 + pret_3
    emprunts = [total_emprunts, 0.0, 0.0]  # Supposons que les emprunts sont en année 1
    
    # Subventions
    subvention_1 = financements.get("Subvention 1", {}).get("montant", 0.0)
    subvention_2 = financements.get("Subvention 2", {}).get("montant", 0.0)
    subventions = subvention_1 + subvention_2
    subventions_list = [subventions, 0.0, 0.0]  # Supposons que les subventions sont en année 1
    
    # Autres financements
    autres_financements = financements.get("Autre financement", 0.0)
    autres_financements_list = [autres_financements, 0.0, 0.0]  # Supposons que c'est en année 1
    
    # Capacité d'auto-financement
    capacite_autofinancement_values = capacite_autofinancement.get("capacite_autofinancement", [0.0, 0.0, 0.0])
    
    # Total des ressources
    total_ressources = [
        apport_personnel_list[0] + emprunts[0] + subventions_list[0] + autres_financements_list[0] + capacite_autofinancement_values[0],
        apport_personnel_list[1] + emprunts[1] + subventions_list[1] + autres_financements_list[1] + capacite_autofinancement_values[1],
        apport_personnel_list[2] + emprunts[2] + subventions_list[2] + autres_financements_list[2] + capacite_autofinancement_values[2]
    ]
    
    # Variation de trésorerie
    variation_tresorerie = [
        total_ressources[0] - total_besoins[0],
        total_ressources[1] - total_besoins[1],
        total_ressources[2] - total_besoins[2]
    ]
    
    # Excédent de trésorerie (cumulatif)
    excedent_tresorerie = []
    cumul_excedent = 0.0
    for i in range(3):
        cumul_excedent += variation_tresorerie[i]
        excedent_tresorerie.append(cumul_excedent)
    
    # Préparation des données pour le tableau
    data_table = {
        "Plan de financement à trois ans": [
            "Immobilisations",
            "Acquisition des stocks",
            "Variation du Besoin en fonds de roulement",
            "Remboursement d'emprunts",
            "Total des besoins",
            "Apport personnel",
            "Emprunts",
            "Subventions",
            "Autres financements",
            "Capacité d'auto-financement",
            "Total des ressources",
            "Variation de trésorerie",
            "Excédent de trésorerie"
        ],
        "Année 1": [
            f"{immobilisations[0]:,.2f} $",
            f"{acquisition_stocks[0]:,.2f} $",
            f"{variation_bfr[0]:,.2f} $",
            f"{remboursements_emprunts[0]:,.2f} $",
            f"{total_besoins[0]:,.2f} $",
            f"{apport_personnel_list[0]:,.2f} $",
            f"{emprunts[0]:,.2f} $",
            f"{subventions_list[0]:,.2f} $",
            f"{autres_financements_list[0]:,.2f} $",
            f"{capacite_autofinancement_values[0]:,.2f} $",
            f"{total_ressources[0]:,.2f} $",
            f"{variation_tresorerie[0]:,.2f} $",
            f"{excedent_tresorerie[0]:,.2f} $"
        ],
        "Année 2": [
            f"{immobilisations[1]:,.2f} $",
            f"{acquisition_stocks[1]:,.2f} $",
            f"{variation_bfr[1]:,.2f} $",
            f"{remboursements_emprunts[1]:,.2f} $",
            f"{total_besoins[1]:,.2f} $",
            f"{apport_personnel_list[1]:,.2f} $",
            f"{emprunts[1]:,.2f} $",
            f"{subventions_list[1]:,.2f} $",
            f"{autres_financements_list[1]:,.2f} $",
            f"{capacite_autofinancement_values[1]:,.2f} $",
            f"{total_ressources[1]:,.2f} $",
            f"{variation_tresorerie[1]:,.2f} $",
            f"{excedent_tresorerie[1]:,.2f} $"
        ],
        "Année 3": [
            f"{immobilisations[2]:,.2f} $",
            f"{acquisition_stocks[2]:,.2f} $",
            f"{variation_bfr[2]:,.2f} $",
            f"{remboursements_emprunts[2]:,.2f} $",
            f"{total_besoins[2]:,.2f} $",
            f"{apport_personnel_list[2]:,.2f} $",
            f"{emprunts[2]:,.2f} $",
            f"{subventions_list[2]:,.2f} $",
            f"{autres_financements_list[2]:,.2f} $",
            f"{capacite_autofinancement_values[2]:,.2f} $",
            f"{total_ressources[2]:,.2f} $",
            f"{variation_tresorerie[2]:,.2f} $",
            f"{excedent_tresorerie[2]:,.2f} $"
        ]
    }
    
    df = pd.DataFrame(data_table)
    st.write("### Tableau du Plan de financement à trois ans")
    st.table(df)
    
    # Stocker les résultats dans les données
    data["plan_financement"] = {
        "immobilisations": immobilisations,
        "acquisition_stocks": acquisition_stocks,
        "variation_bfr": variation_bfr,
        "remboursements_emprunts": remboursements_emprunts,
        "total_besoins": total_besoins,
        "apport_personnel": apport_personnel_list,
        "emprunts": emprunts,
        "subventions": subventions_list,
        "autres_financements": autres_financements_list,
        "capacite_autofinancement": capacite_autofinancement_values,
        "total_ressources": total_ressources,
        "variation_tresorerie": variation_tresorerie,
        "excedent_tresorerie": excedent_tresorerie
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data   
    
    # Préparer les données d'exportation pour Plan de Financement à Trois Ans
    export_table_plan_financement = []
    for idx, label in enumerate(data_table["Plan de financement à trois ans"]):
        export_table_plan_financement.append({
            "Plan de financement à trois ans": label,
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_plan_financement_trois_ans'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_plan_financement
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_plan_financement_trois_ans"):
        telecharger_document_complet()
        


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    export_data_budget_tresorerie_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
    export_data_budget_tresorerie_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})

    st.write(export_data_plan_financement)
    # Vérifiez que toutes les données nécessaires sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data"),
        export_data_budget_tresorerie_part1.get("table_data"),
        export_data_budget_tresorerie_part2.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant ($)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} $"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} $"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} $"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.2f} "
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.2f} "
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.2f} "
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Plan de Financement à Trois Ans ###
    doc.add_heading('Plan de financement à trois ans', level=1)
    
    # Créer le tableau Plan de Financement à Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=4)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Plan de financement à trois ans", "")
        row_cells_plan[1].text = row.get("Année 1", "")
        row_cells_plan[2].text = row.get("Année 2", "")
        row_cells_plan[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 9. Ajouter la section Budget Prévisionnel de Trésorerie Partie 1 ###
    doc.add_heading('Budget prévisionnel de trésorerie - Partie 1', level=1)
    
    # Créer le premier tableau Budget prévisionnel de trésorerie
    table_budget_part1 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part1['table_data'][0]))
    table_budget_part1.style = 'Light List Accent 1'
    table_budget_part1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-têtes
    headers_budget_part1 = export_data_budget_tresorerie_part1['table_data'][0].keys()
    hdr_cells_budget_part1 = table_budget_part1.rows[0].cells
    for i, header in enumerate(headers_budget_part1):
        hdr_cells_budget_part1[i].text = header
        for paragraph in hdr_cells_budget_part1[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part1[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données du premier tableau
    for row in export_data_budget_tresorerie_part1['table_data'][1:]:
        row_cells = table_budget_part1.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les données du budget prévisionnel de trésorerie - Partie 1 sont basées sur les estimations fournies.")
    
    ### 10. Ajouter la section Budget Prévisionnel de Trésorerie Partie 2 ###
    doc.add_heading('Budget prévisionnel de trésorerie - Partie 2', level=1)
    
    # Créer le deuxième tableau Budget prévisionnel de trésorerie
    table_budget_part2 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part2['table_data'][0]))
    table_budget_part2.style = 'Light List Accent 1'
    table_budget_part2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-têtes
    headers_budget_part2 = export_data_budget_tresorerie_part2['table_data'][0].keys()
    hdr_cells_budget_part2 = table_budget_part2.rows[0].cells
    for i, header in enumerate(headers_budget_part2):
        hdr_cells_budget_part2[i].text = header
        for paragraph in hdr_cells_budget_part2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part2[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données du deuxième tableau
    for row in export_data_budget_tresorerie_part2['table_data'][1:]:
        row_cells = table_budget_part2.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les données du budget prévisionnel de trésorerie - Partie 2 sont basées sur les estimations fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")


import streamlit as st
import pandas as pd

def page_budget_previsionnel_tresorerie():
    st.title("Budget prévisionnel de trésorerie")
    
    data = st.session_state.get("data", {})
    
    if not data:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **(Hors TVA)**")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Création du budget prévisionnel pour la première année, mois par mois
    months = [f"Mois {i+1}" for i in range(12)] + ["TOTAL"]
    
    # Initialisation des structures de données
    encaissements = {}
    decaissements = {}
    solde_precedent = [0.0] * 12
    solde_mois = [0.0] * 12
    solde_tresorerie_cumul = [0.0] * 12
    
    # ----------------------------
    # Encaissements
    # ----------------------------
    
    # Récupérer les encaissements depuis "Plan de financement à trois ans"
    plan_financement = data.get("plan_financement", {})
    apport_personnel = plan_financement.get("apport_personnel", [0.0, 0.0, 0.0])[0]
    emprunts = plan_financement.get("emprunts", [0.0, 0.0, 0.0])[0]
    subventions = plan_financement.get("subventions", [0.0, 0.0, 0.0])[0]
    autres_financements = plan_financement.get("autres_financements", [0.0, 0.0, 0.0])[0]
    
    encaissements["Apport personnel"] = [apport_personnel] + [0.0]*11
    encaissements["Emprunts"] = [emprunts] + [0.0]*11
    encaissements["Subventions"] = [subventions] + [0.0]*11
    encaissements["Autres financements"] = [autres_financements] + [0.0]*11
    
    # Récupérer les ventes depuis "Chiffre d'Affaires Prévisionnel"
    chiffre_affaires = data.get("chiffre_affaires", {})
    
    # Initialiser les listes de ventes mensuelles
    vente_marchandises_mensuel = []
    vente_services_mensuel = []
    
    # Remplir les ventes mensuelles de Marchandises
    for i in range(1, 13):
        key_ca = f"Marchandises_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_marchandises_mensuel.append(ca)
    
    # Remplir les ventes mensuelles de Services
    for i in range(1, 13):
        key_ca = f"Services_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_services_mensuel.append(ca)
    
    encaissements["Vente de marchandises"] = vente_marchandises_mensuel
    encaissements["Vente de services"] = vente_services_mensuel
    encaissements["Chiffre d'affaires (total)"] = [vente_marchandises_mensuel[i] + vente_services_mensuel[i] for i in range(12)]
    
    # Total des encaissements
    total_encaissements = []
    for i in range(12):
        total = (
            encaissements["Apport personnel"][i] +
            encaissements["Emprunts"][i] +
            encaissements["Subventions"][i] +
            encaissements["Autres financements"][i] +
            encaissements["Vente de marchandises"][i] +
            encaissements["Vente de services"][i]
        )
        total_encaissements.append(total)
    total_total_encaissements = sum(total_encaissements)
    total_encaissements.append(total_total_encaissements)
    
    # ----------------------------
    # Décaissements
    # ----------------------------
    
    # Récupérer les données nécessaires pour les décaissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    charges_variables = data.get("charges_variables", {})
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    
    # Immobilisations incorporelles et corporelles depuis "besoins_demarrage"
    immobilisations_incorporelles = sum([
        besoins_demarrage.get("Frais d’établissement", 0.0),
        besoins_demarrage.get("Frais d’ouverture de compteurs", 0.0),
        besoins_demarrage.get("Logiciels, formations", 0.0),
        besoins_demarrage.get("Dépôt de marque", 0.0),
        besoins_demarrage.get("Droits d’entrée", 0.0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0.0),
        besoins_demarrage.get("Droit au bail", 0.0),
        besoins_demarrage.get("Caution ou dépôt de garantie", 0.0),
        besoins_demarrage.get("Frais de dossier", 0.0),
        besoins_demarrage.get("Frais de notaire", 0.0),
    ])
    
    immobilisations_corporelles = sum([
        besoins_demarrage.get("Enseigne et éléments de communication", 0.0),
        besoins_demarrage.get("Véhicule", 0.0),
        besoins_demarrage.get("Matériel professionnel", 0.0),
        besoins_demarrage.get("Matériel autre", 0.0),
        besoins_demarrage.get("Matériel de bureau", 0.0)
    ])
    
    immobilisations_total = immobilisations_incorporelles + immobilisations_corporelles
    
    decaissements["Immobilisations incorporelles"] = [immobilisations_incorporelles] + [0.0]*11
    decaissements["Immobilisations corporelles"] = [immobilisations_corporelles] + [0.0]*11
    decaissements["Immobilisations (total)"] = [immobilisations_total] + [0.0]*11
    
    # Acquisition des stocks depuis "Stock de matières et produits"
    acquisition_stocks = besoins_demarrage.get("Stock de matières et produits", 0.0)
    decaissements["Acquisition stocks"] = [acquisition_stocks] + [0.0]*11
    
    # Échéances emprunt : "Principal année 1" divisé par 12
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    principal_annee1 = remboursements_emprunts[0]
    echeances_emprunt_mensuel = principal_annee1 / 12.0 if principal_annee1 > 0 else 0.0
    decaissements["Échéances emprunt"] = [echeances_emprunt_mensuel] * 12
    
    # Achats de marchandises : "Vente de marchandises" * "le coût d'achat de vos marchandises" de "Charges Variables"
    cout_achat_marchandises_pct = charges_variables.get("cout_achat_marchandises_pct", 100.0)
    if cout_achat_marchandises_pct == 0.0:
        cout_achat_marchandises_pct = 100.0  # Supposer 100% si non renseigné
    
    achats_marchandises_mensuel = [vente_marchandises_mensuel[i] * cout_achat_marchandises_pct / 100.0 for i in range(12)]
    decaissements["Achats de marchandises"] = achats_marchandises_mensuel
    
    # Charges externes : Récupérer depuis "soldes_intermediaires_de_gestion"
    charges_externes_annee1 = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])[0]
    charges_externes_mensuel = charges_externes_annee1 / 12.0 if charges_externes_annee1 > 0 else 0.0
    decaissements["Charges externes"] = [charges_externes_mensuel] * 12
    
    # Impôts et taxes
    impots_et_taxes_annee1 = compte_resultat.get("impots_et_taxes", [0.0, 0.0, 0.0])[0]
    impots_et_taxes_mensuel = impots_et_taxes_annee1 / 12.0 if impots_et_taxes_annee1 > 0 else 0.0
    decaissements["Impôts et taxes"] = [impots_et_taxes_mensuel] * 12
    
    # Salaires employés, Charges sociales employés, Prélèvement dirigeant(s), Charges sociales dirigeant(s), Frais bancaires, charges financières
    salaires_employes_annee1 = compte_resultat.get("salaires_employes", [0.0, 0.0, 0.0])[0]
    charges_sociales_employes_annee1 = compte_resultat.get("charges_sociales_employes", [0.0, 0.0, 0.0])[0]
    prelevement_dirigeants_annee1 = compte_resultat.get("salaires_dirigeants", [0.0, 0.0, 0.0])[0]
    charges_sociales_dirigeants_annee1 = compte_resultat.get("charges_sociales_dirigeants", [0.0, 0.0, 0.0])[0]
    frais_bancaires_annuels = compte_resultat.get("total_frais_financiers", [0.0, 0.0, 0.0])[0]
    
    salaires_employes_mensuel = [salaires_employes_annee1 / 12.0] * 12
    charges_sociales_employes_mensuel = [charges_sociales_employes_annee1 / 12.0] * 12
    prelevement_dirigeants_mensuel = [prelevement_dirigeants_annee1 / 12.0] * 12
    charges_sociales_dirigeants_mensuel = [charges_sociales_dirigeants_annee1 / 12.0] * 12
    frais_bancaires_mensuel = [frais_bancaires_annuels / 12.0] * 12 if frais_bancaires_annuels > 0 else [0.0] * 12
    
    decaissements["Salaires employés"] = salaires_employes_mensuel
    decaissements["Charges sociales employés"] = charges_sociales_employes_mensuel
    decaissements["Prélèvement dirigeant(s)"] = prelevement_dirigeants_mensuel
    decaissements["Charges sociales dirigeant(s)"] = charges_sociales_dirigeants_mensuel
    decaissements["Frais bancaires, charges financières"] = frais_bancaires_mensuel
    
    # ----------------------------
    # Total charges de personnel
    # ----------------------------
    # Calculer le total des charges de personnel pour chaque mois
    total_charges_personnel_mensuel = [
        salaires_employes_mensuel[i] + charges_sociales_employes_mensuel[i] +
        prelevement_dirigeants_mensuel[i] + charges_sociales_dirigeants_mensuel[i]
        for i in range(12)
    ]
    decaissements["Total charges de personnel"] = total_charges_personnel_mensuel
    
    # ----------------------------
    # Total des décaissements
    # ----------------------------
    
    # Définir les clés à inclure dans le total des décaissements
    decaissements_keys = [
        "Immobilisations (total)",
        "Acquisition stocks",
        "Échéances emprunt",
        "Achats de marchandises",
        "Charges externes",
        "Impôts et taxes",
        "Total charges de personnel",
        "Frais bancaires, charges financières"
    ]
    
    total_decaissements = []
    for i in range(12):
        total = sum([decaissements[key][i] for key in decaissements_keys])
        total_decaissements.append(total)
    total_total_decaissements = sum(total_decaissements)
    total_decaissements.append(total_total_decaissements)
    
    # ----------------------------
    # Calcul des Soldes
    # ----------------------------
    
    for i in range(12):
        solde_mois[i] = total_encaissements[i] - total_decaissements[i]
        solde_tresorerie_cumul[i] = solde_tresorerie_cumul[i - 1] + solde_mois[i] if i > 0 else solde_mois[i]
        solde_precedent[i] = solde_tresorerie_cumul[i - 1] if i > 0 else 0.0
    
    # Append totals to solde_mois, solde_precedent, solde_tresorerie_cumul
    total_solde_mois = sum(solde_mois)
    solde_mois.append(total_solde_mois)
    
    # Pour solde_precedent, le total n'est pas significatif, on peut ajouter une chaîne vide
    solde_precedent.append("")
    
    # Pour solde_tresorerie_cumul, on peut ajouter la dernière valeur cumulative
    solde_tresorerie_cumul.append(solde_tresorerie_cumul[-1])
    
    # ----------------------------
    # Préparation des données pour le tableau
    # ----------------------------
    
    table_data = {"Description": months}
    
    # Encaissements
    for key in encaissements:
        amounts = encaissements[key]
        total = sum(amounts)
        amounts_with_total = amounts + [total]
        table_data[key] = [f"{value:,.2f} $" if value != 0 else "-" for value in amounts_with_total]
    
    # Décaissements
    for key in decaissements:
        # Inclure toutes les lignes de décaissements
        amounts = decaissements[key]
        total = sum(amounts)
        # Remplacer 0 par '-' si nécessaire
        amounts_with_total = [f"{value:,.2f} $" if value != 0 else "-" for value in amounts] + [f"{total:,.2f} $" if total != 0 else "-"]
        table_data[key] = amounts_with_total
    
    # Totaux et soldes
    table_data["Total des encaissements"] = [f"{value:,.2f} $" if value != 0 else "-" for value in total_encaissements]
    table_data["Total des décaissements"] = [f"{value:,.2f} $" if value != 0 else "-" for value in total_decaissements]
    solde_precedent_formatted = [f"{value:,.2f} $" if isinstance(value, (int, float)) and value != 0 else "-" for value in solde_precedent]
    table_data["Solde précédent"] = solde_precedent_formatted
    table_data["Solde du mois"] = [f"{value:,.2f} $" if value != 0 else "-" for value in solde_mois]
    table_data["Solde de trésorerie (cumul)"] = [f"{value:,.2f} $" if value != 0 else "-" for value in solde_tresorerie_cumul]
    
    # Assurer que toutes les listes ont la même longueur
    max_length = max(len(lst) for lst in table_data.values())
    for key in table_data:
        if len(table_data[key]) < max_length:
            difference = max_length - len(table_data[key])
            table_data[key] += [""] * difference  # Remplir avec des chaînes vides si nécessaire
        elif len(table_data[key]) > max_length:
            table_data[key] = table_data[key][:max_length]  # Tronquer si trop long
    
    # Création du DataFrame complet
    df_full = pd.DataFrame(table_data)
    df_full.set_index("Description", inplace=True)
    df_full = df_full.T  # Transposer pour avoir les mois comme colonnes
    
    # Séparation en deux tableaux
    # Tableau 1 : Mois 1 à Mois 5
    columns_part1 = ["Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5"]
    df_part1 = df_full[columns_part1]
    
    # Tableau 2 : Mois 6 à Mois 12 + TOTAL
    columns_part2 = ["Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"]
    df_part2 = df_full[columns_part2]
    
    ### 3. Ajouter la section Budget Prévisionnel de Trésorerie ###
    # (Les deux tableaux seront ajoutés dans telecharger_document_complet())
    
    ### 4. Affichage des tableaux séparés ###
    st.subheader("Budget prévisionnel de trésorerie")
    st.table(df_part1)
    
    st.subheader("Budget prévisionnel de trésorerie (suite)")
    st.table(df_part2)
    
    # ----------------------------
    # Stockage des résultats dans les données
    # ----------------------------
    
    data["budget_previsionnel_tresorerie"] = {
        "encaissements": encaissements,
        "decaissements": decaissements,
        "total_encaissements": total_encaissements,
        "total_decaissements": total_decaissements,
        "solde_precedent": solde_precedent,
        "solde_mois": solde_mois,
        "solde_tresorerie_cumul": solde_tresorerie_cumul
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data   
    
    # ----------------------------
    # Préparation des données d'exportation pour Budget Prévisionnel de Trésorerie Partie 1
    # ----------------------------
    
    export_table_budget_part1 = []
    headers_part1 = df_part1.columns.tolist()
    export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [""] + headers_part1)))
    for index, row in df_part1.iterrows():
        export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [index] + row.tolist())))
    
    # Stocker les données d'exportation pour Partie 1
    st.session_state['export_data_budget_previsionnel_tresorerie_part1'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part1
    }
    
    # ----------------------------
    # Préparation des données d'exportation pour Budget Prévisionnel de Trésorerie Partie 2
    # ----------------------------
    
    export_table_budget_part2 = []
    headers_part2 = df_part2.columns.tolist()
    export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [""] + headers_part2)))
    for index, row in df_part2.iterrows():
        export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [index] + row.tolist())))
    
    # Stocker les données d'exportation pour Partie 2
    st.session_state['export_data_budget_previsionnel_tresorerie_part2'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part2
    }
    
    # ----------------------------
    # Section Export
    # ----------------------------
    
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_budget_previsionnel_tresorerie"):
        telecharger_document_complet()

# Section 15 : Tableaux d'Analyse Financière
def page_douze_tableaux():
    st.title("Tableaux d'Analyse Financière")
    
    data = st.session_state["data"]
    
    st.markdown("""
    Cette section présente les principaux indicateurs financiers basés sur les données que vous avez saisies.
    """)
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    charges_sociales_dirigeant_annee1 = data.get("charges_sociales", {}).get("dirigeants", {}).get("annee1", 0.0)
    charges_sociales_employes_annee1 = data.get("charges_sociales", {}).get("employes", {}).get("annee1", 0.0)
    amortissements_annee1 = data.get("amortissements", {}).get("total", {}).get("annee1", 0.0)
    
    # Calcul du résultat net
    resultat_net = total_ca_annee1 - total_charges_fixes_annee1 - total_charges_variables - total_salaires_annee1 - charges_sociales_dirigeant_annee1 - charges_sociales_employes_annee1 - amortissements_annee1
    
    # Capacité d'autofinancement (simplifiée)
    capacite_autofinancement = resultat_net + amortissements_annee1  # Les amortissements sont réintégrés
    
    st.write(f"**Résultat Net Année 1 :** {resultat_net:.2f} $")
    st.write(f"**Capacité d'Autofinancement Année 1 :** {capacite_autofinancement:.2f} $")
    
    # Vous pouvez répéter les calculs pour les années 2 et 3 si nécessaire
    
    st.write("---")
    
    st.session_state["data"] = data
    
def load_and_split_documents(file_path):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    raw_documents = PyPDFLoader(file_path).load()
    return text_splitter.split_documents(raw_documents)

def create_faiss_db(documents):
    if not documents:
        raise ValueError("Aucun document trouvé pour créer la base de données FAISS.")
    embeddings = OpenAIEmbeddings(openai_api_key=api_key )
    return FAISS.from_documents(documents, embeddings)

def generate_section(system_message, query, documents, combined_content, tableau_financier, rubriques, business_model,bm_precedent,nom_entreprise,precedent_chain_of, generation=1):
    
    MODEL="gpt-4o"
    MAX_TOKENS_PER_REQUEST = 150
    utilisateur = get_current_user()
    
    # Calculer les tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    # Fusionner tous les textes des documents pour le comptage des tokens
    # Compter les tokens dans la requête (query + full_text)
    full_text = " ".join([doc.page_content for doc in documents])
    tokens_in_input = count_tokens(query + full_text, MODEL)   
    
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    llm = ChatOpenAI(openai_api_key=api_key )
    if generation ==1 :
        if documents:
            db = create_faiss_db(documents)
            if db:
                success, message = consommer_tokens(st.session_state['user_info'], tokens_in_input)
                if not success:
                    return ""
                # Consommer les tokens
            
            
            qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory, verbose=True)
            combined_info = qa_chain.run({'question': query})
             #tokens_utilises = combined_info['usage']['total_tokens']
            tokens_in_ = count_tokens(combined_info +""+query  , MODEL)
            tokens_utilises = tokens_in_+ tokens_in_input + MAX_TOKENS_PER_REQUEST
            # Consommer les tokens
            success, message = consommer_tokens(st.session_state['user_info'], tokens_utilises )
            if not success:
                return ""
            
            full_contents = combined_content + " " + combined_info + " " + query+ " "+ json.dumps(tableau_financier)+ "Voici le nom de l'entreprise:"+ nom_entreprise
            full_content = combined_content + " " + combined_info + " " + query+ " "+ tableau_financier + "Voici le nom de l'entreprise:"+ nom_entreprise
 
        else:
            full_content = combined_content + " " + query+ "Dans ce données où vous allez recuperer les informations generales de l'entreprises "+ tableau_financier+ "utiliser les données financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise. Voici les autres informations à considerer c'est les informations du business model et ca doit etre tenue compte lors de la generation:"+ json.dumps(business_model)+ "Voici le nom de l'entreprise:"+nom_entreprise
            full_contents = combined_content + " " + query+ "Dans ce données où vous allez recuperer les informations generales de l'entreprises "+ json.dumps(tableau_financier)+ "utiliser les données financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise. Voici les autres informations à considerer c'est les informations du business model et ca doit etre tenue compte lors de la generation:"+ json.dumps(business_model)+ "Voici le nom de l'entreprise:"+nom_entreprise
           
    else: 
        if documents:
            db = create_faiss_db(documents)
            if db:
                success, message = consommer_tokens(st.session_state['user_info'], tokens_in_input)
                if not success:
                    return ""
                # Consommer les tokens
            qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory, verbose=True)
            combined_info = qa_chain.run({'question': query})
            
            tokens_in_ = count_tokens(combined_info +""+query  , MODEL)
            tokens_utilises = tokens_in_+ tokens_in_input + MAX_TOKENS_PER_REQUEST
            # Consommer les tokens
            success, message = consommer_tokens(st.session_state['user_info'], tokens_utilises )
            if not success:
                return ""
            
            full_content = combined_content + " " + combined_info + " " + query+ " "+json.dumps(tableau_financier)+ "Voici le nom de l'entreprise:"+ nom_entreprise
            full_contents= combined_content + " " + combined_info + " " + query+ " "+json.dumps(tableau_financier) + "Voici le nom de l'entreprise:"+ nom_entreprise
            prompt = f"""
            voici le prompt à executé pour cette sections:{query}
            Voici le nom de l'entreptise :{nom_entreprise}
            voici les informations que vous avez generer precedement dans d'autres section, combiner des informations entrer par l'utlisateurs: {combined_content}
            Voici les données provenu d'une autres sources qui est un document: {combined_info}
            Voici les données qui nous ont permit de genener le business model: {rubriques}
            Voici le business model generer:{business_model}
            Voici les données des analyses financiers: {tableau_financier}
            Voici les analyses du business model et ce qu'il faudra ameliorer dans le bussiness model: {precedent_chain_of}
            Voici le business plan genener precedement: {bm_precedent}
            """
            full_contents=prompt 
            full_content=prompt 
            
        else:
            full_content = combined_content + " " + query+ "Dans ce données où vous allez recuperer les informations generales de l'entreprises "+ json.dumps(tableau_financier)+ "utiliser les données financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise. Voici les autres informations à considerer c'est les informations du business model et ca doit etre tenue compte lors de la generation:"+ json.dumps(business_model)+ "Voici le nom de l'entreprise:"+""+nom_entreprise
            full_contents = combined_content + " " + query+ "Dans ce données où vous allez recuperer les informations generales de l'entreprises "+ json.dumps(tableau_financier)+ "utiliser les données financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise. Voici les autres informations à considerer c'est les informations du business model et ca doit etre tenue compte lors de la generation:"+ json.dumps(business_model)+ "Voici le nom de l'entreprise:"+""+nom_entreprise
            
            prompt = f"""
                voici le prompt à executé pour cette sections:{query}
                Voici le nom de l'entreptise :{nom_entreprise}
                voici les informations que vous avez generer precedement dans d'autres section, combiner des informations entrer par l'utlisateurs: {combined_content}
                Voici les données qui nous ont permit de genener le business model: {rubriques}
                Voici le business model generer :{business_model}
                Voici les données des analyses financiers: {tableau_financier}
                Voici les analyses du business plan et ce qu'il faudra ameliorer dans le bussiness plan: {precedent_chain_of}
                Voici le business plan genener precedement: {bm_precedent}
                """
            full_content=prompt 
            full_contents=prompt 
    # Calculer les tokens nécessaires (entrée + réponse prévue)
    # Calculer le nombre de tokens dans l'entrée utilisateur
    tokens_in_input = count_tokens(full_contents, MODEL)
    tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
    # Vérifier si l'utilisateur a assez de tokens
    
    if tokens_remaining < tokens_needed:
        st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
        return
    else:    
        completion = openai.ChatCompletion.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": full_content}
            ],
            temperature=0.9
        )
        tokens_utilises = completion['usage']['total_tokens']
        # Consommer les tokens
        success, message = consommer_tokens(st.session_state['user_info'], tokens_utilises)
        if not success:
            return ""
        return completion['choices'][0]['message']['content']

def extract_company_name(text):
    match = re.search(r"(nom de l'entreprise est|Nom de l'entreprise|La vision de) ([\w\s]+)", text, re.IGNORECASE)
    if match:
        return match.group(2).strip()
    return "Nom de l'entreprise non trouvé"

def generate_markdown(results):
    markdown_content = "# Business Plan\n\n"
    for sec_name, content in results.items():
        markdown_content += f"## {sec_name}\n\n"
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.startswith('- '):  # Points de liste
                markdown_content += f"- {paragraph[2:]}\n"
            elif re.match(r'^\d+\.\s', paragraph):  # Points numérotés
                markdown_content += f"{paragraph}\n"
            else:
                markdown_content += f"{paragraph}\n"
        markdown_content += "\n"

    return markdown_content

def convert_table_to_markdown(table_name, table_data):
    """
    Convertit les données d'une table en format Markdown.
    
    Args:
        table_name (str): Nom de la table.
        table_data (list of dict): Données de la table.
    
    Returns:
        str: Table au format Markdown.
    """
    if not table_data:
        return "Aucune donnée disponible."
    
    # Extraire les en-têtes de colonnes
    headers = list(table_data[0].keys())
    markdown_table = "| " + " | ".join(headers) + " |\n"
    markdown_table += "| " + " | ".join(['---'] * len(headers)) + " |\n"
    
    # Ajouter les lignes
    for row in table_data:
        markdown_table += "| " + " | ".join([str(row.get(header, "")) for header in headers]) + " |\n"
    
    return markdown_table

def convert_all_tables_to_markdown(tables):
    """
    Convertit toutes les tables en une seule chaîne de caractères au format Markdown.
    
    Args:
        tables (dict): Dictionnaire contenant les tables financières.
    
    Returns:
        str: Toutes les tables concaténées en Markdown.
    """
    markdown = ""
    for table_name, table_data in tables.items():
        markdown += f"### {table_name}\n\n"
        markdown += convert_table_to_markdown(table_name, table_data) + "\n\n"
    return markdown


def markdown_to_word_via_text(markdown_content):
    # Créer un nouveau document Word
    doc = Document()
    doc.add_heading('Business Plan', 0)

    # Diviser le contenu en lignes
    lines = markdown_content.split('\n')
    table_data = []
    inside_table = False
    plain_text_output = []  # Pour collecter le texte brut

    for line in lines:
        line = line.strip()
        if not line:
            # Si ligne vide et données de table en cours, ajouter le tableau au document
            if table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell.strip()
                table_data = []
                inside_table = False
            continue

        if line.startswith('## '):
            # Sous-titre
            doc.add_heading(line[3:], level=2)
            plain_text_output.append(line[3:])
        elif line.startswith('- '):
            # Liste à puces
            doc.add_paragraph(line[2:], style='List Bullet')
            plain_text_output.append(f"• {line[2:]}")
        elif re.match(r'^\d+\.\s', line):
            # Liste numérotée
            doc.add_paragraph(line, style='List Number')
            plain_text_output.append(line)
        elif line.startswith('|'):
            # Détection des lignes de tableau (évite les lignes de séparation)
            if re.match(r'\|?\s*[-:]+\s*\|', line):
                inside_table = True
                continue  # Ignorer les lignes de séparation
            else:
                inside_table = True
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
        elif line.startswith('**') and line.endswith('**'):
            # Texte en gras
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
            plain_text_output.append(line[2:-2])
        elif not inside_table:
            # Paragraphe normal
            doc.add_paragraph(line)
            plain_text_output.append(line)

    # Traiter les données de table restantes
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell.strip()

    # Sauvegarder le document dans un buffer mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return "\n".join(plain_text_output), buffer

# Fonction pour convertir un dictionnaire en texte formaté
def format_table_data(data, title):
    if not data:
        return f"{title} : Aucune donnée disponible.\n"
    
    text = f"{title} :\n"
    for key, value in data.items():
        if isinstance(value, dict):
            text += f"  {key} :\n"
            for sub_key, sub_value in value.items():
                text += f"    {sub_key} : {sub_value}\n"
        elif isinstance(value, list):
            text += f"  {key} : {', '.join(map(str, value))}\n"
        else:
            text += f"  {key} : {value}\n"
    return text + "\n"

def page_generation_business_plan():
    # Création des colonnes pour le titre et le bouton
    st.title("Générateur de Business Plan")

            
    #st.write(st.session_state["produits_data"])

    uploaded_file = st.file_uploader("Téléchargez votre fichier PDF", type="pdf")
    user_text_input = st.text_area("Entrez des informations supplémentaires ou un texte alternatif:", height=200)
    

    if uploaded_file or user_text_input:
        documents = []
        combined_content = user_text_input  

        if uploaded_file:
            file_path = "uploaded_document.pdf"
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            documents = load_and_split_documents(file_path)



        # Créer un dictionnaire pour stocker les résultats
        results = {}
        
        # Messages système et requêtes pour chaque section
        system_messages = {
            "Couverture": """
                Générer cette section du business plan:
                Voici les textes à afficher sous forme :
                
                # Canevas de Plans d’Affaires

                Nom du projet ou entreprise
                
                 

            """,
            "Sommaire": """
                Générer cette section du business plan:
                Voici les textes à afficher sous forme de liste:
                ## Sommaire
                I. Résumé Exécutif « Executive Summary » / Pitch
                II. Présentation de votre entreprise/projet
                III. Présentation de l’offre de produit(s) et/ou service(s)  
                IV. Étude de marché
                V. Stratégie marketing, communication et politique commerciale
                VI. Moyens de production et organisation 
                VII. Étude des risques/hypothèses  
                VIII. Plan financier 
                
            """,
            "Résumé Exécutif": """
                Générer cette section du business plan:
                
                ## I. Résumé Exécutif « Executive Summary » / Pitch
                Générer deux grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                Attirer l'attention du lecteur en 5 minutes et lui donner envie d'en savoir plus.
                Décrire le projet en quelques phrases simples et impactantes.
                Ne pas essayer de tout couvrir, soyez concis et précis.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Présentation de la PME** : Nom de l’entreprise et brève description du service/produit fourni.
                - **Présentation des porteurs de projet** : Nom, prénom, coordonnées, situation de famille, formation et diplômes, expérience professionnelle, activités extra ou para-professionnelles (Joindre CV en annexe).
                - **Potentiel en termes de taille et de profit** : Démontrez par des calculs simples comment votre PME fera du profit.
                - **Votre besoin financier**.

            """,
            "Présentation de votre entreprise": """
                Générer cette section du business plan:

                ## II. Présentation de votre entreprise/projet

                Générer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de votre entreprise/projet de manière plus détaillée.
                - Présenter l’équipe managériale clé.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Informations générales sur la PME** :
                - Forme juridique : Ets, Sarlu, Sarl, SAS, SA.
                - Siège social : Adresse juridique de l’entreprise.
                - Coordonnées bancaires : Numéro de compte (avec 23 chiffres) de l’entreprise ainsi que la banque où est logé le compte (joindre le Swift Copy).
                - Couverture géographique de l’entreprise et ses activités : lieu d’implantation de l’entreprise et différentes zones couvertes.
                - **Description détaillée de la PME et objectifs de son projet** : Présentez l’entreprise, son origine, introduisez ses atouts/opportunités et enfin décrivez le projet de l’entreprise.
                - **Stade d’avancement de l’entreprise ou du projet** :
                - Décrivez ce qui a été fait et les projets à mener dans le futur.
                - Parlez du niveau de maturité de la PME ou du projet.
                - Lister éventuellement les financements déjà acquis.
                - **Présentation de l’équipe managériale** : Décrivez l’organigramme et l’organisation des ressources humaines, présentez les associés de la PME ainsi que leurs parts sociales.
                - **Analyse SWOT** : Forces, faiblesses, opportunités, contraintes/menaces. de preference ca doit etre presenter sous forme de tableau.
                - **Business Modèle Canevas** : Insérer votre business modèle canevas avec les 9 rubriques bien remplies.

            """,
            "Présentation de l’offre de produit": """
                Générer cette section du business plan :

                ## III. Présentation de l’offre de produit(s) et/ou service(s)
                Générer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de l’offre de produits/services de manière détaillée.
                - Présenter la proposition de valeur différenciante de la PME ou de son offre.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Noms du/des produit(s) ou service(s)**.
                - **Besoins identifiés** sur le marché auxquels répond votre offre.
                - **Description du/des produit(s) ou service(s)** répondant à ces besoins.
                - **Proposition de valeur unique**.
                - **Prise en compte de l’aspect genre** dans le fonctionnement de la PME ou du projet de l’entreprise.
                - **Prise en compte de l’environnement** :
                - Identification des impacts environnementaux et sociaux des activités de la PME.
                - Mise en place de mesures d’atténuation.
                - Existence d’un Plan de Gestion Environnemental et Social.

            """,
            "Étude de marché": """
                Générer cette section du business plan :

                ## IV. Étude de marché

                Générer 8 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Expliquer la méthode utilisée pour la conduite de l’étude de marché.
                - Présenter les résultats de l’étude de marché.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Description des hypothèses et méthodes de l’étude de marché** :
                - Citer le produit ou service pré-ciblé.
                - Préciser le marché pré-ciblé : secteur d’activité dans lequel le produit s’inscrit.
                - Présenter les méthodes choisies pour réaliser l’étude de marché : questionnaire, étude documentaire, étude de concurrence, étude métier, etc.

                2. **Approche générale du marché (précisez les sources à chaque étape)** :
                - Décrire le marché, ses principales caractéristiques, historique et perspectives.
                - Présenter les résultats : marché cible, marché potentiel, marché réel.
                - Présenter les menaces et opportunités du marché.

                3. **Caractéristiques de la demande** :
                - Présenter le volume de la demande, l’évolution de la demande sur le marché ciblé et les tendances de consommation.
                - Détailler les différents types de clientèle (segmentation).
                - Lister les prescripteurs (partenaires qui peuvent apporter des clients).

                4. **Caractéristiques de l’offre** :
                - Présenter la concurrence directe et indirecte : lister les concurrents et décrire leur offre de services/produits.
                - Lister les points forts et les points faibles de la concurrence : avantages concurrentiels de la concurrence sur le marché.
                - Comment vous différenciez-vous de ces concurrents indirects ?

                5. **Caractéristiques de l’environnement** :
                - Décrire l’environnement des affaires relatif au développement de la PME/projet : le cadre légal, réglementaire, les facteurs externes au marché lui-même, l’évolution des technologies.
                - Lister les menaces et opportunités liées à l’environnement.

                6. **Partenariats** :
                - Préciser les partenariats stratégiques noués ou à mettre en place pour faire croître l’entreprise : il peut s’agir des acteurs en amont et en aval de votre chaîne de production/distribution (fournisseurs, distributeurs, partenaires commerciaux, etc.).

                7. **Création d’emplois** :
                - Démontrer l’impact de la PME/projet en termes d’emplois directs déjà créés ou à créer.

                8. **Chiffre d’affaires** :
                - Préciser la part de marché visée et le volume de chiffre d’affaires prévisible à horizon 1 an, 2 ans, 3 ans.

            """,
            "Stratégie Marketing":  """
                Générer cette section du business plan :

                ## V. Stratégie Marketing, Communication et Politique Commerciale

                Générer cette section, l'objectif pour cette section est de :
                - Présenter la stratégie marketing et commerciale à court et moyen terme.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Choix de segments de clientèle** :
                - Expliquer quels segments de clientèle vont constituer la cible de la PME/projet et pourquoi ce choix.
                - Expliquer dans les grandes lignes le positionnement stratégique.

                2. **Marketing-mix (4P : Produit – Prix – Place – Promotion)** :
                - Présenter la politique marketing générale :
                    - Choix du nom, du logo et des couleurs.
                    - Choix du message, du slogan.
                - Tableau synthétique des segments :

                    | Segment de clientèle | Produit proposé | Positionnement en termes de prix | Lieu de distribution | Style et mode de communication |
                    |-----------------------|-----------------|----------------------------------|-----------------------|---------------------------------|
                    | Segment 1            |                 |                                  |                       |                                 |
                    | Segment 2            |                 |                                  |                       |                                 |
                    | Segment 3            |                 |                                  |                       |                                 |

                3. **Plan Marketing et actions commerciales**  :
                - Présenter le plan marketing : lister les actions commerciales et actions de communication prévues ; inscrire leur coût si possible.

                    | Types d’actions       | Janvier | Février | Mars | ... | Décembre |
                    |-----------------------|---------|---------|------|-----|----------|
                    | Action 1             |         |         |      |     |          |
                    | Action 2             |         |         |      |     |          |

                4. **Moyens et partenaires sollicités** :
                - Lister les moyens à mettre en œuvre et les partenaires sollicités pour les actions commerciales et de communication.

            """,
            "Moyens de production et organisation": """
                Générer cette section du business plan:

                ## VI. Moyens de production et organisation

                Générer 4 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Spécifier les moyens humains et matériels à disposition de la PME.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Locaux** :
                - Liste des locaux, bail de location, conditions négociées, coût, utilité.
                - **Matériel** :
                - Liste, mode d’acquisition ou de location, coût, utilité, renouvellement.
                - **Moyens humains** :
                - Personnel, plannings, horaires, coût, charges sociales ; indiquer une répartition claire des tâches.
                - **Fournisseurs et sous-traitants** :
                - Liste des fournisseurs et/ou sous-traitants, devis obtenus, tarifs, conditions négociées.

            """,
            "Étude des risques": """
                Générer cette section du business plan:

                ## VII. Étude des risques/hypothèses

                Générer cette section, l'objectif pour cette section est de :
                - Présenter la synthèse des risques et mesures d’atténuation identifiés quant au développement de la PME/projet.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Tableau des risques** :

                | Nature de risque          | Description              | Stratégie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques liés à l’environnement général |                          |                            |
                | Risques liés au marché    |                          |                            |
                | Risques liés aux outils   |                          |                            |
                | Risques liés aux personnes |                          |                            |
                | Risques liés aux tiers    |                          |                            |
                | Autres risques (spécifiez) |                          |                            |

                Étude des risques/hypothèses:

            """,
            "Annexes": """
                Générer cette section du business plan:
                
                ## VII. Étude des risques/hypothèses

                ### Objectif
                - Présenter la synthèse des risques et mesures d’atténuation identifiés quant au développement de la PME/projet.

                ### Contenu attendu
                - **Tableau des risques** :

                | Nature de risque          | Description              | Stratégie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques liés à l’environnement général |                          |                            |
                | Risques liés au marché    |                          |                            |
                | Risques liés aux outils   |                          |                            |
                | Risques liés aux personnes |                          |                            |
                | Risques liés aux tiers    |                          |                            |
                | Autres risques (spécifiez) |                          |                            |

            """,
            "Annexes": """
                Générer cette section du business plan:

                7 – ANNEXES
                Renvoyer en annexe les documents trop volumineux ou difficiles à lire : - - - -
                étude de marché complète,
                contrats,
                conditions

                Annexes du projet:

            """
        }

        queries = {
            "Couverture": "Afficher seulement le texte fournies",
            "Sommaire": "Afficher seulement le texte fournises",
            "Résumé Exécutif": "Décrire brièvement le projet, son potentiel de profit et les qualifications de l'équipe.",
            "Présentation de votre entreprise": "Fournir une analyse détaillée de l'entreprise, incluant son origine, ses objectifs et son organisation.",
            "Présentation de l’offre de produit": "Décrire les produits ou services, leur proposition de valeur unique, et les besoins du marché qu'ils adressent.",
            "Étude de marché": "Analyser le marché cible, les tendances de consommation, et la concurrence directe et indirecte.",
            "Stratégie Marketing": "Décrire la stratégie marketing, y compris les segments cibles, le positionnement, le mix marketing (Produit, Prix, Place, Promotion) et les actions commerciales prévues.",
            "Moyens de production et organisation": "Décrire les moyens humains et matériels, ainsi que l'organisation opérationnelle de l'entreprise.",
            "Étude des risques": "Identifier les risques potentiels et proposer des stratégies pour les atténuer.",
            "Annexes": "Inclure tous les documents annexes pertinents pour étayer le plan d'affaires."
        }

        # Espaces réservés pour chaque section
        placeholders = {name: st.empty() for name in system_messages.keys()}
        
        data = st.session_state.get("data", {})
        tables = data.get("tables", {})

        section_order = list(system_messages.keys())
        # Définir le point de séparation
        split_section = "Présentation de votre entreprise"
        # Séparer les sections en deux groupes
        first_part = []
        second_part = []
        for section in section_order:
            if section == split_section:
                first_part.append(section)
                second_part = section_order[section_order.index(section)+1:]
                break
            else:
                first_part.append(section)

        
        results_first_part = {}
        results_second_part = {}
        
        

        # Récupérer les données exportées de toutes les sections
        # Récupérer les données exportées de toutes les sections 
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
        export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
        export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
        export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
        export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
        export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
        

        # Concaténer toutes les sections
        final_text = ""
        final_text += format_table_data(export_data_investissements, "Investissements et financements")
        final_text += format_table_data(export_data_salaires, "Salaires et Charges Sociales")
        final_text += format_table_data(export_data_amortissements, "Détail des Amortissements")
        final_text += format_table_data(export_data_compte, "Compte de résultats prévisionnel")
        final_text += format_table_data(export_data_soldes, "Soldes intermédiaires de gestion")
        final_text += format_table_data(export_data_capacite, "Capacité d'autofinancement")
        final_text += format_table_data(export_data_seuil, "Seuil de rentabilité économique")
        final_text += format_table_data(export_data_bfr, "Besoin en fonds de roulement")

        # Ajouter les nouvelles sections
        final_text += format_table_data(export_data_plan_financement, "Plan de financement à trois ans")
        final_text += format_table_data(export_data_budget_part1, "Budget prévisionnel de trésorerie")
        final_text += format_table_data(export_data_budget_part2, "Budget prévisionnel de trésorerie(suite)")


        st.session_state.get('business_plan_markdown_content1') 
        st.session_state.get('business_plan_markdown_content2')
        
        MODEL="gpt-4o"
        MAX_TOKENS_PER_REQUEST = 150
        utilisateur = get_current_user()
        # Calculer les tokens restants
        tokens_purchased = utilisateur.get('tokens_purchased', 0)
        tokens_consumed = utilisateur.get('tokens_consumed', 0)
        tokens_remaining = tokens_purchased - tokens_consumed
        # Fusionner tous les textes des documents pour le comptage des tokens
        # Compter les tokens dans la requête (query + full_text)
        full_text = " ".join([doc.page_content for doc in documents])

    
        # Générer toutes les sections automatiquement
        if st.button("Generer business plan"):
            
            for section_name in first_part:
                with st.spinner(f"Génération de {section_name}..."):
                    system_message = system_messages[section_name]
                    query = queries[section_name]
                            
                    full_contents = combined_content + " " + query+ "" + system_message+ json.dumps(final_text)+ full_text+ ""+ json.dumps(st.session_state["produits_data"])+json.dumps(serialize_membres(st.session_state["membres"]))
                    tokens_in_input = count_tokens(full_contents , MODEL) 
                    tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
                    
                    # Vérifier si l'utilisateur a assez de tokens
                    if tokens_remaining < tokens_needed:
                        st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
                        return
                    
                    try:
                        # Vérifier si la section est "Couverture" ou "Sommaire"
                        if section_name in ["Couverture", "Sommaire"]:
                            results_first_part[section_name] = generate_section(system_message, query, "", "", "", rubriques="", business_model="",bm_precedent="",nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of="", generation=1)

                        
                        else:
                            results_first_part[section_name] = generate_section(system_message, query, documents, combined_content+"Voici la liste de membre de l'entreprise:"+json.dumps(serialize_membres(st.session_state["membres"])), final_text, rubriques="", business_model=st.session_state.business_model_precedent, bm_precedent="",nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of="" ,generation=1)

                    
                    except ValueError as e:
                        results_first_part[section_name] = f"Erreur: {str(e)}"
                    
                    combined_content += " " + results_first_part[section_name]
                    placeholders[section_name].markdown(f"\n\n{results_first_part[section_name]}")
                    
            st.session_state.markdown_content1= generate_markdown(results_first_part)
           
            
            
            # Génération de la seconde partie
            for section_name in second_part:
            
                with st.spinner(f"Génération de {section_name}..."):
                    system_message = system_messages[section_name]
                    query = queries[section_name]
                    full_contents = combined_content + " " + query+ "" + system_message+ json.dumps(final_text)+ full_text+ json.dumps(st.session_state["produits_data"])+json.dumps(serialize_membres(st.session_state["membres"]))
                    tokens_in_input = count_tokens(full_contents , MODEL) 
                    tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
                    
                    # Vérifier si l'utilisateur a assez de tokens
                    if tokens_remaining < tokens_needed:
                        st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
                        return
                    
                    try:
                        results_second_part[section_name] = generate_section(system_message, query, documents, combined_content+"Voici la liste de membre de l'entreprise:"+json.dumps(serialize_membres(st.session_state["membres"])), final_text, rubriques="", business_model=st.session_state.business_model_precedent,bm_precedent="",nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of="", generation=1)
                    except ValueError as e:
                        results_second_part[section_name] = f"Erreur: {str(e)}"
                    
                    combined_content += " " + results_second_part[section_name]  
                    placeholders[section_name].markdown(f"\n\n{results_second_part[section_name]}")

            st.session_state.markdown_content2= generate_markdown(results_second_part)
            
            
            
           
        if st.session_state.get('business_plan_precedent'):
            # Génération du Business Plan et téléchargement des fichiers 
            if st.button("Ameliorer business plan"):
                html_content = chain_of_thougtht("BP",st.session_state.montant_projet,st.session_state.nom_entreprise, st.session_state.business_model_precedent , st.session_state["produits_data"], Analisis=final_text, previousbp=st.session_state.business_plan_precedent, generation=2)
                st.write(html_content)
                if html_content:
                    for section_name in first_part:
                        with st.spinner(f"Génération de {section_name}..."):
                            system_message = system_messages[section_name]
                            query = queries[section_name]
                            full_contents = combined_content + " " + query+ "" + system_message+ json.dumps(final_text)+ full_text + json.dumps(st.session_state["produits_data"])+json.dumps(serialize_membres(st.session_state["membres"]))+html_content
                            tokens_in_input = count_tokens(full_contents , MODEL) 
                            tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
                            
                            # Vérifier si l'utilisateur a assez de tokens
                            if tokens_remaining < tokens_needed:
                                st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
                                return
                            
                            try:
                                # Vérifier si la section est "Couverture" ou "Sommaire"
                                if section_name in ["Couverture", "Sommaire"]:
                                    results_first_part[section_name] = generate_section(system_message, query, documents, combined_content+"Voici la liste de membre de l'entreprise:"+json.dumps(serialize_membres(st.session_state["membres"])), "", rubriques="", business_model="",bm_precedent="",nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of=html_content, generation=2)
                                else:
                                    results_first_part[section_name] = generate_section(system_message, query, documents, combined_content+"Voici la liste de membre de l'entreprise:"+json.dumps(serialize_membres(st.session_state["membres"])), final_text, rubriques=st.session_state["produits_data"], business_model=st.session_state.business_model_precedent,bm_precedent=st.session_state.business_plan_precedent, nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of=html_content, generation=2)
                            except ValueError as e:
                                results_first_part[section_name] = f"Erreur: {str(e)}"
                            
                            combined_content += " " + results_first_part[section_name]
                            placeholders[section_name].markdown(f"\n\n{results_first_part[section_name]}")

                    st.session_state.markdown_content1= generate_markdown(results_first_part)
                    
                    
                    # Génération de la seconde partie
                    
                    for section_name in second_part:
                        with st.spinner(f"Génération de {section_name}..."):
                            system_message = system_messages[section_name]
                            query = queries[section_name]
                            full_contents = combined_content + " " + query+ "" + system_message+ json.dumps(final_text)+ full_text+json.dumps(st.session_state["produits_data"])+json.dumps(serialize_membres(st.session_state["membres"]))
                            tokens_in_input = count_tokens(full_contents  , MODEL) 
                            tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
                            
                            # Vérifier si l'utilisateur a assez de tokens
                            if tokens_remaining < tokens_needed:
                                st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
                                return
                            
                            try:
                                results_second_part[section_name] = generate_section(system_message, query, documents, combined_content+"Voici la liste de membre de l'entreprise:" +json.dumps(serialize_membres(st.session_state["membres"])), final_text, rubriques=st.session_state["produits_data"], business_model=st.session_state.business_model_precedent,bm_precedent=st.session_state.business_plan_precedent, nom_entreprise=st.session_state.nom_entreprise,precedent_chain_of=html_content,generation=2)
                            except ValueError as e:
                                results_second_part[section_name] = f"Erreur: {str(e)}"
                            
                            combined_content += " " + results_second_part[section_name]
                            placeholders[section_name].markdown(f"\n\n{results_second_part[section_name]}")
                    
                    st.session_state.markdown_content2 = generate_markdown(results_second_part)



                

        if st.button("Creer un fichier à telecharger") or st.session_state.get('business_plan_markdown_content1') or st.session_state.get('business_plan_markdown_content2'):
            
            pdf = MarkdownPdf(toc_level=2)
            pdf.add_section(Section(st.session_state.markdown_content2))
            pdf.meta["title"] = "Business Plan" 
            pdf_file_path = "business_plan.pdf"
            pdf.save(pdf_file_path)
            
            
            
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement


            from docx import Document
            import re

            def markdown_to_word_via_text(markdown_contents, doc):
                doc.add_heading('Business Plan', 0)

                # Diviser le contenu en lignes
                lines = markdown_contents.split('\n')
                table_data = []
                inside_table = False

                for line in lines:
                    line = line.strip()
                    if not line:
                        # Si ligne vide et données de table en cours, ajouter le tableau au document
                        if table_data:
                            add_table_with_borders(doc, table_data)
                            table_data = []
                            inside_table = False
                        continue

                    if line.startswith('# '):  # Titre niveau 1
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):  # Titre niveau 2
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):  # Titre niveau 3
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('#### '):  # Titre niveau 4
                        doc.add_heading(line[5:], level=4)
                    elif re.match(r'^\d+\.\s', line):  # Liste numérotée
                        # Vérifier s'il y a du texte en gras dans la liste numérotée
                        match = re.match(r'^(\d+\.\s)(\*\*.+?\*\*)', line)
                        if match:
                            paragraph = doc.add_paragraph(style='List Number')
                            paragraph.add_run(match.group(1))  # Numéro
                            bold_run = paragraph.add_run(match.group(2)[2:-2])  # Texte en gras sans `**`
                            bold_run.bold = True
                        else:
                            doc.add_paragraph(line, style='List Number')
                    elif line.startswith('- ') or line.startswith('•'):  # Liste à puces
                        match = re.match(r'^(•|-)\s\*\*(.+?)\*\*(.*)', line)
                        if match:
                            paragraph = doc.add_paragraph(style='List Bullet')
                            bold_run = paragraph.add_run(match.group(2))  # Texte en gras
                            bold_run.bold = True
                            if match.group(3):  # Texte après le gras
                                paragraph.add_run(match.group(3).strip())
                        else:
                            doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('|'):  # Détection des lignes de tableau
                        if re.match(r'\|?\s*[-:]+\s*\|', line):
                            inside_table = True
                            continue  # Ignorer les lignes de séparation
                        else:
                            inside_table = True
                            table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
                    elif re.match(r'^\*\*.+?\*\*\s*:', line):  # Texte en gras suivi de texte normal
                        match = re.match(r'^\*\*(.+?)\*\*\s*:(.*)', line)
                        if match:
                            paragraph = doc.add_paragraph()
                            bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                            bold_run.bold = True
                            if match.group(2):  # Texte normal après le `:`
                                paragraph.add_run(f":{match.group(2)}")
                    elif re.match(r'^\*\*.+?\*\*$', line):  # Texte entièrement en gras
                        paragraph = doc.add_paragraph()
                        bold_run = paragraph.add_run(line[2:-2])  # Texte sans `**`
                        bold_run.bold = True
                    elif re.match(r'^\*\*.+?\*\*\s[\d.,]+\s?[$$%]$', line):  # Nombres avec symboles monétaires
                        match = re.match(r'^\*\*(.+?)\*\*\s([\d.,]+\s?[$$%])$', line)
                        if match:
                            paragraph = doc.add_paragraph()
                            bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                            bold_run.bold = True
                            paragraph.add_run(f" {match.group(2)}")  # Montant avec symbole
                    elif not inside_table:  # Paragraphe normal
                        doc.add_paragraph(line)

                # Traiter les données de table restantes
                if table_data:
                    add_table_with_borders(doc, table_data)

            def add_table_with_borders(doc, table_data):
                """
                Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
                """
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'  # Appliquer un style de tableau avec bordures

                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        cell_content = table.cell(i, j).paragraphs[0]
                        parts = re.split(r'(\*\*.+?\*\*)', cell)  # Diviser par texte en gras
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):  # Texte en gras
                                run = cell_content.add_run(part[2:-2])
                                run.bold = True
                            else:  # Texte normal
                                cell_content.add_run(part.strip())




            # Ajouter la création et le téléchargement du fichier Word
            
            doc = Document()
            markdown_to_word_via_text(st.session_state.markdown_content1, doc)
            #st.write(st.session_state.markdown_content2)
            produits_data = st.session_state.get("produits_data", [])
                # Boucler à travers chaque produit
            for produit in produits_data:
                nom_produit = produit.get("nom_produit", "Produit")
                business_models = produit.get("business_models", [])
                
                for bm in business_models:
                    description_html = bm.get("description", "")
            
                    generer_docx_business_model(
                        nom_entreprise=st.session_state.nom_entreprise,
                        date_bmc=datetime.date.today(),
                        contenu_business_model=description_html,
                        nom_produit=nom_produit, 
                        doc=doc,
                        value=2)
 

            markdown_to_word_via_text(st.session_state.markdown_content2, doc)
           # st.write(st.session_state.markdown_content2)
            
            st.session_state.business_plan_precedent = st.session_state.markdown_content1+""+ st.session_state.markdown_content2
            #st.write(st.session_state.business_plan_precedent)
            
            # doc.add_paragraph(content_result)
            
            # Vérifier et ajouter le contenu
        

            # Ajouter les sections du Business Plan


            # Récupérer les données des tableaux depuis la session Streamlit
            # Récupérer les données des tableaux depuis la session Streamlit
            export_data_investissements = st.session_state.get('export_data_investissements', {})
            export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
            export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
            export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
            export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
            export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
            export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
            export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
            

            # Ajouter une section pour les tableaux
            doc.add_heading('Résumé des Données Financières', level=1)

            # Fonction pour ajouter un tableau dans le document Word
            def ajouter_tableau(donnees, headers, titre):
                """
                Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
                """
                doc.add_heading(titre, level=2)
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light List Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Ajouter les en-têtes
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Ajouter les données des tableaux
                for row in donnees:
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        cell_value = row.get(header, "")
                        cell_text = str(cell_value)  # Convertir en chaîne de caractères
                        row_cells[i].text = cell_text
                        row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                # Ajouter une note
                doc.add_paragraph()
                doc.add_paragraph("Les résultats sont calculés selon les données fournies.")

            # Ajouter les différents tableaux
            if export_data_investissements.get("table_data"):
                ajouter_tableau(export_data_investissements["table_data"], ["Investissements", "Taux (%)", "Durée (mois)", "Montant ($)"], "Investissements et Financements")
            if export_data_salaires.get("table_data"):
                ajouter_tableau(export_data_salaires["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Salaires et Charges Sociales")
            if export_data_amortissements.get("amortissements"):
                ajouter_tableau(export_data_amortissements["amortissements"], ["Amortissement", "Année 1", "Année 2", "Année 3"], "Détail des Amortissements")
            if export_data_compte.get("table_data"):
                ajouter_tableau(export_data_compte["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Compte de Résultats Prévisionnel")
            if export_data_soldes.get("table_data"):
                ajouter_tableau(export_data_soldes["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Soldes Intermédiaires de Gestion")
            if export_data_capacite.get("table_data"):
                ajouter_tableau(export_data_capacite["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Capacité d'Autofinancement")
            if export_data_seuil.get("table_data"):
                ajouter_tableau(export_data_seuil["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Seuil de Rentabilité Économique")
            if export_data_bfr.get("table_data"):
                ajouter_tableau(export_data_bfr["table_data"], ["Analyse clients / fournisseurs", "Délai jours", "Année 1", "Année 2", "Année 3"], "Besoin en Fonds de Roulement")

            # **Nouvelles sections ajoutées :**

            # Ajouter la section Plan de Financement à Trois Ans
            export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
            if export_data_plan_financement.get("table_data"):
                ajouter_tableau(
                    export_data_plan_financement["table_data"],
                    ["Description", "Année 1", "Année 2", "Année 3"],
                    "Plan de Financement à Trois Ans"
                )

            # Ajouter la section Budget Prévisionnel de Trésorerie Partie 1
            export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
            if export_data_budget_part1.get("table_data"):
                ajouter_tableau(
                    export_data_budget_part1["table_data"],
                    ["Description", "Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5", "TOTAL"],
                    "Budget Prévisionnel de Trésorerie - Partie 1"
                )

            # Ajouter la section Budget Prévisionnel de Trésorerie Partie 2
            export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
            if export_data_budget_part2.get("table_data"):
                ajouter_tableau(
                    export_data_budget_part2["table_data"],
                    ["Description", "Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"],
                    "Budget Prévisionnel de Trésorerie - Partie 2"
                )

            # Enregistrer le document dans un buffer
            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            # Télécharger les fichiers générés
            st.success("Le PDF et le document Word ont été générés avec succès.")
            #with open(pdf_file_path, "rb") as f:
                #st.download_button("Téléchargez le PDF", f, file_name="business_plan.pdf", mime="application/pdf")

            st.download_button("Téléchargez le document Word", word_buffer, file_name="business_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # Sélection du type d'entreprise et nom





# Initialiser les variables dans la session si ce n'est pas déjà fait
if 'business_models' not in st.session_state:
    st.session_state.business_models = {}
    
if 'personas' not in st.session_state:
    st.session_state.personas = {}
    
if 'problem_trees' not in st.session_state:
    st.session_state.problem_trees = {}
    
if 'analyse_marches' not in st.session_state:
    st.session_state.analyse_marches = {}
    
if 'facteurs_limitants' not in st.session_state:
    st.session_state.facteurs_limitants = {}
    
if 'concurrences' not in st.session_state:
    st.session_state.concurrences = {}

if 'date_bmc_generate' not in st.session_state:
    st.session_state.date_bmc_generate = datetime.date.today()

# --- Variable interne nb_products
if "nb_products" not in st.session_state:
    st.session_state["nb_products"] = 1
# --- La liste "produits_data" stockant tous les produits
if "produits_data" not in st.session_state:
    st.session_state["produits_data"] = []
    
if 'selected_idx_produit' not in st.session_state:
    st.session_state['selected_idx_produit'] = 0  # Valeur par défaut (premier produit)
if 'previous_selected_idx_produit' not in st.session_state:
    st.session_state['previous_selected_idx_produit'] = None



# Initialiser les variables dans la session si ce n'est pas déjà fait
    # Initialiser 'business_model_precedent' comme un dictionnaire vide si non présent
if 'business_model_precedent' not in st.session_state:
    st.session_state['business_model_precedent'] = {}

# Initialiser les variables dans la session si ce n'est pas déjà fait
if 'rubriques_initiales' not in st.session_state:
    st.session_state.rubriques_initiales = ""
    
# Initialiser les variables dans la session si ce n'est pas déjà fait
if 'markdown_content2' not in st.session_state:
    st.session_state.markdown_content2 = ""
    
# Initialiser les variables dans la session si ce n'est pas déjà fait
if 'markdown_content1' not in st.session_state:
    st.session_state.markdown_content1 = ""
    
    
    
    
    
# Initialisation des variables dans session_state
if 'type_entreprise' not in st.session_state:
    st.session_state.type_entreprise = "PME"  # Valeur par défaut

if 'montant_projet' not in st.session_state:
    st.session_state.montant_projet = ""

if 'nom_entreprise' not in st.session_state:
    st.session_state.nom_entreprise = ""
 

# Noms des nouveaux onglets du Business Model Canvas
business_model_tab_names = [
    "Collecte des Données",
    "Générer Business Model",

]

# Fonctions correspondantes pour les nouveaux onglets
business_model_sections = [
    page_collecte_donnees,
    page_generer_business_model,
   
]

# Liste des noms d'onglets existants

# Mise à jour des noms d'onglets
tab_names = [
    "Informations Générales", "Besoins de Démarrage", "Financement",
    "Charges Fixes", "Chiffre d'Affaires", "Charges Variables",
    "Fonds de Roulement", "Salaires", "Rentabilité", "Trésorerie","CT","Equipes","Génération du Business Plan",
    "Investissements et Financements", "Salaires et Charges Sociales", "Détail des Amortissements",
    "Compte de Résultats Prévisionnel", "Soldes Intermédiaires de Gestion",
    "Capacité d'Autofinancement", "Seuil de Rentabilité Économique",
    "Besoin en Fonds de Roulement", "Plan de Financement sur 3 Ans",
    "Budget Prévisionnel de Trésorerie"
]

# Mise à jour de la liste des fonctions correspondantes
sections = [
    page_informations_generales, page_besoins_demarrage, page_financement,
    page_charges_fixes, page_chiffre_affaires, page_charges_variables,
    page_fonds_roulement, page_salaires, page_rentabilite, page_tresorerie, 
    ct_model, ajouter_informations_personnel,
    page_generation_business_plan, page_investissements_et_financements,
    page_salaires_charges_sociales, page_detail_amortissements,
    page_compte_resultats_previsionnel, page_soldes_intermediaires_de_gestion,
    page_capacite_autofinancement, page_seuil_rentabilite_economique,
    page_besoin_fonds_roulement, page_plan_financement_trois_ans,
    page_budget_previsionnel_tresorerie
]

# Trouver l'index de "Génération du Business Plan"
try:
    index_generation_bp = tab_names.index("Génération du Business Plan")
except ValueError:
    st.error("L'onglet 'Génération du Business Plan' n'a pas été trouvé dans la liste des onglets.")
    index_generation_bp = len(tab_names)  # Ajouter à la fin si non trouvé

# Insérer les nouveaux onglets avant "Génération du Business Plan"
tab_names =business_model_tab_names + tab_names
sections = business_model_sections + sections


# Fonction de rafraîchissement du Business Plan
def refresh_business_plan():
    """
    Réinitialise toutes les variables de session liées au Business Plan.
    """
    keys_to_reset = [
        'business_plan_markdown_content1',
        'business_plan_markdown_content2',
        'business_plan_precedent',
        'markdown_content1',
        'markdown_content2',
        'business_models',
        'personas',
        'problem_trees',
        'analyse_marches',
        'facteurs_limitants',
        'concurrences',
        'type_entreprise',
        'montant_projet',
        'nom_entreprise',
        'nb_products',
        'produits_data',
        'selected_idx_produit',
        'previous_selected_idx_produit',
        'rubriques_initiales',
        'business_model_precedent',
        'data',
        'tables'
    ]
    
    for key in keys_to_reset:
        if key in st.session_state:
            del st.session_state[key]
    
    # Réinitialiser 'data' et 'tables' à des valeurs par défaut
    st.session_state['data'] = {}
    st.session_state['tables'] = {}
    
    st.success("Le Business Plan a été réinitialisé avec succès.")



# ---------------------- INITIALISATION SESSION ----------------------
if "login_attempts" not in st.session_state:
    st.session_state["login_attempts"] = {}  # {email: nb_tentatives}
if "locked_until" not in st.session_state:
    st.session_state["locked_until"] = {}    # {email: timestamp}
    
def initialize_session():
    """Initialise les variables de session nécessaires."""
    if 'jwt_token' not in st.session_state:
        st.session_state['jwt_token'] = None
    if 'otp_required' not in st.session_state:
        st.session_state['otp_required'] = False
    if 'authenticated' not in st.session_state:
        st.session_state['authenticated'] = False
    if 'user_info' not in st.session_state:
        st.session_state['user_info'] = None
    if 'page' not in st.session_state:
        st.session_state['page'] = 'Connexion'

 
 
 # ---------------------- FONCTIONS UTILITAIRES ----------------------
def format_number(num):
    """
    Formate un nombre en utilisant des suffixes K, M, B, T, Q pour des milliers, millions, milliards,
    trillions, quadrillions respectivement. Pour les nombres supérieurs à un quadrillion, utilise la notation scientifique.

    Args:
        num (int ou float): Le nombre à formater.

    Returns:
        str: Le nombre formaté avec le suffixe approprié ou en notation scientifique.
    """
    try:
        num = float(num)
    except (ValueError, TypeError):
        return str(num)

    abs_num = abs(num)  # Gérer les nombres négatifs

    if abs_num >= 1_000_000_000_000_000:
        # Notation scientifique pour les nombres >= 1 quadrillion
        return f"{num:.2e}"
    elif abs_num >= 1_000_000_000_000:
        return f"{num / 1_000_000_000_000:.1f}Q"  # Quadrillion
    elif abs_num >= 1_000_000_000:
        return f"{num / 1_000_000_000:.1f}B"  # Billion (milliard)
    elif abs_num >= 1_000_000:
        return f"{num / 1_000_000:.1f}M"  # Million
    elif abs_num >= 1_000:
        return f"{num / 1_000:.1f}K"  # Thousand (mille)
    else:
        return f"{int(num)}"
 
 
 


# ---------------------- FONCTIONS JWT ----------------------
def create_jwt_token(user_id, is_admin=False):
    """
    Crée un token JWT avec un champ is_admin, un champ user_id,
    valide pendant JWT_EXP_DELTA_SECONDS.
    """
    payload = {
        "user_id": str(user_id),
        "is_admin": is_admin,
        "exp": datetime.datetime.utcnow() + datetime.timedelta(seconds=JWT_EXP_DELTA_SECONDS),
    }
    token = jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
    return token

def decode_jwt_token(token):
    """Décode un token JWT et renvoie le payload ou None si invalide."""
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        return payload
    except InvalidTokenError as e:
        logging.error(f"Erreur JWT : {e}")
        return None

def get_current_user():
    """
    Récupère l'utilisateur courant à partir des données gérées par streamlit-authenticator.
    Retourne l'objet utilisateur (depuis MongoDB) ou None si pas connecté.
    """
    # Récupère le nom d'utilisateur depuis session_state, géré par streamlit-authenticator
    username = st.session_state.get("username", None)
    if not username:
        return None

    # Requête MongoDB pour obtenir les informations de l'utilisateur
    user = users_collection.find_one({"email": username})
    return user



# ---------------------- FONCTION UTILITAIRE ----------------------
def get_client_ip():
    """
    Dans un déploiement Streamlit, la récupération de l’IP peut être non trivial.
    Ici on renvoie une IP fictive (ou 127.0.0.1).
    En production, adapter via des headers HTTP ou via un reverse proxy.
    """
    return "127.0.0.1"

def envoyer_email(to_email, subject, content):
    message = MIMEMultipart()
    message['From'] = EMAIL_ADDRESS
    message['To'] = to_email
    message['Subject'] = subject
    message.attach(MIMEText(content, 'html'))
    
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.send_message(message)
        server.quit()
        logging.info(f"Email envoyé à {to_email} pour le sujet: {subject}")
        return True
    except Exception as e:
        logging.error(f"Erreur lors de l'envoi de l'email à {to_email} : {e}")
        st.error(f"Erreur lors de l'envoi de l'email: {e}")
        return False
    
    
def envoyer_email_piecejointes(to_email, subject, content, attachments=None):
    """
    Envoie un email avec des pièces jointes (seules les images sont autorisées).

    Args:
        to_email (str): Adresse email du destinataire.
        subject (str): Sujet de l'email.
        content (str): Contenu de l'email (HTML).
        attachments (list of UploadedFile or str): Liste des fichiers à attacher.

    Returns:
        bool: True si l'email a été envoyé avec succès, False sinon.
    """
    try:
        message = MIMEMultipart()
        message['From'] = EMAIL_ADDRESS
        message['To'] = to_email
        message['Subject'] = subject
        message.attach(MIMEText(content, 'html'))
        logging.info(f"Création de l'email avec sujet '{subject}' à '{to_email}'.")

        # Traitement des pièces jointes
        if attachments:
            logging.info(f"Nombre de pièces jointes à traiter: {len(attachments)}")
            for attachment in attachments:
                try:
                    # Déterminer le nom et l'extension
                    if isinstance(attachment, str):
                        filename = os.path.basename(attachment)
                        ext = filename.split('.')[-1].lower()
                        logging.debug(f"Pièce jointe (str): {filename}")
                    else:
                        filename = attachment.name
                        ext = filename.split('.')[-1].lower()
                        logging.debug(f"Pièce jointe (UploadedFile): {filename}")

                    # Vérifier l'extension
                    if ext not in ALLOWED_EXTENSIONS:
                        error_msg = f"Le type de fichier '{ext}' pour '{filename}' n'est pas autorisé. Seules les images sont acceptées."
                        st.error(error_msg)
                        logging.error(error_msg)
                        return False

                    # Vérifier la taille
                    if isinstance(attachment, str):
                        if not os.path.exists(attachment):
                            error_msg = f"Le fichier '{attachment}' n'existe pas."
                            st.error(error_msg)
                            logging.error(error_msg)
                            return False
                        file_size = os.path.getsize(attachment)
                        logging.debug(f"Taille du fichier '{filename}': {file_size} bytes")
                    else:
                        # Si 'size' est disponible
                        if hasattr(attachment, 'size'):
                            file_size = attachment.size
                            logging.debug(f"Taille du fichier '{filename}': {file_size} bytes")
                        else:
                            # Lire le contenu pour obtenir la taille
                            file_content = attachment.read()
                            file_size = len(file_content)
                            attachment.seek(0)  # Réinitialiser le pointeur du fichier après lecture
                            logging.debug(f"Taille du fichier '{filename}': {file_size} bytes")

                    if file_size > MAX_ATTACHMENT_SIZE:
                        error_msg = f"Le fichier '{filename}' dépasse la taille maximale autorisée de 10 Mo."
                        st.error(error_msg)
                        logging.error(error_msg)
                        return False

                    # Création de la pièce jointe
                    if isinstance(attachment, str):
                        with open(attachment, "rb") as f:
                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload(f.read())
                    else:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())

                    # Encodage en Base64
                    encoders.encode_base64(part)

                    # Détermination du type MIME
                    mime_type, _ = mimetypes.guess_type(filename)
                    if mime_type:
                        maintype, subtype = mime_type.split('/', 1)
                        part.set_type(mime_type)
                        logging.debug(f"Type MIME pour '{filename}': {mime_type}")
                    else:
                        maintype, subtype = 'application', 'octet-stream'
                        part.set_type('application/octet-stream')
                        logging.debug(f"Type MIME par défaut pour '{filename}': application/octet-stream")

                    part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
                    message.attach(part)
                    logging.info(f"Pièce jointe ajoutée: {filename}")

                except Exception as e:
                    logging.error(f"Erreur lors de l'ajout de la pièce jointe '{filename}' : {e}")
                    st.error(f"Erreur lors de l'ajout de la pièce jointe '{filename}' : {e}")
                    return False

        # Envoi de l'email
        try:
            server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            server.send_message(message)
            server.quit()
            logging.info(f"Email envoyé à {to_email} pour le sujet: {subject}")
            return True
        except Exception as e:
            logging.error(f"Erreur lors de l'envoi de l'email à {to_email} : {e}")
            st.error(f"Erreur lors de l'envoi de l'email: {e}")
            return False

    except Exception as e:
        logging.error(f"Erreur générale dans envoyer_email_piecejointes : {e}")
        st.error(f"Erreur générale lors de l'envoi de l'email : {e}")
        return False


# ---------------------- FONCTIONS GESTION TENTATIVES ----------------------
def increment_login_attempts(email):
    """
    Incrémente le compteur de tentatives de connexion pour email.
    Gère le blocage si >5 tentatives.
    """
    attempts = st.session_state["login_attempts"].get(email, 0) + 1
    st.session_state["login_attempts"][email] = attempts
    logging.warning(f"Tentative échouée {attempts} pour {email}")

    # Au bout de 5 tentatives -> pause 3 secondes
    if attempts == 5:
        st.warning("Vous avez atteint 5 tentatives de connexion échouées. Attendez 3 secondes avant de réessayer.")
        time.sleep(3)
    elif attempts > 5:
        # Bloque 30 secondes
        st.warning("Trop de tentatives ! Vous êtes bloqué pendant 30 secondes.")
        st.session_state["locked_until"][email] = time.time() + 30

        # Envoi email à l'admin
        ip_address = get_client_ip()
        subject = "[Alerte Système] Trop de tentatives de connexion"
        content = f"""
        <h2>Tentatives de Connexion Suspicious</h2>
        <p>L'utilisateur avec l'email <b>{email}</b> vient de dépasser 5 tentatives de connexion échouées.</p>
        <p>Adresse IP suspecte : <b>{ip_address}</b></p>
        <p>Veuillez prendre les mesures nécessaires (blocage IP, etc.).</p>
        """
        envoyer_email(ADMIN_EMAIL, subject, content)

def is_locked(email):
    """
    Retourne True si l'email est encore bloqué, False sinon.
    """
    if email not in st.session_state["locked_until"]:
        return False
    locked_until_ts = st.session_state["locked_until"][email]
    return time.time() < locked_until_ts


def ajouter_utilisateur(email, nom, mot_de_passe, role="user"):
    """
    Ajoute un utilisateur dans la base avec un rôle ('user', 'admin', ou 'super_admin').
    Par défaut, le rôle est 'user'.
    """
    try:
        # Vérifier si l'email existe déjà
        if users_collection.find_one({"email": email}):
            logging.warning(f"Tentative d'inscription avec un email existant : {email}")
            return False, "L'utilisateur existe déjà."

        # Préparer les identifiants pour le hachage
        credentials = {"usernames": {email: {"password": mot_de_passe}}}
        hashed_credentials = stauth.Hasher.hash_passwords(credentials)

        # Extraire le mot de passe haché
        hashed_password = hashed_credentials["usernames"][email]["password"]

        utilisateur = {
            "email": email,
            "nom": nom,
            "mot_de_passe": hashed_password,
            "tokens_purchased": 100,  # Par exemple : tokens par défaut à l'inscription
            "tokens_consumed": 0,
            "is_blocked": False,
            "created_at": datetime.datetime.utcnow(),
            "role": role,
            "otp_validated": False  # Initialiser à False
        }

        # Insérer l'utilisateur dans la collection MongoDB
        users_collection.insert_one(utilisateur)

        # Envoyer un email de bienvenue (optionnel)
        contenu = f"""
        <h1>Bienvenue {nom}!</h1>
        <p>Merci de vous être inscrit sur notre application.</p>
        <p>Votre rôle est : {role}</p>
        <p>Vous avez reçu 100 000 tokens par pour l'essaie.</p>
        """
        envoyer_email(email, "Bienvenue sur notre application", contenu)

        logging.info(f"Nouvel utilisateur ajouté : {email} avec le rôle {role}")
        return True, "Utilisateur ajouté avec succès."

    except Exception as e:
        logging.error(f"Erreur lors de l'ajout de l'utilisateur {email} : {e}")
        return False, "Une erreur est survenue lors de l'ajout de l'utilisateur."


def verifier_utilisateur(email, mot_de_passe):
    """
    Vérifie le couple email / mot_de_passe manuellement.
    Gère is_blocked, tentatives de connexion, etc.
    Compatible avec un champ 'mot_de_passe' haché au format '$2b$...' stocké en base.
    """
    try:
        if is_locked(email):
            st.error("Vous êtes temporairement bloqué. Veuillez patienter avant de réessayer.")
            return False, None

        utilisateur = users_collection.find_one({"email": email})
        if not utilisateur:
            increment_login_attempts(email)
            return False, None

        if utilisateur.get('is_blocked', False):
            logging.warning(f"Tentative de connexion d'un utilisateur bloqué : {email}")
            st.error("Votre compte a été bloqué. Veuillez contacter l'administrateur.")
            return False, None

        hashed_pw = utilisateur.get('mot_de_passe', "")
        if not hashed_pw:
            # Le compte n'a pas de mot de passe (peut-être Google)
            increment_login_attempts(email)
            return False, None
        
        # bcrypt.checkpw attend des bytes, donc on encode
        if bcrypt.checkpw(mot_de_passe.encode('utf-8'), hashed_pw.encode('utf-8')):
            # succès -> reset attempts
            st.session_state["login_attempts"][email] = 0
            return True, utilisateur
        else:
            increment_login_attempts(email)
            return False, None

    except Exception as e:
        logging.error(f"Erreur lors de la vérification de l'utilisateur {email} : {e}")
        st.error(f"Erreur lors de la vérification de l'utilisateur : {e}")
        return False, None

def reinitialiser_mot_de_passe(email, nouveau_mot_de_passe):
    try:
        utilisateur = users_collection.find_one({"email": email})
        if utilisateur:
            hashed_password = bcrypt.hashpw(nouveau_mot_de_passe.encode('utf-8'), bcrypt.gensalt())
            users_collection.update_one(
                {"email": email},
                {"$set": {"mot_de_passe": hashed_password}}
            )
            # Envoyer un email de confirmation
            contenu = f"""
            <h1>Mot de passe réinitialisé</h1>
            <p>Votre mot de passe a été réinitialisé avec succès.</p>
            """
            envoyer_email(email, "Mot de passe réinitialisé", contenu)
            logging.info(f"Mot de passe réinitialisé pour {email}")
            return True, "Mot de passe réinitialisé avec succès."
        return False, "Utilisateur non trouvé."
    except Exception as e:
        logging.error(f"Erreur lors de la réinitialisation du mot de passe pour {email} : {e}")
        st.error(f"Erreur lors de la réinitialisation du mot de passe : {e}")
        return False, "Une erreur est survenue lors de la réinitialisation du mot de passe."

# ---------------------- FONCTIONS PAIEMENT (EXTRAITES DU CODE ORIGINAL) ----------------------
def creer_session_stripe(customer_id, amount_cents, success_url, cancel_url):
    try:
        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            customer=customer_id,
            line_items=[{
                'price_data': {
                    'currency': 'usd',
                    'product_data': {
                        'name': 'Achat de Tokens',
                    },
                    'unit_amount': amount_cents,
                },
                'quantity': 10000,
            }],
            mode='payment',
            success_url=success_url,
            cancel_url=cancel_url,
        )
        return session.url
    except Exception as e:
        st.error(f"Erreur lors de la création de la session Stripe: {e}")
        return None

def enregistrer_paiement(user_id, amount, tokens, payment_method, uploaded_file_url=None, status="Pending"):
    try:
        paiement = {
            "user_id": user_id,
            "amount": amount,
            "tokens_purchased": tokens,
            "payment_method": payment_method,
            "status": status,
            "timestamp": datetime.datetime.utcnow(),
            "uploaded_file_url": uploaded_file_url
        }
        db['payments'].insert_one(paiement)
        logging.info(f"Paiement enregistré pour l'utilisateur {user_id}, montant: {amount} USD, tokens: {tokens}, méthode: {payment_method}")
        return True, "Paiement enregistré avec succès."
    except Exception as e:
        logging.error(f"Erreur lors de l'enregistrement du paiement pour l'utilisateur {user_id} : {e}")
        st.error(f"Erreur lors de l'enregistrement du paiement : {e}")
        return False, "Erreur lors de l'enregistrement du paiement."

def recuperer_historique_paiements(user_id):
    try:
        paiements = list(db['payments'].find({"user_id": user_id}).sort("timestamp", -1))
        return paiements
    except Exception as e:
        logging.error(f"Erreur lors de la récupération de l'historique des paiements pour l'utilisateur {user_id} : {e}")
        st.error(f"Erreur lors de la récupération des paiements : {e}")
        return []

def recuperer_dernier_paiement(user_id):
    paiements = recuperer_historique_paiements(user_id)
    if paiements:
        # Trier les paiements par timestamp décroissant
        paiements_sorted = sorted(paiements, key=lambda x: x['timestamp'], reverse=True)
        return paiements_sorted[0]
    return None

# ---------------------- FONCTIONS TOKENS & CONSOMMATION ----------------------
def consommer_tokens(utilisateur, tokens_utilises):
    try:
        tokens_purchased = utilisateur.get('tokens_purchased', 0)
        tokens_consumed = utilisateur.get('tokens_consumed', 0)
        tokens_remaining = tokens_purchased - tokens_consumed

        if tokens_remaining >= tokens_utilises:
            users_collection.update_one(
                {"_id": utilisateur['_id']},
                {"$inc": {"tokens_consumed": tokens_utilises}}
            )
            st.session_state['user_info']['tokens_consumed'] += tokens_utilises
            # Enregistrer la consommation dans la base
            enregistrer_consomation(utilisateur['_id'], tokens_utilises)
            logging.info(f"{tokens_utilises} tokens consommés par l'utilisateur {utilisateur['email']}")
            return True, "Tokens consommés avec succès."
        else:
            logging.warning(f"Utilisateur {utilisateur['email']} a tenté de consommer {tokens_utilises} tokens mais n'en a que {tokens_remaining}")
            st.error("Vous n'avez pas assez de tokens.")
            return False, "Tokens insuffisants."
    except Exception as e:
        logging.error(f"Erreur lors de la consommation de tokens pour l'utilisateur {utilisateur['email']} : {e}")
        st.error(f"Erreur lors de la consommation de tokens : {e}")
        return False, "Une erreur est survenue."

def enregistrer_consomation(user_id, tokens_utilises):
    try:
        consommation = {
            "user_id": user_id,
            "date": datetime.datetime.utcnow(),
            "tokens_consumed": tokens_utilises
        }
        consumption_collection.insert_one(consommation)
        logging.info(f"Consommation de {tokens_utilises} tokens enregistrée pour l'utilisateur {user_id}")
        return True
    except Exception as e:
        logging.error(f"Erreur lors de l'enregistrement de la consommation pour l'utilisateur {user_id} : {e}")
        st.error(f"Erreur lors de l'enregistrement de la consommation : {e}")
        return False

def recuperer_consomation(user_id, periode):
    try:
        aujourd_hui = datetime.datetime.utcnow()
        if periode == 'Jour':
            debut = aujourd_hui - datetime.timedelta(days=1)
            date_format = "%Y-%m-%d"
        elif periode == '2 Jours':
            debut = aujourd_hui - datetime.timedelta(days=2)
            date_format = "%Y-%m-%d"
        elif periode == 'Semaine':
            debut = aujourd_hui - datetime.timedelta(weeks=1)
            date_format = "%Y-%m-%d"
        elif periode == 'Mois':
            debut = aujourd_hui - datetime.timedelta(days=30)
            date_format = "%Y-%m"
        elif periode == '3 Mois':
            debut = aujourd_hui - datetime.timedelta(days=90)
            date_format = "%Y-%m"
        elif periode == '6 Mois':
            debut = aujourd_hui - datetime.timedelta(days=180)
            date_format = "%Y-%m"
        elif periode == 'Année':
            debut = aujourd_hui - datetime.timedelta(days=365)
            date_format = "%Y"
        else:
            debut = aujourd_hui - datetime.timedelta(days=30)  # par défaut
            date_format = "%Y-%m"

        pipeline = [
            {"$match": {
                "user_id": user_id,
                "date": {"$gte": debut, "$lte": aujourd_hui}
            }},
            {"$group": {
                "_id": {"$dateToString": {"format": date_format, "date": "$date"}},
                "tokens_consumed": {"$sum": "$tokens_consumed"}
            }},
            {"$sort": {"_id": 1}}
        ]

        result = list(consumption_collection.aggregate(pipeline))
        dates = [doc['_id'] for doc in result]
        tokens = [doc['tokens_consumed'] for doc in result]
        return dates, tokens
    except Exception as e:
        logging.error(f"Erreur lors de la récupération des consommations pour l'utilisateur {user_id} : {e}")
        st.error(f"Erreur lors de la récupération des consommations : {e}")
        return [], []

def afficher_tokens_sidebar(utilisateur):
    st.sidebar.subheader("Consommation de Tokens")
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    st.sidebar.write(f"**Tokens achetés :** {format_number(tokens_purchased)}")
    st.sidebar.write(f"**Tokens consommés :** {format_number(tokens_consumed)}")
    st.sidebar.write(f"**Tokens restants :** {format_number(tokens_remaining)}")
    
    if tokens_purchased > 0:
        progress_value = min(tokens_consumed / tokens_purchased, 1.0)
    else:
        progress_value = 0
    st.sidebar.progress(progress_value)
    st.sidebar.write(f"**{progress_value*100:.2f}%** des tokens ont été consommés.")

# ---------------------- FONCTIONS DE PAIEMENT (ADMIN / MANUEL) ----------------------
def mettre_a_jour_statut_paiement(paiement_id, nouveau_statut):
    try:
        if nouveau_statut not in ["Validated", "Rejected"]:
            st.error("Statut invalide.")
            logging.warning(f"Tentative de mise à jour avec statut invalide : {nouveau_statut}")
            return False, "Statut invalide."
        
        paiement = db['payments'].find_one({"_id": ObjectId(paiement_id)})
        if not paiement:
            st.error("Paiement non trouvé.")
            logging.warning(f"Paiement non trouvé pour l'ID : {paiement_id}")
            return False, "Paiement non trouvé."
        
        db['payments'].update_one(
            {"_id": ObjectId(paiement_id)},
            {"$set": {"status": nouveau_statut}}
        )
        logging.info(f"Statut du paiement {paiement_id} mis à jour à {nouveau_statut}")
        
        # Si validé, ajouter les tokens à l'utilisateur
        if nouveau_statut == "Validated":
            users_collection.update_one(
                {"_id": paiement['user_id']},
                {"$inc": {"tokens_purchased": paiement['tokens_purchased']}}
            )
            # Envoyer un email à l'utilisateur
            user = users_collection.find_one({"_id": paiement['user_id']})
            sujet = "Votre Paiement a été Validé"
            contenu = f"""
            <h2>Votre Paiement a été Validé</h2>
            <p>Merci d'avoir acheté {paiement['tokens_purchased']} tokens.</p>
            <p>Votre solde de tokens a été mis à jour.</p>
            """
            envoyer_email(user['email'], sujet, contenu)
            message = "Paiement validé et tokens ajoutés."
            st.success(message)
            logging.info(f"Tokens ajoutés à l'utilisateur {user['email']} suite à la validation du paiement.")
        elif nouveau_statut == "Rejected":
            # Envoyer un email à l'utilisateur
            user = users_collection.find_one({"_id": paiement['user_id']})
            sujet = "Votre Paiement a été Rejeté"
            contenu = f"""
            <h2>Votre Paiement a été Rejeté</h2>
            <p>Nous n'avons pas pu valider votre paiement de {paiement['amount']} USD.</p>
            <p>Veuillez réessayer ou contacter le support.</p>
            """
            envoyer_email(user['email'], sujet, contenu)
            message = "Paiement rejeté et utilisateur notifié."
            st.warning(message)
            logging.info(f"Utilisateur {user['email']} notifié du rejet du paiement.")
        
        return True, message
    except Exception as e:
        logging.error(f"Erreur lors de la mise à jour du statut du paiement {paiement_id} : {e}")
        st.error(f"Erreur lors de la mise à jour du statut du paiement : {e}")
        return False, "Erreur lors de la mise à jour du statut du paiement."

def traiter_paiement_manuel(user_id, amount, uploaded_file):
    if uploaded_file.type not in ["image/jpeg", "image/png", "image/jpg"]:
        st.error("Veuillez uploader une image valide (JPEG, PNG).")
        logging.warning(f"Utilisateur {user_id} a tenté d'uploader un fichier invalide : {uploaded_file.type}")
        return False, "Type de fichier invalide."
    
    try:
        upload_dir = "uploads/"
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir)
        file_path = os.path.join(upload_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        uploaded_file_url = file_path
    except Exception as e:
        logging.error(f"Erreur lors de la sauvegarde du fichier uploadé par {user_id} : {e}")
        st.error(f"Erreur lors de la sauvegarde du fichier : {e}")
        return False, "Erreur lors de la sauvegarde du fichier."
    
    # 1 USD = 100 tokens (exemple)
    tokens = amount * TOKEN_NUMBER
    
    success, message = enregistrer_paiement(user_id, amount, tokens, "Manual", uploaded_file_url)
    if success:
        sujet = "Nouvelle Demande de Paiement Manuel"
        user = users_collection.find_one({"_id": user_id})
        contenu = f"""
        <h2>Nouvelle Demande de Paiement Manuel</h2>
        <p>Un utilisateur a soumis une demande d'achat de tokens manuellement.</p>
        <p><strong>Email de l'utilisateur :</strong> {user['email']}</p>
        <p><strong>Montant :</strong> {amount} USD</p>
        <p><strong>Nombre de Tokens :</strong> {tokens}</p>
        <p><strong>Fichier Uploadé :</strong> <a href="file://{uploaded_file_url}">Voir le fichier</a></p>
        """
        envoyer_email(ADMIN_EMAIL, sujet, contenu)
        success = envoyer_email_piecejointes(ADMIN_EMAIL, sujet, contenu, attachments=[uploaded_file])
        if not success:
            st.error("Échec de l'envoi de l'email.")
            return
        logging.info(f"Nouvelle demande de paiement manuel enregistrée pour l'utilisateur {user['email']}")
        st.success("Votre demande de paiement a été soumise et est en attente de validation.")
    else:
        st.error(message)
    
    return success, message

def afficher_historique_paiements(utilisateur, items_par_page=6):
    st.subheader("Historique des Paiements")
    paiements = recuperer_historique_paiements(utilisateur['_id'])
    
    if not paiements:
        st.write("Aucun paiement effectué.")
        return

    # Initialiser la page actuelle dans le state si ce n'est pas déjà fait
    if 'page_paiements' not in st.session_state:
        st.session_state.page_paiements = 1

    total_paiements = len(paiements)
    total_pages = ceil(total_paiements / items_par_page)

    # Calculer les indices de début et de fin pour les paiements à afficher
    debut = (st.session_state.page_paiements - 1) * items_par_page
    fin = debut + items_par_page
    paiements_a_afficher = paiements[debut:fin]

    # Afficher les paiements en grille (par exemple, 2 colonnes)
    cols = st.columns(2)
    for idx, paiement in enumerate(paiements_a_afficher):
        col = cols[idx % 2]
        with col:
            st.markdown(f"""
                **Montant :** {paiement['amount']} USD  
                **Tokens :** {format_number(paiement['tokens_purchased'])}  
                **Méthode :** {paiement['payment_method']}  
                **Statut :** {paiement['status']}  
                **Date :** {paiement['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}  
                """)
            if paiement['payment_method'] == "Manual" and paiement.get('uploaded_file_url'):
                st.markdown(f"[Voir le justificatif]({paiement['uploaded_file_url']})")
            st.markdown("---")

    # Afficher la pagination
    st.markdown(f"Page {st.session_state.page_paiements} sur {total_pages}")

    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.session_state.page_paiements > 1:
            if st.button("Précédent"):
                st.session_state.page_paiements -= 1
                st.rerun()
    with col3:
        if st.session_state.page_paiements < total_pages:
            if st.button("Suivant"):
                st.session_state.page_paiements += 1
                st.rerun()

def afficher_historique_paiements_admin():
    st.subheader("Historique des Paiements")
    paiements = list(db['payments'].find().sort("timestamp", -1))
    
    if paiements:
        for paiement in paiements:
            user = users_collection.find_one({"_id": paiement['user_id']})
            st.markdown(f"""
                **Utilisateur :** {user['nom']} ({user['email']})  
                **Montant :** {paiement['amount']} USD  
                **Tokens :** {paiement['tokens_purchased']}  
                **Méthode :** {paiement['payment_method']}  
                **Statut :** {paiement['status']}  
                **Date :** {paiement['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}  
                """)
            if paiement['payment_method'] == "Manual" and paiement['uploaded_file_url']:
                st.markdown(f"[Voir le justificatif]({paiement['uploaded_file_url']})")
            st.markdown("---")
    else:
        st.write("Aucun paiement effectué.")

def filtrer_paiements(paiements, periode):
    aujourd_hui = datetime.datetime.utcnow()
    if periode == 'Jour':
        debut = aujourd_hui - datetime.timedelta(days=1)
    elif periode == '2 Jours':
        debut = aujourd_hui - datetime.timedelta(days=2)
    elif periode == 'Semaine':
        debut = aujourd_hui - datetime.timedelta(weeks=1)
    elif periode == 'Mois':
        debut = aujourd_hui - datetime.timedelta(days=30)
    elif periode == '3 Mois':
        debut = aujourd_hui - datetime.timedelta(days=90)
    elif periode == '6 Mois':
        debut = aujourd_hui - datetime.timedelta(days=180)
    elif periode == 'Année':
        debut = aujourd_hui - datetime.timedelta(days=365)
    else:
        debut = aujourd_hui - datetime.timedelta(days=30)

    paiements_filtrés = [p for p in paiements if debut <= p['timestamp'] <= aujourd_hui]
    return paiements_filtrés

# ---------------------- PAGES ----------------------
def page_connexion():
    st.title("Connexion")

    # Vérifiez si un JWT valide est déjà présent
    jwt_token = st.session_state.get('jwt_token')
    if jwt_token:
        payload = decode_jwt_token(jwt_token)
        if payload:
            st.session_state['page'] = 'Accueil' if payload.get('role') == 'user' else 'page_admin'
            st.rerun()
        else:
            # JWT invalide ou expiré
            st.session_state['jwt_token'] = None
    
    # 1) Charger les credentials depuis Mongo
    credentials = load_credentials_from_mongo()
    #st.write(credentials)

    # Si la base est vide, on affiche un simple avertissement
    if not credentials.get("usernames"):
        st.warning("Aucun utilisateur n'est défini en base. "
                    "Vous pouvez en créer un (via 'Créer un compte') "
                    "ou vérifier la base.")

        # 2) Créer l'instance d'authentification
    authenticator = stauth.Authenticate(
        credentials,
        "my_cookie_name",        # Nom du cookie
        "CLE_SECRETE_STREAMLIT", # Clé secrète pour signer le cookie
        cookie_expiry_days=7     # Durée de validité du cookie (en jours)
        )


    # 3) Afficher le widget de connexion
    authenticator.login(
        location="main",
        fields={
            'Form name': 'Login',
            'Username': 'Username',
            'Password': 'Password',
            'Login': 'Login'
        }
    )

    # 4) Lire l'état de l'authentification dans st.session_state
    authentication_status = st.session_state.get("authentication_status", None)
    name = st.session_state.get("name", None)
    username = st.session_state.get("username", None)
    #st.write(username)

    # 5) Gérer les différents cas
    if authentication_status:
        st.success(f"Bienvenue {name} !")

        # Vérifier si l'utilisateur est bloqué
        user_doc = users_collection.find_one({"email": username})
        if not user_doc:
            st.error("Utilisateur introuvable. Veuillez vérifier vos identifiants.")
            return

        if user_doc.get("is_blocked", False):
            st.error("Votre compte a été bloqué. Veuillez contacter l'administrateur.")
            authenticator.logout("Se déconnecter", "main")
            return

        # Vérifier si l'OTP a déjà été validé
        if user_doc.get('otp_validated', False):
            # OTP déjà validé, rediriger vers la page d'accueil
            st.session_state['page'] = 'Accueil' if user_doc.get('role') == 'user' else 'page_admin'
            st.rerun()
        else:
            # OTP non validé, envoyer l'OTP et rediriger vers la vérification OTP
            st.session_state['temp_user_id'] = str(user_doc['_id'])
            st.info("Votre session nécessite une validation par OTP.")
            envoyer_otp_via_email(username)
            st.session_state['page'] = 'OTP Verification'
            st.rerun()

    elif authentication_status is False:
        st.error("Identifiants incorrects ou compte bloqué.")
    elif authentication_status is None:
        st.warning("Veuillez entrer vos identifiants.")

    # 6) Boutons de navigation
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Créer un compte"):
            st.session_state['page'] = 'Inscription'
            st.rerun()
    with col2:
        if st.button("Mot de passe oublié"):
            st.session_state['page'] = 'Mot de passe oublié'
            st.rerun()


def generate_otp(length=6):
    """Génère un OTP de 'length' chiffres."""
    return ''.join(str(random.randint(0, 9)) for _ in range(length))

def envoyer_otp_via_email(email):
    """Exemple : génère un OTP, le stocke en session + expiration 5 min, et l'envoie par email."""
    otp = generate_otp(6)  # 6 chiffres
    st.session_state['otp'] = otp
    # On fixe une expiration à 5 minutes
    st.session_state['otp_expiration'] = datetime.datetime.utcnow() + datetime.timedelta(minutes=5)

    # Envoi par email (pseudo-code)
    contenu = f"Votre code OTP est : {otp}. Il expirera dans 5 minutes."
    envoyer_email(email, "Votre OTP", contenu)

    st.success("Un code OTP vous a été envoyé par email (valable 5 minutes).")

def page_otp_verification():
    st.title("🔒 Vérification OTP")

    # Récupère l'OTP en session, et son expiration
    otp_stored = st.session_state.get('otp')
    otp_expiration = st.session_state.get('otp_expiration')

    # S'il n'y a pas d'OTP en session, avertir et rediriger
    if not otp_stored or not otp_expiration:
        st.error("Aucun OTP n'a été généré. Veuillez revenir à la page de connexion ou d'inscription.")
        st.session_state['page'] = 'Connexion'
        st.rerun()

    # Vérifie si l'OTP a expiré
    now_utc = datetime.datetime.utcnow()
    if now_utc > otp_expiration:
        st.error("Votre code OTP a expiré. Veuillez redémarrer le processus.")
        del st.session_state['otp']
        del st.session_state['otp_expiration']
        # Proposer un bouton pour renvoyer l'OTP
        st.write("Veuillez renvoyer un nouvel OTP pour continuer.")
        if st.button("Renvoyer l'OTP"):
            # Appeler la fonction pour générer et envoyer un nouvel OTP
            st.session_state['page'] = 'Connexion'
            st.rerun()
        return  # Arrêter l'exécution de la fonction après avoir proposé le renvoi



    st.write("Veuillez entrer le code OTP à 6 chiffres reçu par email.")

    # Affichage de 6 champs sur la même ligne (1 caractère max chacun)
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    with col1:
        d1 = st.text_input("", max_chars=1, key="digit1")
    with col2:
        d2 = st.text_input("", max_chars=1, key="digit2")
    with col3:
        d3 = st.text_input("", max_chars=1, key="digit3")
    with col4:
        d4 = st.text_input("", max_chars=1, key="digit4")
    with col5:
        d5 = st.text_input("", max_chars=1, key="digit5")
    with col6:
        d6 = st.text_input("", max_chars=1, key="digit6")

    if st.button("Valider"):
        # Concaténer les 6 caractères pour former l'OTP saisi
        otp_entered = (d1 + d2 + d3 + d4 + d5 + d6).strip()

        # Vérifiez que l'utilisateur a bien entré 6 chiffres
        if len(otp_entered) < 6 or not otp_entered.isdigit():
            st.error("Veuillez entrer 6 chiffres.")
            return

        # Vérifiez si l'OTP correspond à celui stocké
        if otp_entered == otp_stored:
            st.success("OTP valide. Vous êtes authentifié !")

            # Nettoyer l'OTP après validation
            del st.session_state['otp']
            del st.session_state['otp_expiration']

            # Récupérer l'utilisateur via l'ID temporaire
            user_id = st.session_state.get('temp_user_id')
            if not user_id:
                st.error("Impossible de récupérer l'identifiant utilisateur.")
                st.session_state['page'] = 'Connexion'
                st.rerun()

            # Cherchez l'utilisateur dans la base de données
            utilisateur = users_collection.find_one({"_id": ObjectId(user_id)})
            if not utilisateur:
                st.error("Utilisateur introuvable.")
                st.session_state['page'] = 'Connexion'
                st.rerun()

            # Générez un JWT
            user_role = utilisateur.get("role", "user")
            jwt_token = create_jwt_token(user_id, user_role)  # Remplacez par votre fonction JWT

            # Stockez les informations utilisateur et le JWT dans la session
            st.session_state['jwt_token'] = jwt_token
            st.session_state['user_info'] = utilisateur
            st.session_state['authenticated'] = True

            # Mettre à jour otp_validated dans MongoDB
            users_collection.update_one(
                {"_id": ObjectId(user_id)},
                {"$set": {"otp_validated": True}}
            )

            # Redirigez l'utilisateur en fonction de son rôle
            st.session_state['page'] = 'Accueil' if user_role == "user" else 'page_admin'
            st.rerun()

        else:
            st.error("Le code OTP saisi est incorrect.")


            
def page_inscription():
    st.title("📝 Inscription")
    nom = st.text_input("Nom")
    email = st.text_input("Email")
    mot_de_passe = st.text_input("Mot de passe", type="password")
    mot_de_passe_conf = st.text_input("Confirmer le mot de passe", type="password")

    if st.button("S'inscrire"):
        if mot_de_passe != mot_de_passe_conf:
            st.error("Les mots de passe ne correspondent pas.")
        else:
            # FORCER le rôle = "user"
            #success, message = ajouter_utilisateur(email, nom, mot_de_passe, role="admin")
            success, message = ajouter_utilisateur(email, nom, mot_de_passe, role="user")

            if success:
                st.success(message)
                logging.info(f"Nouvel utilisateur inscrit : {email}")

                # Récupérer l’utilisateur pour générer et envoyer l’OTP
                utilisateur = users_collection.find_one({"email": email})
                #otp = generate_otp()
                #st.session_state['otp'] = otp
                st.session_state['temp_user_id'] = str(utilisateur['_id'])

                #contenu_otp = f"Votre code OTP d'inscription est : {otp}"
                #envoyer_email(email, "Code OTP d'Inscription", contenu_otp)
                envoyer_otp_via_email(email)

                # Rediriger vers la page OTP
                st.session_state['page'] = 'OTP Verification'
                st.rerun()
            else:
                st.error(message)

    st.markdown("---")
    if st.button("Retour à la connexion"):
        st.session_state['page'] = 'Connexion'
        st.rerun()

def page_reinitialisation():
    st.title("🔒 Réinitialiser le mot de passe")
    email = st.text_input("Email")
    if st.button("Envoyer OTP"):
        utilisateur = users_collection.find_one({"email": email})
        if utilisateur:
            otp = generate_otp()
            st.session_state['reset_otp'] = otp
            st.session_state['reset_user_id'] = str(utilisateur['_id'])
            envoyer_email(email, "Votre OTP de Réinitialisation", f"Votre code OTP pour réinitialiser le mot de passe est : {otp}")
            logging.info(f"OTP de réinitialisation envoyé à {email}")
            st.session_state['page'] = 'Reset OTP Verification'
            st.rerun()
        else:
            st.error("Utilisateur non trouvé.")
    
    st.markdown("---")
    if st.button("Retour à la connexion"):
        st.session_state['page'] = 'Connexion'
        st.rerun()

def page_reset_otp_verification():
    st.title("🔑 Vérification OTP de Réinitialisation")
    otp_input = st.text_input("Entrez votre OTP")
    nouveau_mot_de_passe = st.text_input("Nouveau mot de passe", type="password")
    nouveau_mot_de_passe_conf = st.text_input("Confirmer le nouveau mot de passe", type="password")
    if st.button("Réinitialiser le mot de passe"):
        if 'reset_otp' in st.session_state and otp_input == st.session_state['reset_otp']:
            if nouveau_mot_de_passe != nouveau_mot_de_passe_conf:
                st.error("Les mots de passe ne correspondent pas.")
            else:
                user_id = st.session_state['reset_user_id']
                email = users_collection.find_one({"_id": ObjectId(user_id)})['email']
                reinitialiser_mot_de_passe(email, nouveau_mot_de_passe)
                logging.info(f"Mot de passe réinitialisé pour l'utilisateur {email}")
                st.success("Mot de passe réinitialisé avec succès.")
                st.session_state['page'] = 'Connexion'
                st.rerun()
        else:
            st.error("OTP invalide.")

def page_google_login():
    st.title("🔗 Connexion avec Google")
    
    oauth = OAuth2Session(
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scope='openid email profile',
        redirect_uri=f"{os.getenv('BASE_URL')}/?page=Google Callback"
    )
    
    authorization_url, state = oauth.create_authorization_url(
        GOOGLE_AUTHORIZATION_ENDPOINT,
        access_type='offline',
        prompt='select_account'
    )
    
    st.markdown(f"[Cliquez ici pour vous connecter avec Google]({authorization_url})")
    
    query_params = st.experimental_get_query_params()
    if 'code' in query_params:
        code = query_params['code'][0]
        try:
            token = oauth.fetch_token(
                GOOGLE_TOKEN_ENDPOINT,
                code=code,
                client_secret=GOOGLE_CLIENT_SECRET
            )
            user_info = oauth.get(GOOGLE_USERINFO_ENDPOINT).json()
            email = user_info['email']
            nom = user_info['name']
            google_id = user_info['id']
            profile_picture_url = user_info.get('picture', "https://via.placeholder.com/100")
            
            # Vérifier si l'utilisateur existe
            utilisateur = users_collection.find_one({"email": email})
            if utilisateur:
                if utilisateur.get('is_blocked', False):
                    st.error("Votre compte a été bloqué. Veuillez contacter l'administrateur.")
                    logging.warning(f"Utilisateur bloqué via Google : {email}")
                    return
                # Connecter l'utilisateur
                jwt_token = create_jwt_token(utilisateur['_id'], utilisateur.get("is_admin", False))
                st.session_state['jwt_token'] = jwt_token
                st.session_state['user_info'] = utilisateur
                logging.info(f"Utilisateur connecté via Google : {email}")
                st.success("Connexion réussie!")
            else:
                # Créer un nouvel utilisateur
                success, message = ajouter_utilisateur(email, nom, google_id, role="admin")
                if success:
                    utilisateur = users_collection.find_one({"email": email})
                    jwt_token = create_jwt_token(utilisateur['_id'], utilisateur.get("is_admin", False))
                    st.session_state['jwt_token'] = jwt_token
                    st.session_state['user_info'] = utilisateur
                    logging.info(f"Nouvel utilisateur inscrit et connecté via Google : {email}")
                    st.success("Inscription et connexion réussies!")
                else:
                    st.error(message)
            st.session_state['page'] = 'Accueil'
            st.rerun()
        except Exception as e:
            logging.error(f"Erreur lors de la connexion via Google : {e}")
            st.error(f"Erreur lors de la connexion via Google ")

def page_accueilles():
    st.title('Business Plan')
    utilisateur = get_current_user()
    #if not utilisateur:
       #st.warning("Vous n'êtes pas connecté !")
       # return
    
    st.sidebar.header("Configuration Initiale pour le business model")
    st.sidebar.selectbox(
        "Type d'entreprise",
        ["PME", "Startup"],
        key="type_entreprise"
    ) 

    st.sidebar.text_input(
        "Montant limite pour le projet",
        value=st.session_state.montant_projet,
        key="montant_projet"
    )

    st.sidebar.text_input(
        "Nom de l'entreprise",
        value=st.session_state.nom_entreprise,
        key="nom_entreprise"
    )

    if not nom_entreprise:
        st.sidebar.warning("Veuillez entrer le nom de votre entreprise.")
    
    tab_names = ["Section 1", "Section 2", "Section 3"]
    
    # Onglets
    tabs = st.tabs(tab_names)
    # On n'a pas les fonctions sections[i], on peut juste faire un for
    for i, tab in enumerate(tabs):
        with tab:
            st.write(f"Contenu de {tab_names[i]} - À personnaliser...")

    # Bouton de déconnexion
    if st.button("Se déconnecter"):
        st.session_state.pop('jwt_token', None)
        st.session_state.pop('user_info', None)
        st.session_state['page'] = 'Connexion'
        st.rerun()

    # Afficher sidebar tokens
    afficher_tokens_sidebar(utilisateur)
    
    
    
def page_accueil() :
    col1, col2 = st.columns([4, 1])  # Ajustez les ratios selon vos besoins
    with col1:
        st.title("Business Plan")  # Titre dans la première colonne

    with col2:
        if st.button("🔄 Actualiser"):
            st.rerun()
            
    #st.title('Business Plan')
    utilisateur = get_current_user()
    st.session_state['user_info'] = utilisateur
    
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté !")
        return

    # Définir les pages pour la navigationdc
    pages = [
        st.Page(page_accueil_content, title="🏠 Accueil"),
        st.Page(page_chat, title="💬 Chat"),
        st.Page(page_achat, title="💰 Acheter des Tokens"),
        st.Page(page_profil, title="👤 Profil"),
        st.Page(page_deconnexion, title="❌ Déconnexion"),
    ]
    # Créer et exécuter la navigation
    pg = st.navigation(pages)
    pg.run()

def page_accueil_content():
    utilisateur = get_current_user()
    
    #st.subheader(f"Bienvenue {utilisateur['nom']} !")
    #st.write("Bienvenue sur votre tableau de bord. Utilisez la navigation ci-dessus pour accéder aux différentes sections.")   
    st.sidebar.header("Configuration Initiale")
    type_entreprise = st.sidebar.selectbox("Type d'entreprise", ["PME", "Startup"], key="type_entreprise")
    montant_projet = st.sidebar.text_input("Montant limite pour le projet", value="", key="montant_projet")
    nom_entreprise = st.sidebar.text_input("Nom de l'entreprise", value="", key="nom_entreprise")

    if not nom_entreprise:
        st.sidebar.warning("Veuillez entrer le nom de votre entreprise.")
    
    afficher_tokens_sidebar(utilisateur)
    
    # Création des onglets
    tabs = st.tabs(tab_names)

    # Parcours des onglets
    for i, tab in enumerate(tabs):
        with tab:
            sections[i]()
        
    
def page_chatss():
    st.title("💬 Chat Interface")
    utilisateur = get_current_user()
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté.")
        return
   
    MODEL = "gpt-4"
    MAX_TOKENS_PER_REQUEST = 150
    
    # Calculer les tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    st.session_state['user_info'] = utilisateur  # Assurez-vous que la session est à jour
    
    afficher_tokens_sidebar(utilisateur)
    
    # Initialiser l'historique des messages dans la session
    if 'messages' not in st.session_state:
        st.session_state['messages'] = [
            {"role": "system", "content": "Tu es un assistant expert en génération de business et business plan."}
        ]

    # Afficher l'historique des messages
    for msg in st.session_state['messages']:
        if msg['role'] == 'user':
            st.chat_message("user").write(msg['content'])
        elif msg['role'] == 'assistant':
            st.chat_message("assistant").write(msg['content'])

    # Entrée utilisateur via le composant de chat
    user_input = st.chat_input("Entrez votre question")

    if user_input:
        # Ajouter le message de l'utilisateur à l'historique
        st.session_state['messages'].append({"role": "user", "content": user_input})
        st.chat_message("user").write(user_input)
        
        # Calculer les tokens nécessaires (entrée + réponse prévue)
        tokens_in_input = count_tokens(user_input, MODEL)
        tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
        st.write(f"Tokens nécessaires : {tokens_needed}")
        
        # Vérifier si l'utilisateur a assez de tokens
        if tokens_remaining < tokens_needed:
            st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
            return
        else:
            try:
                # Créer un conteneur vide pour la réponse de l'assistant une seule fois
                assistant_placeholder = st.chat_message("assistant")
                
                # Initialiser la réponse de l'assistant
                assistant_reply = ""
                
                # Appel à l'API OpenAI ChatCompletion avec streaming
                response = openai.ChatCompletion.create(
                    model=MODEL,  # Assurez-vous que le nom du modèle est correct
                    messages=st.session_state['messages'],
                    max_tokens=1500,  # Ajustez selon vos besoins
                    temperature=0.7,
                    stream=True  # Activer le streaming
                )
                
                # Itérer sur les fragments de réponse
                for chunk in response:
                    if 'choices' in chunk and len(chunk['choices']) > 0:
                        delta = chunk['choices'][0].get('delta', {})
                        content = delta.get('content', '')
                        if content:
                            assistant_reply += content
                            # Mettre à jour le même conteneur avec le contenu actuel
                            assistant_placeholder.write(assistant_reply)
                
                # Une fois la réponse complète, ajouter le message à l'historique
                st.session_state['messages'].append({"role": "assistant", "content": assistant_reply})
                
                # Calculer les tokens utilisés
                tokens_utilises = response['usage']['total_tokens'] if 'usage' in response else tokens_needed
                success, message = consommer_tokens(utilisateur, tokens_utilises)

                # Afficher les tokens consommés
                if success:
                    st.success(f"{tokens_utilises} tokens consommés.")
                else:
                    st.error(message)
            
            except Exception as e:
                logging.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
                st.error("Erreur lors de l'appel à l'API OpenAI. Veuillez réessayer plus tard.")


def page_chatssss():
    st.title("💬 Interface de Chat")

    utilisateur = get_current_user()
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté.")
        return
   
    MODEL = "gpt-4"  # Assurez-vous que le modèle est correct et supporte le streaming
    MAX_TOKENS_PER_REQUEST = 150  # Ajustez selon vos besoins
    
    # Calculer les tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    st.session_state['user_info'] = utilisateur  # Assurez-vous que la session est à jour
    
    afficher_tokens_sidebar(utilisateur)
    
    # Initialiser l'historique des messages dans la session
    if 'messages' not in st.session_state:
        st.session_state['messages'] = [
            {"role": "system", "content": "Tu es un assistant expert en génération de business et business plan."}
        ]

    # Afficher l'historique des messages
    for msg in st.session_state['messages']:
        if msg['role'] == 'user':
            st.chat_message("user").write(msg['content'])
        elif msg['role'] == 'assistant':
            st.chat_message("assistant").write(msg['content'])

    # Entrée utilisateur via le composant de chat
    user_input = st.chat_input("Entrez votre question")

    if user_input:
        # Ajouter le message de l'utilisateur à l'historique
        st.session_state['messages'].append({"role": "user", "content": user_input})
        st.chat_message("user").write(user_input)
        
        # Calculer les tokens nécessaires (entrée + réponse prévue)
        tokens_in_input = count_tokens(user_input, MODEL)
        tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
        st.write(f"Tokens nécessaires : {tokens_needed}")
        
        # Vérifier si l'utilisateur a assez de tokens
        if tokens_remaining < tokens_needed:
            st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
            return
        else:
            # Vérifier si une réponse est déjà en cours
            if 'response_pending' in st.session_state and st.session_state['response_pending']:
                st.warning("Une réponse est déjà en cours de génération. Veuillez patienter.")
                return
            else:
                st.session_state['response_pending'] = True  # Marquer comme réponse en cours
                
                try:
                    # Créer un conteneur unique pour la réponse de l'assistant
                    if 'assistant_placeholder' not in st.session_state:
                        st.session_state['assistant_placeholder'] = st.empty()
                    
                    assistant_placeholder = st.session_state['assistant_placeholder']
                    
                    # Initialiser la réponse de l'assistant
                    assistant_reply = ""
                    
                    # Appel à l'API OpenAI ChatCompletion avec streaming
                    response = openai.ChatCompletion.create(
                        model=MODEL,  # Assurez-vous que le nom du modèle est correct
                        messages=st.session_state['messages'],
                        max_tokens=1500,  # Ajustez selon vos besoins
                        temperature=0.7,
                        stream=True  # Activer le streaming
                    )
                    
                    # Itérer sur les fragments de réponse
                    for chunk in response:
                        if 'choices' in chunk and len(chunk['choices']) > 0:
                            delta = chunk['choices'][0].get('delta', {})
                            content = delta.get('content', '')
                            if content:
                                assistant_reply += content
                                # Mettre à jour le conteneur avec le contenu accumulé
                                assistant_placeholder.markdown(f"**Assistant:** {assistant_reply}")
                    
                    # Une fois la réponse complète, ajouter le message à l'historique
                    st.session_state['messages'].append({"role": "assistant", "content": assistant_reply})
                    
                    # Calculer les tokens utilisés
                    tokens_utilises = response['usage']['total_tokens'] if 'usage' in response else tokens_needed
                    success, message = consommer_tokens(utilisateur, tokens_utilises)

                    # Afficher les tokens consommés
                    if success:
                        st.success(f"{tokens_utilises} tokens consommés.")
                    else:
                        st.error(message)
                
                except Exception as e:
                    logging.error(f"Erreur lors de l'appel à l'API OpenAI : {e}")
                    st.error("Erreur lors de l'appel à l'API OpenAI. Veuillez réessayer plus tard.")
                
                finally:
                    st.session_state['response_pending'] = False  # Réinitialiser le drapeau
                    
import time                  
# Fonction principale de l'application de chat
# Définir la limite de contexte
CONTEXT_LIMIT = 4100  # Ajustez selon votre modèle
CONTEXT_LIMIT_INPUT = 3590  # Ajustez selon votre modèle


def trim_message_history(messages, user_input, model="gpt-4", context_limit=CONTEXT_LIMIT):
    """
    Garde les messages les plus récents jusqu'à ce que la somme des tokens soit inférieure à la limite de contexte.
    Inclut toujours le user_input et préserve le message système.
    """
    total_tokens = 0
    trimmed_messages = []
    
    # Inclure le message système en premier
    system_messages = [msg for msg in messages if msg['role'] == 'system']
    if system_messages:
        trimmed_messages.extend(system_messages)
        total_tokens += count_tokens(system_messages[0]['content'], model)
    
    # Calculer les tokens de user_input
    user_input_tokens = count_tokens(user_input, model)
    
    # Parcourir les messages en ordre inverse (du plus récent au plus ancien)
    for msg in reversed(messages):
        if msg['role'] == 'system':
            continue  # Déjà inclus
        msg_tokens = count_tokens(msg['content'], model)
        if total_tokens + msg_tokens + user_input_tokens > context_limit + 200:
            break
        trimmed_messages.insert(1, msg)  # Insérer après le message système
        total_tokens += msg_tokens
    
    # Ajouter le user_input à la fin de l'historique trimé
    #trimmed_messages.append({"role": "user", "content": user_input})
    total_tokens += user_input_tokens
    return trimmed_messages


def page_chat():
    st.title("💬 Interface de Chat")

    utilisateur = get_current_user()
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté.")
        return
    
    MODEL = "gpt-4"
    MAX_TOKENS_PER_REQUEST = 150
    
    # Calcul des tokens restants
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed

    afficher_tokens_sidebar(utilisateur)
    
    # Initialiser l'historique des messages dans la session
    if 'messages' not in st.session_state:
        st.session_state['messages'] = [
            {"role": "system", "content": "Tu es un assistant expert en génération de business model et business plan. aussi un expert en business"}
        ]
    
    # Initialiser le drapeau de réponse en cours
    if 'response_pending' not in st.session_state:
        st.session_state['response_pending'] = False

    # Afficher tous les messages dans l'ordre chronologique
    for msg in st.session_state['messages']:
        if msg['role'] == 'user':
            st.chat_message("user").write(msg['content'])
        elif msg['role'] == 'assistant':
            st.chat_message("assistant").write(msg['content'])

    # Placeholder pour la réponse en cours de l'assistant
    assistant_placeholder = None
    if st.session_state['response_pending']:
        assistant_placeholder = st.empty()
        assistant_placeholder.markdown("**Assistant:** En cours de génération...")
    
    # Entrée utilisateur via le composant de chat
    user_input = st.chat_input("Entrez votre question")

    if user_input:
        # Ajouter le message de l'utilisateur à l'historique
        if count_tokens(user_input, MODEL) > CONTEXT_LIMIT_INPUT :
            st.error("❌ Votre message est trop long et dépasse la limite de contexte autorisée. Veuillez réduire la longueur de votre message et réessayer.")
            return

        st.session_state['messages'].append({"role": "user", "content": user_input})
        st.chat_message("user").write(user_input)
        
        # Calculer les tokens nécessaires (entrée + réponse prévue)
        tokens_in_input = count_tokens(user_input+ json.dumps(st.session_state['messages']), MODEL)
        tokens_needed = tokens_in_input + MAX_TOKENS_PER_REQUEST
        st.write(f"**Tokens nécessaires :** {tokens_needed}")
        
        # Vérifier si l'utilisateur a assez de tokens
        if tokens_remaining < tokens_needed:
            st.error("Vous n'avez pas assez de tokens pour effectuer cette action. Veuillez acheter plus de tokens.")
            return
        else:
            # Vérifier si une réponse est déjà en cours
            if st.session_state['response_pending']:
                st.warning("Une réponse est déjà en cours de génération. Veuillez patienter.")
                return
            else:
                # Trimmer l'historique des messages pour respecter la limite de contexte
                trimmed_messages = trim_message_history(st.session_state['messages'],user_input, MODEL, CONTEXT_LIMIT)
                
                # Mettre à jour l'historique des messages avec les messages trimés
                st.session_state['messages'] = trimmed_messages
                
                st.session_state['response_pending'] = True  # Marquer comme réponse en cours

                try:
                    # Créer un conteneur unique pour la réponse de l'assistant
                    assistant_placeholder = st.empty()
                    assistant_placeholder.markdown("**Assistant:** En cours de génération...")
                    
                    # Initialiser la réponse de l'assistant
                    assistant_reply = ""
                    
                    # Appel à l'API OpenAI ChatCompletion avec streaming
                    response = openai.ChatCompletion.create(
                        model=MODEL,
                        messages=st.session_state['messages'], 
                        max_tokens=4000,
                        temperature=0.7,
                        stream=True  # Activer le streaming
                    )
                    
                    # Itérer sur les fragments de réponse
                    for chunk in response:
                        if 'choices' in chunk and len(chunk['choices']) > 0:
                            delta = chunk['choices'][0].get('delta', {})
                            content = delta.get('content', '')
                            if content:
                                assistant_reply += content
                                # Mettre à jour le conteneur avec le contenu accumulé
                                assistant_placeholder.markdown(f"**Assistant:** {assistant_reply}")
                                time.sleep(0.05)  # Petit délai pour permettre à Streamlit de mettre à jour l'interface
                    
                    # Une fois la réponse complète, ajouter le message à l'historique
                    st.session_state['messages'].append({"role": "assistant", "content": assistant_reply})
                    
                    # Calculer les tokens utilisés
                    tokens_utilises = tokens_needed+count_tokens(assistant_reply, MODEL)  # Simplification si 'usage' n'est pas disponible
                    if 'usage' in response:
                        tokens_utilises = response['usage']['total_tokens']
                    
                    success, message = consommer_tokens(utilisateur, tokens_utilises)

                    # Afficher les tokens consommés séparément
                    if success:
                        st.success(f"{tokens_utilises} tokens consommés.")
                    else:
                        st.error(message)
                
                except openai.error.OpenAIError as e:
                    logging.error(f"Erreur OpenAI: {e}")
                    st.error(f"Erreur lors de l'appel à l'API OpenAI: {e}")
                
                except Exception as e:
                    logging.error(f"Erreur inattendue: {e}")
                    st.error(f"Erreur inattendue: {e}")
                
                finally:
                    st.session_state['response_pending'] = False





def page_achat():
    st.title("💰 Acheter des Tokens")
    utilisateur = get_current_user()
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté.")
        return
    
    afficher_tokens_sidebar(utilisateur)
    
    mode_paiement = st.radio("Choisissez une méthode de paiement :", ["Stripe", "Manuel"])
    
    if mode_paiement == "Stripe":
        montant = st.number_input("Montant en USD", min_value=1, step=1)
        if st.button("Acheter avec Stripe"):
            if montant < 1:
                st.error("Le montant doit être d'au moins 1 USD.")
            else:
                if not utilisateur.get('stripe_customer_id'):
                    try:
                        customer = stripe.Customer.create(email=utilisateur['email'])
                        users_collection.update_one(
                            {"_id": utilisateur['_id']},
                            {"$set": {"stripe_customer_id": customer['id']}}
                        )
                        utilisateur['stripe_customer_id'] = customer['id']
                        st.session_state['user_info']['stripe_customer_id'] = customer['id']
                        logging.info(f"Client Stripe créé pour {utilisateur['email']}")
                    except Exception as e:
                        logging.error(f"Erreur lors de la création du client Stripe ")
                        st.error(f"Erreur lors de la création du client Stripe")
                        return
                
                success_url = f"{os.getenv('BASE_URL')}/?page=Accueil&success=true"
                cancel_url = f"{os.getenv('BASE_URL')}/?page=Accueil&cancel=true"
                session_url = creer_session_stripe(
                    utilisateur['stripe_customer_id'],
                    montant * 100,
                    success_url,
                    cancel_url
                )
                if session_url:
                    st.markdown(f"[Cliquez ici pour payer]({session_url})")
    
    elif mode_paiement == "Manuel":
        with st.form("paiement_manuel"):
            montant = st.number_input("Montant en USD", min_value=1, step=1)
            uploaded_file = st.file_uploader("Uploader une image de paiement (JPEG, JPG, PNG)", type=["jpg", "jpeg", "png"])
            submit_button = st.form_submit_button("Soumettre le Paiement")
        
        if submit_button:
            if montant < 1:
                st.error("Le montant doit être d'au moins 1 USD.")
            elif not uploaded_file:
                st.error("Veuillez uploader une image de paiement.")
            else:
                success, message = traiter_paiement_manuel(utilisateur['_id'], montant, uploaded_file)
                if success:
                    st.success("Votre paiement a été soumis et est en attente de validation.")
                else:
                    st.error(message)
    
    st.subheader("Dernier Paiement")
    st.write("🔍 **Astuce :** Pour voir l'historique complet de vos paiements, rendez-vous sur votre")

    dernier_paiement = recuperer_dernier_paiement(utilisateur['_id'])
    if dernier_paiement:
        st.markdown(f"""
            **Montant :** {dernier_paiement['amount']} USD  
            **Tokens :** {dernier_paiement['tokens_purchased']}  
            **Méthode :** {dernier_paiement['payment_method']}  
            **Statut :** {dernier_paiement['status']}  
            **Date :** {dernier_paiement['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}  
            """)
        if dernier_paiement['payment_method'] == "Manual" and dernier_paiement.get('uploaded_file_url'):
            st.markdown(f"[Voir le justificatif]({dernier_paiement['uploaded_file_url']})")

    else:
        st.write("Aucun paiement effectué.")


def page_profil():
    st.title("🧑 Profil Utilisateur")
    utilisateur = get_current_user()
    if not utilisateur:
        st.warning("Vous n'êtes pas connecté.")
        return

    afficher_tokens_sidebar(utilisateur)
    
    st.markdown(f"""
        <style>
        .profile-container {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }}
        .profile-header {{
            display: flex;
            align-items: center;
            margin-bottom: 20px;
        }}
        .profile-header img {{
            border-radius: 50%;
            margin-right: 20px;
            width: 100px;
            height: 100px;
            object-fit: cover;
        }}
        .profile-header .profile-name {{
            font-size: 28px;
            font-weight: bold;
        }}
        .profile-info {{
            font-size: 18px;
            margin-bottom: 10px;
        }}
        .profile-actions button {{
            background-color: #4CAF50;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 16px;
        }}
        .profile-actions button:hover {{
            background-color: #45a049;
        }}
        .cards-container {{
            display: flex;
            justify-content: space-between;
            flex-wrap: wrap;
            margin-bottom: 20px;
        }}
        .card {{
            background-color: #ffffff;
            border: 1px solid #dddddd;
            border-radius: 10px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            padding: 20px;
            width: 30%;
            min-width: 200px;
            text-align: center;
            margin-bottom: 20px;
        }}
        .card h3 {{
            margin-bottom: 10px;
            font-size: 22px;
            color: #333333;
        }}
        .card p {{
            font-size: 20px;
            color: #555555;
        }}
        </style>
        <div class="profile-container">
            <div class="profile-header">
                <img src="https://www.gravatar.com/avatar/00000000000000000000000000000000?d=mp&f=y" width="100" height="100">
                <div class="profile-name">{utilisateur['nom']}</div>
            </div>
            <div class="profile-info"><strong>Email :</strong> {utilisateur['email']}</div>
            <div class="profile-actions">
                <button onclick="window.location.reload()">Mettre à jour le profil</button>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    tokens_purchased = utilisateur.get('tokens_purchased', 0)
    tokens_consumed = utilisateur.get('tokens_consumed', 0)
    tokens_remaining = tokens_purchased - tokens_consumed
    
    st.markdown(f"""
        <div class="cards-container">
            <div class="card">
                <h3>Tokens Achetés</h3>
                <p>{format_number(tokens_purchased)}</p>
            </div>
            <div class="card">
                <h3>Tokens Consommés</h3>
                <p>{format_number(tokens_consumed)}</p>
            </div>
            <div class="card">
                <h3>Tokens Restants</h3>
                <p>{format_number(tokens_remaining)}</p>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    st.subheader("Consommation de Tokens")
    
    periode = st.selectbox("Filtrer par :", ["Jour", "2 Jours", "Semaine", "Mois", "3 Mois", "6 Mois", "Année"])
    dates, tokens = recuperer_consomation(utilisateur['_id'], periode)
    
    if dates and tokens:
        df = pd.DataFrame({'Date': dates, 'Tokens Consommés': tokens})
        fig = px.bar(
            df, 
            x='Date', 
            y='Tokens Consommés',
            title=f"Consommation de Tokens - {periode}",
            labels={'Tokens Consommés': 'Tokens Consommés', 'Date': 'Date'},
            color='Tokens Consommés',
            color_continuous_scale='Viridis'
        )
        
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.write("Aucune donnée de consommation disponible pour la période sélectionnée.")
    
    afficher_historique_paiements(utilisateur)

def page_deconnexion():
    st.session_state.pop("jwt_token", None)
    st.session_state.pop("user_info", None)
    st.session_state['page'] = "Connexion"


    user_info = get_current_user()
    if not user_info:
        st.error("Utilisateur non connecté.")
        return None

    role = user_info.get("role", "user")  # Par défaut, le rôle est "user"
    if role == "admin":
        st.session_state['authentication_status'] = None
        #authenticator.logout("Déconnexion", "main")
    else :
        users_collection.update_one(
            {"_id": ObjectId(user_info['_id'])},
            {"$set": {"otp_validated": False}}
            )
    st.session_state['page'] = "Connexion"
    st.success("Vous êtes déconnecté.")
    logging.info("Utilisateur déconnecté.")
    st.rerun()
    
    
def page_admin():
    st.title("🛠️ Administration - Gestion des Paiements Manuels")
    utilisateur = get_current_user()
    
    # 1) Vérifier si l'utilisateur est connecté
    if not utilisateur:
        st.error("Vous n'êtes pas connecté.")
        return  # On arrête la fonction pour éviter l'erreur
    
    # 2) Vérifier le rôle
    if utilisateur.get("role") not in ["admin", "super_admin"]:
        st.error("Vous n'avez pas les droits administrateur.")
        return
    
    
    # ---------------------- Dashboard Admin ----------------------
    
    
    # Dashboard Admin
    col1, col2 , col3= st.columns([6, 1, 2])  # Utilisation de st.columns au lieu de col
    with col1:
        st.header("Dashboard Administrateur")
    with col3:
        if st.button("Déconnexion"):  # Correction de st.bouton en st.button et capitalisation
           page_deconnexion()  # Assurez-vous que cette fonction est définie

    periode = st.selectbox(
        "Filtrer par :", 
        ["Jour", "2 Jours", "Semaine", "Mois", "3 Mois", "6 Mois", "Année"], 
        key="admin_periode"
    )
    
    # Récupérer tous les paiements et les trier par date décroissante
    paiements = list(db['payments'].find().sort("timestamp", -1))
    paiements_filtrés = filtrer_paiements(paiements, periode)
    
    # Séparer les paiements par statut
    validated_paiements = [p for p in paiements_filtrés if p['status'] == "Validated"]
    pending_paiements = [p for p in paiements_filtrés if p['status'] == "Pending"]
    rejected_paiements = [p for p in paiements_filtrés if p['status'] == "Rejected"]
    
    # Calculer les métriques financières
    money_earned = sum([p['amount'] for p in validated_paiements])
    money_pending = sum([p['amount'] for p in pending_paiements])
    money_rejected = sum([p['amount'] for p in rejected_paiements])
    money_potential = money_pending
    
    # Calculer les métriques utilisateur
    total_users = users_collection.count_documents({})
    if paiements_filtrés:
        debut_nouveaux = paiements_filtrés[-1]['timestamp']
    else:
        debut_nouveaux = datetime.datetime.utcnow() - datetime.timedelta(days=1)
    
    new_users = users_collection.count_documents({"created_at": {"$gte": debut_nouveaux}})
    
    # Afficher les métriques principales
    cols = st.columns(3)
    with cols[0]:
        st.metric(label="Paiements Validés", value=format_number(len(validated_paiements)))
    with cols[1]:
        st.metric(label="Paiements en Attente", value=format_number(len(pending_paiements)))
    with cols[2]:
        st.metric(label="Paiements Rejetés", value=format_number(len(rejected_paiements)))
    
    cols = st.columns(3)
    with cols[0]:
        st.metric(label="Argent Gagné (USD)", value=f"${format_number(money_earned)}")
    with cols[1]:
        st.metric(label="Argent Rejeté (USD)", value=f"${format_number(money_rejected)}")
    with cols[2]:
        st.metric(label="Argent à Gagner (USD)", value=f"${format_number(money_potential)}")
    
    cols = st.columns(3)
    with cols[0]:
        st.metric(label="Nombre Total d'Utilisateurs", value=total_users)
    with cols[1]:
        st.metric(label="Nouveaux Utilisateurs", value=new_users)
    with cols[2]:
        st.metric(label="Tokens Vendus", value=format_number(sum([p['tokens_purchased'] for p in validated_paiements])))
    
    st.markdown("---")
    st.header("Gestion des Utilisateurs")
    
    # Pagination pour les utilisateurs
    page_size = 10
    total_users_count = users_collection.count_documents({})
    total_pages = (total_users_count // page_size) + 1
    if 'admin_user_page' not in st.session_state:
        st.session_state['admin_user_page'] = 1
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("◀️ Précédent Utilisateurs") and st.session_state['admin_user_page'] > 1:
            st.session_state['admin_user_page'] -= 1
            st.rerun()
    with col2:
        st.write(f"Page {st.session_state['admin_user_page']} de {total_pages}")
    with col3:
        if st.button("Suivant ▶️ Utilisateurs") and st.session_state['admin_user_page'] < total_pages:
            st.session_state['admin_user_page'] += 1
            st.rerun()
    
    utilisateurs = list(
        users_collection.find()
        .skip((st.session_state['admin_user_page'] - 1) * page_size)
        .limit(page_size)
    )
    
    st.subheader("Liste des Utilisateurs")
    cols = st.columns(2)
    for i, utilisateur_ in enumerate(utilisateurs):
        with cols[i % 2]:
            status = "Bloqué" if utilisateur_.get('is_blocked', False) else "Autorisé"
            block_key = f"block_{utilisateur_['_id']}"
            unblock_key = f"unblock_{utilisateur_['_id']}"
    
            st.markdown(f"""
                <div style="border:1px solid #ddd; padding:10px; border-radius:5px; margin-bottom:10px;">
                    <h3>{utilisateur_['nom']}</h3>
                    <p><strong>Email :</strong> {utilisateur_['email']}</p>
                    <p><strong>Tokens Achetés :</strong> {utilisateur_.get('tokens_purchased',0)}</p>
                    <p><strong>Tokens Consommés :</strong> {utilisateur_.get('tokens_consumed',0)}</p>
                    <p><strong>Status :</strong> {status}</p>
                </div>
            """, unsafe_allow_html=True)
            
            if not utilisateur_.get('is_blocked', False):
                if st.button("Bloquer", key=block_key):
                    users_collection.update_one(
                        {"_id": ObjectId(utilisateur_['_id'])},
                        {"$set": {"is_blocked": True}}
                    )
                    st.success(f"Utilisateur {utilisateur_['nom']} bloqué avec succès.")
                    logging.info(f"Utilisateur bloqué : {utilisateur_['_id']}")
                    st.rerun()
            else:
                if st.button("Débloquer", key=unblock_key):
                    users_collection.update_one(
                        {"_id": ObjectId(utilisateur_['_id'])},
                        {"$set": {"is_blocked": False}}
                    )
                    st.success(f"Utilisateur {utilisateur_['nom']} débloqué avec succès.")
                    logging.info(f"Utilisateur débloqué : {utilisateur_['_id']}")
                    st.rerun()
    
    st.markdown("---")
    st.header("Demandes de Validation de Paiement")
    
    # Filtrer les paiements en attente
    paiement_periode = st.selectbox(
        "Filtrer par :", 
        ["Jour", "2 Jours", "Semaine", "Mois", "3 Mois", "6 Mois", "Année"], 
        key="paiement_periode"
    )
    paiements_attente = list(
        db['payments'].find({"payment_method": "Manual", "status": "Pending"})
        .sort("timestamp", -1)
    )
    paiements_attente_filtrés = filtrer_paiements(paiements_attente, paiement_periode)
    
    # Pagination pour les paiements en attente
    page_size_paiement = 6
    total_paiements_count = len(paiements_attente_filtrés)
    total_pages_paiement = (total_paiements_count // page_size_paiement) + 1
    if 'admin_paiement_page' not in st.session_state:
        st.session_state['admin_paiement_page'] = 1
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("◀️ Précédent Paiements") and st.session_state['admin_paiement_page'] > 1:
            st.session_state['admin_paiement_page'] -= 1
            st.rerun()
    with col2:
        st.write(f"Page {st.session_state['admin_paiement_page']} de {total_pages_paiement}")
    with col3:
        if st.button("Suivant ▶️ Paiements") and st.session_state['admin_paiement_page'] < total_pages_paiement:
            st.session_state['admin_paiement_page'] += 1
            st.rerun()
    
    paiements_page = paiements_attente_filtrés[
        (st.session_state['admin_paiement_page'] - 1) * page_size_paiement : 
        st.session_state['admin_paiement_page'] * page_size_paiement
    ]
    
    if paiements_page:
        st.subheader("Demandes de Validation")
        cols = st.columns(3)
        for i, paiement in enumerate(paiements_page):
            with cols[i % 3]:
                user_ = users_collection.find_one({"_id": paiement['user_id']})
                validate_key = f"validate_{paiement['_id']}"
                reject_key = f"reject_{paiement['_id']}"
                
                st.markdown(f"""
                    <div style="border:1px solid #ddd; padding:10px; border-radius:5px; margin-bottom:10px;">
                        <h3>{user_['nom']}</h3>
                        <p><strong>Email :</strong> {user_['email']}</p>
                        <p><strong>Montant :</strong> {paiement['amount']} USD</p>
                        <p><strong>Tokens :</strong> {paiement['tokens_purchased']}</p>
                        <p><strong>Date :</strong> {paiement['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}</p>
                        <p><strong>Justificatif :</strong> <a href="{paiement['uploaded_file_url']}" target="_blank">Voir le fichier</a></p>
                    </div>
                """, unsafe_allow_html=True)
                
                col_validate, col_reject = st.columns(2)
                with col_validate:
                    if st.button("Valider", key=validate_key):
                        success, message = mettre_a_jour_statut_paiement(paiement['_id'], "Validated")
                        if success:
                            st.success(f"Paiement {paiement['_id']} validé avec succès.")
                            st.rerun()
                        else:
                            st.error("Erreur lors de la validation du paiement.")
                with col_reject:
                    if st.button("Rejeter", key=reject_key):
                        success, message = mettre_a_jour_statut_paiement(paiement['_id'], "Rejected")
                        if success:
                            st.warning(f"Paiement {paiement['_id']} rejeté avec succès.")
                            st.rerun()
                        else:
                            st.error("Erreur lors du rejet du paiement.")
    else:
        st.write("Aucune demande de validation de paiement pour cette période.")
    
    # ---------------------- Nouvelle Section : Paiements Déjà Validés ----------------------
    st.markdown("---")
    st.header("Paiements Déjà Validés")
    
    # Filtrer les paiements validés par période
    periode_valides = st.selectbox(
        "Filtrer par :", 
        ["Jour", "2 Jours", "Semaine", "Mois", "3 Mois", "6 Mois", "Année"], 
        key="periode_valides"
    )
    paiements_valides = list(
        db['payments'].find({"status": "Validated"})
        .sort("timestamp", -1)
    )
    paiements_valides_filtrés = filtrer_paiements(paiements_valides, periode_valides)
    
    # Pagination pour les paiements validés
    page_size_valides = 10
    total_paiements_valides = len(paiements_valides_filtrés)
    total_pages_valides = (total_paiements_valides // page_size_valides) + 1
    if 'admin_validated_paiement_page' not in st.session_state:
        st.session_state['admin_validated_paiement_page'] = 1
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("◀️ Précédent Paiements", key="prev_validated_paiements") and st.session_state['admin_validated_paiement_page'] > 1:
            st.session_state['admin_validated_paiement_page'] -= 1
            st.rerun()
    with col2:
        st.write(f"Page {st.session_state['admin_validated_paiement_page']} de {total_pages_valides}")
    with col3:
        if st.button("Suivant ▶️ Paiements", key="next_validated_paiements") and st.session_state['admin_validated_paiement_page'] < total_pages_valides:
            st.session_state['admin_validated_paiement_page'] += 1
            st.rerun()

    
    paiements_valides_page = paiements_valides_filtrés[
        (st.session_state['admin_validated_paiement_page'] - 1) * page_size_valides : 
        st.session_state['admin_validated_paiement_page'] * page_size_valides
    ]
    
    if paiements_valides_page:
        st.subheader("Liste des Paiements Validés")
        cols_valides = st.columns(3)
        for i, paiement in enumerate(paiements_valides_page):
            with cols_valides[i % 3]:
                user_ = users_collection.find_one({"_id": paiement['user_id']})
                st.markdown(f"""
                    <div style="border:1px solid #ddd; padding:10px; border-radius:5px; margin-bottom:10px;">
                        <h3>{user_['nom']}</h3>
                        <p><strong>Email :</strong> {user_['email']}</p>
                        <p><strong>Montant s:</strong> {paiement['amount']} USD</p>
                        <p><strong>Tokens Achetés :</strong> {paiement['tokens_purchased']}</p>
                        <p><strong>Date :</strong> {paiement['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}</p>
                        <p><strong>Méthode :</strong> {paiement['payment_method']}</p>
                        <p><strong>Justificatif :</strong> <a href="{paiement['uploaded_file_url']}" target="_blank">Voir le fichier</a></p>
                    </div>
                """, unsafe_allow_html=True)
    else:
        st.write("Aucun paiement validé pour cette période.")


# ---------------------- FONCTION PRINCIPALE ----------------------
def main():
    if 'page' not in st.session_state:
        st.session_state['page'] = 'Connexion'

    page_courante = st.session_state['page']

    # On route vers la page correspondante
    if page_courante == "Connexion":
        page_connexion()
    elif page_courante == "Inscription":
        page_inscription()
    elif page_courante == "Mot de passe oublié":
        page_reinitialisation()
    elif page_courante == "OTP Verification":
        page_otp_verification()
    elif page_courante == "Reset OTP Verification":
        page_reset_otp_verification()
    elif page_courante == "Google Login":
        page_google_login()
    elif page_courante == "Accueil":
        page_accueil()
    elif page_courante == "page_admin":
        page_admin()
    elif page_courante == "Chat":
        page_chat()
    elif page_courante == "Achat":
        page_achat()
    elif page_courante == "Profil":
        page_profil()
    elif page_courante == "Deconnexion":
        page_deconnexion()
    else:
        page_connexion()

if __name__ == "__main__":
    initialize_session()
    main()
